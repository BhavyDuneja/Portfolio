"""
build_04_msa.py

Generates the AnantaSutra Master Service Agreement (MSA) as a
professionally formatted .docx at:
  c:/Users/duneja8515/Desktop/bhavya/prsnl/sutranet/legal-documents/04_Master_Service_Agreement.docx

Drafted against Indian law anchors: Indian Contract Act 1872, Sale of Goods
Act 1930, IT Act 2000, DPDP Act 2023, CGST Act 2017, Income-tax Act 1961,
Arbitration & Conciliation Act 1996 (as amended), Specific Relief Act 1963,
Indian Stamp Act 1899 (as applicable to NCT of Delhi), Companies Act 2013.

Uses python-docx v1.2.0.
"""

from __future__ import annotations

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Output target
# ---------------------------------------------------------------------------
OUTPUT_PATH = (
    r"c:/Users/duneja8515/Desktop/bhavya/prsnl/sutranet/legal-documents/"
    r"04_Master_Service_Agreement.docx"
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(11)



# ---------------------------------------------------------------------------
# Brand header (AnantaSutra) — injected by rebrand
# ---------------------------------------------------------------------------
LOGO_PATH = r"c:/Users/duneja8515/Desktop/bhavya/prsnl/sutranet/public/Favicon/logo-nobg.png"


def add_brand_header(doc):
    """Insert centered logo + brand lines at top of page 1. Never crashes."""
    try:
        from docx.shared import Inches as _Inches, Pt as _Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH as _CENTERED
    except Exception:
        return
    p = doc.add_paragraph()
    p.alignment = _CENTERED.CENTER
    try:
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=_Inches(1.2))
    except Exception:
        pass
    p2 = doc.add_paragraph()
    p2.alignment = _CENTERED.CENTER
    r2 = p2.add_run("AnantaSutra — अनन्तसूत्र")
    r2.italic = True
    r2.font.size = _Pt(10)
    p3 = doc.add_paragraph()
    p3.alignment = _CENTERED.CENTER
    r3 = p3.add_run("https://anantasutra.com  •  contact@anantasutra.com")
    r3.italic = True
    r3.font.size = _Pt(9)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:color"), "000000")
            tcBorders.append(border)
    tcPr.append(tcBorders)


def apply_body_format(paragraph, justify=True, first_line_indent=False):
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    if justify:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_line_indent:
        pf.first_line_indent = Cm(0.5)


def add_body_paragraph(doc, text, bold=False, italic=False, justify=True):
    p = doc.add_paragraph()
    apply_body_format(p, justify=justify)
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE
    run.bold = bold
    run.italic = italic
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    run = p.add_run(text)
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    return p


def add_clause(doc, number, heading, body_text=None):
    """Adds a numbered clause with heading; optional body text."""
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    run = p.add_run(f"{number}. {heading}")
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    if body_text:
        add_body_paragraph(doc, body_text)


def add_sub(doc, number, text, bold_lead=None):
    """Adds a sub-clause like '1.1 Something.' with justified body."""
    p = doc.add_paragraph()
    apply_body_format(p)
    p.paragraph_format.left_indent = Cm(0.6)
    lead_run = p.add_run(f"{number}  ")
    lead_run.font.name = BODY_FONT
    lead_run.font.size = BODY_SIZE
    lead_run.bold = True
    if bold_lead:
        lr = p.add_run(f"{bold_lead} ")
        lr.font.name = BODY_FONT
        lr.font.size = BODY_SIZE
        lr.bold = True
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    apply_body_format(p)
    p.paragraph_format.left_indent = Cm(1.0 + 0.5 * level)
    # Remove bullet style run formatting and set our font
    for r in p.runs:
        r.text = ""
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(6)  # WD_BREAK.PAGE == 6 via integer fallback
    # Use the proper enumeration safely
    from docx.enum.text import WD_BREAK
    # Some versions require explicit:
    # Already added a break above; remove and re-add via WD_BREAK.PAGE if needed
    # We'll just rely on docx.add_page_break:
    doc.add_page_break()


def set_page_and_margins(doc):
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def set_default_style(doc):
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Confidential — AnantaSutra   |   Page ")
        run.font.name = BODY_FONT
        run.font.size = Pt(9)

        # PAGE field
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_begin)

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        run._r.append(instr)

        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        run._r.append(fld_sep)

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_end)

        run2 = p.add_run(" of ")
        run2.font.name = BODY_FONT
        run2.font.size = Pt(9)

        # NUMPAGES field
        fld_begin2 = OxmlElement("w:fldChar")
        fld_begin2.set(qn("w:fldCharType"), "begin")
        run2._r.append(fld_begin2)
        instr2 = OxmlElement("w:instrText")
        instr2.set(qn("xml:space"), "preserve")
        instr2.text = "NUMPAGES"
        run2._r.append(instr2)
        fld_sep2 = OxmlElement("w:fldChar")
        fld_sep2.set(qn("w:fldCharType"), "separate")
        run2._r.append(fld_sep2)
        fld_end2 = OxmlElement("w:fldChar")
        fld_end2.set(qn("w:fldCharType"), "end")
        run2._r.append(fld_end2)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(16)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    return p


def add_spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)


# ---------------------------------------------------------------------------
# Build the document
# ---------------------------------------------------------------------------
def build():
    doc = Document()
    add_brand_header(doc)
    set_page_and_margins(doc)
    set_default_style(doc)
    add_footer(doc)

    # -----------------------------------------------------------------
    # Title
    # -----------------------------------------------------------------
    add_title(doc, "MASTER SERVICE AGREEMENT")

    # Intro line
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = intro.add_run(
        "This Master Service Agreement is made and executed at New Delhi, "
        "Delhi, India, on this [DAY] day of [MONTH], [YEAR] (the "
        "\"Effective Date\")."
    )
    r.font.name = BODY_FONT
    r.font.size = BODY_SIZE
    r.italic = True
    intro.paragraph_format.space_after = Pt(10)

    # -----------------------------------------------------------------
    # Parties
    # -----------------------------------------------------------------
    add_h1(doc, "BETWEEN")
    add_body_paragraph(
        doc,
        "ANANTASUTRA, a business concern operated and represented by its "
        "Founder & CEO, Mr. Himanshu Mishra, carrying on business under the "
        "trade name \"AnantaSutra\", having its principal place of business "
        "at Delhi, India, with contact address contact@anantasutra.com "
        "(hereinafter referred to as \"AnantaSutra\" or the \"Service "
        "Provider\", which expression shall, unless repugnant to the "
        "context or meaning thereof, include its successors, assigns, and "
        "the person(s) for the time being in control of the business), "
        "acting through its Founder & CEO, Mr. Himanshu Mishra, "
        "of the ONE PART;"
    )
    add_h1(doc, "AND")
    add_body_paragraph(
        doc,
        "[CLIENT LEGAL NAME], a company incorporated under the Companies "
        "Act, 2013 / a [PARTNERSHIP FIRM/LLP/SOLE PROPRIETORSHIP/FOREIGN "
        "ENTITY], bearing [CIN/LLPIN/REGISTRATION NO.], PAN [PAN], GSTIN "
        "[GSTIN], having its registered office at [CLIENT REGISTERED "
        "OFFICE ADDRESS] (hereinafter referred to as the \"Client\", which "
        "expression shall, unless repugnant to the context or meaning "
        "thereof, include its successors and permitted assigns) of the "
        "OTHER PART."
    )
    add_body_paragraph(
        doc,
        "AnantaSutra and the Client are hereinafter individually referred to "
        "as a \"Party\" and collectively as the \"Parties\".",
    )

    # -----------------------------------------------------------------
    # Recitals
    # -----------------------------------------------------------------
    add_h1(doc, "RECITALS")
    add_body_paragraph(
        doc,
        "A. AnantaSutra is, inter alia, engaged in the business of providing "
        "information technology, software development, technology "
        "consulting, data engineering, and allied professional services.",
    )
    add_body_paragraph(
        doc,
        "B. The Client is desirous of engaging AnantaSutra, from time to "
        "time, to perform certain services as more particularly described "
        "in one or more Statements of Work (as defined hereinafter) to be "
        "executed under this Agreement.",
    )
    add_body_paragraph(
        doc,
        "C. The Parties intend to set out in this Agreement the master "
        "commercial, legal, and operational framework that shall govern "
        "each such Statement of Work, without prejudice to project-"
        "specific commercial terms contained therein.",
    )
    add_body_paragraph(
        doc,
        "D. The Parties have negotiated the terms of this Agreement at "
        "arm's length and are satisfied that it represents the lawful, "
        "bargained-for exchange between them under the Indian Contract "
        "Act, 1872.",
    )
    add_body_paragraph(
        doc,
        "NOW, THEREFORE, in consideration of the mutual covenants, "
        "representations, warranties, and agreements hereinafter set "
        "forth, and for other good and valuable consideration, the receipt "
        "and sufficiency of which is hereby acknowledged, the Parties "
        "agree as follows:",
        bold=True,
    )

    # -----------------------------------------------------------------
    # 1. Definitions
    # -----------------------------------------------------------------
    add_clause(
        doc,
        1,
        "DEFINITIONS AND INTERPRETATION",
        "In this Agreement, unless the context otherwise requires, the "
        "following capitalised terms shall have the meanings ascribed to "
        "them below. Terms defined in the singular shall include the "
        "plural and vice versa. References to statutes include any "
        "amendment, re-enactment, or subordinate legislation made under "
        "them.",
    )

    defs = [
        (
            "Affiliate",
            "means, in relation to a Party, any entity that directly or "
            "indirectly controls, is controlled by, or is under common "
            "control with such Party, where \"control\" means beneficial "
            "ownership of fifty per cent (50%) or more of the voting "
            "equity or the power to direct the management and policies of "
            "such entity.",
        ),
        (
            "Applicable Laws",
            "means all statutes, rules, regulations, notifications, "
            "circulars, orders, directions, and judgments of any "
            "governmental, statutory, judicial, quasi-judicial, or "
            "regulatory authority of the Republic of India (and, where "
            "relevant to a Statement of Work, any other jurisdiction) "
            "which are in force and are applicable to the performance of "
            "the obligations of the Parties under this Agreement.",
        ),
        (
            "Background IP",
            "means all Intellectual Property that is (a) owned, licensed, "
            "or controlled by a Party prior to the Effective Date; or (b) "
            "developed or acquired by a Party independently of, and "
            "outside the scope of, this Agreement; together with any "
            "modifications, enhancements, or derivatives thereof made "
            "independently of this Agreement.",
        ),
        (
            "Business Day",
            "means a day other than a Saturday, Sunday, or a gazetted "
            "public holiday in New Delhi, Delhi, India, on which "
            "commercial banks are ordinarily open for business.",
        ),
        (
            "Confidential Information",
            "has the meaning set out in Clause 10 and shall be construed "
            "consistently with the definition thereof in any separately "
            "executed non-disclosure agreement between the Parties.",
        ),
        (
            "Deliverables",
            "means any report, software, source code, object code, "
            "documentation, design, specification, data, analysis, "
            "material, or other work product that AnantaSutra is required to "
            "deliver to the Client under a Statement of Work.",
        ),
        (
            "DPDP Act",
            "means the Digital Personal Data Protection Act, 2023 and the "
            "rules, regulations, and directions made thereunder, as "
            "amended from time to time.",
        ),
        (
            "Effective Date",
            "means the later of (a) the date of last signature to this "
            "Agreement, or (b) the date on which the stamp duty payable "
            "on this Agreement, if any, is paid.",
        ),
        (
            "Fees",
            "means the fees, charges, and expenses payable by the Client "
            "to AnantaSutra under a Statement of Work, exclusive of "
            "applicable Taxes.",
        ),
        (
            "Foreground IP",
            "means all Intellectual Property first created, conceived, "
            "developed, or reduced to practice by or on behalf of "
            "AnantaSutra in the course of, or as a result of, performance of "
            "the Services under a Statement of Work, other than "
            "Background IP.",
        ),
        (
            "Intellectual Property\" or \"IP",
            "means all intellectual property rights, whether registered "
            "or unregistered and wherever subsisting, including copyright "
            "and related rights (under the Copyright Act, 1957), patents "
            "and patent applications (under the Patents Act, 1970), "
            "trademarks, service marks, trade names, logos, and get-up "
            "(under the Trade Marks Act, 1999), designs (under the "
            "Designs Act, 2000), semiconductor layout designs, trade "
            "secrets, know-how, database rights, moral rights (subject "
            "to Section 57 of the Copyright Act, 1957), and all similar "
            "or equivalent rights anywhere in the world.",
        ),
        (
            "Personal Data",
            "has the meaning given to the term \"personal data\" under "
            "Section 2(t) of the DPDP Act.",
        ),
        (
            "Data Fiduciary\" and \"Data Processor",
            "have the meanings respectively given to \"Data Fiduciary\" "
            "under Section 2(i) and \"Data Processor\" under Section 2(k) "
            "of the DPDP Act.",
        ),
        (
            "Services",
            "means the professional services to be performed by AnantaSutra "
            "as described in each Statement of Work executed hereunder.",
        ),
        (
            "Statement of Work\" or \"SOW",
            "means a written document executed by the authorised "
            "representatives of both Parties, substantially in the form "
            "of Schedule 1, which describes a specific engagement under "
            "this Agreement, including the scope, Deliverables, "
            "timelines, resources, and Fees.",
        ),
        (
            "Taxes",
            "means all indirect and withholding taxes imposed under "
            "Applicable Laws, including goods and services tax under the "
            "CGST Act, 2017, IGST Act, 2017, and applicable State GST "
            "Acts; customs duty; and tax deducted at source under the "
            "Income-tax Act, 1961.",
        ),
        (
            "Term",
            "has the meaning given in Clause 18.",
        ),
    ]

    for term, meaning in defs:
        p = doc.add_paragraph()
        apply_body_format(p)
        p.paragraph_format.left_indent = Cm(0.6)
        r1 = p.add_run(f"\"{term}\" ")
        r1.bold = True
        r1.font.name = BODY_FONT
        r1.font.size = BODY_SIZE
        r2 = p.add_run(meaning)
        r2.font.name = BODY_FONT
        r2.font.size = BODY_SIZE

    add_sub(
        doc,
        "1.2",
        "References to clauses, sub-clauses, and Schedules are to "
        "clauses, sub-clauses, and Schedules of this Agreement. The "
        "Schedules form an integral and operative part of this Agreement.",
    )
    add_sub(
        doc,
        "1.3",
        "Headings are for convenience only and shall not affect the "
        "interpretation of this Agreement.",
    )
    add_sub(
        doc,
        "1.4",
        "The expression \"including\" shall be construed as \"including "
        "without limitation\".",
    )

    # -----------------------------------------------------------------
    # 2. Structure of Engagement
    # -----------------------------------------------------------------
    add_clause(doc, 2, "STRUCTURE OF ENGAGEMENT")
    add_sub(
        doc,
        "2.1",
        "This Agreement establishes the master legal and commercial "
        "framework pursuant to which AnantaSutra shall provide Services to "
        "the Client. No Services shall be deemed to be rendered under "
        "this Agreement unless they are covered by a duly executed "
        "Statement of Work.",
    )
    add_sub(
        doc,
        "2.2",
        "Each Statement of Work shall be in substantially the form set "
        "out in Schedule 1 and shall, upon execution by the authorised "
        "representatives of both Parties, be deemed to incorporate, and "
        "be governed by, the terms of this Agreement.",
    )
    add_sub(
        doc,
        "2.3  Order of Precedence.",
        "In the event of any conflict or inconsistency between this "
        "Agreement and a Statement of Work: (a) on matters of commercial "
        "scope (namely scope of Services, Deliverables, timelines, Fees, "
        "milestones, personnel, and acceptance criteria), the Statement "
        "of Work shall prevail; and (b) on matters of limitation of "
        "liability, indemnity, intellectual property, confidentiality, "
        "data protection, insurance, and dispute resolution, this "
        "Agreement shall prevail, unless the Statement of Work expressly "
        "references the clause of this Agreement being varied and both "
        "Parties expressly agree in writing to such variation.",
    )
    add_sub(
        doc,
        "2.4",
        "No purchase order, acknowledgement, invoice, or similar "
        "document issued by either Party shall have the effect of "
        "modifying this Agreement, and any pre-printed terms on such "
        "documents are hereby expressly rejected.",
    )

    # -----------------------------------------------------------------
    # 3. Services and SOWs
    # -----------------------------------------------------------------
    add_clause(doc, 3, "SERVICES AND STATEMENTS OF WORK")
    add_sub(
        doc,
        "3.1",
        "AnantaSutra shall perform the Services described in each Statement "
        "of Work with the standard of care, skill, and diligence "
        "reasonably expected of a competent professional services firm "
        "providing services of a similar nature in the Republic of India.",
    )
    add_sub(
        doc,
        "3.2",
        "AnantaSutra shall perform the Services through personnel who are "
        "appropriately qualified and experienced. AnantaSutra remains the "
        "sole employer of its personnel, and nothing in this Agreement "
        "or any Statement of Work shall be construed to create any "
        "employer-employee relationship between the Client and any "
        "person deployed by AnantaSutra. All statutory contributions under "
        "the Employees' Provident Funds and Miscellaneous Provisions "
        "Act, 1952, the Employees' State Insurance Act, 1948, the "
        "Payment of Gratuity Act, 1972, and the Code on Social Security, "
        "2020 (as and when notified), in respect of such personnel shall "
        "be the exclusive responsibility of AnantaSutra.",
    )
    add_sub(
        doc,
        "3.3  Acceptance.",
        "Unless otherwise specified in a Statement of Work, the Client "
        "shall review each Deliverable within fifteen (15) Business Days "
        "of delivery (the \"Acceptance Period\") and shall notify "
        "AnantaSutra in writing of acceptance or of any material "
        "non-conformity with the applicable specifications. If the "
        "Client fails to notify AnantaSutra within the Acceptance Period, "
        "the Deliverable shall be deemed accepted. AnantaSutra shall, "
        "within a reasonable period, correct any notified material "
        "non-conformity and resubmit the Deliverable for re-review.",
    )
    add_sub(
        doc,
        "3.4  Warranty.",
        "AnantaSutra warrants that the Deliverables shall, for a period of "
        "ninety (90) days from acceptance (or such longer period as is "
        "set out in the applicable Statement of Work), materially "
        "conform to the specifications agreed in the Statement of Work. "
        "AnantaSutra's sole obligation, and the Client's sole and "
        "exclusive remedy, for breach of this warranty shall be the "
        "re-performance of the relevant Services or correction of the "
        "Deliverable, at AnantaSutra's option and cost.",
    )

    # -----------------------------------------------------------------
    # 4. Change Control
    # -----------------------------------------------------------------
    add_clause(doc, 4, "CHANGE CONTROL")
    add_sub(
        doc,
        "4.1",
        "Either Party may at any time propose a change to the scope, "
        "Deliverables, timelines, resources, or Fees under a Statement "
        "of Work by issuing a written change request to the other Party.",
    )
    add_sub(
        doc,
        "4.2",
        "AnantaSutra shall respond with a written change order setting out "
        "the impact of the proposed change on scope, timelines, Fees, "
        "and any other relevant matter. No change shall take effect "
        "unless and until the change order is signed by the authorised "
        "representatives of both Parties.",
    )
    add_sub(
        doc,
        "4.3",
        "Pending agreement on a change order, AnantaSutra shall continue "
        "to perform in accordance with the then-current Statement of "
        "Work.",
    )

    # -----------------------------------------------------------------
    # 5. Fees, Invoicing and Payment
    # -----------------------------------------------------------------
    add_clause(doc, 5, "FEES, INVOICING AND PAYMENT")
    add_sub(
        doc,
        "5.1  Fees.",
        "In consideration of the Services, the Client shall pay the Fees "
        "specified in the applicable Statement of Work. Unless otherwise "
        "stated, Fees are exclusive of Taxes and reasonable pre-approved "
        "out-of-pocket expenses (including travel, boarding, and lodging "
        "at actuals supported by receipts), which shall be reimbursed at "
        "cost.",
    )
    add_sub(
        doc,
        "5.2  Currency.",
        "Fees shall be invoiced in Indian Rupees (INR) for domestic "
        "engagements and in United States Dollars (USD) or such other "
        "freely convertible foreign currency as may be agreed for "
        "cross-border engagements, subject always to the Foreign "
        "Exchange Management Act, 1999 and the regulations made "
        "thereunder.",
    )
    add_sub(
        doc,
        "5.3  Invoicing.",
        "AnantaSutra shall issue a tax invoice in accordance with Rule 46 "
        "of the Central Goods and Services Tax Rules, 2017, together "
        "with such supporting documentation as is reasonably required. "
        "Unless otherwise specified in the Statement of Work, invoices "
        "shall be raised monthly for time-and-materials engagements and "
        "upon milestone achievement for fixed-price engagements.",
    )
    add_sub(
        doc,
        "5.4  Payment Terms.",
        "Undisputed invoices shall be paid by the Client within thirty "
        "(30) days (or such longer period, not exceeding forty-five "
        "(45) days, as may be specified in the Statement of Work) of "
        "the date of invoice, by electronic funds transfer to the bank "
        "account designated by AnantaSutra in writing.",
    )
    add_sub(
        doc,
        "5.5  Late Payment.",
        "Without prejudice to any other right or remedy, any undisputed "
        "amount not paid when due shall bear simple interest at the "
        "rate of one and a half per cent (1.5%) per month, or the "
        "highest rate permissible under Applicable Laws, whichever is "
        "lower, from the due date until the date of actual payment.",
    )
    add_sub(
        doc,
        "5.6  GST.",
        "All Fees are exclusive of goods and services tax chargeable "
        "under the CGST Act, 2017, the IGST Act, 2017, and the "
        "applicable State GST legislation, which shall be levied by "
        "AnantaSutra at the prevailing rate and paid by the Client in "
        "addition to the Fees.",
    )
    add_sub(
        doc,
        "5.7  TDS.",
        "The Client shall be entitled to withhold tax at source at the "
        "rates prescribed under the Income-tax Act, 1961 (including, as "
        "applicable, Sections 194C, 194J, 194Q, or 195). The Client "
        "shall remit the tax so withheld to the credit of the Central "
        "Government and shall furnish to AnantaSutra the applicable TDS "
        "certificate in Form 16A (or, in respect of non-resident "
        "payments, Form 16B) within the statutory period. There shall "
        "be no gross-up unless expressly agreed in the Statement of "
        "Work.",
    )
    add_sub(
        doc,
        "5.8  Disputed Invoices.",
        "If the Client in good faith disputes any portion of an invoice, "
        "it shall, within ten (10) Business Days of receipt, notify "
        "AnantaSutra in writing of the disputed portion, detailing the "
        "reasons. The undisputed portion shall be paid in accordance "
        "with Clause 5.4. The Parties shall endeavour to resolve the "
        "dispute within fifteen (15) Business Days of such notice, "
        "failing which Clause 20 shall apply.",
    )

    # -----------------------------------------------------------------
    # 6. Personnel and Non-Solicit
    # -----------------------------------------------------------------
    add_clause(doc, 6, "PERSONNEL")
    add_sub(
        doc,
        "6.1",
        "AnantaSutra shall be responsible for the selection, engagement, "
        "training, supervision, and discharge of its personnel. The "
        "Client shall have no authority to direct, discipline, or "
        "terminate any AnantaSutra personnel.",
    )
    add_sub(
        doc,
        "6.2",
        "AnantaSutra shall, at the Client's reasonable request and cost, "
        "replace any personnel whose performance is materially "
        "unsatisfactory, subject to reasonable transition and "
        "knowledge-transfer arrangements.",
    )
    add_sub(
        doc,
        "6.3  Background Checks.",
        "AnantaSutra shall conduct reasonable pre-employment background "
        "verification on its personnel performing the Services, "
        "consistent with Applicable Laws and AnantaSutra's standard "
        "policies.",
    )

    # -----------------------------------------------------------------
    # 7. Subcontracting
    # -----------------------------------------------------------------
    add_clause(doc, 7, "SUBCONTRACTING")
    add_sub(
        doc,
        "7.1",
        "AnantaSutra shall not subcontract the performance of any material "
        "part of the Services without the prior written consent of the "
        "Client, such consent not to be unreasonably withheld or "
        "delayed.",
    )
    add_sub(
        doc,
        "7.2",
        "Where subcontracting is permitted, AnantaSutra shall (a) flow "
        "down written obligations on confidentiality, intellectual "
        "property, and data protection that are no less onerous than "
        "those set out in this Agreement; and (b) remain primarily "
        "liable to the Client for the acts and omissions of its "
        "subcontractors as if they were its own.",
    )

    # -----------------------------------------------------------------
    # 8. Intellectual Property
    # -----------------------------------------------------------------
    add_clause(doc, 8, "INTELLECTUAL PROPERTY")
    add_sub(
        doc,
        "8.1  Background IP.",
        "Each Party shall retain all right, title, and interest in and "
        "to its Background IP. Save as expressly granted herein or in a "
        "Statement of Work, nothing in this Agreement shall operate to "
        "transfer or license any Background IP of a Party to the other "
        "Party.",
    )
    add_sub(
        doc,
        "8.2  Foreground IP.",
        "Subject to full and final payment of the Fees payable in "
        "respect of the applicable Statement of Work, and subject to "
        "Clause 8.3, AnantaSutra hereby assigns to the Client, with effect "
        "from the date of such payment, all right, title, and interest "
        "(including copyright under Section 18 of the Copyright Act, "
        "1957) in and to the Foreground IP comprised in the "
        "Deliverables specifically created for the Client under the "
        "applicable Statement of Work, for the full term of protection "
        "of such IP and for the whole world. AnantaSutra shall, at the "
        "Client's cost, execute such further deeds and documents as "
        "may be reasonably necessary to perfect such assignment in any "
        "jurisdiction, including any short-form assignments required "
        "for recordal with the Registrar of Copyrights, Controller of "
        "Patents, Registrar of Trade Marks, or Controller of Designs.",
    )
    add_sub(
        doc,
        "8.3  Carve-back; AnantaSutra Retained IP.",
        "Notwithstanding Clause 8.2, AnantaSutra shall retain all right, "
        "title, and interest in and to (a) its Background IP; (b) any "
        "generic tools, libraries, utilities, frameworks, algorithms, "
        "methodologies, know-how, and pre-existing components used in "
        "the creation of the Deliverables; and (c) any residual "
        "knowledge retained unaided in the memory of its personnel. "
        "AnantaSutra hereby grants the Client a perpetual, irrevocable "
        "(save for termination for the Client's material breach or "
        "insolvency), worldwide, non-exclusive, royalty-free licence to "
        "use, copy, modify, and create derivative works of the "
        "AnantaSutra Retained IP solely to the extent embedded in the "
        "Deliverables and solely for the Client's internal business "
        "purposes.",
    )
    add_sub(
        doc,
        "8.4  Open Source.",
        "AnantaSutra shall, prior to incorporating any open-source software "
        "into a Deliverable, disclose such intended use to the Client, "
        "identify the applicable open-source licence, and, where the "
        "licence is of a \"copyleft\" nature (including the GNU General "
        "Public License, Affero GPL, or Lesser GPL), obtain the "
        "Client's prior written consent.",
    )
    add_sub(
        doc,
        "8.5  Moral Rights.",
        "To the extent permissible under Section 57 of the Copyright "
        "Act, 1957, AnantaSutra waives, and shall procure that its "
        "personnel waive, all moral rights in the Foreground IP, save "
        "to the extent relating to false attribution or distortion "
        "prejudicial to the author's honour or reputation.",
    )
    add_sub(
        doc,
        "8.6  Feedback.",
        "If the Client provides AnantaSutra with any suggestions, "
        "comments, or feedback in relation to the Services or "
        "Deliverables (\"Feedback\"), the Client grants AnantaSutra a "
        "perpetual, irrevocable, worldwide, royalty-free, "
        "sub-licensable licence to use and exploit such Feedback for "
        "any purpose without obligation to the Client.",
    )

    # -----------------------------------------------------------------
    # 9. Confidentiality
    # -----------------------------------------------------------------
    add_clause(doc, 9, "CONFIDENTIALITY")
    add_sub(
        doc,
        "9.1",
        "Each Party (the \"Recipient\") acknowledges that in the course "
        "of performance of this Agreement, it may receive Confidential "
        "Information from the other Party (the \"Discloser\"). "
        "\"Confidential Information\" means all non-public information "
        "disclosed by or on behalf of the Discloser, whether orally, in "
        "writing, electronically, or visually, and whether or not "
        "marked as confidential, which a reasonable person would "
        "consider confidential in the circumstances, including "
        "technical, commercial, financial, customer, employee, "
        "source-code, algorithmic, and product information, together "
        "with all derivatives, notes, analyses, and extracts.",
    )
    add_sub(
        doc,
        "9.2",
        "The Recipient shall (a) use Confidential Information solely "
        "for the performance of this Agreement; (b) protect it with the "
        "same degree of care as it uses to protect its own confidential "
        "information of like importance, but in no event less than a "
        "reasonable standard of care; and (c) not disclose it to any "
        "third party save to its personnel, advisors, and permitted "
        "subcontractors on a strict need-to-know basis, each of whom is "
        "bound by written confidentiality obligations no less onerous "
        "than those contained herein.",
    )
    add_sub(
        doc,
        "9.3  Exclusions.",
        "The obligations in this Clause 9 shall not apply to information "
        "which (a) is or becomes publicly available otherwise than "
        "through breach of this Agreement; (b) was lawfully in the "
        "Recipient's possession prior to disclosure, free from any "
        "obligation of confidence; (c) is lawfully received from a "
        "third party free of any obligation of confidence; or (d) is "
        "independently developed by the Recipient without reference to "
        "the Discloser's Confidential Information.",
    )
    add_sub(
        doc,
        "9.4  Compelled Disclosure.",
        "If the Recipient is required by Applicable Laws or by a "
        "competent court or regulator to disclose Confidential "
        "Information, it shall, to the extent permitted by law, "
        "promptly notify the Discloser so as to enable the Discloser to "
        "seek a protective order and shall disclose only such portion "
        "of the Confidential Information as is legally required.",
    )
    add_sub(
        doc,
        "9.5  Survival.",
        "The obligations in this Clause 9 shall survive termination of "
        "this Agreement for a period of five (5) years, save in respect "
        "of Confidential Information constituting a trade secret, for "
        "which the obligations shall subsist for so long as the "
        "information remains a trade secret.",
    )
    add_sub(
        doc,
        "9.6",
        "Where the Parties have executed a separate non-disclosure "
        "agreement, the provisions of this Clause 9 shall be read "
        "consistently with that non-disclosure agreement; in the event "
        "of conflict, the provisions more protective of the Discloser "
        "shall prevail.",
    )

    # -----------------------------------------------------------------
    # 10. Data Protection (DPDP Act)
    # -----------------------------------------------------------------
    add_clause(doc, 10, "DATA PROTECTION AND PRIVACY")
    add_sub(
        doc,
        "10.1  Roles.",
        "Where, in the course of performance of the Services, AnantaSutra "
        "processes Personal Data on behalf of the Client, the Client "
        "shall be the Data Fiduciary and AnantaSutra shall be the Data "
        "Processor within the meaning of Section 2 of the DPDP Act. "
        "Where AnantaSutra independently determines the purpose and means "
        "of processing Personal Data in the performance of the "
        "Services, it shall act as an independent Data Fiduciary in "
        "respect of such Personal Data, subject to its separate "
        "compliance obligations under the DPDP Act.",
    )
    add_sub(
        doc,
        "10.2  Processor Obligations.",
        "In its capacity as Data Processor, AnantaSutra shall: (a) process "
        "Personal Data only on the documented instructions of the "
        "Client and as required for the performance of the Services; "
        "(b) implement reasonable security safeguards consistent with "
        "Section 8(5) of the DPDP Act, Rule 8 of the Information "
        "Technology (Reasonable Security Practices and Procedures and "
        "Sensitive Personal Data or Information) Rules, 2011, and, "
        "where applicable, the standards set out in IS/ISO/IEC 27001; "
        "(c) ensure that personnel authorised to process Personal Data "
        "are bound by written confidentiality obligations; (d) assist "
        "the Client, to the extent reasonably possible and at the "
        "Client's cost, in responding to requests from Data Principals "
        "exercising their rights under Sections 11 to 14 of the DPDP "
        "Act; and (e) maintain records of processing activities to the "
        "extent required by Applicable Laws.",
    )
    add_sub(
        doc,
        "10.3  Sub-processors.",
        "AnantaSutra shall not engage any sub-processor to process "
        "Personal Data without the prior written consent of the "
        "Client, which shall not be unreasonably withheld. AnantaSutra "
        "shall impose on each approved sub-processor, by written "
        "contract, obligations no less onerous than those set out in "
        "this Clause 10, and shall remain liable to the Client for the "
        "acts and omissions of its sub-processors.",
    )
    add_sub(
        doc,
        "10.4  Breach Notification.",
        "In the event of any personal data breach (as defined in the "
        "DPDP Act) affecting Personal Data processed under this "
        "Agreement, AnantaSutra shall notify the Client without undue "
        "delay and in any event within forty-eight (48) hours of "
        "becoming aware of the breach, providing such information as "
        "is reasonably required to enable the Client to discharge its "
        "obligations under Section 8(6) of the DPDP Act, including "
        "notification to the Data Protection Board of India and "
        "affected Data Principals within statutory timelines.",
    )
    add_sub(
        doc,
        "10.5  Cross-Border Transfer.",
        "AnantaSutra shall not transfer Personal Data outside the "
        "territory of India save in accordance with Section 16 of the "
        "DPDP Act and any notifications issued by the Central "
        "Government restricting transfer to specified countries or "
        "territories, and subject to the Client's prior written "
        "consent and the execution of such additional safeguards "
        "(including model contractual clauses equivalent to Module 2 "
        "of the EU Standard Contractual Clauses, where required by the "
        "Client) as the Client may reasonably require.",
    )
    add_sub(
        doc,
        "10.6  Audit.",
        "Once in every twelve (12) month period, and more often if "
        "required by a competent regulator or following a personal "
        "data breach, the Client (or an independent auditor appointed "
        "by it and bound by confidentiality) may, on not less than "
        "thirty (30) days' prior written notice, audit AnantaSutra's "
        "compliance with this Clause 10. Such audits shall be conducted "
        "during business hours, shall not unreasonably disrupt "
        "AnantaSutra's operations, and shall not extend to the data or "
        "information of any other customer of AnantaSutra. The cost of "
        "any such audit shall be borne by the Client unless the audit "
        "discloses a material breach by AnantaSutra.",
    )
    add_sub(
        doc,
        "10.7  Return and Deletion.",
        "On termination or expiry of this Agreement or the relevant "
        "Statement of Work, AnantaSutra shall, at the Client's option, "
        "return or securely destroy all Personal Data in its "
        "possession or control, and shall certify such destruction in "
        "writing, save to the extent that AnantaSutra is required to "
        "retain Personal Data to comply with Applicable Laws.",
    )
    add_sub(
        doc,
        "10.8  Full Data Processing Addendum.",
        "A summary of the data processing particulars is set out in "
        "Schedule 2. Where a Statement of Work involves processing of "
        "Personal Data of any significant volume or sensitivity, the "
        "Parties shall execute a separate Data Processing Agreement, "
        "which shall be a Schedule to this Agreement and, as regards "
        "the processing of Personal Data, shall prevail over the "
        "provisions of this Clause 10 in the event of conflict.",
    )

    # -----------------------------------------------------------------
    # 11. Representations and Warranties
    # -----------------------------------------------------------------
    add_clause(doc, 11, "REPRESENTATIONS AND WARRANTIES")
    add_sub(
        doc,
        "11.1  Mutual.",
        "Each Party represents and warrants to the other that (a) it is "
        "duly incorporated and validly existing under Applicable Laws; "
        "(b) it has full corporate power and authority to execute and "
        "perform this Agreement, and the signatory hereto is duly "
        "authorised, including by way of a resolution of its board of "
        "directors where required under the Companies Act, 2013; (c) "
        "the execution and performance of this Agreement does not "
        "conflict with or breach any other contract, judgment, or "
        "Applicable Law binding on it; and (d) it shall comply with "
        "all Applicable Laws in the performance of its obligations "
        "hereunder.",
    )
    add_sub(
        doc,
        "11.2  AnantaSutra-Specific.",
        "AnantaSutra further represents and warrants that (a) to the best "
        "of its knowledge, the Services and Deliverables (other than "
        "Client-supplied materials and Client-directed modifications) "
        "will not infringe the intellectual property rights of any "
        "third party; (b) it has the necessary tools, personnel, and "
        "expertise to perform the Services; and (c) it shall comply "
        "with its obligations under the Prevention of Corruption Act, "
        "1988, and, in relation to cross-border engagements, the US "
        "Foreign Corrupt Practices Act, 1977 and the UK Bribery Act, "
        "2010.",
    )
    add_sub(
        doc,
        "11.3  Disclaimer.",
        "Save as expressly set out in this Agreement, and to the "
        "maximum extent permitted by Applicable Laws, AnantaSutra "
        "disclaims all other conditions, warranties, and "
        "representations, whether express or implied, statutory or "
        "otherwise, including the implied conditions and warranties "
        "under the Sale of Goods Act, 1930 in so far as they may "
        "apply.",
    )

    # -----------------------------------------------------------------
    # 12. Service Levels and Acceptance
    # -----------------------------------------------------------------
    add_clause(doc, 12, "SERVICE LEVELS")
    add_sub(
        doc,
        "12.1",
        "Where the Services are subject to service levels, those "
        "service levels, the measurement methodology, the applicable "
        "service credits (if any), and any exclusions shall be as set "
        "out in the applicable Statement of Work, read with the "
        "framework set out in Schedule 3.",
    )
    add_sub(
        doc,
        "12.2",
        "Any service credits payable shall be the Client's sole and "
        "exclusive financial remedy for breach of the relevant service "
        "level, save where the breach also constitutes a material "
        "breach of this Agreement.",
    )

    # -----------------------------------------------------------------
    # 13. Indemnities
    # -----------------------------------------------------------------
    add_clause(doc, 13, "INDEMNITIES")
    add_sub(
        doc,
        "13.1  IP Infringement Indemnity.",
        "AnantaSutra shall, at its cost, defend, indemnify, and hold "
        "harmless the Client against any third-party claim alleging "
        "that the Deliverables, as delivered by AnantaSutra, infringe the "
        "intellectual property rights of such third party, and shall "
        "pay any damages finally awarded against the Client by a court "
        "of competent jurisdiction, or agreed by AnantaSutra in "
        "settlement, in respect of such claim. If any Deliverable is, "
        "or in AnantaSutra's reasonable opinion is likely to be, the "
        "subject of such a claim, AnantaSutra may, at its option and "
        "cost, (i) procure for the Client the right to continue using "
        "the Deliverable; (ii) modify or replace the Deliverable so as "
        "to avoid infringement without materially diminishing its "
        "functionality; or (iii) if neither of the foregoing is "
        "commercially reasonable, accept return of the Deliverable and "
        "refund the Fees paid for it, depreciated on a straight-line "
        "basis over thirty-six (36) months from delivery.",
    )
    add_sub(
        doc,
        "13.2  Carve-outs.",
        "Clause 13.1 shall not apply to any claim arising out of (a) "
        "materials, specifications, or directions supplied by the "
        "Client; (b) modifications to the Deliverable not made by "
        "AnantaSutra; (c) combination of the Deliverable with products or "
        "services not supplied by AnantaSutra, where the claim would not "
        "have arisen but for such combination; (d) use of the "
        "Deliverable outside the scope of the licence or permitted "
        "use; or (e) continued use of an allegedly infringing version "
        "after AnantaSutra has made a non-infringing version available.",
    )
    add_sub(
        doc,
        "13.3  Data-Breach Indemnity.",
        "AnantaSutra shall indemnify the Client against direct losses, "
        "third-party claims, and monetary penalties imposed by the "
        "Data Protection Board of India under Section 33 of, and the "
        "Schedule to, the DPDP Act to the extent arising out of "
        "AnantaSutra's breach of Clause 10 of this Agreement.",
    )
    add_sub(
        doc,
        "13.4  Client Indemnity.",
        "The Client shall indemnify AnantaSutra against any third-party "
        "claim to the extent arising out of (a) Client-supplied "
        "materials, data, or specifications that infringe third-party "
        "rights or breach Applicable Laws; or (b) the Client's use of "
        "the Deliverables otherwise than in accordance with this "
        "Agreement or any documentation provided by AnantaSutra.",
    )
    add_sub(
        doc,
        "13.5  Indemnity Procedure.",
        "The Party seeking indemnity (the \"Indemnified Party\") "
        "shall (a) promptly notify the other Party (the \"Indemnifying "
        "Party\") in writing of the claim; (b) grant the Indemnifying "
        "Party sole control of the defence and settlement of the "
        "claim, save that no settlement admitting liability or "
        "imposing non-monetary obligations on the Indemnified Party "
        "shall be agreed without the Indemnified Party's prior "
        "written consent (not to be unreasonably withheld); and (c) "
        "provide reasonable co-operation at the Indemnifying Party's "
        "cost. Failure to provide prompt notice shall relieve the "
        "Indemnifying Party of its obligations only to the extent it "
        "is materially prejudiced thereby.",
    )
    add_sub(
        doc,
        "13.6",
        "The indemnities in this Clause 13 state the sole and "
        "exclusive liability of each Party, and the other Party's sole "
        "and exclusive remedy, in respect of the matters covered by "
        "those indemnities.",
    )

    # -----------------------------------------------------------------
    # 14. Limitation of Liability
    # -----------------------------------------------------------------
    add_clause(doc, 14, "LIMITATION OF LIABILITY")
    add_sub(
        doc,
        "14.1  Exclusion of Indirect Damages.",
        "To the maximum extent permitted by Applicable Laws, neither "
        "Party shall be liable to the other, whether in contract, tort "
        "(including negligence), under statute, or otherwise, for any "
        "indirect, consequential, incidental, special, punitive, or "
        "exemplary damages, or for any loss of profits, loss of "
        "revenue, loss of goodwill, loss of business opportunity, loss "
        "of data (subject to Clause 14.3), or loss of anticipated "
        "savings, even if such Party has been advised of the "
        "possibility of such damages.",
    )
    add_sub(
        doc,
        "14.2  General Cap.",
        "Subject to Clause 14.3, the aggregate liability of each "
        "Party under or in connection with this Agreement or any "
        "Statement of Work, whether in contract, tort (including "
        "negligence), under statute, or otherwise, shall not exceed "
        "the total Fees actually paid and payable by the Client to "
        "AnantaSutra under the Statement of Work giving rise to the "
        "claim in the twelve (12) months immediately preceding the "
        "first event giving rise to liability.",
    )
    add_sub(
        doc,
        "14.3  Carve-outs from Cap and Exclusion.",
        "The exclusions and cap in Clauses 14.1 and 14.2 shall not "
        "apply to: (a) AnantaSutra's liability under the IP infringement "
        "indemnity in Clause 13.1; (b) breach of the confidentiality "
        "obligations in Clause 9, in respect of which the aggregate "
        "liability shall be capped at two (2) times the cap in Clause "
        "14.2 (the \"Confidentiality Super-Cap\"); (c) AnantaSutra's "
        "liability under the Data-Breach Indemnity in Clause 13.3 and "
        "its breach of Clause 10, in respect of which the aggregate "
        "liability shall be capped at two (2) times the cap in Clause "
        "14.2 or such higher amount as may be expressly agreed in the "
        "Statement of Work having regard to the Parties' respective "
        "exposure under Section 33 of the DPDP Act (the \"Data "
        "Protection Super-Cap\"); (d) fraud or fraudulent "
        "misrepresentation; (e) gross negligence or wilful misconduct; "
        "(f) death or personal injury caused by the negligence of a "
        "Party; (g) the Client's obligation to pay undisputed Fees; "
        "and (h) any liability which cannot, as a matter of Indian "
        "public policy under Section 23 of the Indian Contract Act, "
        "1872, be lawfully excluded or limited.",
    )
    add_sub(
        doc,
        "14.4",
        "Each Party shall use reasonable endeavours to mitigate its "
        "losses. The limitations and exclusions in this Clause 14 "
        "reflect the commercial bargain between the Parties and the "
        "Fees paid under this Agreement.",
    )

    # -----------------------------------------------------------------
    # 15. Insurance
    # -----------------------------------------------------------------
    add_clause(doc, 15, "INSURANCE")
    add_sub(
        doc,
        "15.1",
        "AnantaSutra shall, throughout the Term, maintain with reputable "
        "insurers the following insurance cover, in each case with "
        "limits no less than those set out below (or such higher "
        "limits as may be specified in a Statement of Work): (a) "
        "professional indemnity insurance of not less than INR 2 "
        "(two) crore per claim; (b) cyber-liability insurance of not "
        "less than INR 1 (one) crore per claim; (c) commercial general "
        "liability insurance of not less than INR 1 (one) crore per "
        "occurrence; and (d) employees' compensation insurance as "
        "required under the Employees' Compensation Act, 1923, and "
        "the Employees' State Insurance Act, 1948, as applicable.",
    )
    add_sub(
        doc,
        "15.2",
        "AnantaSutra shall, upon reasonable request, furnish the Client "
        "with certificates of insurance evidencing the cover required "
        "under Clause 15.1.",
    )

    # -----------------------------------------------------------------
    # 16. Term and Termination
    # -----------------------------------------------------------------
    add_clause(doc, 16, "TERM AND TERMINATION")
    add_sub(
        doc,
        "16.1  Term.",
        "This Agreement shall commence on the Effective Date and shall "
        "continue in force for an initial period of two (2) years, "
        "and shall thereafter renew for successive periods of one (1) "
        "year each unless either Party gives the other not less than "
        "sixty (60) days' prior written notice of non-renewal (the "
        "\"Term\"). Termination or expiry of this Agreement shall "
        "automatically terminate all then-subsisting Statements of "
        "Work, save as the Parties may otherwise agree in writing.",
    )
    add_sub(
        doc,
        "16.2  Termination for Convenience.",
        "Either Party may terminate this Agreement or any Statement of "
        "Work for convenience, without cause, by giving the other not "
        "less than sixty (60) days' prior written notice, such notice "
        "not to be given during the first twelve (12) months of a "
        "Statement of Work except as expressly permitted therein.",
    )
    add_sub(
        doc,
        "16.3  Termination for Cause.",
        "Either Party may terminate this Agreement or any Statement "
        "of Work with immediate effect by written notice to the other "
        "if the other (a) commits a material breach of this Agreement "
        "which, if capable of cure, is not cured within thirty (30) "
        "days of written notice requiring cure; or (b) commits a "
        "series of persistent breaches which, taken together, "
        "constitute a material breach.",
    )
    add_sub(
        doc,
        "16.4  Insolvency.",
        "Either Party may terminate this Agreement with immediate "
        "effect by written notice to the other if the other (a) is "
        "unable to pay its debts as they fall due within the meaning "
        "of Section 4 of the Insolvency and Bankruptcy Code, 2016; "
        "(b) has a corporate insolvency resolution process initiated "
        "against it under the Insolvency and Bankruptcy Code, 2016; "
        "(c) enters into a scheme of arrangement with its creditors "
        "under Sections 230–232 of the Companies Act, 2013 otherwise "
        "than for solvent reconstruction; or (d) has a receiver, "
        "liquidator, or resolution professional appointed over the "
        "whole or any material part of its assets.",
    )
    add_sub(
        doc,
        "16.5  Change of Control.",
        "The Client may terminate this Agreement by written notice to "
        "AnantaSutra if AnantaSutra undergoes a change of control in favour "
        "of a direct competitor of the Client, such notice to be "
        "given within sixty (60) days of the Client becoming aware of "
        "such change of control.",
    )
    add_sub(
        doc,
        "16.6  Consequences of Termination.",
        "On termination or expiry of this Agreement or any Statement "
        "of Work: (a) the Client shall pay all undisputed Fees for "
        "Services performed, and out-of-pocket expenses incurred, up "
        "to the effective date of termination; (b) each Party shall, "
        "at the other's option, return or destroy all Confidential "
        "Information of the other in its possession, subject to "
        "Clause 10.7 in relation to Personal Data; (c) subject to "
        "payment in accordance with Clause 16.6(a), the IP assignment "
        "in Clause 8.2 shall operate in respect of Deliverables "
        "actually delivered and paid for; and (d) AnantaSutra shall, at "
        "the Client's reasonable request and on reasonable terms, "
        "provide transition assistance for a period of up to ninety "
        "(90) days at its then-prevailing rates.",
    )
    add_sub(
        doc,
        "16.7  Survival.",
        "Clauses 1, 5 (in respect of accrued Fees), 8, 9, 10, 13, 14, "
        "16.6, 16.7, 17, 19, 20, 21, and 22 shall survive termination "
        "or expiry of this Agreement.",
    )

    # -----------------------------------------------------------------
    # 17. Force Majeure
    # -----------------------------------------------------------------
    add_clause(doc, 17, "FORCE MAJEURE")
    add_sub(
        doc,
        "17.1",
        "Neither Party shall be liable for any failure or delay in "
        "performing its obligations under this Agreement (other than "
        "payment obligations) to the extent such failure or delay is "
        "caused by an event beyond its reasonable control, including "
        "acts of God, natural disasters, war, terrorism, civil "
        "commotion, riots, strikes, lockouts (other than by its own "
        "workforce), governmental orders or lockdowns, epidemics or "
        "pandemics, failure of public infrastructure (including "
        "internet or electricity), or any event recognised under "
        "Section 56 of the Indian Contract Act, 1872 as frustrating "
        "the contract (a \"Force Majeure Event\").",
    )
    add_sub(
        doc,
        "17.2",
        "The affected Party shall promptly notify the other of the "
        "Force Majeure Event, use reasonable endeavours to mitigate "
        "its effects, and resume performance as soon as reasonably "
        "practicable. If the Force Majeure Event continues for more "
        "than ninety (90) days, either Party may terminate this "
        "Agreement or the affected Statement of Work by written "
        "notice, without liability save for accrued rights and "
        "obligations.",
    )

    # -----------------------------------------------------------------
    # 18. Non-Solicitation
    # -----------------------------------------------------------------
    add_clause(doc, 18, "NON-SOLICITATION")
    add_sub(
        doc,
        "18.1",
        "During the Term and for a period of twelve (12) months "
        "thereafter, neither Party shall directly or indirectly "
        "solicit for employment or engagement any person who is, or "
        "has within the preceding six (6) months been, an employee or "
        "consultant of the other Party and who has been materially "
        "involved in the performance of the Services. This Clause "
        "shall not prohibit (a) general advertising or recruitment "
        "campaigns not specifically targeted at the other Party's "
        "personnel; or (b) the employment of a person who "
        "independently responds to such a campaign.",
    )
    add_sub(
        doc,
        "18.2",
        "The Parties acknowledge that this Clause 18 is reasonable "
        "and does not operate as a restraint of trade under Section "
        "27 of the Indian Contract Act, 1872, but is a legitimate "
        "measure to protect each Party's investment in its workforce.",
    )

    # -----------------------------------------------------------------
    # 19. Anti-Bribery and Anti-Corruption
    # -----------------------------------------------------------------
    add_clause(doc, 19, "ANTI-BRIBERY AND ANTI-CORRUPTION")
    add_sub(
        doc,
        "19.1",
        "Each Party shall, and shall procure that its Affiliates, "
        "personnel, and subcontractors shall, comply with all "
        "applicable anti-bribery and anti-corruption laws, including "
        "the Prevention of Corruption Act, 1988, and, in relation to "
        "cross-border performance, the US Foreign Corrupt Practices "
        "Act, 1977 and the UK Bribery Act, 2010.",
    )
    add_sub(
        doc,
        "19.2",
        "Each Party warrants that it has not and shall not, directly "
        "or indirectly, offer, promise, give, request, or accept any "
        "bribe, kickback, or other improper payment or advantage in "
        "connection with the Services. Breach of this Clause 19 shall "
        "be a material breach, entitling the other Party to terminate "
        "this Agreement for cause with immediate effect and to claim "
        "all direct losses suffered.",
    )

    # -----------------------------------------------------------------
    # 20. Dispute Resolution
    # -----------------------------------------------------------------
    add_clause(doc, 20, "DISPUTE RESOLUTION")
    add_sub(
        doc,
        "20.1  Good-Faith Negotiation.",
        "In the event of any dispute, controversy, or claim arising "
        "out of or in connection with this Agreement or any Statement "
        "of Work (a \"Dispute\"), the Parties shall first attempt to "
        "resolve the Dispute by good-faith negotiation between their "
        "respective project managers within fifteen (15) Business "
        "Days of written notice of the Dispute. Failing resolution, "
        "the Dispute shall be escalated to senior executives (of the "
        "rank of Director or equivalent) of each Party for a further "
        "period of fifteen (15) Business Days.",
    )
    add_sub(
        doc,
        "20.2  Arbitration.",
        "If the Dispute is not resolved under Clause 20.1, it shall "
        "be referred to and finally resolved by arbitration under "
        "the Arbitration and Conciliation Act, 1996, as amended "
        "(including the Arbitration and Conciliation (Amendment) "
        "Acts of 2015, 2019, and 2021). Where the amount in dispute "
        "is equal to or less than INR 5 (five) crore, the arbitral "
        "tribunal shall comprise a sole arbitrator to be mutually "
        "appointed by the Parties, failing which appointed in "
        "accordance with Section 11 of the Arbitration and "
        "Conciliation Act, 1996. Where the amount in dispute exceeds "
        "INR 5 (five) crore, the tribunal shall comprise three "
        "arbitrators, one to be appointed by each Party and the "
        "third (who shall be the presiding arbitrator) by the two "
        "Party-appointed arbitrators. The Parties may, by mutual "
        "written agreement, elect institutional arbitration under "
        "the rules of the Mumbai Centre for International Arbitration "
        "(MCIA) or the Delhi International Arbitration Centre "
        "(DIAC).",
    )
    add_sub(
        doc,
        "20.3  Seat, Venue and Language.",
        "The seat of arbitration shall be New Delhi, Delhi, "
        "India, and the courts at New Delhi shall accordingly have "
        "exclusive supervisory jurisdiction over the arbitration in "
        "accordance with the principles in BALCO v. Kaiser Aluminium "
        "Technical Services, (2012) 9 SCC 552 and Indus Mobile "
        "Distribution v. Datawind Innovations, (2017) 7 SCC 678. The "
        "venue of arbitral hearings may, for convenience, be held "
        "elsewhere as the tribunal may determine in consultation "
        "with the Parties. The language of arbitration shall be "
        "English.",
    )
    add_sub(
        doc,
        "20.4  Interim Relief.",
        "Notwithstanding Clause 20.2, either Party may at any time "
        "apply to a court of competent jurisdiction for interim, "
        "protective, or conservatory relief (including injunctive "
        "relief under the Specific Relief Act, 1963) under Section 9 "
        "of the Arbitration and Conciliation Act, 1996, and such "
        "application shall not be treated as a waiver of the "
        "arbitration agreement.",
    )
    add_sub(
        doc,
        "20.5",
        "The award shall be final and binding on the Parties. The "
        "costs of arbitration shall be borne as the tribunal may "
        "determine.",
    )

    # -----------------------------------------------------------------
    # 21. Governing Law and Jurisdiction
    # -----------------------------------------------------------------
    add_clause(doc, 21, "GOVERNING LAW AND JURISDICTION")
    add_body_paragraph(
        doc,
        "This Agreement shall be governed by and construed in "
        "accordance with the laws of the Republic of India. Subject "
        "to Clause 20, the courts at New Delhi, Delhi shall have "
        "exclusive jurisdiction over any matter arising out of or in "
        "connection with this Agreement.",
    )

    # -----------------------------------------------------------------
    # 22. Miscellaneous
    # -----------------------------------------------------------------
    add_clause(doc, 22, "MISCELLANEOUS")
    add_sub(
        doc,
        "22.1  Notices.",
        "Any notice under this Agreement shall be in writing and "
        "shall be delivered by hand, by recognised courier, or by "
        "electronic mail with delivery and read-receipt requested, "
        "to the addresses set out below (or to such other address as "
        "a Party may notify to the other):",
    )
    add_body_paragraph(
        doc,
        "If to AnantaSutra: Attn: Mr. Himanshu Mishra, Founder & CEO, "
        "AnantaSutra, Delhi, India; email: contact@anantasutra.com.",
    )
    add_body_paragraph(
        doc,
        "If to the Client: Attention: [NAME, DESIGNATION], [CLIENT "
        "LEGAL NAME], [CLIENT ADDRESS]; email: [CLIENT LEGAL NOTICES "
        "EMAIL].",
    )
    add_sub(
        doc,
        "22.2  Assignment.",
        "Neither Party shall assign, novate, transfer, or charge any "
        "of its rights or obligations under this Agreement without "
        "the prior written consent of the other Party, save that "
        "either Party may, on prior written notice, assign its "
        "rights and obligations to an Affiliate or to a successor "
        "entity arising out of a bona fide internal reorganisation, "
        "merger, or sale of substantially all assets, provided that "
        "such assignee is not a direct competitor of the other Party.",
    )
    add_sub(
        doc,
        "22.3  Subcontracting.",
        "Subcontracting is governed by Clause 7.",
    )
    add_sub(
        doc,
        "22.4  Publicity.",
        "Neither Party shall use the name, logo, or trademarks of "
        "the other Party in any marketing, press release, website, "
        "or case study without the prior written consent of that "
        "other Party, such consent not to be unreasonably withheld "
        "for routine identification of customer relationships.",
    )
    add_sub(
        doc,
        "22.5  Entire Agreement.",
        "This Agreement, together with the Schedules and all "
        "Statements of Work executed hereunder, constitutes the "
        "entire agreement between the Parties in relation to its "
        "subject matter and supersedes all prior agreements, "
        "understandings, and representations, whether oral or "
        "written, including any non-disclosure agreement previously "
        "executed between the Parties (which shall be deemed "
        "subsumed into this Agreement save to the extent of any "
        "obligations expressly preserved therein).",
    )
    add_sub(
        doc,
        "22.6  Amendment.",
        "No amendment to this Agreement shall be effective unless "
        "made in writing and signed by the authorised "
        "representatives of both Parties.",
    )
    add_sub(
        doc,
        "22.7  Severability.",
        "If any provision of this Agreement is held by a court or "
        "arbitral tribunal of competent jurisdiction to be invalid, "
        "illegal, or unenforceable, that provision shall be "
        "severed, and the remainder of this Agreement shall "
        "continue in full force and effect. The Parties shall "
        "negotiate in good faith to replace the severed provision "
        "with a valid provision that most closely reflects their "
        "original intention.",
    )
    add_sub(
        doc,
        "22.8  Waiver.",
        "No failure or delay by a Party in exercising any right "
        "under this Agreement shall operate as a waiver thereof, "
        "and no single or partial exercise of any right shall "
        "preclude any further exercise of that right or any other "
        "right.",
    )
    add_sub(
        doc,
        "22.9  Counterparts and Electronic Execution.",
        "This Agreement may be executed in any number of "
        "counterparts, each of which shall be deemed an original "
        "and all of which together shall constitute one and the "
        "same instrument. The Parties agree that this Agreement may "
        "be executed by electronic means, including by Aadhaar "
        "e-Sign, digital signature under Section 5 of the "
        "Information Technology Act, 2000, or by affixing scanned "
        "signatures and exchanging executed PDF counterparts, and "
        "such electronic execution shall be deemed valid and "
        "binding in accordance with Section 10A of the Information "
        "Technology Act, 2000.",
    )
    add_sub(
        doc,
        "22.10  No Third-Party Rights.",
        "This Agreement does not confer any rights on any person or "
        "party who is not a Party to it.",
    )
    add_sub(
        doc,
        "22.11  Stamp Duty.",
        "The Parties acknowledge that this Agreement shall be "
        "stamped in accordance with the Indian Stamp Act, 1899 (as applicable to the NCT of Delhi) "
        "read with the Indian Stamp Act, 1899, and that each "
        "Statement of Work executed hereunder shall, where "
        "required, bear separate stamp duty in accordance with the "
        "applicable State stamp legislation. The stamp duty payable "
        "on this Agreement shall be borne by [ANANTASUTRA/THE CLIENT/"
        "EQUALLY BY THE PARTIES]. The Parties acknowledge the "
        "position of law as set out by the Hon'ble Supreme Court in "
        "the 2023 Curative ruling concerning the interplay between "
        "arbitration agreements and the Indian Stamp Act, and shall "
        "procure that this Agreement is duly stamped before being "
        "tendered in evidence.",
    )
    add_sub(
        doc,
        "22.12  Compliance with Laws.",
        "Each Party shall comply with all Applicable Laws in the "
        "performance of this Agreement, including, as applicable, "
        "the Companies Act, 2013 (including Section 188 relating to "
        "related-party transactions), the Income-tax Act, 1961, the "
        "CGST Act, 2017, and all labour, environmental, and data "
        "protection laws.",
    )

    # -----------------------------------------------------------------
    # Execution block
    # -----------------------------------------------------------------
    add_spacer(doc, 10)
    add_h1(doc, "IN WITNESS WHEREOF")
    add_body_paragraph(
        doc,
        "The Parties have caused this Master Service Agreement to be "
        "executed by their duly authorised representatives on the "
        "date first above written.",
    )

    add_spacer(doc, 6)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_border(cell, top=True, bottom=True, left=True, right=True)
    sig_rows = [
        ("For and on behalf of", "For and on behalf of"),
        ("ANANTASUTRA", "[CLIENT LEGAL NAME]"),
        ("", ""),
        ("Signature: ____________________", "Signature: ____________________"),
        ("Name: Mr. Himanshu Mishra", "Name: [NAME]"),
        ("Designation: Founder & CEO", "Designation: [DESIGNATION]"),
        ("Email: contact@anantasutra.com", ""),
        ("Date: [DATE]", "Date: [DATE]"),
        ("Place: Delhi", "Place: [PLACE]"),
    ]
    # Populate first row
    first_row = table.rows[0]
    first_row.cells[0].text = ""
    first_row.cells[1].text = ""
    for i, (l, r) in enumerate(sig_rows):
        if i == 0:
            row = first_row
        else:
            row = table.add_row()
            for cell in row.cells:
                set_cell_border(cell, top=True, bottom=True, left=True, right=True)
        for idx, text in enumerate((l, r)):
            cell = row.cells[idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.name = BODY_FONT
            run.font.size = Pt(11)
            if i in (0, 1):
                run.bold = True

    add_spacer(doc, 6)
    add_body_paragraph(
        doc,
        "WITNESSES:",
        bold=True,
    )
    witness_table = doc.add_table(rows=4, cols=2)
    witness_table.autofit = False
    for row in witness_table.rows:
        for cell in row.cells:
            set_cell_border(cell, top=True, bottom=True, left=True, right=True)
    witnesses = [
        ("1. Signature: ____________________", "2. Signature: ____________________"),
        ("    Name: [WITNESS 1 NAME]", "    Name: [WITNESS 2 NAME]"),
        ("    Address: [WITNESS 1 ADDRESS]", "    Address: [WITNESS 2 ADDRESS]"),
        ("    Date: [DATE]", "    Date: [DATE]"),
    ]
    for i, (l, r) in enumerate(witnesses):
        for idx, text in enumerate((l, r)):
            cell = witness_table.rows[i].cells[idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.name = BODY_FONT
            run.font.size = Pt(11)

    # -----------------------------------------------------------------
    # Schedule 1 — Template SOW
    # -----------------------------------------------------------------
    doc.add_page_break()
    add_title(doc, "SCHEDULE 1")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TEMPLATE STATEMENT OF WORK")
    r.bold = True
    r.font.name = BODY_FONT
    r.font.size = Pt(14)
    p.paragraph_format.space_after = Pt(12)

    add_body_paragraph(
        doc,
        "This Statement of Work No. [SOW-NO.] dated [DATE] (\"this "
        "SOW\") is issued under and subject to the Master Service "
        "Agreement dated [MSA EFFECTIVE DATE] between AnantaSutra "
        "(\"AnantaSutra\") and [CLIENT LEGAL NAME] "
        "(\"the Client\") (the \"MSA\"). Capitalised terms used but "
        "not defined in this SOW have the meanings given to them in "
        "the MSA. In the event of conflict, Clause 2.3 of the MSA "
        "shall determine the order of precedence.",
    )
    sow_items = [
        ("1. Project Name", "[PROJECT NAME]"),
        (
            "2. Scope of Services",
            "[DESCRIBE SCOPE OF SERVICES, DELIVERABLES, AND OUT-OF-"
            "SCOPE ITEMS]",
        ),
        (
            "3. Deliverables",
            "[LIST DELIVERABLES, SPECIFICATIONS, ACCEPTANCE "
            "CRITERIA, AND FORMAT]",
        ),
        (
            "4. Timelines / Milestones",
            "[LIST MILESTONES, DATES, DEPENDENCIES]",
        ),
        (
            "5. Personnel / Project Team",
            "[LIST KEY PERSONNEL, ROLES, AND RATES (IF T&M)]",
        ),
        (
            "6. Commercial Model",
            "[FIXED-PRICE / TIME-AND-MATERIALS / MILESTONE-BASED / "
            "OUTCOME-BASED]",
        ),
        (
            "7. Fees",
            "[STATE FEES IN INR/USD, INCLUSIVE/EXCLUSIVE OF TAXES; "
            "MILESTONE PAYMENT SCHEDULE]",
        ),
        (
            "8. Out-of-Pocket Expenses",
            "[POLICY FOR REIMBURSEMENT, ADVANCE APPROVAL "
            "THRESHOLDS]",
        ),
        (
            "9. Invoicing Schedule",
            "[MONTHLY / MILESTONE / UPON COMPLETION]",
        ),
        (
            "10. Payment Terms",
            "[NET 30 / 45 DAYS FROM INVOICE DATE]",
        ),
        (
            "11. Acceptance Criteria",
            "[DETAILED ACCEPTANCE CRITERIA; UAT PROCEDURE; "
            "ACCEPTANCE PERIOD]",
        ),
        (
            "12. Service Levels",
            "[SERVICE LEVELS, SERVICE CREDITS, EXCLUSIONS; OR "
            "\"NOT APPLICABLE\"]",
        ),
        (
            "13. Personal Data",
            "[DESCRIBE CATEGORIES OF DATA PRINCIPALS, PERSONAL "
            "DATA, PURPOSE, DURATION; OR \"NOT APPLICABLE\"]",
        ),
        (
            "14. Sub-processors",
            "[LIST APPROVED SUB-PROCESSORS / PROCESS FOR APPROVAL]",
        ),
        (
            "15. Change of MSA Terms",
            "[ANY EXPRESS VARIATION OF MSA CLAUSES; DEFAULT: "
            "\"NONE\"]",
        ),
        (
            "16. Special Terms",
            "[ANY PROJECT-SPECIFIC TERMS]",
        ),
    ]
    for heading, body in sow_items:
        add_h2(doc, heading)
        add_body_paragraph(doc, body)

    add_spacer(doc, 8)
    add_body_paragraph(
        doc,
        "Executed by the duly authorised representatives of the "
        "Parties:",
    )
    sow_sig = doc.add_table(rows=5, cols=2)
    sow_sig.autofit = False
    for row in sow_sig.rows:
        for cell in row.cells:
            set_cell_border(cell, top=True, bottom=True, left=True, right=True)
    sow_sig_rows = [
        ("For ANANTASUTRA", "For [CLIENT LEGAL NAME]"),
        ("Signature: ____________________", "Signature: ____________________"),
        ("Name: Mr. Himanshu Mishra", "Name: [NAME]"),
        ("Designation: Founder & CEO", "Designation: [DESIGNATION]"),
        ("Date: [DATE]", "Date: [DATE]"),
    ]
    for i, (l, r) in enumerate(sow_sig_rows):
        for idx, text in enumerate((l, r)):
            cell = sow_sig.rows[i].cells[idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.name = BODY_FONT
            run.font.size = Pt(11)
            if i == 0:
                run.bold = True

    # -----------------------------------------------------------------
    # Schedule 2 — DPA Summary
    # -----------------------------------------------------------------
    doc.add_page_break()
    add_title(doc, "SCHEDULE 2")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DATA PROCESSING ADDENDUM — SUMMARY")
    r.bold = True
    r.font.name = BODY_FONT
    r.font.size = Pt(14)
    p.paragraph_format.space_after = Pt(12)

    add_body_paragraph(
        doc,
        "This Schedule 2 is a summary of the data processing "
        "particulars applicable where AnantaSutra processes Personal "
        "Data on behalf of the Client under a Statement of Work. It "
        "is read together with Clause 10 of the MSA. Where the "
        "scope and volume of processing so warrant, the Parties "
        "shall execute a separate, fuller Data Processing Agreement "
        "(the \"DPA\"), which shall, upon execution, prevail over "
        "this Schedule 2 and, as regards processing of Personal "
        "Data, over Clause 10 of the MSA.",
    )

    rows = [
        ("Subject-matter of processing", "[DESCRIBE THE SERVICES INVOLVING PERSONAL DATA]"),
        ("Duration of processing", "Term of the applicable Statement of Work, and thereafter as required for return/destruction under Clause 10.7 of the MSA."),
        ("Nature and purpose of processing", "[DESCRIBE, e.g. storage, analytics, helpdesk, machine learning, development/testing]"),
        ("Types of Personal Data", "[DESCRIBE — e.g. name, contact details, employee IDs, financial data, health data, etc.]"),
        ("Categories of Data Principals", "[DESCRIBE — e.g. employees, customers, job applicants]"),
        ("Sub-processors (approved)", "[LIST APPROVED SUB-PROCESSORS WITH LOCATION]"),
        ("Cross-border transfer locations", "[LIST COUNTRIES; NOTE ANY RESTRICTIONS NOTIFIED UNDER SECTION 16 DPDP ACT]"),
        ("Security measures", "Technical and organisational measures aligned with IS/ISO/IEC 27001, including access controls, encryption in transit and at rest, logging, vulnerability management, and incident response."),
        ("Breach notification timeline", "Within forty-eight (48) hours of AnantaSutra becoming aware of a personal data breach, in accordance with Clause 10.4 of the MSA."),
        ("Data Principal rights", "AnantaSutra shall assist the Client in responding to requests under Sections 11-14 of the DPDP Act within the statutory timelines."),
        ("Audit rights", "Once per 12 months and upon a breach, as per Clause 10.6 of the MSA."),
        ("Deletion / return", "On termination, in accordance with Clause 10.7 of the MSA."),
    ]
    dpa_table = doc.add_table(rows=1, cols=2)
    dpa_table.autofit = False
    header = dpa_table.rows[0].cells
    for idx, text in enumerate(("Particulars", "Details")):
        cell = header[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.name = BODY_FONT
        run.font.size = Pt(11)
        set_cell_border(cell, top=True, bottom=True, left=True, right=True)
    for particulars, details in rows:
        r = dpa_table.add_row()
        for idx, text in enumerate((particulars, details)):
            cell = r.cells[idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.name = BODY_FONT
            run.font.size = Pt(11)
            if idx == 0:
                run.bold = True
            set_cell_border(cell, top=True, bottom=True, left=True, right=True)

    # -----------------------------------------------------------------
    # Schedule 3 — SLA Framework
    # -----------------------------------------------------------------
    doc.add_page_break()
    add_title(doc, "SCHEDULE 3")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SERVICE LEVEL FRAMEWORK — TEMPLATE")
    r.bold = True
    r.font.name = BODY_FONT
    r.font.size = Pt(14)
    p.paragraph_format.space_after = Pt(12)

    add_body_paragraph(
        doc,
        "Where a Statement of Work specifies that Services are "
        "subject to service levels, this Schedule 3 sets out the "
        "framework within which such service levels are to be "
        "measured, reported, and remedied.",
    )

    add_h2(doc, "1. Availability Service Level")
    sla_table = doc.add_table(rows=1, cols=4)
    sla_table.autofit = False
    header = sla_table.rows[0].cells
    for idx, text in enumerate((
        "Service Level", "Target", "Measurement Window", "Service Credit",
    )):
        cell = header[idx]
        cell.text = ""
        pp = cell.paragraphs[0]
        rn = pp.add_run(text)
        rn.bold = True
        rn.font.name = BODY_FONT
        rn.font.size = Pt(11)
        set_cell_border(cell, top=True, bottom=True, left=True, right=True)

    sla_rows = [
        ("System Availability", "99.5% per calendar month", "Calendar month; excluding scheduled maintenance", "[X]% of monthly Fees per 0.5% shortfall, capped at [Y]%"),
        ("P1 Incident Response", "Within 1 hour of logging", "24x7", "[X]% of monthly Fees per breach, capped at [Y]%"),
        ("P1 Incident Resolution", "Within 8 business hours", "24x7", "[X]% of monthly Fees per breach, capped at [Y]%"),
        ("P2 Incident Response", "Within 4 business hours", "Business hours", "[X]% of monthly Fees per breach, capped at [Y]%"),
        ("P2 Incident Resolution", "Within 2 Business Days", "Business hours", "[X]% of monthly Fees per breach, capped at [Y]%"),
    ]
    for row_data in sla_rows:
        r = sla_table.add_row()
        for idx, text in enumerate(row_data):
            cell = r.cells[idx]
            cell.text = ""
            pp = cell.paragraphs[0]
            rn = pp.add_run(text)
            rn.font.name = BODY_FONT
            rn.font.size = Pt(11)
            set_cell_border(cell, top=True, bottom=True, left=True, right=True)

    add_h2(doc, "2. Exclusions")
    add_body_paragraph(
        doc,
        "The following shall be excluded from service-level "
        "measurement: (a) scheduled maintenance notified in advance; "
        "(b) downtime attributable to Client-side systems, networks, "
        "or third-party services not controlled by AnantaSutra; (c) "
        "Force Majeure Events; (d) failure by the Client to comply "
        "with its obligations; and (e) causes outside AnantaSutra's "
        "reasonable control.",
    )

    add_h2(doc, "3. Reporting")
    add_body_paragraph(
        doc,
        "AnantaSutra shall furnish monthly service-level reports within "
        "ten (10) Business Days of the end of each calendar month, "
        "setting out measured performance against each service level "
        "and any service credits accrued.",
    )

    add_h2(doc, "4. Service Credits as Sole Financial Remedy")
    add_body_paragraph(
        doc,
        "Service credits shall be applied against the Client's next "
        "invoice following the month of accrual and shall be the "
        "Client's sole and exclusive financial remedy for breach of "
        "service levels, save where the breach also constitutes a "
        "material breach of the MSA entitling the Client to "
        "terminate under Clause 16.3.",
    )

    add_h2(doc, "5. Chronic Failure")
    add_body_paragraph(
        doc,
        "The occurrence of any of the following shall constitute a "
        "material breach of the MSA entitling the Client to "
        "terminate the affected Statement of Work for cause: (a) "
        "failure to meet the same service level in three (3) "
        "consecutive calendar months; or (b) failure to meet any "
        "service level in any six (6) out of twelve (12) "
        "consecutive calendar months.",
    )

    # -----------------------------------------------------------------
    # Save
    # -----------------------------------------------------------------
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    size = os.path.getsize(OUTPUT_PATH)
    print(f"Generated: {OUTPUT_PATH}")
    print(f"Size: {size} bytes ({size/1024:.1f} KB)")


if __name__ == "__main__":
    build()
