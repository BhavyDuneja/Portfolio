"""
build_08_service_agreement_awish.py

Generates a SHORT-FORM, PLAIN-LANGUAGE but STRONGLY PRO-ANANTASUTRA
Service Agreement between:

    AnantaSutra (Service Provider)  <->  Awish Clinic (Client)

Output: legal-documents/08_Service_Agreement_Awish_Clinic.docx

This is a deliberately simpler alternative to the full-form Master Service
Agreement in build_04_msa.py. It is tailored to a single, specific
engagement (2 engineers + 1 freelance designer) and is drafted to protect
AnantaSutra: payment by the 20th of each month, suspension/termination for
non-payment, IP passing ONLY on full payment, low liability cap,
non-solicitation with liquidated damages, no refunds, and client-side
indemnity for its own (medical/advertising) content.

Indian-law anchors referenced: Indian Contract Act 1872 (incl. s.27),
Copyright Act 1957, Information Technology Act 2000, DPDP Act 2023,
CGST Act 2017 / Income-tax Act 1961 (GST & TDS), Arbitration &
Conciliation Act 1996, Specific Relief Act 1963, Drugs and Magic Remedies
(Objectionable Advertisements) Act 1954 (clinic advertising compliance).

Uses python-docx v1.2.0.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# -------------------------------------------------------------------
# Constants
# -------------------------------------------------------------------
OUTPUT_PATH = r"c:/Users/duneja8515/Desktop/bhavya/prsnl/sutranet/legal-documents/08_Service_Agreement_Awish_Clinic.docx"
LOGO_PATH = r"c:/Users/duneja8515/Desktop/bhavya/prsnl/sutranet/public/Favicon/logo-nobg.png"

BASE_FONT = "Calibri"
BODY_SIZE = Pt(11)
HEADING1_SIZE = Pt(13)
TITLE_SIZE = Pt(16)

# Service Provider (AnantaSutra) — hardcoded
SP_BRAND = "AnantaSutra"
SP_DESCRIPTION = (
    "AnantaSutra, a business concern owned and represented by its "
    "Co-Founder and Owner, Bhavya Duneja, having its principal "
    "place of business at Delhi, India"
)
SP_CONTACT = "contact@anantasutra.com"
SP_SIGNATORY_NAME = "Bhavya Duneja"
SP_SIGNATORY_DESIGNATION = "Co-Founder & Owner"

# Client (Awish Clinic) — owner Dr. Vijay Kumar, CEO Abhishek Pandey
CLIENT_BRAND = "Awish Clinic"
CLIENT_DESCRIPTION = (
    "Awish Clinic, a healthcare establishment owned by Dr. Vijay Kumar and "
    "represented for the purposes of this Agreement by its Chief Executive "
    "Officer, Mr. Abhishek Pandey, having its principal place of business at "
    "L-5, 5, Pocket L, Sarita Vihar, New Delhi, Delhi 110076, India"
)
CLIENT_SIGNATORY_NAME = "Mr. Abhishek Pandey"
CLIENT_SIGNATORY_DESIGNATION = "Chief Executive Officer"
CLIENT_OWNER_NAME = "Dr. Vijay Kumar"

# Commercials
ENGINEER_COUNT = "two (2)"
ENGINEER_FEE = "Rs. 30,000/- (Rupees Thirty Thousand only)"
ENGINEER_TOTAL = "Rs. 60,000/- (Rupees Sixty Thousand only)"
PAYMENT_DAY = "20th (twentieth)"


# -------------------------------------------------------------------
# Helpers
# -------------------------------------------------------------------
def add_brand_header(doc):
    """Insert centered logo + brand lines at top of page 1. Never crashes."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=Inches(1.2))
    except Exception:
        pass
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("AnantaSutra — अनन्तसूत्र")
    r2.italic = True
    r2.font.size = Pt(10)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("https://anantasutra.com  •  contact@anantasutra.com")
    r3.italic = True
    r3.font.size = Pt(9)


def _simple_field(paragraph, instr):
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = instr
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_end)


def configure_styles(doc):
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = BASE_FONT
    normal.font.size = BODY_SIZE
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    h1 = styles['Heading 1']
    h1.font.name = BASE_FONT
    h1.font.size = HEADING1_SIZE
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(4)
    h1.paragraph_format.keep_with_next = True


def set_page_layout(doc):
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Confidential — AnantaSutra & Awish Clinic    |    Page ")
        run.font.name = BASE_FONT
        run.font.size = Pt(9)
        _simple_field(p, "PAGE")
        mid = p.add_run(" of ")
        mid.font.name = BASE_FONT
        mid.font.size = Pt(9)
        _simple_field(p, "NUMPAGES")


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = TITLE_SIZE
    run.font.name = BASE_FONT
    p.paragraph_format.space_after = Pt(4)
    return p


def add_centered(doc, text, italic=False, size=None, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.name = BASE_FONT
    run.font.size = size or BODY_SIZE
    p.paragraph_format.space_after = Pt(6)
    return p


def add_h1(doc, text):
    return doc.add_paragraph(text, style='Heading 1')


def add_para(doc, text, bold=False, italic=False, justify=True):
    p = doc.add_paragraph()
    p.alignment = (WD_ALIGN_PARAGRAPH.JUSTIFY if justify
                   else WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = BASE_FONT
    run.font.size = BODY_SIZE
    return p


def add_sub(doc, lead, text):
    """A lettered/numbered sub-point, slightly indented, bold lead."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.7)
    lr = p.add_run(lead + "  ")
    lr.bold = True
    lr.font.name = BASE_FONT
    lr.font.size = BODY_SIZE
    r = p.add_run(text)
    r.font.name = BASE_FONT
    r.font.size = BODY_SIZE
    return p


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "000000")
        tcBorders.append(border)
    tcPr.append(tcBorders)


# -------------------------------------------------------------------
# Document construction
# -------------------------------------------------------------------
def build_document():
    doc = Document()
    add_brand_header(doc)
    configure_styles(doc)
    set_page_layout(doc)
    add_footer(doc)

    # ---------------- Title ----------------
    add_title(doc, "SERVICE AGREEMENT")
    add_centered(
        doc,
        "(Software Development, Technical Support & Creative Design Services)",
        italic=True,
    )
    add_centered(
        doc,
        "Between AnantaSutra (Service Provider) and Awish Clinic (Client)",
        italic=True,
    )

    # ---------------- Date / execution ----------------
    add_para(
        doc,
        "This Service Agreement (the \"Agreement\") is made and executed at "
        "Delhi, India on this ______ day of ____________, 20____ (the "
        "\"Effective Date\"),",
    )

    # ---------------- Parties ----------------
    add_h1(doc, "BETWEEN")
    add_para(
        doc,
        SP_DESCRIPTION
        + " (hereinafter referred to as \"AnantaSutra\" or the \"Service "
        "Provider\", which expression shall, unless repugnant to the context, "
        "include its successors and permitted assigns), of the ONE PART;",
    )
    add_h1(doc, "AND")
    add_para(
        doc,
        CLIENT_DESCRIPTION
        + " (hereinafter referred to as \"Awish Clinic\" or the \"Client\", "
        "which expression shall, unless repugnant to the context, include its "
        "successors and permitted assigns), of the OTHER PART.",
    )
    add_para(
        doc,
        "AnantaSutra and the Client are each referred to as a \"Party\" and "
        "together as the \"Parties\".",
    )

    # ---------------- Background ----------------
    add_h1(doc, "BACKGROUND")
    add_para(
        doc,
        "A. AnantaSutra provides software development, website development, "
        "technical support, and creative design (social media, print, and "
        "video) services.",
    )
    add_para(
        doc,
        "B. The Client wishes to engage AnantaSutra to provide a dedicated "
        "team comprising " + ENGINEER_COUNT + " engineers and one (1) "
        "freelance designer, on the terms set out in this Agreement.",
    )
    add_para(
        doc,
        "C. The Parties have agreed the commercial and legal terms recorded "
        "below, which they intend to be binding under the Indian Contract "
        "Act, 1872.",
    )
    add_para(
        doc,
        "NOW THEREFORE, in consideration of the mutual promises below, the "
        "Parties agree as follows:",
        bold=True,
    )

    # ---------------- 1. Definitions ----------------
    add_h1(doc, "1. Definitions")
    add_para(
        doc,
        "In this Agreement: \"Services\" means the services described in "
        "Clause 2; \"Deliverables\" means any work product (including code, "
        "websites, web sections, social media posts, banners, pamphlets, and "
        "videos) created by AnantaSutra for the Client under this Agreement; "
        "\"Fees\" means the amounts payable to AnantaSutra under Clause 4 and "
        "Annexure A; \"Monthly Engagement Fee\" means the fixed monthly fee "
        "for the engineers under Clause 4.1; and \"Rate Card\" means the "
        "designer charges set out in Annexure A. Headings are for convenience "
        "only. \"Including\" means \"including without limitation\".",
    )

    # ---------------- 2. Services and Team ----------------
    add_h1(doc, "2. Services and Team")
    add_sub(
        doc,
        "2.1",
        "AnantaSutra shall make available to the Client a team comprising "
        + ENGINEER_COUNT + " engineers (the \"Engineers\") on a dedicated "
        "monthly basis, and one (1) designer (the \"Designer\") engaged on a "
        "freelance, per-deliverable basis.",
    )
    add_sub(
        doc,
        "2.2",
        "The Engineers shall provide software development, website "
        "development, maintenance, and technical support during AnantaSutra's "
        "standard working hours and days. The Designer shall provide creative "
        "design Deliverables (social media posts, banners, pamphlets, "
        "website sections, and videos) as and when requested by the Client "
        "and accepted by AnantaSutra.",
    )
    add_sub(
        doc,
        "2.3",
        "AnantaSutra shall perform the Services with reasonable skill and "
        "care. AnantaSutra may, at its discretion, substitute or replace any "
        "member of the team with another person of comparable skill, and the "
        "Client shall have no right to select, direct, discipline, or demand "
        "any particular individual.",
    )
    add_sub(
        doc,
        "2.4",
        "The scope of the Services is limited to what is expressly described "
        "in this Agreement and any written brief accepted by AnantaSutra. Any "
        "work outside that scope (including additional engineers, additional "
        "Deliverables beyond the agreed brief, or services requested outside "
        "standard working hours) shall be treated as a separate, additional "
        "engagement and charged separately.",
    )
    add_sub(
        doc,
        "2.5",
        "AnantaSutra does not guarantee any particular business result, "
        "ranking, reach, traffic, conversion, or revenue outcome from the "
        "Services or Deliverables.",
    )

    # ---------------- 3. Term ----------------
    add_h1(doc, "3. Term")
    add_sub(
        doc,
        "3.1",
        "This Agreement shall commence on the Effective Date and continue on "
        "a rolling monthly basis until terminated in accordance with Clause "
        "16.",
    )
    add_sub(
        doc,
        "3.2",
        "The engagement of the Engineers is on a full-calendar-month basis. "
        "Any part-month at commencement or on termination shall be charged "
        "pro-rata only if AnantaSutra agrees in writing; otherwise the full "
        "Monthly Engagement Fee for the relevant month shall be payable.",
    )

    # ---------------- 4. Fees and Charges ----------------
    add_h1(doc, "4. Fees and Charges")
    add_sub(
        doc,
        "4.1  Engineers (fixed monthly).",
        "The Client shall pay a Monthly Engagement Fee of " + ENGINEER_FEE
        + " per Engineer per calendar month, i.e. " + ENGINEER_TOTAL
        + " per calendar month for " + ENGINEER_COUNT + " Engineers. This "
        "fee is a flat retainer payable for each month in full, irrespective "
        "of the number of hours actually used, public holidays, weekly offs, "
        "or reasonable leave taken by the Engineers.",
    )
    add_sub(
        doc,
        "4.2  Designer (per deliverable).",
        "The Designer is charged on a per-deliverable basis at the rates set "
        "out in the Rate Card at Annexure A. There is no minimum monthly "
        "commitment for the Designer; charges accrue only as Deliverables "
        "are produced. Charges for a Deliverable accrue, and become payable, "
        "upon delivery of that Deliverable to the Client, whether or not the "
        "Client subsequently publishes or uses it.",
    )
    add_sub(
        doc,
        "4.3  Taxes.",
        "As at the Effective Date, AnantaSutra is not registered for Goods "
        "and Services Tax (GST), and accordingly no GST is charged on the "
        "Fees. If AnantaSutra becomes liable to register for or to charge GST "
        "or any other tax in respect of the Services, such tax shall be "
        "charged additionally and borne by the Client with effect from the "
        "date such liability arises.",
    )
    add_sub(
        doc,
        "4.4  Third-party costs.",
        "Any third-party or paid assets required for a Deliverable "
        "(including paid stock images, fonts, music, plugins, hosting, "
        "domain, or advertising spend) are not included in the Fees and "
        "shall be paid by the Client, either directly or as a reimbursement "
        "at actuals to AnantaSutra, subject to the Client's prior approval.",
    )
    add_sub(
        doc,
        "4.5  Price revision.",
        "AnantaSutra may revise the Fees and the Rate Card by giving the "
        "Client not less than thirty (30) days' prior written notice. Revised "
        "rates shall apply to all Services and Deliverables provided after "
        "the effective date of the revision.",
    )

    # ---------------- 5. Invoicing and Payment ----------------
    add_h1(doc, "5. Invoicing and Payment")
    add_sub(
        doc,
        "5.1  Payment date.",
        "At the end of each calendar month, AnantaSutra shall raise an "
        "invoice covering (a) the Monthly Engagement Fee for that month, and "
        "(b) the Designer's charges for Deliverables produced during that "
        "month. The Client shall pay each invoice in full on or before the "
        "10th (tenth) day of the following calendar month, by bank transfer "
        "to the account notified by AnantaSutra.",
    )
    add_sub(
        doc,
        "5.2  No deductions.",
        "All payments shall be made in full without any set-off or deduction, "
        "save for any tax the Client is required by law to deduct at source "
        "(TDS), for which the Client shall furnish the relevant certificate "
        "within the statutory time.",
    )
    add_sub(
        doc,
        "5.3  Late payment.",
        "Without prejudice to any other right, any amount not paid by the due "
        "date shall carry interest at one and a half per cent (1.5%) per "
        "month (or part month) from the due date until actual payment.",
    )
    add_sub(
        doc,
        "5.4  Suspension for non-payment.",
        "If any undisputed amount remains unpaid for more than five (5) days "
        "after the due date, AnantaSutra may, without liability and without "
        "prejudice to its other rights, suspend all or part of the Services "
        "(including withdrawing the Engineers and Designer and pausing "
        "access to in-progress Deliverables) until all outstanding amounts, "
        "together with interest, are paid in full. Time lost due to such "
        "suspension shall not reduce the Fees payable.",
    )
    add_sub(
        doc,
        "5.5  Advance.",
        "AnantaSutra may require the first month's Monthly Engagement Fee to "
        "be paid in advance before commencing the Services, and may require "
        "a reasonable advance for any large or asset-heavy Deliverable.",
    )

    # ---------------- 6. Client Responsibilities ----------------
    add_h1(doc, "6. Client Responsibilities")
    add_sub(
        doc,
        "6.1",
        "The Client shall, in a timely manner, provide all content, "
        "materials, information, brand assets, access, approvals, and "
        "feedback that AnantaSutra reasonably requires to perform the "
        "Services, and shall nominate a single point of contact authorised "
        "to approve briefs, Deliverables, and charges.",
    )
    add_sub(
        doc,
        "6.2",
        "AnantaSutra shall not be liable for any delay, additional cost, or "
        "deficiency caused by the Client's failure to provide the above, by "
        "incorrect or incomplete information supplied by the Client, or by "
        "delays in the Client's review or approval. Any timelines are "
        "indicative and conditional on the Client's timely cooperation.",
    )
    add_sub(
        doc,
        "6.3  Content and advertising compliance.",
        "The Client is solely responsible for the accuracy, legality, and "
        "regulatory compliance of all content, claims, and information it "
        "supplies or approves for use in any Deliverable, including all "
        "medical, clinical, treatment, qualification, and advertising claims "
        "relating to Awish Clinic. The Client warrants that such content "
        "complies with all applicable laws, including the Drugs and Magic "
        "Remedies (Objectionable Advertisements) Act, 1954, the Consumer "
        "Protection Act, 2019, applicable medical-council advertising norms, "
        "and the Information Technology Act, 2000. AnantaSutra acts only as a "
        "service provider giving effect to the Client's instructions and "
        "does not verify or endorse such content.",
    )

    # ---------------- 7. Deliverables, Approvals & Revisions ----------------
    add_h1(doc, "7. Deliverables, Approvals and Revisions")
    add_sub(
        doc,
        "7.1  Sections and units.",
        "For website work, a \"section\" means a distinct page or major "
        "content block as identified in the agreed brief, and the per-section "
        "charge in Annexure A is inclusive of images placed within that "
        "section (but not of paid third-party assets under Clause 4.4). For "
        "video work, each video is charged per Annexure A and must not exceed "
        "three (3) minutes; any video longer than three (3) minutes shall be "
        "charged as additional video unit(s) or separately quoted.",
    )
    add_sub(
        doc,
        "7.2  Revisions.",
        "Each Deliverable includes up to two (2) rounds of reasonable "
        "revisions within the originally agreed scope. Further revisions, or "
        "any change to the brief after work has begun, shall be charged "
        "separately at the applicable Rate Card or on a quotation basis.",
    )
    add_sub(
        doc,
        "7.3  Deemed acceptance.",
        "The Client shall review each Deliverable and notify AnantaSutra of "
        "any material non-conformity within five (5) days of delivery. If the "
        "Client does not do so, or uses or publishes the Deliverable, the "
        "Deliverable shall be deemed accepted and the corresponding charge "
        "shall be payable.",
    )

    # ---------------- 8. Intellectual Property ----------------
    add_h1(doc, "8. Intellectual Property")
    add_sub(
        doc,
        "8.1  Ownership on full payment only.",
        "Title to, and intellectual property rights in, a Deliverable shall "
        "pass to the Client only upon AnantaSutra's receipt of full and final "
        "payment of all Fees and charges relating to that Deliverable and of "
        "all other amounts then due under this Agreement. Until such full "
        "payment, all rights in the Deliverables remain the sole property of "
        "AnantaSutra, and the Client has no right or licence to use, copy, "
        "publish, or modify them.",
    )
    add_sub(
        doc,
        "8.2  AnantaSutra retained IP.",
        "Notwithstanding Clause 8.1, AnantaSutra retains all right, title, and "
        "interest in its own pre-existing materials and in any tools, "
        "libraries, frameworks, templates, source-code components, design "
        "elements, methodologies, and know-how used in creating the "
        "Deliverables (\"AnantaSutra Materials\"). To the extent any "
        "AnantaSutra Materials are embedded in a Deliverable, AnantaSutra "
        "grants the Client, on full payment, a non-exclusive, "
        "non-transferable licence to use them solely as part of that "
        "Deliverable for the Client's own business.",
    )
    add_sub(
        doc,
        "8.3  Portfolio rights.",
        "AnantaSutra may, after publication by the Client, display the "
        "non-confidential Deliverables and reference the Client's name and "
        "logo in its portfolio, website, and marketing as an example of its "
        "work, unless the Client objects in writing on reasonable grounds.",
    )

    # ---------------- 9. Confidentiality ----------------
    add_h1(doc, "9. Confidentiality")
    add_para(
        doc,
        "Each Party shall keep confidential all non-public information of the "
        "other Party disclosed in connection with this Agreement, use it only "
        "for performing this Agreement, and not disclose it to third parties "
        "except to its personnel on a need-to-know basis. This obligation "
        "survives for three (3) years after termination, and indefinitely in "
        "respect of trade secrets and patient or personal data. Each Party "
        "shall comply with the Digital Personal Data Protection Act, 2023 in "
        "respect of any personal data it handles.",
    )

    # ---------------- 10. Independent Contractor ----------------
    add_h1(doc, "10. Independent Contractor; No Employment")
    add_para(
        doc,
        "AnantaSutra is an independent service provider. Nothing in this "
        "Agreement creates any employer-employee, partnership, agency, or "
        "joint-venture relationship between the Client and AnantaSutra or any "
        "of its personnel. AnantaSutra remains the sole employer/principal of "
        "its Engineers and Designer and is solely responsible for their "
        "wages, statutory contributions, and taxes. The Client shall not "
        "exercise any employer-style control over AnantaSutra's personnel.",
    )

    # ---------------- 11. Non-Solicitation ----------------
    add_h1(doc, "11. Non-Solicitation of Personnel")
    add_sub(
        doc,
        "11.1",
        "During the Term and for twelve (12) months thereafter, the Client "
        "shall not, directly or indirectly, solicit, hire, engage, or "
        "contract (whether as employee, consultant, or freelancer) any "
        "Engineer, Designer, or other personnel of AnantaSutra who was "
        "involved in providing the Services, whether or not approached by the "
        "Client first.",
    )
    add_sub(
        doc,
        "11.2  Liquidated damages.",
        "The Parties agree that a breach of Clause 11.1 would cause "
        "AnantaSutra loss that is difficult to quantify, and that a genuine "
        "pre-estimate of that loss is an amount equal to twelve (12) times "
        "the monthly fee applicable to the relevant person (and, for the "
        "Designer, twelve (12) times the average monthly charges billed for "
        "that Designer over the preceding three (3) months). The Client shall "
        "pay this amount to AnantaSutra as liquidated damages for each person "
        "so solicited or engaged, in accordance with Section 74 of the Indian "
        "Contract Act, 1872. This Clause does not amount to a restraint of "
        "trade under Section 27 of that Act, being a protection of "
        "AnantaSutra's investment in its workforce.",
    )

    # ---------------- 12. Warranties and Disclaimer ----------------
    add_h1(doc, "12. Warranties and Disclaimer")
    add_sub(
        doc,
        "12.1",
        "AnantaSutra warrants only that the Services will be performed with "
        "reasonable skill and care. Its sole obligation, and the Client's "
        "sole remedy, for breach of this warranty is to re-perform the "
        "affected Services or correct the affected Deliverable at "
        "AnantaSutra's cost.",
    )
    add_sub(
        doc,
        "12.2",
        "Except as expressly stated in this Agreement, and to the maximum "
        "extent permitted by law, AnantaSutra gives no other warranties, "
        "conditions, or representations of any kind, whether express or "
        "implied, including any implied warranty of fitness for a particular "
        "purpose or merchantability, and disclaims all of them.",
    )

    # ---------------- 13. Limitation of Liability ----------------
    add_h1(doc, "13. Limitation of Liability")
    add_sub(
        doc,
        "13.1",
        "To the maximum extent permitted by law, AnantaSutra shall not be "
        "liable for any indirect, incidental, special, or consequential loss, "
        "or for any loss of profit, revenue, goodwill, business, data, or "
        "anticipated savings, even if advised of the possibility.",
    )
    add_sub(
        doc,
        "13.2",
        "AnantaSutra's total aggregate liability arising out of or in "
        "connection with this Agreement, whether in contract, tort, or "
        "otherwise, shall not exceed the total Fees actually paid by the "
        "Client to AnantaSutra in the one (1) month immediately preceding the "
        "event giving rise to the claim.",
    )
    add_sub(
        doc,
        "13.3",
        "Nothing in this Agreement limits any liability that cannot lawfully "
        "be limited, including liability for fraud.",
    )

    # ---------------- 14. Indemnity ----------------
    add_h1(doc, "14. Indemnity by the Client")
    add_para(
        doc,
        "The Client shall indemnify, defend, and hold harmless AnantaSutra "
        "and its personnel against all claims, losses, penalties, costs, and "
        "expenses (including reasonable legal fees) arising out of (a) "
        "content, materials, instructions, claims, or data supplied or "
        "approved by the Client, including any medical, treatment, or "
        "advertising claim relating to Awish Clinic; (b) the Client's breach "
        "of this Agreement or of any applicable law; or (c) the Client's use "
        "of the Deliverables other than as permitted under this Agreement.",
    )

    # ---------------- 15. Suspension ----------------
    add_h1(doc, "15. Suspension")
    add_para(
        doc,
        "In addition to Clause 5.4, AnantaSutra may suspend the Services with "
        "immediate effect on written notice if the Client is in material "
        "breach of this Agreement. Suspension does not relieve the Client of "
        "its payment obligations for the period of suspension where the "
        "suspension arises from the Client's breach.",
    )

    # ---------------- 16. Termination ----------------
    add_h1(doc, "16. Termination")
    add_sub(
        doc,
        "16.1  For convenience.",
        "Either Party may terminate this Agreement by giving the other not "
        "less than thirty (30) days' prior written notice. The Client "
        "remains liable for the full Monthly Engagement Fee for the month in "
        "which the notice period falls and for all Designer charges accrued "
        "up to the effective date of termination.",
    )
    add_sub(
        doc,
        "16.2  For cause.",
        "AnantaSutra may terminate this Agreement immediately by written "
        "notice if the Client (a) fails to pay any undisputed amount within "
        "ten (10) days after the due date; (b) commits any other material "
        "breach not cured within fifteen (15) days of notice; or (c) becomes "
        "insolvent or unable to pay its debts. The Client may terminate "
        "immediately if AnantaSutra commits a material breach not cured "
        "within thirty (30) days of written notice.",
    )

    # ---------------- 17. Effect of Termination ----------------
    add_h1(doc, "17. Effect of Termination")
    add_sub(
        doc,
        "17.1",
        "On termination, the Client shall immediately pay all Fees and "
        "charges accrued or payable up to the effective date of termination, "
        "including the amounts described in Clause 16.1.",
    )
    add_sub(
        doc,
        "17.2  No refunds.",
        "All amounts paid are non-refundable. No credit or refund shall be "
        "due for any unused portion of any month, for any Deliverable already "
        "produced, or for Services already performed.",
    )
    add_sub(
        doc,
        "17.3",
        "Termination does not transfer any intellectual property in respect "
        "of which full payment has not been received (Clause 8.1). Clauses 8, "
        "9, 11, 12, 13, 14, 17, 18, and 19 survive termination.",
    )

    # ---------------- 18. Force Majeure ----------------
    add_h1(doc, "18. Force Majeure")
    add_para(
        doc,
        "Neither Party shall be liable for any failure or delay in "
        "performance (other than a payment obligation) caused by events "
        "beyond its reasonable control, including acts of God, natural "
        "disasters, epidemics or pandemics, war, civil unrest, governmental "
        "orders, strikes, or failure of internet or power infrastructure. If "
        "such an event continues for more than thirty (30) days, either Party "
        "may terminate this Agreement on written notice, without prejudice to "
        "amounts already accrued.",
    )

    # ---------------- 19. Governing Law and Dispute Resolution ----------------
    add_h1(doc, "19. Governing Law and Dispute Resolution")
    add_para(
        doc,
        "This Agreement is governed by the laws of India. The Parties shall "
        "first attempt to resolve any dispute amicably. Failing resolution "
        "within fifteen (15) days, the dispute shall be referred to and "
        "finally resolved by arbitration by a sole arbitrator mutually "
        "appointed by the Parties under the Arbitration and Conciliation Act, "
        "1996. The seat and venue of arbitration shall be Delhi, and the "
        "language shall be English. Subject to arbitration, the courts at "
        "Delhi shall have exclusive jurisdiction. Either Party may seek "
        "interim relief from a court under Section 9 of that Act and under "
        "the Specific Relief Act, 1963.",
    )

    # ---------------- 20. Notices ----------------
    add_h1(doc, "20. Notices")
    add_para(
        doc,
        "Any notice under this Agreement shall be in writing and sent by "
        "email or recognised courier to the contact details of the other "
        "Party set out in this Agreement or notified in writing. Notice by "
        "email is deemed received on the next business day after sending. "
        "Notices to AnantaSutra shall be marked to Bhavya Duneja, "
        "Co-Founder & Owner, at " + SP_CONTACT + ".",
    )

    # ---------------- 21. General ----------------
    add_h1(doc, "21. General")
    add_sub(
        doc,
        "21.1  Entire agreement.",
        "This Agreement (with its Annexure) is the entire agreement between "
        "the Parties on its subject matter and supersedes all prior "
        "discussions and understandings.",
    )
    add_sub(
        doc,
        "21.2  Amendment.",
        "No change to this Agreement is valid unless in writing and signed by "
        "both Parties. Pre-printed terms on any purchase order or other "
        "Client document are rejected and have no effect.",
    )
    add_sub(
        doc,
        "21.3  Assignment.",
        "The Client shall not assign or transfer this Agreement without "
        "AnantaSutra's prior written consent. AnantaSutra may subcontract or "
        "assign to an affiliate or successor.",
    )
    add_sub(
        doc,
        "21.4  Severability and waiver.",
        "If any provision is held invalid, the rest of the Agreement "
        "continues in effect. A delay or failure to enforce any right is not "
        "a waiver of it.",
    )
    add_sub(
        doc,
        "21.5  Electronic execution.",
        "This Agreement may be signed in counterparts and executed "
        "electronically (including by scanned or digital signatures), which "
        "shall be valid and binding under the Information Technology Act, "
        "2000.",
    )

    # ---------------- Signature block ----------------
    doc.add_paragraph("")
    add_para(
        doc,
        "IN WITNESS WHEREOF, the Parties have signed this Agreement on the "
        "date first written above.",
        bold=True,
    )
    doc.add_paragraph("")

    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.autofit = True
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    company_cell = sig_table.rows[0].cells[0]
    client_cell = sig_table.rows[0].cells[1]

    def fill_sig_cell(cell, heading_lines, block_lines):
        cell.text = ""
        for i, line in enumerate(heading_lines):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(line)
            r.bold = True
            r.font.name = BASE_FONT
            r.font.size = BODY_SIZE
        cell.add_paragraph("")
        cell.add_paragraph("_______________________________")
        for line in block_lines:
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(line)
            r.font.name = BASE_FONT
            r.font.size = BODY_SIZE

    fill_sig_cell(
        company_cell,
        ["FOR AND ON BEHALF OF ANANTASUTRA"],
        [
            "Name: " + SP_SIGNATORY_NAME,
            "Designation: " + SP_SIGNATORY_DESIGNATION,
            "Email: " + SP_CONTACT,
            "Place: Delhi, India",
            "Date: ______________________",
        ],
    )
    fill_sig_cell(
        client_cell,
        ["FOR AND ON BEHALF OF AWISH CLINIC"],
        [
            "Name: " + CLIENT_SIGNATORY_NAME,
            "Designation: " + CLIENT_SIGNATORY_DESIGNATION,
            "Signature: ______________________",
            "Place: ______________________",
            "Date: ______________________",
        ],
    )

    doc.add_paragraph("")
    add_para(
        doc,
        "Confirmed and acknowledged by the Owner of Awish Clinic:",
        bold=True,
    )
    add_para(doc, "Name: " + CLIENT_OWNER_NAME + " (Owner, Awish Clinic)")
    add_para(doc, "Signature: ______________________      Date: ______________")

    # =================================================================
    # ANNEXURE A — Rate Card
    # =================================================================
    doc.add_page_break()
    add_title(doc, "ANNEXURE A")
    add_centered(doc, "FEES AND DESIGNER RATE CARD", bold=True, size=Pt(13))

    add_h1(doc, "A. Engineers — Fixed Monthly Fee")
    eng_table = doc.add_table(rows=1, cols=3)
    eng_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    eng_headers = ("Resource", "Rate", "Basis")
    for idx, text in enumerate(eng_headers):
        cell = eng_table.rows[0].cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.name = BASE_FONT
        r.font.size = BODY_SIZE
        set_cell_border(cell)
    eng_rows = [
        ("Engineer (each)", "Rs. 30,000 / month", "Flat monthly retainer, payable in full per month"),
        ("Engineers x 2 (total)", "Rs. 60,000 / month", "Combined monthly fee for the dedicated team"),
    ]
    for row_data in eng_rows:
        row = eng_table.add_row()
        for idx, text in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = BASE_FONT
            r.font.size = BODY_SIZE
            set_cell_border(cell)

    add_h1(doc, "B. Designer — Per-Deliverable Charges")
    des_table = doc.add_table(rows=1, cols=3)
    des_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    des_headers = ("Deliverable", "Charge", "Notes")
    for idx, text in enumerate(des_headers):
        cell = des_table.rows[0].cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.name = BASE_FONT
        r.font.size = BODY_SIZE
        set_cell_border(cell)
    des_rows = [
        ("Post (social media post / banner / pamphlet)",
         "Rs. 200 each",
         "Per individual design unit"),
        ("Website",
         "Rs. 500 per section",
         "Inclusive of images placed within that section"),
        ("Video",
         "Rs. 1,000 per video",
         "Up to 3 (three) minutes; longer videos charged as additional units / separately quoted"),
    ]
    for row_data in des_rows:
        row = des_table.add_row()
        for idx, text in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = BASE_FONT
            r.font.size = BODY_SIZE
            set_cell_border(cell)

    doc.add_paragraph("")
    add_h1(doc, "C. Notes on Charges")
    for note in [
        "AnantaSutra is not currently registered for GST, so no GST is charged on the Fees. If GST becomes applicable to AnantaSutra in future, it will be added (Clause 4.3).",
        "Designer charges accrue per Deliverable upon delivery, whether or not the Deliverable is later published or used (Clause 4.2).",
        "Each Deliverable includes up to two (2) rounds of revisions; further revisions are charged separately (Clause 7.2).",
        "Paid third-party assets (stock images, fonts, music, plugins, hosting, domains, ad spend) are not included and are billed at actuals (Clause 4.4).",
        "Invoices are raised at the end of each month and are due on or before the 10th of the following calendar month (Clause 5.1). Late payment carries interest at 1.5% per month (Clause 5.3).",
        "AnantaSutra may revise these rates on 30 days' written notice (Clause 4.5).",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.9)
        for r0 in p.runs:
            r0.text = ""
        r = p.add_run(note)
        r.font.name = BASE_FONT
        r.font.size = BODY_SIZE

    # Worked example
    doc.add_paragraph("")
    add_h1(doc, "D. Illustrative Monthly Invoice (Example)")
    add_para(
        doc,
        "Example only — actual amounts depend on Deliverables produced in the "
        "month:",
        italic=True,
    )
    ex_table = doc.add_table(rows=1, cols=3)
    ex_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    ex_headers = ("Item", "Calculation", "Amount (Rs.)")
    for idx, text in enumerate(ex_headers):
        cell = ex_table.rows[0].cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.name = BASE_FONT
        r.font.size = BODY_SIZE
        set_cell_border(cell)
    ex_rows = [
        ("Engineers (2)", "2 x 30,000", "60,000"),
        ("Posts (10)", "10 x 200", "2,000"),
        ("Website sections (4)", "4 x 500", "2,000"),
        ("Videos (2)", "2 x 1,000", "2,000"),
        ("Total payable (no GST)", "", "66,000"),
    ]
    for row_data in ex_rows:
        row = ex_table.add_row()
        for idx, text in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.font.name = BASE_FONT
            r.font.size = BODY_SIZE
            if row_data[0].startswith(("Sub-total", "Total")):
                r.bold = True
            set_cell_border(cell)

    # Save
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    out = build_document()
    size_kb = os.path.getsize(out) / 1024.0
    print("AWISH CLINIC SERVICE AGREEMENT COMPLETE")
    print("Output: " + out)
    print("Size: %.1f KB" % size_kb)
