"""
Build Script: Employee Confidentiality, Non-Solicitation & IP-Assignment Agreement
Entity:       AnantaSutra (Delhi, India)
Jurisdiction: Republic of India
Output:       02_Employee_NDA_and_Confidentiality.docx

Generates an execution-ready, India-law-compliant Employee NDA using python-docx.

Design notes (do not remove):
- Post-employment NON-COMPETE is intentionally OMITTED per s.27 Indian Contract
  Act 1872 (Niranjan Shankar Golikari v. Century Spinning; Superintendence Co.
  v. Krishan Murgai; Percept D'Mark v. Zaheer Khan).
- Non-solicitation is retained with a reasonable 12-month tail.
- IP assignment is PRESENT-TENSE ("hereby irrevocably assigns") to overcome the
  s.17(c) Copyright Act 1957 default vis-a-vis independent-contractor scenarios
  and to capture patentable subject matter which is not covered by s.17(c).
- Moral rights waiver follows s.57 Copyright Act 1957 (with the statutory
  limitation that false-attribution / distortion rights cannot be fully waived).
- Data protection aligns with DPDP Act 2023 and s.43A / s.72A IT Act 2000.
- Arbitration seated at New Delhi under the Arbitration & Conciliation Act 1996.
- Stamp duty acknowledgment per Indian Stamp Act 1899 (as applicable to NCT of Delhi), Article 5(j).
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# -----------------------------------------------------------------------------
# Constants
# -----------------------------------------------------------------------------
OUT_PATH = Path(
    r"c:/Users/duneja8515/Desktop/bhavya/prsnl/sutranet/legal-documents/"
    r"02_Employee_NDA_and_Confidentiality.docx"
)

BODY_FONT = "Calibri"
BODY_SIZE = Pt(11)
TITLE_SIZE = Pt(16)

DOC_TITLE = (
    "EMPLOYEE CONFIDENTIALITY, NON-SOLICITATION "
    "& IP-ASSIGNMENT AGREEMENT"
)


# -----------------------------------------------------------------------------
# Low-level formatting helpers
# -----------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Brand header (AnantaSutra) — injected by rebrand
# ---------------------------------------------------------------------------
LOGO_PATH = r"c:/Users/duneja8515/Desktop/bhavya/prsnl/sutranet/public/Favicon/logo-nobg.png"


def add_brand_header(doc):
    """Insert centered logo + brand lines at top of page 1. Never crashes."""
    try:
        from docx.shared import Inches as _Inches, Pt as _Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH as _CENTERED
    except Exception:
        return
    p = doc.add_paragraph()
    p.alignment = _CENTERED.CENTER
    try:
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=_Inches(1.2))
    except Exception:
        pass
    p2 = doc.add_paragraph()
    p2.alignment = _CENTERED.CENTER
    r2 = p2.add_run("AnantaSutra — अनन्तसूत्र")
    r2.italic = True
    r2.font.size = _Pt(10)
    p3 = doc.add_paragraph()
    p3.alignment = _CENTERED.CENTER
    r3 = p3.add_run("https://anantasutra.com  •  contact@anantasutra.com")
    r3.italic = True
    r3.font.size = _Pt(9)

def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:color"), "000000")
            tcBorders.append(border)
    tcPr.append(tcBorders)


def configure_page(document: Document) -> None:
    """A4 with 1-inch margins on all sides."""
    for section in document.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.orientation = WD_ORIENT.PORTRAIT


def configure_base_styles(document: Document) -> None:
    """Set 11pt Calibri, justified, 1.15 line spacing on the Normal style."""
    style = document.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    pf = style.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Heading 1
    h1 = document.styles["Heading 1"]
    h1.font.name = BODY_FONT
    h1.font.size = Pt(13)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True

    # Heading 2
    h2 = document.styles["Heading 2"]
    h2.font.name = BODY_FONT
    h2.font.size = Pt(11)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True


def add_footer(document: Document) -> None:
    """Footer: 'Confidential — AnantaSutra' (left) + page number (right)."""
    for section in document.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.text = ""

        run_left = p.add_run("Confidential \u2014 AnantaSutra")
        run_left.font.size = Pt(9)
        run_left.font.name = BODY_FONT

        # Tab-to-right then PAGE field
        run_tab = p.add_run("\t\t")
        run_tab.font.name = BODY_FONT

        run_page_label = p.add_run("Page ")
        run_page_label.font.size = Pt(9)
        run_page_label.font.name = BODY_FONT

        # Insert PAGE field
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")

        run_page_num = p.add_run()
        run_page_num.font.size = Pt(9)
        run_page_num.font.name = BODY_FONT
        run_page_num._r.append(fld_begin)
        run_page_num._r.append(instr)
        run_page_num._r.append(fld_sep)
        run_page_num._r.append(fld_end)

        run_of = p.add_run(" of ")
        run_of.font.size = Pt(9)
        run_of.font.name = BODY_FONT

        fld_begin2 = OxmlElement("w:fldChar")
        fld_begin2.set(qn("w:fldCharType"), "begin")
        instr2 = OxmlElement("w:instrText")
        instr2.set(qn("xml:space"), "preserve")
        instr2.text = " NUMPAGES "
        fld_sep2 = OxmlElement("w:fldChar")
        fld_sep2.set(qn("w:fldCharType"), "separate")
        fld_end2 = OxmlElement("w:fldChar")
        fld_end2.set(qn("w:fldCharType"), "end")

        run_num_pages = p.add_run()
        run_num_pages.font.size = Pt(9)
        run_num_pages.font.name = BODY_FONT
        run_num_pages._r.append(fld_begin2)
        run_num_pages._r.append(instr2)
        run_num_pages._r.append(fld_sep2)
        run_num_pages._r.append(fld_end2)


# -----------------------------------------------------------------------------
# Content helpers
# -----------------------------------------------------------------------------
def add_title(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = TITLE_SIZE
    run.font.name = BODY_FONT
    p.paragraph_format.space_after = Pt(12)


def add_h1(document: Document, text: str) -> None:
    p = document.add_paragraph(text, style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_h2(document: Document, text: str) -> None:
    p = document.add_paragraph(text, style="Heading 2")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_para(document: Document, text: str, bold: bool = False) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE
    if bold:
        run.bold = True


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE


def add_defined_term(document: Document, term: str, meaning: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_term = p.add_run(f"\u201C{term}\u201D ")
    run_term.bold = True
    run_term.font.name = BODY_FONT
    run_term.font.size = BODY_SIZE
    run_body = p.add_run(f"means {meaning}")
    run_body.font.name = BODY_FONT
    run_body.font.size = BODY_SIZE


def page_break(document: Document) -> None:
    document.add_page_break()


# -----------------------------------------------------------------------------
# Document build
# -----------------------------------------------------------------------------
def build_document() -> Document:
    doc = Document()
    add_brand_header(doc)
    configure_page(doc)
    configure_base_styles(doc)
    add_footer(doc)

    add_title(doc, DOC_TITLE)

    # ------------------------------------------------------------------
    # PARTIES
    # ------------------------------------------------------------------
    add_h1(doc, "1. PARTIES")

    add_para(
        doc,
        "This Employee Confidentiality, Non-Solicitation and Intellectual "
        "Property Assignment Agreement (this \u201CAgreement\u201D) is made "
        "and executed at New Delhi, Delhi, on this [DAY] day of [MONTH], "
        "[YEAR] (the \u201CEffective Date\u201D) by and between:"
    )

    add_para(
        doc,
        "(1) ANANTASUTRA, a business concern operated and represented by its "
        "Founder & CEO, Mr. Himanshu Mishra, carrying on business under the "
        "trade name \u201CAnantaSutra\u201D, having its principal place of "
        "business at Delhi, India, with contact address "
        "contact@anantasutra.com, acting through its Founder & CEO, "
        "Mr. Himanshu Mishra "
        "(hereinafter referred to as the \u201CCompany\u201D or "
        "\u201CAnantaSutra\u201D, which expression shall, unless repugnant "
        "to the context or meaning thereof, include its successors, "
        "assigns, and the person(s) for the time being in control of the "
        "business) of the ONE PART;"
    )

    add_para(
        doc,
        "AND"
    )

    add_para(
        doc,
        "(2) [EMPLOYEE FULL NAME], son/daughter/spouse of [PARENT/SPOUSE "
        "NAME], aged [AGE] years, holding Permanent Account Number [PAN] "
        "and Aadhaar Number [LAST 4 DIGITS MASKED], residing at [RESIDENTIAL "
        "ADDRESS \u2013 PIN], engaged by the Company in the capacity of "
        "[DESIGNATION] in the [DEPARTMENT] department, bearing Employee "
        "Code [EMPLOYEE CODE], reporting to [REPORTING MANAGER NAME, "
        "DESIGNATION], with effect from [DATE OF JOINING] (hereinafter "
        "referred to as the \u201CEmployee\u201D, which expression shall, "
        "unless repugnant to the context, include his/her heirs, legal "
        "representatives and permitted assigns) of the OTHER PART."
    )

    add_para(
        doc,
        "The Company and the Employee are hereinafter referred to "
        "individually as a \u201CParty\u201D and collectively as the "
        "\u201CParties\u201D."
    )

    # ------------------------------------------------------------------
    # RECITALS
    # ------------------------------------------------------------------
    add_h1(doc, "2. RECITALS")

    add_para(
        doc,
        "WHEREAS, the Company is engaged in the business of developing, "
        "licensing and providing software, information technology, data "
        "analytics, product engineering and allied professional services "
        "(the \u201CBusiness\u201D), and in the ordinary course of its "
        "Business develops, uses and holds proprietary technical, "
        "commercial, financial and customer information of significant "
        "value;"
    )
    add_para(
        doc,
        "AND WHEREAS, the Company has offered employment to the Employee "
        "pursuant to an offer letter / employment agreement dated [OFFER "
        "LETTER DATE] (the \u201CEmployment Agreement\u201D), in terms of "
        "which the Employee will, in the course of employment, have access "
        "to Confidential Information and will create, conceive or contribute "
        "to Company Inventions (as hereinafter defined);"
    )
    add_para(
        doc,
        "AND WHEREAS, execution of this Agreement is a condition precedent "
        "to the Employee\u2019s engagement, continuation in employment and "
        "payment of compensation by the Company, and the Employee has "
        "read, understood and voluntarily agreed to be bound by the terms "
        "of this Agreement;"
    )
    add_para(
        doc,
        "NOW, THEREFORE, in consideration of the mutual covenants "
        "contained herein, the offer of employment by the Company, the "
        "compensation, access to Confidential Information and training to "
        "be provided to the Employee, the sufficiency of which is hereby "
        "acknowledged, the Parties hereby agree as follows:"
    )

    # ------------------------------------------------------------------
    # DEFINITIONS
    # ------------------------------------------------------------------
    add_h1(doc, "3. DEFINITIONS AND INTERPRETATION")

    add_para(
        doc,
        "In this Agreement, unless the context otherwise requires, the "
        "following words and expressions shall have the meanings set out "
        "below:"
    )

    add_defined_term(
        doc, "Affiliate",
        "in relation to the Company, any entity that, directly or "
        "indirectly, controls, is controlled by, or is under common control "
        "with the Company, where \u201Ccontrol\u201D means the beneficial "
        "ownership of fifty per cent (50%) or more of the voting equity of "
        "such entity or the ability to direct its management and policies.",
    )

    add_defined_term(
        doc, "Applicable Law",
        "all statutes, enactments, acts of the legislature, ordinances, "
        "rules, regulations, notifications, guidelines, directions, "
        "directives, circulars and orders of any Governmental Authority in "
        "the Republic of India, including without limitation the Indian "
        "Contract Act, 1872; the Specific Relief Act, 1963; the Copyright "
        "Act, 1957; the Patents Act, 1970; the Trade Marks Act, 1999; the "
        "Designs Act, 2000; the Information Technology Act, 2000; the "
        "Digital Personal Data Protection Act, 2023; the Companies Act, "
        "2013; the Delhi Shops and Commercial Establishments Act, "
        "1961; the Arbitration and Conciliation Act, 1996; and the Indian "
        "Stamp Act, 1899 read with the Indian Stamp Act, 1899 (as applicable to the NCT of Delhi).",
    )

    add_defined_term(
        doc, "Company Group",
        "the Company together with each of its Affiliates, and their "
        "respective directors, officers, employees, agents, clients and "
        "representatives.",
    )

    add_defined_term(
        doc, "Company Inventions",
        "any and all inventions, discoveries, improvements, developments, "
        "designs, software, source code, object code, algorithms, "
        "architectures, flowcharts, database structures, schematics, "
        "technical or business methods, know-how, trade secrets, processes, "
        "formulae, works of authorship, documentation, reports, "
        "presentations, data, datasets and other creations of any kind, "
        "whether or not patentable, copyrightable or protectable under any "
        "law, that are conceived, created, developed, reduced to practice, "
        "authored, made or first fixed in any tangible medium by the "
        "Employee, solely or jointly with others, (a) during the course of "
        "and in connection with the employment; (b) using any resources, "
        "facilities, equipment, Confidential Information, time or "
        "materials of the Company Group; or (c) that relate, at the time "
        "of conception or creation, to the Business, research or "
        "development activities of the Company Group or to work performed "
        "by the Employee for the Company Group.",
    )

    add_defined_term(
        doc, "Confidential Information",
        "all non-public information of, or in the possession or control of, "
        "the Company Group or any of its clients, suppliers or business "
        "partners, whether disclosed before, on or after the Effective "
        "Date, in any form (written, oral, visual, electronic, graphic, "
        "digital or otherwise) and whether or not marked or identified as "
        "\u201Cconfidential\u201D, which a reasonable person would "
        "consider to be confidential in the circumstances, including "
        "without limitation: (a) technical information, source code, "
        "object code, architectures, algorithms, models, datasets, test "
        "cases, cloud configurations, infrastructure designs, API designs "
        "and Company Inventions; (b) commercial information, customer and "
        "prospect lists, pricing, discounts, margins, bids, proposals, "
        "responses to requests for proposal, sales and marketing "
        "strategies; (c) financial information, budgets, forecasts, "
        "business plans, payroll, costs and profitability data; (d) "
        "employee, consultant and contractor information, compensation "
        "data and performance evaluations; (e) client-supplied data, "
        "including Personal Data; and (f) all copies, reproductions, "
        "analyses, compilations, studies, derivatives and summaries "
        "thereof.",
    )

    add_defined_term(
        doc, "DPDP Act",
        "the Digital Personal Data Protection Act, 2023, together with "
        "the rules, regulations and notifications framed thereunder, as "
        "amended from time to time.",
    )

    add_defined_term(
        doc, "Moral Rights",
        "any and all rights of paternity, integrity, attribution and "
        "similar rights under Section 57 of the Copyright Act, 1957 and "
        "analogous rights available under Applicable Law or under the law "
        "of any other jurisdiction.",
    )

    add_defined_term(
        doc, "Permitted Purpose",
        "the performance by the Employee of his/her lawful duties and "
        "responsibilities for and on behalf of the Company Group in "
        "accordance with the Employment Agreement, the Company\u2019s "
        "policies and Applicable Law.",
    )

    add_defined_term(
        doc, "Personal Data",
        "\u201Cpersonal data\u201D as defined in Section 2(t) of the DPDP "
        "Act and includes any data about an individual who is identifiable "
        "by or in relation to such data.",
    )

    add_defined_term(
        doc, "Prior Inventions",
        "any inventions, original works of authorship, developments, "
        "improvements, trade secrets, know-how and other intellectual "
        "property that were made, created, conceived, developed or "
        "reduced to practice by the Employee, alone or with others, prior "
        "to the Effective Date and listed in Schedule A to this "
        "Agreement.",
    )

    add_defined_term(
        doc, "Restricted Period",
        "the period commencing on the Effective Date and expiring on the "
        "date falling twelve (12) months after the effective date of "
        "cessation of the Employee\u2019s employment with the Company, "
        "howsoever arising.",
    )

    add_defined_term(
        doc, "Territory",
        "the Republic of India and each other jurisdiction in which the "
        "Company Group conducts or has, during the last twelve (12) "
        "months of the Employee\u2019s employment, actively pursued the "
        "Business.",
    )

    add_defined_term(
        doc, "Third-Party Information",
        "Confidential Information of any client, customer, supplier, "
        "vendor or other third party that has been received by or made "
        "available to the Company Group subject to an obligation of "
        "confidentiality or restricted use.",
    )

    add_defined_term(
        doc, "Trade Secrets",
        "that subset of Confidential Information which derives "
        "independent economic value, actual or potential, from not being "
        "generally known to, and not being readily ascertainable by "
        "proper means by, other persons who can obtain economic value "
        "from its disclosure or use, and which is the subject of "
        "reasonable efforts by the Company Group to maintain its "
        "secrecy.",
    )

    add_h2(doc, "3.2 Interpretation")
    add_para(
        doc,
        "(a) The headings are for convenience only and shall not affect "
        "the construction of this Agreement. (b) Words in the singular "
        "include the plural and vice versa. (c) References to statutes "
        "include subordinate legislation and amendments from time to "
        "time. (d) \u201CIncluding\u201D and \u201Cincludes\u201D are "
        "without limitation. (e) References to clauses and Schedules are "
        "to clauses of, and Schedules to, this Agreement. (f) \u201CINR\u201D "
        "means Indian Rupees, the lawful currency of the Republic of India."
    )

    # ------------------------------------------------------------------
    # CONFIDENTIALITY
    # ------------------------------------------------------------------
    add_h1(doc, "4. CONFIDENTIALITY OBLIGATIONS")

    add_h2(doc, "4.1 Acknowledgment of Proprietary Nature")
    add_para(
        doc,
        "The Employee acknowledges and agrees that the Confidential "
        "Information is the sole and exclusive property of the Company "
        "Group (or of the relevant third party in the case of Third-Party "
        "Information), that it has been developed or acquired at "
        "considerable expense, that it is of substantial commercial value "
        "to the Company Group, and that the unauthorised disclosure or "
        "use thereof would cause irreparable harm to the Company Group."
    )

    add_h2(doc, "4.2 Obligations During Employment")
    add_para(
        doc,
        "During the term of employment, the Employee shall:"
    )
    add_bullet(
        doc,
        "(a) use Confidential Information solely for the Permitted Purpose "
        "and in accordance with the Company\u2019s policies, including "
        "its Information Technology and Acceptable Use Policy, Data "
        "Protection Policy and Information Security Policy (each as "
        "amended from time to time);"
    )
    add_bullet(
        doc,
        "(b) hold all Confidential Information in strict confidence and "
        "exercise at least the same degree of care to protect it from "
        "unauthorised disclosure, copying or misuse as a reasonable person "
        "would exercise to protect his/her own confidential information of "
        "like importance, and in any event no less than a reasonable "
        "standard of care;"
    )
    add_bullet(
        doc,
        "(c) not copy, reproduce, store, transmit, upload, forward or "
        "transfer any Confidential Information to any personal email "
        "account, personal cloud storage (including personal Google "
        "Drive, Gmail, Dropbox, OneDrive or iCloud accounts), personal "
        "code repository (including personal GitHub, GitLab or Bitbucket "
        "accounts) or personally-owned removable media, except with the "
        "prior written consent of the Company;"
    )
    add_bullet(
        doc,
        "(d) not disclose Confidential Information to any person within "
        "the Company Group who does not have a legitimate need-to-know "
        "for the Permitted Purpose, nor to any person outside the Company "
        "Group, except as expressly authorised in writing by the Company;"
    )
    add_bullet(
        doc,
        "(e) promptly notify the Company, in writing, of any actual, "
        "suspected or threatened unauthorised access, use, disclosure, "
        "loss or theft of Confidential Information, in any event within "
        "twenty-four (24) hours of becoming aware of the same, to enable "
        "the Company to comply with its statutory obligations including "
        "under Section 8(6) of the DPDP Act;"
    )
    add_bullet(
        doc,
        "(f) use only Company-provisioned devices, accounts and systems "
        "for processing Confidential Information, and comply with "
        "mandatory mobile device management (MDM) enrolment where "
        "personal devices are permitted under any bring-your-own-device "
        "policy;"
    )
    add_bullet(
        doc,
        "(g) not undertake any external engagement, consultancy, moonlighting "
        "or commercial activity that conflicts with the Employee\u2019s "
        "duties, makes use of Confidential Information or Company "
        "resources, or competes with the Business, except with the prior "
        "written consent of the Company; and"
    )
    add_bullet(
        doc,
        "(h) observe all reasonable physical, technical and administrative "
        "security measures prescribed by the Company, including "
        "clean-desk, least-privilege, multi-factor authentication and "
        "encryption requirements."
    )

    add_h2(doc, "4.3 Obligations After Cessation of Employment")
    add_para(
        doc,
        "Upon and following the cessation of employment (for any reason "
        "whatsoever, whether by resignation, termination, retirement, "
        "abandonment or otherwise), the Employee shall:"
    )
    add_bullet(
        doc,
        "(a) continue to hold all Confidential Information in strict "
        "confidence and not use or disclose any Confidential Information "
        "for any purpose whatsoever;"
    )
    add_bullet(
        doc,
        "(b) not, directly or indirectly, induce, permit or assist any "
        "third party to access, extract, use or disclose any Confidential "
        "Information; and"
    )
    add_bullet(
        doc,
        "(c) comply with the return, non-solicitation and cooperation "
        "obligations set out elsewhere in this Agreement."
    )

    add_h2(doc, "4.4 Survival of Confidentiality")
    add_para(
        doc,
        "The obligations of confidentiality under this Clause 4 shall "
        "survive the cessation of employment and shall continue (a) for "
        "Trade Secrets, in perpetuity, for so long as the information "
        "retains the character of a trade secret; and (b) for all other "
        "Confidential Information, for a period of three (3) years from "
        "the date of cessation of employment."
    )

    # ------------------------------------------------------------------
    # CARVE-OUTS
    # ------------------------------------------------------------------
    add_h1(doc, "5. CARVE-OUTS AND PERMITTED DISCLOSURES")

    add_h2(doc, "5.1 Exclusions from Confidential Information")
    add_para(
        doc,
        "The obligations set out in Clause 4 shall not apply to "
        "information that the Employee can demonstrate by contemporaneous "
        "written evidence:"
    )
    add_bullet(
        doc,
        "(a) is or becomes publicly known and generally available through "
        "no wrongful act, breach or omission of the Employee;"
    )
    add_bullet(
        doc,
        "(b) was rightfully in the Employee\u2019s possession, free of "
        "any obligation of confidentiality, prior to receipt from the "
        "Company Group, and has been expressly disclosed in Schedule A to "
        "this Agreement;"
    )
    add_bullet(
        doc,
        "(c) is lawfully received by the Employee from a third party who, "
        "to the Employee\u2019s knowledge, is not under any duty of "
        "confidentiality to the Company Group; or"
    )
    add_bullet(
        doc,
        "(d) is independently developed by the Employee outside the scope "
        "of employment, without use of or reference to any Confidential "
        "Information, as evidenced by contemporaneous written records."
    )

    add_h2(doc, "5.2 Legally Compelled Disclosure")
    add_para(
        doc,
        "If the Employee is required by any Governmental Authority, court "
        "of competent jurisdiction or subpoena of comparable authority to "
        "disclose any Confidential Information, the Employee shall, to "
        "the extent legally permissible, (a) promptly notify the Company "
        "in writing prior to making any such disclosure; (b) cooperate "
        "with the Company, at the Company\u2019s cost, in seeking a "
        "protective order or other appropriate remedy; and (c) disclose "
        "only that portion of the Confidential Information that is "
        "legally required to be disclosed."
    )

    add_h2(doc, "5.3 Whistleblower and Public-Policy Carve-Out")
    add_para(
        doc,
        "Nothing in this Agreement shall restrict or be construed to "
        "restrict the Employee from: (a) making any disclosure required "
        "or permitted by Applicable Law to any regulatory, investigative "
        "or judicial authority; (b) reporting any suspected violation of "
        "law, including corruption, fraud, money-laundering, tax evasion "
        "or securities-market abuse, to the appropriate Governmental "
        "Authority or through the Company\u2019s whistleblower / vigil "
        "mechanism; or (c) otherwise exercising any right which cannot be "
        "lawfully waived under Applicable Law."
    )

    add_h2(doc, "5.4 POSH Act Carve-Out")
    add_para(
        doc,
        "Notwithstanding anything to the contrary, nothing in this "
        "Agreement shall operate to prevent or restrict any person from "
        "making a complaint, or participating in any proceeding, under "
        "the Sexual Harassment of Women at Workplace (Prevention, "
        "Prohibition and Redressal) Act, 2013 or from providing "
        "information to the Internal Committee constituted under the "
        "said Act."
    )

    # ------------------------------------------------------------------
    # RETURN OF PROPERTY
    # ------------------------------------------------------------------
    add_h1(doc, "6. RETURN OF CONFIDENTIAL INFORMATION AND COMPANY PROPERTY")

    add_para(
        doc,
        "Upon cessation of employment (for any reason) or earlier on the "
        "written request of the Company, the Employee shall promptly, and "
        "in any event within three (3) Working Days:"
    )
    add_bullet(
        doc,
        "(a) return to the Company all Confidential Information, Company "
        "Inventions and documents, files, notes, designs, diagrams, code, "
        "datasets, reports and other materials (in whatever form, "
        "including electronic) containing or derived from Confidential "
        "Information;"
    )
    add_bullet(
        doc,
        "(b) return all Company property in the Employee\u2019s "
        "possession or control, including laptops, desktops, mobile "
        "devices, SIM cards, tokens, access cards, identification badges, "
        "keys, removable media, manuals, credit cards and any other "
        "equipment or materials belonging to the Company Group;"
    )
    add_bullet(
        doc,
        "(c) permanently delete or destroy all electronic copies of "
        "Confidential Information from all personal devices, personal "
        "email accounts, personal cloud-storage accounts and personal "
        "code repositories, and certify such deletion in writing;"
    )
    add_bullet(
        doc,
        "(d) disclose and transfer to the Company all usernames, "
        "passwords, access keys, API keys, tokens, SaaS credentials, "
        "GitHub repository access, AWS / Azure / GCP IAM users, "
        "third-party-tool credentials and any other access credentials "
        "used in the course of employment; and"
    )
    add_bullet(
        doc,
        "(e) execute and deliver the Exit Declaration and Acknowledgment "
        "set out in Schedule B."
    )

    # ------------------------------------------------------------------
    # IP
    # ------------------------------------------------------------------
    add_h1(doc, "7. INTELLECTUAL PROPERTY")

    add_h2(doc, "7.1 Statutory Vesting under Section 17(c) Copyright Act")
    add_para(
        doc,
        "The Employee acknowledges that, in accordance with Section 17(c) "
        "of the Copyright Act, 1957, copyright in any work made by the "
        "Employee in the course of employment under a contract of service "
        "shall, in the absence of any agreement to the contrary, vest in "
        "the Company as first owner. This Clause 7 reinforces, and is "
        "supplemental to, such statutory vesting and extends the same to "
        "all forms of intellectual property, whether or not subsisting in "
        "copyright."
    )

    add_h2(doc, "7.2 Present-Tense Assignment of Company Inventions")
    add_para(
        doc,
        "The Employee hereby irrevocably, unconditionally and absolutely "
        "assigns, transfers and conveys to the Company, with effect from "
        "the date of creation thereof, the entire right, title and "
        "interest, throughout the Territory and for the full term of "
        "protection (including any extensions and renewals) and in "
        "perpetuity where no such term is prescribed, in and to each and "
        "every Company Invention, including all associated intellectual "
        "property rights of any description, including copyright, rights "
        "in computer software, patent rights, rights to apply for patents, "
        "design rights, trade mark rights, database rights, rights in "
        "know-how, rights in confidential information and all other "
        "rights of a similar character, whether registered, unregistered "
        "or capable of registration, and whether existing now or arising "
        "in the future (collectively, the \u201CAssigned Rights\u201D)."
    )

    add_h2(doc, "7.3 Consideration")
    add_para(
        doc,
        "The Parties acknowledge and agree that the salary, allowances, "
        "benefits, training, access to Confidential Information and other "
        "consideration paid or provided by the Company to the Employee "
        "pursuant to the Employment Agreement constitute full, fair and "
        "adequate consideration for the assignment of the Assigned "
        "Rights, and the Employee shall have no further claim, whether "
        "pecuniary or otherwise, in respect thereof."
    )

    add_h2(doc, "7.4 Waiver of Moral Rights")
    add_para(
        doc,
        "To the fullest extent permitted under Section 57 of the Copyright "
        "Act, 1957 and any other Applicable Law, the Employee hereby "
        "irrevocably and unconditionally waives all Moral Rights in and "
        "to the Company Inventions in favour of the Company and its "
        "successors, licensees and assigns. The Employee acknowledges "
        "that the statutory right to restrain false attribution and "
        "derogatory treatment prejudicial to honour or reputation, to "
        "the limited extent such right is not capable of being waived, "
        "shall continue to subsist, but the Employee agrees not to "
        "exercise any such right against any bona fide commercial "
        "exploitation of the Company Inventions by the Company Group or "
        "its licensees."
    )

    add_h2(doc, "7.5 Further Assurance and Assistance")
    add_para(
        doc,
        "The Employee shall, at the Company\u2019s cost and upon the "
        "Company\u2019s request (whether during or after employment), "
        "promptly execute and deliver all further documents (including "
        "deeds of assignment, powers of attorney and declarations), "
        "provide all information, and do all acts and things, as may be "
        "necessary or desirable to: (a) perfect, record or register the "
        "Assigned Rights in the name of the Company or its nominee in "
        "any jurisdiction; (b) apply for, prosecute, obtain, maintain, "
        "defend and enforce any patent, copyright, design, trade mark or "
        "other registration relating to any Company Invention; and (c) "
        "otherwise give full effect to the intent of this Clause 7. If "
        "the Employee fails to execute any such document within fifteen "
        "(15) Working Days of request, the Employee hereby irrevocably "
        "appoints the Company as his/her attorney-in-fact, coupled with "
        "interest, to execute and deliver such documents on his/her "
        "behalf."
    )

    add_h2(doc, "7.6 Prior Inventions")
    add_para(
        doc,
        "The Employee represents and warrants that Schedule A sets out a "
        "true, correct and complete list of all Prior Inventions and of "
        "all pre-existing obligations of confidentiality owed to any "
        "third party. If no such Prior Inventions or obligations are "
        "listed, the Employee represents that none exist. The Employee "
        "shall not incorporate any Prior Invention into any Company "
        "Invention without the prior written consent of the Company; "
        "and if any Prior Invention is so incorporated, the Employee "
        "hereby grants to the Company a non-exclusive, perpetual, "
        "irrevocable, worldwide, royalty-free, fully paid-up, "
        "sub-licensable licence to use, reproduce, modify, distribute and "
        "commercially exploit such Prior Invention solely as part of the "
        "relevant Company Invention."
    )

    add_h2(doc, "7.7 Third-Party Information")
    add_para(
        doc,
        "The Employee shall not, in the course of employment, bring onto "
        "the premises or into the systems of the Company Group any "
        "confidential information or intellectual property of any former "
        "employer or other third party, and shall not use any such "
        "information in the performance of duties hereunder, save with "
        "the written authorisation of both the relevant third party and "
        "the Company."
    )

    # ------------------------------------------------------------------
    # NON-SOLICIT
    # ------------------------------------------------------------------
    add_h1(doc, "8. NON-SOLICITATION")

    add_h2(doc, "8.1 Non-Solicitation of Clients and Customers")
    add_para(
        doc,
        "The Employee agrees that, during the Restricted Period, the "
        "Employee shall not, directly or indirectly, whether on his/her "
        "own account or as an employee, partner, consultant, agent, "
        "shareholder or otherwise, actively solicit, induce or attempt to "
        "induce any person or entity who was a client or customer of the "
        "Company Group, and with whom the Employee had material dealings "
        "or about whom the Employee obtained Confidential Information in "
        "the twelve (12) months immediately preceding the cessation of "
        "employment, to: (a) terminate, reduce or adversely modify its "
        "relationship with the Company Group; or (b) transfer to any "
        "other provider any business that would otherwise reasonably be "
        "expected to be placed with the Company Group."
    )

    add_h2(doc, "8.2 Non-Solicitation of Personnel")
    add_para(
        doc,
        "During the Restricted Period, the Employee shall not, directly "
        "or indirectly, actively solicit, entice or seek to induce any "
        "person who is, or who was at any time during the six (6) months "
        "preceding the cessation of employment, a director, officer, "
        "employee or exclusive consultant of the Company Group, to "
        "terminate his/her engagement with the Company Group. The mere "
        "publication of general recruitment advertisements not "
        "specifically targeted at such persons, and the bona fide hire "
        "of a person who responds to such general advertisement, shall "
        "not constitute a breach of this Clause 8.2."
    )

    add_h2(doc, "8.3 Reasonableness")
    add_para(
        doc,
        "The Employee expressly acknowledges that the restrictions in "
        "this Clause 8 are reasonable in scope, duration and geographic "
        "extent, are necessary for the protection of the legitimate "
        "business interests of the Company Group, and do not amount to a "
        "restraint of trade within the meaning of Section 27 of the "
        "Indian Contract Act, 1872. Each restriction is severable and "
        "independently enforceable."
    )

    # ------------------------------------------------------------------
    # NO NON-COMPETE (drafting notice)
    # ------------------------------------------------------------------
    add_h1(doc, "9. NO POST-EMPLOYMENT NON-COMPETE")

    add_para(
        doc,
        "The Parties acknowledge that, in view of Section 27 of the "
        "Indian Contract Act, 1872, and the settled jurisprudence of the "
        "Hon\u2019ble Supreme Court of India (including in Niranjan "
        "Shankar Golikari v. Century Spinning & Manufacturing Co. Ltd., "
        "AIR 1967 SC 1098; Superintendence Company of India v. Krishan "
        "Murgai, AIR 1980 SC 1717; and Percept D\u2019Mark (India) Pvt. "
        "Ltd. v. Zaheer Khan, (2006) 4 SCC 227), no restraint is imposed "
        "upon the Employee from engaging in any trade, profession or "
        "business of any kind after the cessation of employment. "
        "Nothing in this Agreement shall be construed as a post-"
        "employment covenant against competition. The Employee\u2019s "
        "only post-employment restraints are those of confidentiality "
        "(Clause 4) and non-solicitation (Clause 8), each of which is "
        "narrowly tailored and reasonable."
    )

    # ------------------------------------------------------------------
    # GARDEN LEAVE
    # ------------------------------------------------------------------
    add_h1(doc, "10. GARDEN LEAVE")

    add_para(
        doc,
        "The Company may, at its sole option, by written notice to the "
        "Employee, require the Employee, during all or any part of the "
        "applicable notice period specified in the Employment Agreement, "
        "to: (a) cease performing any or all of his/her duties; (b) not "
        "attend the Company\u2019s premises; (c) not contact any client, "
        "supplier, employee or consultant of the Company Group, save as "
        "expressly authorised by the Company; and (d) otherwise remain "
        "available to the Company (\u201CGarden Leave\u201D). During "
        "Garden Leave, the Employee shall (i) continue to receive full "
        "salary and contractual benefits; (ii) remain an employee of the "
        "Company and bound by all obligations under this Agreement and "
        "the Employment Agreement, including the duties of fidelity and "
        "confidentiality; and (iii) not commence any other employment or "
        "engagement. Garden Leave is, for the avoidance of doubt, a "
        "restraint applicable during, and not after, employment."
    )

    # ------------------------------------------------------------------
    # DATA PROTECTION
    # ------------------------------------------------------------------
    add_h1(doc, "11. DATA PROTECTION (DPDP ACT 2023)")

    add_h2(doc, "11.1 Handling of Personal Data")
    add_para(
        doc,
        "The Employee acknowledges that, in the course of employment, "
        "he/she may Process Personal Data on behalf of the Company (in "
        "its capacity as a Data Fiduciary or a Data Processor, as "
        "applicable) under the DPDP Act. The Employee shall:"
    )
    add_bullet(
        doc,
        "(a) Process Personal Data only for the Permitted Purpose and "
        "only in accordance with the written instructions of the Company "
        "and Applicable Law;"
    )
    add_bullet(
        doc,
        "(b) implement and maintain reasonable security safeguards to "
        "protect Personal Data against unauthorised or unlawful Processing, "
        "accidental loss, destruction, damage, alteration or disclosure, "
        "including those prescribed by the Company under its Information "
        "Security Policy and by reference to the \u201Creasonable security "
        "practices and procedures\u201D contemplated by Section 43A of the "
        "Information Technology Act, 2000;"
    )
    add_bullet(
        doc,
        "(c) not transfer, export or permit any onward disclosure of "
        "Personal Data to any third party or to any personal device, "
        "account or location outside the Company\u2019s approved systems;"
    )
    add_bullet(
        doc,
        "(d) promptly, and in any event within twenty-four (24) hours of "
        "becoming aware, notify the Company in writing of any Personal "
        "Data Breach (as defined under the DPDP Act) or any suspected "
        "breach, to enable the Company to discharge its reporting "
        "obligation to the Data Protection Board of India under Section "
        "8(6) of the DPDP Act; and"
    )
    add_bullet(
        doc,
        "(e) cooperate with the Company in responding to requests from "
        "Data Principals, regulators and law-enforcement agencies."
    )

    add_h2(doc, "11.2 Criminal and Civil Liability")
    add_para(
        doc,
        "The Employee acknowledges that wrongful disclosure of Personal "
        "Data or Confidential Information in breach of a lawful contract "
        "may constitute an offence under Section 72A of the Information "
        "Technology Act, 2000, and may also give rise to civil liability "
        "for the Company under Section 43A thereof and under the DPDP "
        "Act."
    )

    add_h2(doc, "11.3 Processing of Employee\u2019s Own Personal Data")
    add_para(
        doc,
        "The Employee consents to the collection and Processing by the "
        "Company of the Employee\u2019s Personal Data (including "
        "identity, contact, financial, health, performance and background-"
        "verification data) for purposes of human-resources administration, "
        "compliance with Applicable Law, statutory filings, payroll, "
        "benefits administration, internal audit, and legitimate business "
        "operations, in accordance with the Company\u2019s Privacy Notice "
        "communicated separately under the DPDP Act."
    )

    # ------------------------------------------------------------------
    # REMEDIES
    # ------------------------------------------------------------------
    add_h1(doc, "12. REMEDIES")

    add_h2(doc, "12.1 Acknowledgment of Irreparable Harm")
    add_para(
        doc,
        "The Employee acknowledges that any breach or threatened breach "
        "of Clauses 4 (Confidentiality), 6 (Return of Property), 7 "
        "(Intellectual Property) or 8 (Non-Solicitation) would cause "
        "immediate, substantial and irreparable harm to the Company "
        "Group for which monetary damages alone would not be an adequate "
        "remedy."
    )

    add_h2(doc, "12.2 Injunctive Relief and Specific Relief")
    add_para(
        doc,
        "Accordingly, the Company shall be entitled, without the need to "
        "furnish security and without prejudice to any other right or "
        "remedy, to seek and obtain interim, interlocutory and perpetual "
        "injunctive relief, specific performance and other equitable "
        "relief under Sections 10, 14 and 36 to 42 of the Specific "
        "Relief Act, 1963, from any court of competent jurisdiction or "
        "from the arbitral tribunal, to restrain any such breach or "
        "threatened breach, in addition to monetary damages."
    )

    add_h2(doc, "12.3 Liquidated Damages")
    add_para(
        doc,
        "Without prejudice to the Company\u2019s right to claim actual "
        "damages, in the event of a material breach of Clause 4 or Clause "
        "8, the Employee shall pay to the Company a sum equivalent to "
        "three (3) months\u2019 gross fixed salary last drawn as "
        "liquidated damages, being a genuine pre-estimate of loss within "
        "the meaning of Section 74 of the Indian Contract Act, 1872, and "
        "not a penalty. The Parties acknowledge that, where actual loss "
        "can be proved, the Company may elect to claim such actual loss "
        "in lieu of, or in addition to (but without double recovery in "
        "respect of the same loss), the liquidated sum."
    )

    add_h2(doc, "12.4 Set-Off against Final Settlement")
    add_para(
        doc,
        "The Company reserves the right to set off, against the "
        "Employee\u2019s full and final settlement dues (subject to "
        "statutory protections), any amounts due from the Employee on "
        "account of breach of this Agreement, recovery of training "
        "costs actually incurred, unreturned Company property, or "
        "damages, without prejudice to any further right of recovery."
    )

    # ------------------------------------------------------------------
    # TERM & SURVIVAL
    # ------------------------------------------------------------------
    add_h1(doc, "13. TERM AND SURVIVAL")

    add_para(
        doc,
        "This Agreement shall take effect on the Effective Date and "
        "shall continue in force during the term of employment and "
        "thereafter in accordance with this Clause 13. The obligations "
        "set out in Clauses 4 (Confidentiality), 6 (Return of Property), "
        "7 (Intellectual Property), 8 (Non-Solicitation), 11 (Data "
        "Protection), 12 (Remedies), 14 (Dispute Resolution) and 15 "
        "(Governing Law) shall survive the cessation of employment for "
        "the periods specified therein or, where no period is specified, "
        "in perpetuity. The assignment of Company Inventions under "
        "Clause 7 is irrevocable and shall survive termination "
        "indefinitely."
    )

    # ------------------------------------------------------------------
    # DISPUTE RESOLUTION
    # ------------------------------------------------------------------
    add_h1(doc, "14. DISPUTE RESOLUTION")

    add_h2(doc, "14.1 Amicable Resolution")
    add_para(
        doc,
        "The Parties shall first endeavour, in good faith, to resolve "
        "any dispute, controversy or claim arising out of or in "
        "connection with this Agreement (a \u201CDispute\u201D) through "
        "mutual discussion between the Employee and the Head of Human "
        "Resources of the Company, to be completed within thirty (30) "
        "days of a written notice of Dispute."
    )

    add_h2(doc, "14.2 Arbitration")
    add_para(
        doc,
        "Any Dispute not so resolved shall be referred to and finally "
        "resolved by arbitration under the Arbitration and Conciliation "
        "Act, 1996 (as amended from time to time). The arbitral tribunal "
        "shall consist of a sole arbitrator, to be appointed by the "
        "Company and consented to in writing by the Employee (such "
        "consent not to be unreasonably withheld), failing which by the "
        "competent court in accordance with Section 11 of the said "
        "Act. The seat of arbitration shall be New Delhi, Delhi, "
        "India, and the venue shall also be New Delhi unless otherwise "
        "agreed. The language of the arbitration shall be English. The "
        "award shall be final and binding on the Parties."
    )

    add_h2(doc, "14.3 Interim Relief (Section 9) Carve-Out")
    add_para(
        doc,
        "Notwithstanding Clause 14.2, either Party may, at any time "
        "before or during the arbitration proceedings, apply to the "
        "competent court in New Delhi under Section 9 of the Arbitration "
        "and Conciliation Act, 1996 for interim measures of protection, "
        "including injunctive relief, without any such application "
        "being deemed a waiver of or inconsistent with the agreement to "
        "arbitrate."
    )

    # ------------------------------------------------------------------
    # GOVERNING LAW
    # ------------------------------------------------------------------
    add_h1(doc, "15. GOVERNING LAW AND JURISDICTION")

    add_para(
        doc,
        "This Agreement shall be governed by, and construed in "
        "accordance with, the laws of the Republic of India. Subject to "
        "Clause 14 (Dispute Resolution), the courts at New Delhi, "
        "Delhi shall have exclusive jurisdiction to the exclusion of "
        "all other courts in respect of any matter arising out of, or in "
        "connection with, this Agreement."
    )

    # ------------------------------------------------------------------
    # BOILERPLATE
    # ------------------------------------------------------------------
    add_h1(doc, "16. GENERAL PROVISIONS")

    add_h2(doc, "16.1 Notices")
    add_para(
        doc,
        "Any notice under this Agreement shall be in writing and shall "
        "be served by hand delivery, registered post with acknowledgment "
        "due, recognised courier, or email (with read-receipt or "
        "confirmation of delivery) to the address of the receiving "
        "Party set out above (or to such other address as is notified "
        "in writing). Notice by email shall be deemed received on "
        "successful transmission; by hand or courier, on the day of "
        "delivery; and by registered post, on the fifth (5th) Working "
        "Day after dispatch."
    )

    add_h2(doc, "16.2 Assignment")
    add_para(
        doc,
        "The Employee shall not assign, transfer or sub-contract any of "
        "his/her rights or obligations under this Agreement, the "
        "services under it being personal in nature. The Company may "
        "assign this Agreement to any of its Affiliates or to any "
        "successor in a business transfer, merger, amalgamation or "
        "similar transaction, without the Employee\u2019s consent."
    )

    add_h2(doc, "16.3 Waiver")
    add_para(
        doc,
        "No waiver of any provision of this Agreement shall be "
        "effective unless in writing and signed by the waiving Party. "
        "No failure or delay by the Company in exercising any right or "
        "remedy shall operate as a waiver thereof, nor shall any single "
        "or partial exercise preclude any further exercise of the same "
        "or any other right or remedy."
    )

    add_h2(doc, "16.4 Severability")
    add_para(
        doc,
        "If any provision of this Agreement is held by a court of "
        "competent jurisdiction or arbitral tribunal to be invalid, "
        "void or unenforceable (including by reason of Section 27 of "
        "the Indian Contract Act, 1872), such provision shall be "
        "severed and the remaining provisions shall continue in full "
        "force and effect. The Parties shall negotiate in good faith "
        "to replace the severed provision with a valid and enforceable "
        "provision that most closely reflects the original commercial "
        "intent."
    )

    add_h2(doc, "16.5 Entire Agreement")
    add_para(
        doc,
        "This Agreement, together with the Employment Agreement and "
        "the Company\u2019s policies referred to herein, constitutes "
        "the entire agreement between the Parties with respect to its "
        "subject matter and supersedes all prior negotiations, "
        "understandings and agreements (whether oral or written). In "
        "the event of any conflict between this Agreement and the "
        "Employment Agreement, the provision more protective of "
        "Confidential Information, Company Inventions and Company "
        "Group interests shall prevail."
    )

    add_h2(doc, "16.6 Amendments")
    add_para(
        doc,
        "No amendment or modification of this Agreement shall be "
        "effective unless made in writing and signed by both Parties."
    )

    add_h2(doc, "16.7 Counterparts and Electronic Execution")
    add_para(
        doc,
        "This Agreement may be executed in any number of counterparts, "
        "each of which when executed shall constitute an original, and "
        "all of which taken together shall constitute one and the same "
        "instrument. The Parties agree that execution by electronic "
        "signature, including Aadhaar e-Sign and equivalent mechanisms "
        "recognised under the Information Technology Act, 2000, shall "
        "be deemed valid and binding, and a scanned or electronic copy "
        "of an executed counterpart shall have the same evidentiary "
        "value as the original."
    )

    add_h2(doc, "16.8 Relationship of Parties")
    add_para(
        doc,
        "Nothing in this Agreement shall be construed to create any "
        "partnership, joint venture or agency relationship between the "
        "Parties beyond the employer-employee relationship established "
        "by the Employment Agreement."
    )

    add_h2(doc, "16.9 Subject to Labour Legislation")
    add_para(
        doc,
        "The provisions of this Agreement shall apply subject to, and "
        "shall not be construed to waive or dilute, any rights or "
        "protections conferred on the Employee by applicable labour "
        "legislation, including the Delhi Shops and Commercial "
        "Establishments Act, 1961; the Industrial Disputes Act, 1947 "
        "(to the extent the Employee is a \u201Cworkman\u201D "
        "thereunder); and the Code on Wages, 2019, the Industrial "
        "Relations Code, 2020, the Code on Social Security, 2020 and "
        "the Occupational Safety, Health and Working Conditions Code, "
        "2020 (as and when notified)."
    )

    # ------------------------------------------------------------------
    # STAMP DUTY
    # ------------------------------------------------------------------
    add_h1(doc, "17. STAMP DUTY")

    add_para(
        doc,
        "The Parties acknowledge that this Agreement has been executed "
        "in New Delhi, Delhi and has been duly stamped under Article "
        "5(j) of the Schedule to the Indian Stamp Act, 1899 (as applicable to the NCT of Delhi) (being "
        "the article applicable to \u201Cagreements not otherwise "
        "provided for\u201D). The cost of stamp duty shall be borne by "
        "the Company. The Parties acknowledge that an insufficiently "
        "stamped instrument is inadmissible in evidence under Section "
        "35 of the Indian Stamp Act, 1899 until duty and any penalty "
        "are paid, and shall ensure that any required uplift is "
        "discharged promptly."
    )

    # ------------------------------------------------------------------
    # EMPLOYEE REPRESENTATIONS
    # ------------------------------------------------------------------
    add_h1(doc, "18. EMPLOYEE REPRESENTATIONS AND WARRANTIES")

    add_para(
        doc,
        "The Employee represents and warrants to the Company that:"
    )
    add_bullet(
        doc,
        "(a) the Employee has full legal capacity to enter into and "
        "perform this Agreement, and the execution hereof does not "
        "breach any obligation (contractual or otherwise) owed to any "
        "third party, including any former employer;"
    )
    add_bullet(
        doc,
        "(b) the Employee has carefully read and understood this "
        "Agreement, has had the opportunity to take independent legal "
        "advice, and is signing voluntarily and without any coercion, "
        "undue influence or misrepresentation;"
    )
    add_bullet(
        doc,
        "(c) Schedule A contains a complete list of all Prior Inventions "
        "and pre-existing confidentiality obligations; and"
    )
    add_bullet(
        doc,
        "(d) the Employee has provided true, complete and accurate "
        "information in all onboarding and background-verification "
        "documentation and consents to the Company verifying such "
        "information in accordance with the DPDP Act."
    )

    # ------------------------------------------------------------------
    # EXECUTION BLOCK
    # ------------------------------------------------------------------
    add_h1(doc, "19. EXECUTION")

    add_para(
        doc,
        "IN WITNESS WHEREOF, the Parties hereto have set their hands on "
        "this Agreement on the day, month and year first hereinabove "
        "written, in the presence of the witnesses named below."
    )

    # Signature 2-column table for Company and Employee
    table = doc.add_table(rows=6, cols=2)
    table.autofit = True

    # Row 0 - headers
    table.cell(0, 0).text = "For and on behalf of\nANANTASUTRA"
    table.cell(0, 1).text = "EMPLOYEE"

    # Row 1 - signature space
    table.cell(1, 0).text = "\n\nSignature: _______________________"
    table.cell(1, 1).text = "\n\nSignature: _______________________"

    # Row 2 - name
    table.cell(2, 0).text = "Name: Mr. Himanshu Mishra"
    table.cell(2, 1).text = "Name: [EMPLOYEE FULL NAME]"

    # Row 3 - designation
    table.cell(3, 0).text = "Designation: Founder & CEO"
    table.cell(3, 1).text = "Employee Code: [EMPLOYEE CODE]"

    # Row 4 - date/place
    table.cell(4, 0).text = "Date: [DATE]   Place: Delhi"
    table.cell(4, 1).text = "Date: [DATE]   Place: Delhi, India"

    # Row 5 - contact email / PAN
    table.cell(5, 0).text = "Email: contact@anantasutra.com"
    table.cell(5, 1).text = "PAN: [PAN NUMBER]"

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = BODY_FONT
                    run.font.size = BODY_SIZE

    add_para(doc, " ")

    # HR Countersign
    add_h2(doc, "19.1 HR Countersignature")
    add_para(
        doc,
        "Countersigned by the Human Resources department of the Company "
        "for record and onboarding compliance:"
    )
    add_para(
        doc,
        "Signature: _______________________    Name: [HR REPRESENTATIVE "
        "NAME]    Designation: [HR DESIGNATION]    Date: [DATE]"
    )

    # Witnesses
    add_h2(doc, "19.2 Witnesses")

    wit = doc.add_table(rows=4, cols=2)
    wit.cell(0, 0).text = "WITNESS 1"
    wit.cell(0, 1).text = "WITNESS 2"
    wit.cell(1, 0).text = "\n\nSignature: _______________________"
    wit.cell(1, 1).text = "\n\nSignature: _______________________"
    wit.cell(2, 0).text = "Name: [WITNESS 1 NAME]"
    wit.cell(2, 1).text = "Name: [WITNESS 2 NAME]"
    wit.cell(3, 0).text = "Address: [WITNESS 1 ADDRESS]"
    wit.cell(3, 1).text = "Address: [WITNESS 2 ADDRESS]"
    for row in wit.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = BODY_FONT
                    run.font.size = BODY_SIZE

    # ------------------------------------------------------------------
    # SCHEDULE A
    # ------------------------------------------------------------------
    page_break(doc)
    add_title(doc, "SCHEDULE A")
    add_h1(
        doc,
        "PRIOR INVENTIONS AND PRE-EXISTING CONFIDENTIALITY OBLIGATIONS",
    )

    add_para(
        doc,
        "This Schedule forms an integral part of the Employee "
        "Confidentiality, Non-Solicitation and IP-Assignment Agreement "
        "between AnantaSutra and the Employee dated "
        "[EFFECTIVE DATE]. It is to be completed by the Employee before "
        "signature of the Agreement. If no disclosures are made, the "
        "Employee shall write \u201CNIL\u201D in each section and "
        "initial."
    )

    add_h2(doc, "Part A \u2014 Prior Inventions")

    pia = doc.add_table(rows=5, cols=4)
    headers_a = ["S. No.", "Title / Description of Prior Invention",
                 "Date of Creation", "Rights Retained / Third Parties Involved"]
    for i, h in enumerate(headers_a):
        cell = pia.cell(0, i)
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = BODY_FONT
                run.font.size = BODY_SIZE
    for r in range(1, 5):
        pia.cell(r, 0).text = str(r)
        for c in range(1, 4):
            pia.cell(r, c).text = ""

    add_h2(doc, "Part B \u2014 Pre-Existing Confidentiality Obligations")

    pib = doc.add_table(rows=5, cols=3)
    headers_b = [
        "S. No.",
        "Counterparty (Former Employer / Client / Other)",
        "Nature and Duration of Obligation",
    ]
    for i, h in enumerate(headers_b):
        cell = pib.cell(0, i)
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = BODY_FONT
                run.font.size = BODY_SIZE
    for r in range(1, 5):
        pib.cell(r, 0).text = str(r)
        for c in range(1, 3):
            pib.cell(r, c).text = ""

    add_para(doc, " ")
    add_para(
        doc,
        "I certify that the information provided in this Schedule A is "
        "true, complete and correct to the best of my knowledge and "
        "belief, and I acknowledge that any inaccuracy or omission may "
        "amount to a material breach of the Agreement."
    )
    add_para(
        doc,
        "Employee Signature: _______________________"
    )
    add_para(
        doc,
        "Name: [EMPLOYEE FULL NAME]    Date: [DATE]"
    )

    # ------------------------------------------------------------------
    # SCHEDULE B
    # ------------------------------------------------------------------
    page_break(doc)
    add_title(doc, "SCHEDULE B")
    add_h1(doc, "EXIT DECLARATION AND ACKNOWLEDGMENT")

    add_para(
        doc,
        "This Exit Declaration is to be executed by the Employee on the "
        "last working day with AnantaSutra, pursuant to "
        "Clause 6 of the Employee Confidentiality, Non-Solicitation and "
        "IP-Assignment Agreement dated [EFFECTIVE DATE]."
    )

    add_h2(doc, "Part A \u2014 Declaration by the Employee")

    add_para(
        doc,
        "I, [EMPLOYEE FULL NAME], Employee Code [EMPLOYEE CODE], "
        "[DESIGNATION], [DEPARTMENT], whose last working day with "
        "AnantaSutra is [LAST WORKING DAY], hereby declare "
        "and certify as follows:"
    )
    add_bullet(
        doc,
        "(a) I have returned all Confidential Information, documents, "
        "files, source code, datasets and other materials (in any form "
        "whatsoever) belonging to or relating to the Company Group or "
        "its clients that are or were in my possession or control;"
    )
    add_bullet(
        doc,
        "(b) I have returned all Company property, including laptops, "
        "mobile devices, SIM cards, access cards, keys, tokens, "
        "credit/charge cards and other equipment listed in the "
        "attached clearance checklist;"
    )
    add_bullet(
        doc,
        "(c) I have permanently deleted all Confidential Information "
        "from my personal devices, personal email accounts, personal "
        "cloud-storage accounts, personal code repositories and any "
        "other location outside the Company\u2019s systems, and no "
        "copies (whether in whole or in part, whether digital or "
        "physical) remain in my possession or control;"
    )
    add_bullet(
        doc,
        "(d) I have disclosed to the Company all usernames, passwords, "
        "API keys, SaaS credentials, repository access, cloud IAM "
        "users and other access credentials used by me in the course "
        "of employment;"
    )
    add_bullet(
        doc,
        "(e) I acknowledge that my obligations of confidentiality "
        "under Clause 4, non-solicitation under Clause 8, intellectual "
        "property assignment under Clause 7, and data protection under "
        "Clause 11 of the Agreement continue to bind me after the "
        "cessation of my employment in accordance with their "
        "respective terms; and"
    )
    add_bullet(
        doc,
        "(f) I have not misappropriated, and will not disclose or use, "
        "any Confidential Information or Trade Secrets of the Company "
        "Group for any purpose whatsoever."
    )

    add_h2(doc, "Part B \u2014 Clearance Checklist")

    check = doc.add_table(rows=9, cols=3)
    chk_headers = ["S. No.", "Item", "Returned / Confirmed (Y/N)"]
    for i, h in enumerate(chk_headers):
        cell = check.cell(0, i)
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = BODY_FONT
                run.font.size = BODY_SIZE
    items = [
        "Laptop / desktop and charger",
        "Mobile device / SIM card",
        "ID card / access card / biometric enrolment",
        "All physical documents, notebooks and drawings",
        "Credentials / SaaS logins / repository access",
        "Deletion of Confidential Information from personal devices",
        "Cooperation in transition / handover",
        "Outstanding advances / expense reports settled",
    ]
    for i, it in enumerate(items, start=1):
        check.cell(i, 0).text = str(i)
        check.cell(i, 1).text = it
        check.cell(i, 2).text = ""

    add_para(doc, " ")

    add_h2(doc, "Part C \u2014 Signatures")

    sig = doc.add_table(rows=4, cols=2)
    sig.cell(0, 0).text = "EMPLOYEE"
    sig.cell(0, 1).text = "FOR ANANTASUTRA"
    sig.cell(1, 0).text = "\n\nSignature: _______________________"
    sig.cell(1, 1).text = "\n\nSignature: _______________________"
    sig.cell(2, 0).text = "Name: [EMPLOYEE FULL NAME]"
    sig.cell(2, 1).text = "Name: Mr. Himanshu Mishra"
    sig.cell(3, 0).text = "Date: [LAST WORKING DAY]"
    sig.cell(3, 1).text = "Designation: Founder & CEO    Date: [DATE]"

    for row in sig.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = BODY_FONT
                    run.font.size = BODY_SIZE

    return doc


def main() -> None:
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(str(OUT_PATH))
    size_kb = OUT_PATH.stat().st_size / 1024.0
    print(f"Wrote: {OUT_PATH}")
    print(f"Size : {size_kb:.2f} KB")


if __name__ == "__main__":
    main()
