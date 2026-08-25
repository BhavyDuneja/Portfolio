"""
Build a SHORT-FORM Employment Letter for Mr. Lakshay Chauhan (AnantaSutra).

Generates: legal-documents/05a_Simple_Employment_Lakshay_Chauhan.docx

This is a deliberately lightweight ~2-4 page letter appropriate for a junior
Rs. 10,000/month hire at AnantaSutra (a pre-incorporation business concern
operated by Mr. Himanshu Mishra, Founder & CEO). It is NOT a replacement for
the full-form Employment Agreement in build_05_employment.py; it is a plain,
short alternative.

Key compliance touchpoints referenced (not exhaustive):
- Indian Contract Act 1872 (incl. s.27, hence NO post-employment non-compete)
- Delhi Shops and Establishments Act 1954 (as and when applicable)
- Payment of Wages Act 1936 / Code on Wages 2019
- Copyright Act 1957 (work-for-hire / moral rights waiver scope)
- POSH Act 2013 (acknowledgment only)
- Arbitration & Conciliation Act 1996 (disputes)
- EPF Act 1952 / ESI Act 1948 / Gratuity Act 1972 (applicable when
  statutory thresholds/registrations are triggered)
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# -------------------------------------------------------------------
# Constants
# -------------------------------------------------------------------

OUTPUT_PATH = r"c:/Users/duneja8515/Desktop/bhavya/prsnl/sutranet/legal-documents/05a_Simple_Employment_Lakshay_Chauhan.docx"
LOGO_PATH = r"c:/Users/duneja8515/Desktop/bhavya/prsnl/sutranet/public/Favicon/logo-nobg.png"

BASE_FONT = "Calibri"
BODY_SIZE = Pt(11)
HEADING1_SIZE = Pt(13)
TITLE_SIZE = Pt(16)

# Hardcoded employee details
EMPLOYEE_NAME = "Mr. Lakshay Chauhan"
# NOTE: Default suggestion is "Associate"; user did not confirm, so kept as placeholder.
EMPLOYEE_DESIGNATION = "[DESIGNATION]"  # suggested default: "Associate"
EMPLOYEE_START_DATE = "[DATE OF JOINING]"
EMPLOYEE_GROSS_MONTHLY = "Rs. 10,000/- (Rupees Ten Thousand only)"
EMPLOYEE_ANNUAL_CTC = "Rs. 1,20,000/- (Rupees One Lakh Twenty Thousand only)"
PLACE_OF_WORK = "Delhi, India"

# Employer side (hardcoded — no placeholders)
EMPLOYER_BRAND = "AnantaSutra"
EMPLOYER_DESCRIPTION = (
    "AnantaSutra, a business concern operated and represented by its Founder "
    "& Chief Executive Officer, Mr. Himanshu Mishra, having its principal "
    "place of business at Delhi, India"
)
EMPLOYER_CONTACT = "contact@anantasutra.com"
EMPLOYER_SIGNATORY_NAME = "Mr. Himanshu Mishra"
EMPLOYER_SIGNATORY_DESIGNATION = "Founder & CEO"


# -------------------------------------------------------------------
# Helpers (logo header reused in style from build_05_employment.py)
# -------------------------------------------------------------------


def add_brand_header(doc):
    """Insert centered logo + brand lines at top of page 1. Never crashes."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=Inches(1.2))
    except Exception:
        pass
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("AnantaSutra — अनन्तसूत्र")
    r2.italic = True
    r2.font.size = Pt(10)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("https://anantasutra.com  •  contact@anantasutra.com")
    r3.italic = True
    r3.font.size = Pt(9)


def _simple_field(paragraph, instr):
    """Insert a Word field (e.g. PAGE, NUMPAGES) into a paragraph."""
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = instr
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_end)


def configure_styles(doc):
    styles = doc.styles

    normal = styles['Normal']
    normal.font.name = BASE_FONT
    normal.font.size = BODY_SIZE
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    h1 = styles['Heading 1']
    h1.font.name = BASE_FONT
    h1.font.size = HEADING1_SIZE
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(4)
    h1.paragraph_format.keep_with_next = True


def set_page_layout(doc):
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Confidential — AnantaSutra    |    Page ")
        run.font.name = BASE_FONT
        run.font.size = Pt(9)
        _simple_field(p, "PAGE")
        mid = p.add_run(" of ")
        mid.font.name = BASE_FONT
        mid.font.size = Pt(9)
        _simple_field(p, "NUMPAGES")


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = TITLE_SIZE
    run.font.name = BASE_FONT
    p.paragraph_format.space_after = Pt(10)


def add_h1(doc, text):
    return doc.add_paragraph(text, style='Heading 1')


def add_para(doc, text, bold=False, italic=False, justify=True):
    p = doc.add_paragraph()
    p.alignment = (WD_ALIGN_PARAGRAPH.JUSTIFY if justify
                   else WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = BASE_FONT
    run.font.size = BODY_SIZE
    return p


# -------------------------------------------------------------------
# Document construction
# -------------------------------------------------------------------

def build_document():
    doc = Document()
    add_brand_header(doc)
    configure_styles(doc)
    set_page_layout(doc)
    add_footer(doc)

    # ---------------- Title ----------------
    add_title(doc, "EMPLOYMENT LETTER")
    add_para(
        doc,
        "(Short-Form Employment Agreement)",
        italic=True,
        justify=False,
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---------------- Date & Parties ----------------
    add_para(
        doc,
        "Date: " + EMPLOYEE_START_DATE,
    )
    add_para(
        doc,
        "To,",
    )
    add_para(
        doc,
        EMPLOYEE_NAME + ",",
    )
    add_para(
        doc,
        "Delhi, India.",
    )

    add_h1(doc, "1. Parties")
    add_para(
        doc,
        "This Employment Letter (\"Letter\" or \"Agreement\") is issued by "
        + EMPLOYER_DESCRIPTION
        + " (the \"Employer\" or \"AnantaSutra\"), and is addressed to "
        + EMPLOYEE_NAME
        + ", residing in Delhi, India (the \"Employee\"). The Employer and "
        "the Employee are referred to individually as a \"Party\" and "
        "collectively as the \"Parties\"."
    )

    # ---------------- Appointment ----------------
    add_h1(doc, "2. Appointment")
    add_para(
        doc,
        "AnantaSutra is pleased to offer, and the Employee accepts, the "
        "position of " + EMPLOYEE_DESIGNATION + " with AnantaSutra. The "
        "Employee shall report to the Founder & CEO or such other person "
        "as the Founder & CEO may nominate from time to time. The Employee "
        "shall devote his full working time, attention, and skill to "
        "AnantaSutra on an exclusive basis during working hours, and shall "
        "not engage in any other employment, business, or consulting "
        "arrangement during working hours without the prior written consent "
        "of AnantaSutra."
    )

    # ---------------- Commencement & Probation ----------------
    add_h1(doc, "3. Commencement and Probation")
    add_para(
        doc,
        "The Employee's employment shall commence on " + EMPLOYEE_START_DATE
        + " (\"Start Date\"). The first three (3) months from the Start "
        "Date shall be a probationary period, during which either Party "
        "may terminate this Letter by giving seven (7) days' written "
        "notice, without assigning any reason. Upon successful completion "
        "of probation, the Employee shall be deemed confirmed in writing "
        "by AnantaSutra. AnantaSutra may, at its sole discretion, extend "
        "the probationary period once by up to three (3) additional months."
    )

    # ---------------- Place of Work ----------------
    add_h1(doc, "4. Place of Work")
    add_para(
        doc,
        "The Employee's place of work shall be " + PLACE_OF_WORK
        + ". AnantaSutra may, at its discretion, permit the Employee to "
        "work remotely (work-from-home) or on a hybrid basis, and may "
        "also require the Employee to travel within India for business "
        "purposes with reasonable prior notice."
    )

    # ---------------- Working Hours ----------------
    add_h1(doc, "5. Working Hours")
    add_para(
        doc,
        "The Employee's standard working hours shall be nine (9) hours "
        "per day, six (6) days per week, as per the working schedule "
        "communicated by AnantaSutra. Working hours, weekly offs, and "
        "overtime (if any) shall be regulated in accordance with the "
        "Delhi Shops and Establishments Act, 1954, and any other "
        "applicable labour law, as and when such laws become applicable "
        "to AnantaSutra."
    )

    # ---------------- Compensation ----------------
    add_h1(doc, "6. Compensation")
    add_para(
        doc,
        "The Employee shall be paid a gross salary of "
        + EMPLOYEE_GROSS_MONTHLY + " per month, aggregating to an annual "
        "Cost-to-Company (CTC) of " + EMPLOYEE_ANNUAL_CTC + ". Salary "
        "shall be payable on or before the seventh (7th) day of the "
        "following calendar month, by bank transfer to the Employee's "
        "designated bank account. Tax Deducted at Source (TDS) and any "
        "other statutory deductions, if and when applicable, shall be "
        "deducted at source. No separate variable pay, bonus, or "
        "incentive shall be payable unless specifically communicated to "
        "the Employee in writing by AnantaSutra."
    )
    add_para(
        doc,
        "Note: At the current annual CTC of Rs. 1,20,000/-, no income "
        "tax is payable under the Income-tax Act, 1961, as the amount is "
        "below the basic exemption limit of Rs. 2,50,000/- per annum. "
        "The Employee shall, however, be solely responsible for filing "
        "any applicable personal tax returns.",
        italic=True,
    )
    add_para(
        doc,
        "Statutory benefits such as Employees' Provident Fund (EPF), "
        "Employees' State Insurance (ESI), and Gratuity shall become "
        "applicable as and when the relevant statutory thresholds are "
        "triggered and AnantaSutra obtains the necessary registrations "
        "under the applicable laws."
    )

    # ---------------- Leave ----------------
    add_h1(doc, "7. Leave and Holidays")
    add_para(
        doc,
        "The Employee shall be entitled to one (1) paid casual leave per "
        "completed calendar month of service, accruing on the last day "
        "of each month, aggregating to twelve (12) paid casual leaves "
        "per annum. Unavailed leaves shall lapse at the end of each "
        "calendar year and shall not be encashable. The Employee shall "
        "additionally be entitled to public holidays as per the "
        "AnantaSutra holiday calendar notified from time to time. Any "
        "additional leave, paid or unpaid, shall be granted at the sole "
        "discretion of the Founder & CEO."
    )

    # ---------------- Confidentiality ----------------
    add_h1(doc, "8. Confidentiality")
    add_para(
        doc,
        "The Employee shall, during and after the term of employment, "
        "keep strictly confidential all non-public information of "
        "AnantaSutra, including without limitation business plans, "
        "customer and supplier lists, pricing, source code, technology, "
        "financial information, and any other information designated as "
        "confidential (\"Confidential Information\"). The Employee shall "
        "not use or disclose any Confidential Information except in the "
        "proper performance of duties or with the prior written consent "
        "of AnantaSutra. This obligation shall continue for a period of "
        "three (3) years from the date of cessation of employment, and "
        "shall continue in perpetuity in respect of trade secrets."
    )

    # ---------------- IP ----------------
    add_h1(doc, "9. Intellectual Property")
    add_para(
        doc,
        "All inventions, works of authorship, designs, source code, "
        "documentation, data, and other materials created, conceived, or "
        "first reduced to practice by the Employee in the course of "
        "employment, whether during or outside working hours and whether "
        "alone or with others (\"Work Product\"), shall be the exclusive "
        "property of AnantaSutra. The Employee hereby assigns, with "
        "present-tense effect, all right, title, and interest (including "
        "all intellectual property rights) in the Work Product to "
        "AnantaSutra, and waives all moral rights in the Work Product to "
        "the fullest extent permitted under the Copyright Act, 1957. The "
        "Employee shall execute such further documents as AnantaSutra "
        "may reasonably request to give effect to this clause."
    )

    # ---------------- Non-Solicitation ----------------
    add_h1(doc, "10. Non-Solicitation")
    add_para(
        doc,
        "For a period of twelve (12) months following the cessation of "
        "employment for any reason, the Employee shall not, directly or "
        "indirectly, (a) solicit, entice, or attempt to solicit any "
        "employee, contractor, or consultant of AnantaSutra to leave "
        "AnantaSutra, or (b) solicit or divert the business of any "
        "customer, client, or prospective client of AnantaSutra with "
        "whom the Employee had material dealings during the twelve (12) "
        "months preceding cessation. For the avoidance of doubt, this "
        "Agreement does not impose any post-employment non-compete "
        "restriction, in view of Section 27 of the Indian Contract Act, "
        "1872."
    )

    # ---------------- Code of Conduct ----------------
    add_h1(doc, "11. Code of Conduct and Compliance")
    add_para(
        doc,
        "The Employee shall at all times comply with (a) AnantaSutra's "
        "internal policies and code of conduct, as amended from time to "
        "time; (b) the Sexual Harassment of Women at Workplace "
        "(Prevention, Prohibition and Redressal) Act, 2013 (\"POSH "
        "Act\") and AnantaSutra's POSH policy; (c) anti-bribery and "
        "anti-corruption norms, including the Prevention of Corruption "
        "Act, 1988; and (d) all other applicable laws. Any breach of "
        "this clause shall constitute cause for termination."
    )

    # ---------------- Termination ----------------
    add_h1(doc, "12. Termination")
    add_para(
        doc,
        "After confirmation, either Party may terminate this Agreement "
        "by giving thirty (30) days' prior written notice to the other "
        "Party, or by paying/forfeiting one (1) month's gross salary in "
        "lieu of notice. AnantaSutra may terminate this Agreement with "
        "immediate effect, without notice or payment in lieu, for cause, "
        "including but not limited to misconduct, fraud, breach of "
        "confidentiality, insubordination, unauthorised absence, or "
        "breach of any material term of this Agreement. Upon termination "
        "for any reason, the Employee shall promptly return all "
        "AnantaSutra property and shall receive full and final settlement "
        "of dues within the timelines prescribed by applicable law."
    )

    # ---------------- Governing Law ----------------
    add_h1(doc, "13. Governing Law and Dispute Resolution")
    add_para(
        doc,
        "This Agreement shall be governed by and construed in accordance "
        "with the laws of India. Any dispute, controversy, or claim "
        "arising out of or in connection with this Agreement shall be "
        "referred to and finally resolved by arbitration by a sole "
        "arbitrator, to be mutually appointed by the Parties, under the "
        "Arbitration and Conciliation Act, 1996. The seat and venue of "
        "arbitration shall be New Delhi, and the language of arbitration "
        "shall be English. Nothing in this clause shall preclude either "
        "Party from approaching the statutory forums or authorities "
        "having exclusive jurisdiction over employment disputes under "
        "applicable Indian labour law."
    )

    # ---------------- Acknowledgment ----------------
    add_h1(doc, "14. Acknowledgment and Acceptance")
    add_para(
        doc,
        "By signing below, the Employee acknowledges that he has read "
        "and understood the terms of this Agreement, has had an "
        "opportunity to seek independent advice, and accepts the terms "
        "of employment set out herein."
    )

    # ---------------- Signature block ----------------
    doc.add_paragraph("")
    add_para(doc, "IN WITNESS WHEREOF, the Parties have executed this "
                  "Employment Letter on the date first written above.")
    doc.add_paragraph("")

    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.autofit = True
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    company_cell = sig_table.rows[0].cells[0]
    employee_cell = sig_table.rows[0].cells[1]

    def fill_sig_cell(cell, heading_lines, block_lines):
        cell.text = ""
        for i, line in enumerate(heading_lines):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(line)
            r.bold = True
            r.font.name = BASE_FONT
            r.font.size = BODY_SIZE
        cell.add_paragraph("")
        cell.add_paragraph("_______________________________")
        for line in block_lines:
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(line)
            r.font.name = BASE_FONT
            r.font.size = BODY_SIZE

    fill_sig_cell(
        company_cell,
        ["FOR AND ON BEHALF OF ANANTASUTRA"],
        [
            "Name: " + EMPLOYER_SIGNATORY_NAME,
            "Designation: " + EMPLOYER_SIGNATORY_DESIGNATION,
            "Email: " + EMPLOYER_CONTACT,
            "Place: Delhi, India",
            "Date: ______________________",
        ],
    )

    fill_sig_cell(
        employee_cell,
        ["ACCEPTED BY THE EMPLOYEE"],
        [
            "Name: " + EMPLOYEE_NAME,
            "Signature: ______________________",
            "Place: Delhi, India",
            "Date: ______________________",
        ],
    )

    # Save
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    out = build_document()
    size_kb = os.path.getsize(out) / 1024.0
    assert size_kb > 100, (
        "Output DOCX is smaller than expected (%.1f KB); logo may have failed to embed."
        % size_kb
    )
    print("SIMPLE LAKSHAY LETTER COMPLETE")
    print("Output: " + out)
    print("Size: %.1f KB" % size_kb)
