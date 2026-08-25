"""
Build script for Document 07 — Data Processing Agreement (DPDP Act 2023).

Generates an execution-ready DPA for AnantaSutra (New Delhi, Delhi)
compliant with the Digital Personal Data Protection Act, 2023 and aligned with
residual obligations under the IT Act 2000 (s.43A, s.72A) and the SPDI Rules
2011.

Primary structure: PROCESSOR-SIDE DPA (AnantaSutra = Data Processor; Client =
Data Fiduciary), with explicit role-mapping notes so the same template can be
executed standalone where AnantaSutra is the Data Fiduciary and the counterparty
is a Sub-Processor / Vendor.

Output: 07_Data_Processing_Agreement_DPDP.docx
"""

from __future__ import annotations

import os
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, Inches, RGBColor


# --------------------------------------------------------------------------- #
# Configuration
# --------------------------------------------------------------------------- #

OUTPUT_DIR = r"c:/Users/duneja8515/Desktop/bhavya/prsnl/sutranet/legal-documents"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "07_Data_Processing_Agreement_DPDP.docx")

BODY_FONT = "Calibri"
BODY_SIZE = Pt(11)
LINE_SPACING = 1.15


# --------------------------------------------------------------------------- #
# Low-level helpers
# --------------------------------------------------------------------------- #


# ---------------------------------------------------------------------------
# Brand header (AnantaSutra) — injected by rebrand
# ---------------------------------------------------------------------------
LOGO_PATH = r"c:/Users/duneja8515/Desktop/bhavya/prsnl/sutranet/public/Favicon/logo-nobg.png"


def add_brand_header(doc):
    """Insert centered logo + brand lines at top of page 1. Never crashes."""
    try:
        from docx.shared import Inches as _Inches, Pt as _Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH as _CENTERED
    except Exception:
        return
    p = doc.add_paragraph()
    p.alignment = _CENTERED.CENTER
    try:
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=_Inches(1.2))
    except Exception:
        pass
    p2 = doc.add_paragraph()
    p2.alignment = _CENTERED.CENTER
    r2 = p2.add_run("AnantaSutra — अनन्तसूत्र")
    r2.italic = True
    r2.font.size = _Pt(10)
    p3 = doc.add_paragraph()
    p3.alignment = _CENTERED.CENTER
    r3 = p3.add_run("https://anantasutra.com  •  contact@anantasutra.com")
    r3.italic = True
    r3.font.size = _Pt(9)

def _set_page(doc: Document) -> None:
    """A4 paper with 1-inch margins on every section."""
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def _set_base_style(doc: Document) -> None:
    """Configure the Normal style as 11pt Calibri, justified, 1.15 spacing."""
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    pf = normal.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(6)

    # Heading styles
    for lvl, size, bold in (("Heading 1", 14, True), ("Heading 2", 12, True),
                            ("Heading 3", 11, True)):
        style = doc.styles[lvl]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True


def _add_page_number_field(paragraph) -> None:
    """Inject a PAGE field into a paragraph."""
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def _build_footer(doc: Document) -> None:
    """Footer: 'Confidential — AnantaSutra' left, page number right."""
    for section in doc.sections:
        footer = section.footer
        # Clear any default paragraph content
        para = footer.paragraphs[0]
        para.text = ""
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        run = para.add_run("Confidential — AnantaSutra")
        run.font.name = BODY_FONT
        run.font.size = Pt(9)
        run.font.italic = True

        para.add_run("\t\t")
        pg_run = para.add_run("Page ")
        pg_run.font.name = BODY_FONT
        pg_run.font.size = Pt(9)
        _add_page_number_field(para)


def _p(doc: Document, text: str, *, bold: bool = False,
       italic: bool = False, align=None, size: int | None = None,
       space_after: int | None = None):
    """Add a body paragraph."""
    para = doc.add_paragraph()
    para.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if space_after is not None:
        para.paragraph_format.space_after = Pt(space_after)
    return para


def _h1(doc: Document, text: str):
    h = doc.add_heading(text, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def _h2(doc: Document, text: str):
    h = doc.add_heading(text, level=2)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def _h3(doc: Document, text: str):
    h = doc.add_heading(text, level=3)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def _bullet(doc: Document, text: str):
    p = doc.add_paragraph(text, style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def _sub_clause(doc: Document, number: str, text: str):
    """Numbered sub-clause like '4.1 ...' indented and justified."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.left_indent = Cm(0.6)
    run_num = para.add_run(f"{number}  ")
    run_num.bold = True
    para.add_run(text)
    return para


def _page_break(doc: Document) -> None:
    doc.add_page_break()


def _shade_cell(cell, color_hex: str = "D9D9D9") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, *, bold: bool = False,
                   italic: bool = False, size: int = 11):
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def _two_col_table(doc: Document, rows: list[tuple[str, str]],
                   header: tuple[str, str] | None = None,
                   col_widths: tuple[float, float] = (5.5, 10.5)):
    """Render a two-column table with optional header."""
    total_rows = len(rows) + (1 if header else 0)
    table = doc.add_table(rows=total_rows, cols=2)
    table.style = "Light Grid Accent 1"
    table.autofit = False

    for row in table.rows:
        row.cells[0].width = Cm(col_widths[0])
        row.cells[1].width = Cm(col_widths[1])

    idx = 0
    if header:
        _set_cell_text(table.rows[0].cells[0], header[0], bold=True)
        _set_cell_text(table.rows[0].cells[1], header[1], bold=True)
        _shade_cell(table.rows[0].cells[0])
        _shade_cell(table.rows[0].cells[1])
        idx = 1

    for i, (left, right) in enumerate(rows):
        _set_cell_text(table.rows[idx + i].cells[0], left, bold=True)
        _set_cell_text(table.rows[idx + i].cells[1], right)
    return table


# --------------------------------------------------------------------------- #
# Content — Title, Parties, Recitals
# --------------------------------------------------------------------------- #

def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DATA PROCESSING AGREEMENT (DPDP Act 2023)")
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = BODY_FONT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(
        "Executed under the Digital Personal Data Protection Act, 2023 "
        "and the Indian Contract Act, 1872"
    )
    sub_run.italic = True
    sub_run.font.size = Pt(11)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    n_run = note.add_run(
        "Primary structure: Processor-side DPA "
        "(AnantaSutra as Data Processor). "
        "Role-mapping notes included for use as a standalone Fiduciary↔Sub-Processor DPA."
    )
    n_run.italic = True
    n_run.font.size = Pt(9)
    note.paragraph_format.space_after = Pt(12)


def add_parties_and_recitals(doc: Document) -> None:
    _h1(doc, "1.  Parties")

    _p(doc,
       "This Data Processing Agreement (\"DPA\" or \"Agreement\") is made and "
       "entered into at New Delhi, Delhi, on the [DATE OF EXECUTION] "
       "(\"Effective Date\"),")

    _p(doc,
       "BY AND BETWEEN:",
       bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)

    _p(doc,
       "(1)  ANANTASUTRA, a business concern operated and represented by "
       "its Founder & CEO, Mr. Himanshu Mishra, carrying on business under "
       "the trade name \"AnantaSutra\", having its principal place of "
       "business at Delhi, India, with contact address "
       "contact@anantasutra.com, represented by Mr. Himanshu Mishra, "
       "Founder & CEO (hereinafter referred to as \"AnantaSutra\" or the "
       "\"Company\", which expression shall, unless repugnant to the "
       "context or meaning thereof, include its successors, assigns, and "
       "the person(s) for the time being in control of the business);")

    _p(doc, "AND", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4)

    _p(doc,
       "(2)  [COUNTERPARTY FULL LEGAL NAME], a [COMPANY / LLP / PARTNERSHIP "
       "/ PROPRIETORSHIP] incorporated/registered under the laws of [INDIA / "
       "JURISDICTION], having its [registered/principal] office at [ADDRESS], "
       "represented by [AUTHORISED SIGNATORY NAME, DESIGNATION] (hereinafter "
       "referred to as the \"Counterparty\", which expression shall, unless "
       "repugnant to the context or meaning thereof, include its successors "
       "and permitted assigns).")

    _p(doc,
       "AnantaSutra and the Counterparty are hereinafter referred to individually "
       "as a \"Party\" and collectively as the \"Parties\".")

    _h2(doc, "1A.  Role Allocation")

    _p(doc,
       "For the purposes of this Agreement and the Digital Personal Data "
       "Protection Act, 2023 (\"DPDP Act\"):")

    _bullet(doc,
            "Where this DPA is executed as an addendum to the Master Services "
            "Agreement dated [MSA DATE] under which AnantaSutra provides services "
            "to the Counterparty (\"Underlying Agreement\"), the Counterparty "
            "shall be the \"Data Fiduciary\" and AnantaSutra shall be the "
            "\"Data Processor\" in relation to the Personal Data processed "
            "under the Underlying Agreement. [DEFAULT — select this box: ☒]")

    _bullet(doc,
            "Where this DPA is executed standalone between AnantaSutra and a "
            "vendor or sub-processor appointed by AnantaSutra, AnantaSutra shall be "
            "the \"Data Fiduciary\" and the Counterparty shall be the "
            "\"Data Processor\" or \"Sub-Processor\" (as the case may be). "
            "[ALTERNATIVE — select this box: ☐]")

    _p(doc,
       "For GDPR cross-reference (where the Counterparty is established in or "
       "processes personal data of Data Principals located in the European "
       "Economic Area / United Kingdom), the term \"Data Fiduciary\" "
       "corresponds to \"Controller\" and \"Data Processor\" corresponds to "
       "\"Processor\" under Regulation (EU) 2016/679 (\"GDPR\") and the UK "
       "GDPR. Where EU Standard Contractual Clauses are required, they shall "
       "be executed as a separate schedule to this DPA and shall prevail in "
       "respect of transfers to which they apply.",
       italic=True)

    _h1(doc, "2.  Recitals")

    _p(doc,
       "A.  The Parties have entered into, or propose to enter into, the "
       "Underlying Agreement pursuant to which one Party (the \"Fiduciary\") "
       "engages the other Party (the \"Processor\") to perform services "
       "involving the processing of Personal Data of identifiable "
       "individuals (\"Data Principals\") on the Fiduciary's behalf.")

    _p(doc,
       "B.  Section 8(5) of the DPDP Act requires that a Data Fiduciary may "
       "engage, appoint, use or otherwise involve a Data Processor only "
       "under a valid contract. The Parties intend this Agreement to satisfy "
       "that statutory requirement and to record their respective rights, "
       "duties, and liabilities under the DPDP Act and allied law.")

    _p(doc,
       "C.  The Parties further intend that this Agreement shall, to the "
       "extent applicable, address residual obligations under Section 43A "
       "and Section 72A of the Information Technology Act, 2000, the "
       "Information Technology (Reasonable Security Practices and Procedures "
       "and Sensitive Personal Data or Information) Rules, 2011 (\"SPDI "
       "Rules 2011\"), and such sectoral regulatory requirements as may "
       "apply to the Parties.")

    _p(doc,
       "D.  The Parties acknowledge that the DPDP Rules, 2025 (\"DPDP "
       "Rules\") are being notified in phases by the Central Government and "
       "agree that, where specific operational timelines, forms or "
       "procedures are prescribed under the DPDP Rules, the Parties shall "
       "comply with such timelines, forms and procedures as they stand "
       "notified from time to time.")

    _p(doc,
       "NOW, THEREFORE, in consideration of the mutual covenants and promises "
       "herein contained, and for other good and valuable consideration the "
       "receipt and adequacy of which are hereby acknowledged, the Parties "
       "agree as follows:",
       bold=True)


# --------------------------------------------------------------------------- #
# Definitions
# --------------------------------------------------------------------------- #

def add_definitions(doc: Document) -> None:
    _h1(doc, "3.  Definitions and Interpretation")

    _p(doc,
       "3.1  Unless the context otherwise requires, capitalised terms used in "
       "this Agreement shall have the meanings set out below. Terms used and "
       "not defined herein shall have the meanings ascribed to them in the "
       "DPDP Act and, where applicable, the DPDP Rules.")

    defs = [
        ("\"Applicable Law\"",
         "means all statutes, rules, regulations, notifications, orders, "
         "directions, judgments and decrees of the Republic of India, "
         "including the DPDP Act, the DPDP Rules, the Information Technology "
         "Act, 2000 and the SPDI Rules 2011, and any sectoral regulatory "
         "requirements applicable to a Party."),
        ("\"Consent Manager\"",
         "has the meaning given in Section 2(g) read with Section 6(7) of "
         "the DPDP Act."),
        ("\"Cross-Border Transfer\"",
         "means the transfer of Personal Data by the Processor to any "
         "location outside the territory of India, including transfer to an "
         "Affiliate, a Sub-Processor, a cloud infrastructure provider, or "
         "any third party situated outside India."),
        ("\"Data Fiduciary\"",
         "has the meaning given in Section 2(i) of the DPDP Act, being any "
         "person who, alone or in conjunction with other persons, determines "
         "the purpose and means of Processing of Personal Data."),
        ("\"Data Principal\"",
         "has the meaning given in Section 2(j) of the DPDP Act, being the "
         "individual to whom the Personal Data relates, and where such "
         "individual is a child, includes the parents or lawful guardian, "
         "and where such individual is a person with disability, includes "
         "her lawful guardian, acting on her behalf."),
        ("\"Data Processor\"",
         "has the meaning given in Section 2(k) of the DPDP Act, being any "
         "person who processes Personal Data on behalf of a Data Fiduciary."),
        ("\"Data Protection Board\" or \"DPB\"",
         "means the Data Protection Board of India established under "
         "Section 18 of the DPDP Act."),
        ("\"DPDP Act\"",
         "means the Digital Personal Data Protection Act, 2023 (Act 22 of "
         "2023), as amended, modified, re-enacted or supplemented from "
         "time to time."),
        ("\"DPDP Rules\"",
         "means the rules made under Section 40 of the DPDP Act, including "
         "the Digital Personal Data Protection Rules, 2025 (as notified "
         "from time to time, in whole or in part)."),
        ("\"Effective Date\"",
         "means the date first written above, or the date of last signature "
         "of the Parties, whichever is later."),
        ("\"Personal Data\"",
         "has the meaning given in Section 2(t) of the DPDP Act, being any "
         "data about an individual who is identifiable by or in relation to "
         "such data."),
        ("\"Personal Data Breach\"",
         "has the meaning given in Section 2(u) of the DPDP Act, being any "
         "unauthorised processing of Personal Data or accidental disclosure, "
         "acquisition, sharing, use, alteration, destruction or loss of "
         "access to Personal Data, that compromises the confidentiality, "
         "integrity or availability of Personal Data."),
        ("\"Processing\"",
         "has the meaning given in Section 2(x) of the DPDP Act, being a "
         "wholly or partly automated operation or set of operations "
         "performed on digital Personal Data, and includes operations such "
         "as collection, recording, organisation, structuring, storage, "
         "adaptation, retrieval, use, alignment or combination, indexing, "
         "sharing, disclosure by transmission, dissemination or otherwise "
         "making available, restriction, erasure or destruction."),
        ("\"Processor Personnel\"",
         "means employees, officers, directors, contractors, agents, and "
         "Sub-Processors of the Processor, and employees, officers and "
         "agents of any Sub-Processor, who are authorised by the Processor "
         "to Process Personal Data under this Agreement."),
        ("\"Significant Data Fiduciary\" or \"SDF\"",
         "has the meaning given in Section 10 of the DPDP Act, being a "
         "Data Fiduciary or a class of Data Fiduciaries notified as such by "
         "the Central Government."),
        ("\"Sub-Processor\"",
         "means any third party (including any Affiliate of the Processor) "
         "engaged by the Processor to Process Personal Data on behalf of "
         "the Fiduciary under this Agreement."),
        ("\"Underlying Agreement\"",
         "means the Master Services Agreement, Statement of Work, vendor "
         "agreement or other commercial agreement between the Parties to "
         "which this DPA relates."),
    ]

    _two_col_table(doc, defs, header=("Term", "Definition"))

    _p(doc,
       "3.2  Interpretation. Headings are for convenience only. Singular "
       "includes plural and vice versa. References to statutes include "
       "amendments and re-enactments. \"Including\" means \"including without "
       "limitation\". References to Clauses, Sections and Schedules are to "
       "those of this Agreement unless otherwise specified.")


# --------------------------------------------------------------------------- #
# Scope, Subject-matter, Roles
# --------------------------------------------------------------------------- #

def add_scope(doc: Document) -> None:
    _h1(doc, "4.  Scope, Subject-Matter and Roles")

    _sub_clause(doc, "4.1",
                "Subject-matter, duration, nature and purpose of Processing. "
                "The subject-matter, duration, nature, purpose, categories of "
                "Personal Data and categories of Data Principals in respect "
                "of the Processing under this Agreement are set out in "
                "Schedule 1 (Processing Details). The Processor shall Process "
                "Personal Data strictly in accordance with Schedule 1 and the "
                "documented instructions of the Fiduciary.")

    _sub_clause(doc, "4.2",
                "Fiduciary's lawful basis. The Fiduciary represents and "
                "warrants that it has established a lawful basis for the "
                "Processing under Section 4 of the DPDP Act (namely, consent "
                "under Section 6 or a legitimate use under Section 7, as "
                "applicable), that it has issued the Notice required under "
                "Section 5 of the DPDP Act, and that the Processor is "
                "entitled to rely on such representation. The Fiduciary "
                "shall indemnify and hold harmless the Processor against all "
                "loss arising from a breach of this representation.")

    _sub_clause(doc, "4.3",
                "No independent purpose determination. The Processor shall "
                "not determine the purpose or means of Processing, and shall "
                "not Process Personal Data for its own purposes, save for "
                "anonymised or aggregated data lawfully derived in accordance "
                "with Clause 6.11 below.")

    _sub_clause(doc, "4.4",
                "Co-terminus. Unless earlier terminated in accordance with "
                "this Agreement, this DPA shall remain in force for so long "
                "as the Processor Processes Personal Data on behalf of the "
                "Fiduciary under the Underlying Agreement.")


# --------------------------------------------------------------------------- #
# Processor obligations
# --------------------------------------------------------------------------- #

def add_processor_obligations(doc: Document) -> None:
    _h1(doc, "5.  Processor Obligations")

    _sub_clause(doc, "5.1",
                "Documented instructions. The Processor shall Process "
                "Personal Data only on the documented written instructions "
                "of the Fiduciary (including in respect of Cross-Border "
                "Transfers), as set out in this Agreement, the Underlying "
                "Agreement, Schedule 1, and any further written instructions "
                "issued by the Fiduciary from time to time. The Processor "
                "shall promptly inform the Fiduciary if, in its opinion, an "
                "instruction infringes Applicable Law.")

    _sub_clause(doc, "5.2",
                "Confidentiality of Processor Personnel. The Processor shall "
                "ensure that all Processor Personnel authorised to Process "
                "Personal Data: (a) are bound by written obligations of "
                "confidentiality no less stringent than those in this "
                "Agreement; (b) have received appropriate training in data "
                "protection and information security; and (c) are accessed "
                "on a strict need-to-know basis.")

    _sub_clause(doc, "5.3",
                "Security Measures. The Processor shall implement and "
                "maintain reasonable technical and organisational security "
                "safeguards to protect Personal Data in accordance with "
                "Section 8(4) of the DPDP Act and the DPDP Rules. Such "
                "measures shall, at a minimum, include those set out in "
                "Schedule 2 (Technical and Organisational Measures) and "
                "shall be benchmarked to ISO/IEC 27001:2022 (or successor) "
                "controls, encompassing encryption in transit (TLS 1.2 or "
                "higher) and at rest (AES-256 or equivalent), role-based "
                "access control, multi-factor authentication, logging and "
                "monitoring, vulnerability management, secure software "
                "development life-cycle, backup and business continuity, "
                "physical security, personnel training and vendor risk "
                "management.")

    _sub_clause(doc, "5.4",
                "Sub-Processor engagement. The Fiduciary grants the Processor "
                "[general / specific] authorisation to engage Sub-Processors "
                "subject to this Clause 5.4. The initial list of approved "
                "Sub-Processors is set out in Schedule 3. The Processor "
                "shall: (a) give the Fiduciary at least thirty (30) days' "
                "prior written notice of any intended addition or "
                "replacement of a Sub-Processor, during which period the "
                "Fiduciary may object on reasonable data-protection grounds; "
                "(b) enter into a written agreement with each Sub-Processor "
                "imposing data-protection obligations no less onerous than "
                "those in this DPA (flow-down); and (c) remain fully liable "
                "to the Fiduciary for acts and omissions of its "
                "Sub-Processors as if they were its own.")

    _sub_clause(doc, "5.5",
                "Assistance with Data Principal rights. Taking into account "
                "the nature of Processing, the Processor shall assist the "
                "Fiduciary by appropriate technical and organisational "
                "measures, insofar as possible, for the fulfilment of the "
                "Fiduciary's obligation to respond to requests by Data "
                "Principals exercising rights under Sections 11 to 14 of the "
                "DPDP Act, including the rights to access information, "
                "correction, completion, updating and erasure, grievance "
                "redressal, and nomination. The Processor shall provide such "
                "assistance within five (5) working days of receipt of a "
                "written request from the Fiduciary (or such shorter period "
                "as may be required to enable the Fiduciary to comply with "
                "its statutory response timeline).")

    _sub_clause(doc, "5.6",
                "Assistance with Fiduciary's obligations. The Processor "
                "shall assist the Fiduciary in ensuring compliance with the "
                "Fiduciary's obligations under Sections 8(4) to 8(9) of the "
                "DPDP Act, including in respect of: (a) security of "
                "Processing; (b) notification of a Personal Data Breach to "
                "the Data Protection Board and to affected Data Principals; "
                "(c) where applicable to the Fiduciary, the performance of "
                "Data Protection Impact Assessments and independent audits "
                "pursuant to Section 10 of the DPDP Act; and (d) "
                "consultation with the Data Protection Board or any other "
                "competent authority.")

    _sub_clause(doc, "5.7",
                "Erasure on purpose-completion or consent-withdrawal. Upon "
                "written direction of the Fiduciary, and in any event upon "
                "completion of the purpose for which Personal Data was "
                "collected or upon withdrawal of consent by the Data "
                "Principal, the Processor shall erase the relevant Personal "
                "Data in accordance with Section 8(7) of the DPDP Act and "
                "this Agreement, save for such copies as are required to be "
                "retained by Applicable Law (which shall continue to be "
                "secured in accordance with this Agreement).")

    _sub_clause(doc, "5.8",
                "Return or deletion on termination. The provisions of Clause "
                "11 (Termination and Consequences) shall apply.")

    _sub_clause(doc, "5.9",
                "Records of Processing. The Processor shall maintain "
                "complete and accurate records of Processing carried out on "
                "behalf of the Fiduciary, including records of "
                "Sub-Processors, categories of Personal Data, locations of "
                "Processing, and security incidents. Such records shall be "
                "made available to the Fiduciary and, as required, to the "
                "Data Protection Board.")

    _sub_clause(doc, "5.10",
                "Demonstrable compliance and audit. The Processor shall make "
                "available to the Fiduciary all information reasonably "
                "necessary to demonstrate compliance with this Agreement and "
                "shall allow for, and contribute to, audits, including "
                "inspections, conducted by the Fiduciary or an independent "
                "auditor mandated by the Fiduciary, subject to the "
                "conditions in Clause 7.")

    _sub_clause(doc, "5.11",
                "No onward disclosure. The Processor shall not disclose "
                "Personal Data to any third party (including any government "
                "authority), except as: (a) expressly permitted under this "
                "Agreement; (b) required by Applicable Law (in which case "
                "the Processor shall, where not prohibited by law, promptly "
                "notify the Fiduciary); or (c) otherwise pre-authorised in "
                "writing by the Fiduciary.")

    _sub_clause(doc, "5.12",
                "Cooperation with Data Protection Board. The Processor "
                "shall cooperate with, and respond to lawful enquiries and "
                "directions of, the Data Protection Board and, to the extent "
                "applicable, the relevant sectoral regulator, and shall "
                "promptly notify the Fiduciary of any such enquiry or "
                "direction affecting Personal Data Processed hereunder "
                "(unless prohibited from doing so by law).")


# --------------------------------------------------------------------------- #
# Personal Data Breach
# --------------------------------------------------------------------------- #

def add_breach(doc: Document) -> None:
    _h1(doc, "6.  Personal Data Breach")

    _sub_clause(doc, "6.1",
                "Notification by Processor. The Processor shall notify the "
                "Fiduciary in writing without undue delay, and in any event "
                "no later than twenty-four (24) hours, after becoming aware "
                "of any Personal Data Breach affecting Personal Data "
                "Processed under this Agreement. The notification shall be "
                "made to the Fiduciary's notice address specified in "
                "Clause 15 and, in addition, to [DESIGNATED CONTACT EMAIL].")

    _sub_clause(doc, "6.2",
                "Information content. The notification shall, to the extent "
                "then known, include the information set out in Schedule 4 "
                "(Breach Notification Template), namely: (a) the nature of "
                "the Personal Data Breach including, where possible, the "
                "categories and approximate number of Data Principals and "
                "Personal Data records concerned; (b) the likely "
                "consequences of the Personal Data Breach; (c) the measures "
                "taken or proposed to address the Personal Data Breach, "
                "including, where appropriate, measures to mitigate its "
                "possible adverse effects; and (d) the contact details of a "
                "point of contact from whom further information may be "
                "obtained.")

    _sub_clause(doc, "6.3",
                "Supplementation. Where and insofar as it is not possible to "
                "provide the information at the same time, the information "
                "may be provided in phases without further undue delay. The "
                "Processor shall provide a root-cause analysis and "
                "remediation plan to the Fiduciary within seventy-two (72) "
                "hours of initial notification or such other period as may "
                "be agreed in writing.")

    _sub_clause(doc, "6.4",
                "Fiduciary notification to DPB and Data Principals. The "
                "Fiduciary shall, in accordance with Section 8(6) of the "
                "DPDP Act and within the timelines and in the form "
                "prescribed under the DPDP Rules, notify the Data "
                "Protection Board and each affected Data Principal of the "
                "Personal Data Breach. The Processor shall provide the "
                "Fiduciary with all reasonable assistance required to make "
                "such notifications, including provision of information and "
                "draft notifications where requested.")

    _sub_clause(doc, "6.5",
                "Concurrent regulatory reporting. Where concurrent reporting "
                "obligations apply (including under the CERT-In Directions "
                "dated 28 April 2022 for cyber-incident reporting within "
                "six (6) hours, and any applicable RBI, SEBI, IRDAI or "
                "other sectoral directions), each Party shall discharge its "
                "own reporting obligations and shall cooperate with the "
                "other Party in doing so.")

    _sub_clause(doc, "6.6",
                "No admission. A notification under this Clause 6 shall not "
                "constitute an admission of fault or liability by either "
                "Party.")


# --------------------------------------------------------------------------- #
# Audits
# --------------------------------------------------------------------------- #

def add_audits(doc: Document) -> None:
    _h1(doc, "7.  Audits and Inspections")

    _sub_clause(doc, "7.1",
                "Audit right. Not more than once in any twelve (12) month "
                "period (and in addition, without such limit, following a "
                "confirmed Personal Data Breach or at the lawful direction "
                "of the Data Protection Board), the Fiduciary may, on at "
                "least thirty (30) days' prior written notice, carry out, "
                "or mandate an independent third-party auditor bound by "
                "written confidentiality to carry out, an audit of the "
                "Processor's compliance with this Agreement.")

    _sub_clause(doc, "7.2",
                "Audit conditions. Audits shall be conducted during normal "
                "business hours, shall not unreasonably interfere with the "
                "Processor's business, shall be subject to the Processor's "
                "reasonable security and confidentiality requirements, and "
                "shall not extend to Personal Data or confidential "
                "information of other customers of the Processor. Costs of "
                "audit shall be borne by the Fiduciary, save where the "
                "audit discloses a material breach of this Agreement by the "
                "Processor, in which case the Processor shall bear the "
                "Fiduciary's reasonable audit costs.")

    _sub_clause(doc, "7.3",
                "Alternative assurance. The Processor may discharge its "
                "audit obligations by providing to the Fiduciary: (a) "
                "current SOC 2 Type II or equivalent audit reports; (b) "
                "ISO/IEC 27001 certification; or (c) an independent audit "
                "report conducted pursuant to Section 10(2)(c) of the DPDP "
                "Act (where applicable), provided such reports "
                "substantively cover the Processing under this Agreement.")


# --------------------------------------------------------------------------- #
# Cross-Border Transfer
# --------------------------------------------------------------------------- #

def add_cross_border(doc: Document) -> None:
    _h1(doc, "8.  Cross-Border Transfer of Personal Data")

    _sub_clause(doc, "8.1",
                "Statutory framework. The Parties acknowledge that Section "
                "16 of the DPDP Act empowers the Central Government to "
                "restrict the transfer of Personal Data by a Data Fiduciary "
                "for Processing to any country or territory outside India "
                "as may be notified (the \"negative-list\" framework). The "
                "Parties further acknowledge the extra-territorial "
                "application of the DPDP Act in relation to the offering of "
                "goods or services to Data Principals within the territory "
                "of India.")

    _sub_clause(doc, "8.2",
                "Compliance obligation. The Processor shall not effect any "
                "Cross-Border Transfer of Personal Data except: (a) to such "
                "destinations as are set out in Schedule 1 or Schedule 3 "
                "or as are subsequently approved by the Fiduciary in "
                "writing; and (b) in compliance with Section 16 and "
                "Section 17 of the DPDP Act, the DPDP Rules, any "
                "restrictions notified by the Central Government, and any "
                "applicable sectoral data-localisation requirements "
                "(including the Reserve Bank of India's Payment Data "
                "Storage Direction dated 6 April 2018, and applicable SEBI, "
                "IRDAI and other sectoral circulars).")

    _sub_clause(doc, "8.3",
                "Safeguards. Every Cross-Border Transfer shall be secured "
                "by appropriate safeguards, including encryption in transit "
                "and at rest, access controls, and contractual obligations "
                "on the recipient no less stringent than those set out in "
                "this Agreement.")

    _sub_clause(doc, "8.4",
                "EU / UK transfers. Where the Processor or any Sub-Processor "
                "is established in, or Processes Personal Data originating "
                "from, the European Economic Area or the United Kingdom, "
                "the Parties shall execute the applicable European "
                "Commission Standard Contractual Clauses (Module 2 or "
                "Module 3, as appropriate) and the UK International Data "
                "Transfer Addendum, which shall be annexed to this DPA and "
                "shall prevail over this DPA in respect of such transfers "
                "to the extent of any conflict.")

    _sub_clause(doc, "8.5",
                "Localisation flow-down for SDFs. If either Party is "
                "notified as a Significant Data Fiduciary under Section 10 "
                "of the DPDP Act and is directed to localise certain "
                "categories of Personal Data, the other Party shall "
                "cooperate with such localisation, including by adjusting "
                "hosting locations and Sub-Processor engagement, at the "
                "localising Party's reasonable cost.")


# --------------------------------------------------------------------------- #
# Data Principal Rights
# --------------------------------------------------------------------------- #

def add_rights(doc: Document) -> None:
    _h1(doc, "9.  Data Principal Rights; Grievance Redressal")

    _sub_clause(doc, "9.1",
                "Primary responsibility. The Fiduciary shall bear primary "
                "responsibility for responding to Data Principal requests "
                "exercising rights under Sections 11 to 14 of the DPDP Act, "
                "namely the right to access information about Personal Data, "
                "the right to correction and erasure of Personal Data, the "
                "right of grievance redressal, and the right to nominate.")

    _sub_clause(doc, "9.2",
                "Processor's role. If the Processor receives a request "
                "directly from a Data Principal, the Processor shall not "
                "respond substantively (except to acknowledge receipt and "
                "signpost the Fiduciary) and shall forward the request to "
                "the Fiduciary within two (2) working days. The Processor "
                "shall provide the assistance described in Clause 5.5 to "
                "enable the Fiduciary to respond within the timeline "
                "prescribed under the DPDP Rules.")

    _sub_clause(doc, "9.3",
                "Erasure and backups. Where a request for erasure is to be "
                "honoured, the Processor shall erase the relevant Personal "
                "Data from its working production systems within seven (7) "
                "working days, and from backup systems at the next scheduled "
                "backup-rotation cycle, during which period such backup "
                "copies shall be inaccessible to live Processing and shall "
                "be protected in accordance with Schedule 2.")

    _sub_clause(doc, "9.4",
                "Grievance officer. Each Party shall appoint a grievance "
                "officer or other point of contact for Data Principal "
                "grievances, and shall inform the other Party of any change "
                "in such contact. The initial contacts are set out in "
                "Schedule 1.")


# --------------------------------------------------------------------------- #
# Significant Data Fiduciary, Children's Data, Sectoral Overlay
# --------------------------------------------------------------------------- #

def add_sdf_and_sectoral(doc: Document) -> None:
    _h1(doc, "10.  Significant Data Fiduciary, Children, and Sectoral Overlay")

    _sub_clause(doc, "10.1",
                "Significant Data Fiduciary. As at the Effective Date, "
                "AnantaSutra is not a Significant Data Fiduciary (\"SDF\") "
                "and has not been so notified by the Central Government "
                "under Section 10 of the DPDP Act. If either Party is "
                "subsequently notified as a Significant Data Fiduciary "
                "under Section 10 of the DPDP Act, that Party shall "
                "promptly inform the other Party and shall undertake the "
                "additional obligations prescribed from the date of such "
                "notification, including the appointment of a Data "
                "Protection Officer based in India, the performance of "
                "periodic Data Protection Impact Assessments and "
                "independent audits, and such other measures as the "
                "Central Government may prescribe. The other Party shall "
                "provide reasonable cooperation, including access to "
                "information and systems reasonably required for the "
                "performance of such assessments and audits. Until such "
                "notification, SDF-specific obligations (DPO appointment, "
                "DPIA, independent audit) shall not apply to AnantaSutra.")

    _sub_clause(doc, "10.2",
                "Children's Personal Data. Where the Processing involves "
                "Personal Data of children (individuals under 18 years of "
                "age) or persons with disability in accordance with Section "
                "9 of the DPDP Act: (a) the Fiduciary warrants that "
                "verifiable parental (or lawful guardian) consent has been "
                "obtained; (b) the Processor shall not undertake tracking, "
                "behavioural monitoring or targeted advertising directed at "
                "such Data Principals; and (c) the Parties shall observe "
                "the heightened penalty exposure of up to INR 200 crore for "
                "non-compliance.")

    _sub_clause(doc, "10.3",
                "Sensitive, financial, health and biometric data. While the "
                "DPDP Act does not prescribe a separate category of "
                "\"sensitive personal data\", the Parties acknowledge that "
                "Section 43A of the Information Technology Act, 2000 and "
                "the SPDI Rules 2011 continue to apply to financial, health, "
                "biometric, sexual-orientation, medical and Aadhaar-related "
                "information Processed by bodies corporate, and that "
                "sectoral regulators (including the Reserve Bank of India, "
                "the Insurance Regulatory and Development Authority of "
                "India, and the Securities and Exchange Board of India) "
                "impose heightened standards. The Processor shall apply "
                "such heightened standards where applicable to the "
                "Processing.")


# --------------------------------------------------------------------------- #
# Termination
# --------------------------------------------------------------------------- #

def add_termination(doc: Document) -> None:
    _h1(doc, "11.  Term, Termination and Consequences")

    _sub_clause(doc, "11.1",
                "Term. This Agreement shall commence on the Effective Date "
                "and shall continue in force for so long as the Processor "
                "Processes Personal Data on behalf of the Fiduciary. "
                "Termination of the Underlying Agreement for any reason "
                "shall automatically terminate this DPA.")

    _sub_clause(doc, "11.2",
                "Return or deletion. Within thirty (30) days of termination "
                "or expiry, the Processor shall, at the Fiduciary's written "
                "option, either: (a) return to the Fiduciary all Personal "
                "Data in a structured, commonly used and machine-readable "
                "format; or (b) securely delete or destroy all Personal "
                "Data (and all copies thereof) in its possession or "
                "control, including copies held by Sub-Processors.")

    _sub_clause(doc, "11.3",
                "Certificate of deletion. On completion of deletion or "
                "return, the Processor shall furnish to the Fiduciary a "
                "signed certificate of deletion / return, substantially in "
                "the form required by the Fiduciary.")

    _sub_clause(doc, "11.4",
                "Permitted retention. The Processor may retain Personal "
                "Data only to the extent, and for such period as, "
                "Applicable Law requires such retention. In such case, "
                "the Processor shall: (a) notify the Fiduciary of the "
                "nature and duration of retention; (b) continue to apply "
                "the security measures set out in Schedule 2; and (c) "
                "Process such retained Personal Data only for the purpose "
                "required by Applicable Law.")

    _sub_clause(doc, "11.5",
                "Survival. Clauses 1 (Parties), 3 (Definitions), 5.2 "
                "(Confidentiality), 5.10 (Audit — for one year), 6 (Breach), "
                "11 (Termination), 12 (Liability), 13 (Confidentiality), "
                "15 (Notices), 16 (Governing Law), 17 (Dispute Resolution) "
                "and any other Clause that by its nature is intended to "
                "survive, shall survive termination.")


# --------------------------------------------------------------------------- #
# Liability
# --------------------------------------------------------------------------- #

def add_liability(doc: Document) -> None:
    _h1(doc, "12.  Liability, Indemnity and DPB Penalty Exposure")

    _sub_clause(doc, "12.1",
                "Acknowledgement. The Parties acknowledge that non-compliance "
                "with the DPDP Act may result in monetary penalties imposed "
                "by the Data Protection Board of up to INR 250 crore per "
                "breach under the Schedule to Section 33 of the DPDP Act, "
                "including INR 200 crore for breaches involving Personal "
                "Data of children or SDF obligations and INR 150 crore for "
                "failure to notify a Personal Data Breach.")

    _sub_clause(doc, "12.2",
                "Indemnity by Processor. The Processor shall indemnify, "
                "defend and hold harmless the Fiduciary from and against all "
                "direct losses, damages, fines, penalties (including DPB "
                "penalties under Section 33 of the DPDP Act), regulatory "
                "charges, reasonable legal costs and third-party claims "
                "arising out of or attributable to: (a) any breach by the "
                "Processor (or its Sub-Processors or Processor Personnel) of "
                "this Agreement; (b) any negligent, wilful or unauthorised "
                "act or omission by the Processor in the Processing of "
                "Personal Data; or (c) a Personal Data Breach caused by "
                "the Processor.")

    _sub_clause(doc, "12.3",
                "Indemnity by Fiduciary. The Fiduciary shall indemnify, "
                "defend and hold harmless the Processor from and against "
                "all direct losses, damages, fines, penalties (including "
                "DPB penalties), regulatory charges, reasonable legal costs "
                "and third-party claims arising out of or attributable to: "
                "(a) the Fiduciary's breach of its representations under "
                "Clause 4.2 (lawful basis and notice); (b) instructions "
                "given by the Fiduciary that infringe Applicable Law; or "
                "(c) any act or omission of the Fiduciary in its capacity "
                "as Data Fiduciary that gives rise to liability under the "
                "DPDP Act.")

    _sub_clause(doc, "12.4",
                "Super-cap for DPDP exposure. Notwithstanding any limitation "
                "or exclusion of liability in the Underlying Agreement, the "
                "aggregate liability of each Party under this DPA (including "
                "in respect of indemnities for DPB penalties) shall be "
                "capped at [two (2) / three (3)] times the annual fees "
                "payable under the Underlying Agreement, or INR "
                "[AMOUNT IN WORDS] (INR [AMOUNT IN FIGURES]), whichever is "
                "higher. The general cap in the Underlying Agreement shall "
                "not apply to claims under this DPA, and the limitations "
                "herein shall not apply to: (i) gross negligence, wilful "
                "misconduct or fraud; (ii) breach of Clause 5.11 "
                "(no onward disclosure); (iii) breach of confidentiality "
                "obligations; or (iv) indemnified third-party claims.")

    _sub_clause(doc, "12.5",
                "Insurance. Each Party shall maintain cyber-liability "
                "insurance with a reputable insurer for a sum assured of "
                "not less than INR [AMOUNT] crore, and shall produce "
                "certificates of insurance on request.")


# --------------------------------------------------------------------------- #
# Confidentiality, Order of Precedence, Boilerplate
# --------------------------------------------------------------------------- #

def add_misc(doc: Document) -> None:
    _h1(doc, "13.  Confidentiality")

    _p(doc,
       "13.1  The Parties shall treat this Agreement, the Personal Data "
       "Processed hereunder, and any information exchanged in connection "
       "with this Agreement as confidential information, and shall apply "
       "the confidentiality undertakings in the Underlying Agreement (or, "
       "failing such undertakings, customary market-standard undertakings) "
       "with effect from the Effective Date.")

    _p(doc,
       "13.2  The Processor shall ensure that each member of Processor "
       "Personnel is bound by written confidentiality undertakings, "
       "including in respect of obligations surviving termination of "
       "their engagement for a period of not less than five (5) years.")

    _h1(doc, "14.  Order of Precedence")

    _p(doc,
       "In the event of any conflict or inconsistency between this DPA and "
       "the Underlying Agreement in relation to the Processing of Personal "
       "Data, this DPA shall prevail. In the event of any conflict between "
       "this DPA and the European Commission Standard Contractual Clauses "
       "(where executed), the Standard Contractual Clauses shall prevail in "
       "respect of transfers to which they apply. In the event of any "
       "conflict between this DPA and a mandatory provision of the DPDP "
       "Act or DPDP Rules, the statutory provision shall prevail.")

    _h1(doc, "15.  Notices")

    _p(doc,
       "15.1  All notices, consents, instructions and other communications "
       "required or permitted under this Agreement shall be in writing and "
       "shall be delivered by hand, by registered post with acknowledgement "
       "due, or by email to the addresses set out below (or to such other "
       "address as a Party may notify to the other from time to time):")

    _two_col_table(doc,
                   [
                       ("To AnantaSutra",
                        "Attn: Mr. Himanshu Mishra, Founder & CEO\n"
                        "Place: Delhi, India\n"
                        "Email: contact@anantasutra.com"),
                       ("To Counterparty",
                        "Attention: [CONTACT NAME]\n"
                        "Address: [ADDRESS]\n"
                        "Email: [EMAIL]\n"
                        "Copy: [EMAIL]"),
                   ])

    _p(doc,
       "15.2  Notices sent by email shall be deemed received on the first "
       "working day after transmission, provided no bounce-back is received.")

    _h1(doc, "16.  Governing Law")

    _p(doc,
       "This Agreement and any dispute or claim arising out of or in "
       "connection with it (including non-contractual disputes or claims) "
       "shall be governed by and construed in accordance with the laws of "
       "the Republic of India.")

    _h1(doc, "17.  Dispute Resolution")

    _sub_clause(doc, "17.1",
                "Amicable resolution. The Parties shall first attempt to "
                "resolve any dispute amicably through discussions between "
                "senior representatives for a period of thirty (30) days "
                "from written notice of dispute.")

    _sub_clause(doc, "17.2",
                "Arbitration. Failing amicable resolution, the dispute shall "
                "be finally resolved by arbitration under the Arbitration "
                "and Conciliation Act, 1996 (as amended) by a sole "
                "arbitrator mutually appointed by the Parties (and failing "
                "agreement, appointed in accordance with the said Act). "
                "The seat and venue of arbitration shall be New Delhi, "
                "Delhi. The language of arbitration shall be English. "
                "The arbitral award shall be final and binding.")

    _sub_clause(doc, "17.3",
                "Regulatory carve-out. Nothing in this Clause 17 shall "
                "restrict the jurisdiction of the Data Protection Board, "
                "the Telecom Disputes Settlement and Appellate Tribunal "
                "(TDSAT) under Section 29 of the DPDP Act, or any "
                "sectoral regulator, in respect of matters falling within "
                "their statutory jurisdiction. The courts at New Delhi "
                "shall have exclusive jurisdiction in respect of any "
                "interim or urgent relief under Section 9 of the "
                "Arbitration and Conciliation Act, 1996 and any matters "
                "not capable of being resolved by arbitration.")

    _h1(doc, "18.  Miscellaneous")

    _sub_clause(doc, "18.1",
                "Entire agreement. This DPA, together with the Underlying "
                "Agreement and the Schedules hereto, constitutes the entire "
                "agreement between the Parties relating to the Processing of "
                "Personal Data and supersedes all prior understandings, "
                "whether oral or written.")

    _sub_clause(doc, "18.2",
                "Severability. If any provision of this Agreement is held "
                "to be invalid, illegal or unenforceable, the remaining "
                "provisions shall remain in full force and effect, and the "
                "Parties shall negotiate in good faith to replace the "
                "invalid provision with a valid provision that most closely "
                "reflects the Parties' original intention.")

    _sub_clause(doc, "18.3",
                "Amendment. No amendment or modification of this Agreement "
                "shall be effective unless made in writing and signed by "
                "both Parties.")

    _sub_clause(doc, "18.4",
                "Assignment. Neither Party may assign this Agreement "
                "without the prior written consent of the other Party, "
                "save to an Affiliate or in connection with a merger, "
                "reorganisation or sale of substantially all of its assets.")

    _sub_clause(doc, "18.5",
                "Counterparts and electronic execution. This Agreement may "
                "be executed in counterparts, each of which shall be deemed "
                "an original and all of which together shall constitute one "
                "instrument. The Parties may execute this Agreement by "
                "Aadhaar e-Sign or Digital Signature Certificate in "
                "accordance with Sections 3, 3A and 5 of the Information "
                "Technology Act, 2000, and such execution shall be legally "
                "valid and binding.")

    _sub_clause(doc, "18.6",
                "Stamp duty. This Agreement, where executed standalone, "
                "shall be adequately stamped under the Delhi Stamp Act, "
                "1957 (or the Stamp Act of the state of execution, as "
                "applicable), as an \"agreement not otherwise provided "
                "for\". Where executed as an addendum to the Underlying "
                "Agreement, the Parties record that the stamp duty paid on "
                "the Underlying Agreement covers this DPA. Under-stamping "
                "shall be cured by impounding and payment of the "
                "deficient duty together with any penalty under Section 35 "
                "of the Indian Stamp Act, 1899.")

    _sub_clause(doc, "18.7",
                "No waiver. A failure or delay by a Party in exercising any "
                "right shall not operate as a waiver of that right.")

    _sub_clause(doc, "18.8",
                "No partnership. Nothing in this Agreement shall create a "
                "partnership, joint venture or agency between the Parties.")


# --------------------------------------------------------------------------- #
# Execution
# --------------------------------------------------------------------------- #

def add_execution(doc: Document) -> None:
    _h1(doc, "19.  Execution")

    _p(doc,
       "IN WITNESS WHEREOF, the Parties hereto have caused this Data "
       "Processing Agreement to be executed by their duly authorised "
       "representatives as of the Effective Date.",
       bold=True)

    sign_table = doc.add_table(rows=7, cols=2)
    sign_table.style = "Table Grid"
    sign_table.autofit = False
    for row in sign_table.rows:
        row.cells[0].width = Cm(8)
        row.cells[1].width = Cm(8)

    _set_cell_text(sign_table.rows[0].cells[0],
                   "For and on behalf of ANANTASUTRA", bold=True)
    _set_cell_text(sign_table.rows[0].cells[1],
                   "For and on behalf of [COUNTERPARTY]", bold=True)
    _shade_cell(sign_table.rows[0].cells[0])
    _shade_cell(sign_table.rows[0].cells[1])

    ananta_values = [
        "Signature:\n\n_____________________________________",
        "Name: Mr. Himanshu Mishra",
        "Designation: Founder & CEO",
        "Date: [DATE]",
        "Place: Delhi",
        "Witness (Name, Signature):\n\n_____________________________________",
    ]
    counter_values = [
        "Signature:\n\n_____________________________________",
        "Name: [NAME]",
        "Designation: [DESIGNATION]",
        "Date: [DATE]",
        "Place: [PLACE]",
        "Witness (Name, Signature):\n\n_____________________________________",
    ]
    for i, (l, r) in enumerate(zip(ananta_values, counter_values), start=1):
        _set_cell_text(sign_table.rows[i].cells[0], l)
        _set_cell_text(sign_table.rows[i].cells[1], r)


# --------------------------------------------------------------------------- #
# Schedules
# --------------------------------------------------------------------------- #

def add_schedule_1(doc: Document) -> None:
    _page_break(doc)
    _h1(doc, "Schedule 1 — Processing Details")

    _p(doc,
       "This Schedule 1 is issued pursuant to Clause 4.1 of the DPA and "
       "sets out the details of the Processing.",
       italic=True)

    rows = [
        ("Subject-matter of Processing",
         "[Brief description — e.g., operation, hosting and support of the "
         "[PRODUCT / SERVICE] through which the Fiduciary delivers "
         "services to its customers.]"),
        ("Duration of Processing",
         "For the term of the Underlying Agreement dated [DATE] and any "
         "permitted extension, plus any permitted retention period under "
         "Clause 11.4."),
        ("Nature of Processing",
         "Collection, recording, organisation, structuring, storage, "
         "retrieval, use, alignment, indexing, sharing, transmission, "
         "restriction, erasure and destruction — as reasonably necessary to "
         "provide the services under the Underlying Agreement."),
        ("Purpose of Processing",
         "[Describe — e.g., customer account management; delivery of the "
         "Fiduciary's services to its end customers; analytics; support.]"),
        ("Categories of Data Principals",
         "[Select all that apply:\n"
         "☐ Employees of the Fiduciary\n"
         "☐ Contractors / consultants of the Fiduciary\n"
         "☐ End customers / users of the Fiduciary's services\n"
         "☐ Prospects / leads\n"
         "☐ Vendors / suppliers\n"
         "☐ Minors (<18) — trigger Clause 10.2\n"
         "☐ Other: [specify]]"),
        ("Categories of Personal Data",
         "[Select all that apply:\n"
         "☐ Identity data (name, date of birth, identifiers)\n"
         "☐ Contact data (email, phone, address)\n"
         "☐ Authentication data (hashed passwords, tokens)\n"
         "☐ Transaction / payment data\n"
         "☐ Usage and device / log data\n"
         "☐ Location data\n"
         "☐ Financial data (subject to sectoral overlay)\n"
         "☐ Health data (subject to sectoral overlay)\n"
         "☐ Biometric data\n"
         "☐ Aadhaar or other government-issued ID\n"
         "☐ Other: [specify]]"),
        ("Processing locations",
         "[Primary: India — specify data centre region.\n"
         "Secondary / disaster-recovery: [specify]\n"
         "Any other locations (for Cross-Border Transfer): [specify]]"),
        ("Retention period",
         "[Specify per data category — default: for the term of the "
         "Underlying Agreement plus [X] years for tax, companies-law and "
         "audit compliance.]"),
        ("Fiduciary's grievance officer / DPO",
         "Name: [NAME]\nEmail: [EMAIL]\nPhone: [PHONE]"),
        ("Processor's grievance officer (AnantaSutra)",
         "Name: Mr. Himanshu Mishra\n"
         "Designation: Founder & CEO\n"
         "Email: contact@anantasutra.com\n"
         "Note: A Data Protection Officer (DPO) shall be appointed only if "
         "AnantaSutra is notified as a Significant Data Fiduciary under "
         "Section 10 of the DPDP Act."),
        ("Lawful basis relied upon by Fiduciary",
         "☐ Consent under Section 6 DPDP Act\n"
         "☐ Legitimate use under Section 7 DPDP Act (specify sub-clause)\n"
         "☐ Employment (Section 7(i))\n"
         "☐ Other: [specify]"),
    ]
    _two_col_table(doc, rows, header=("Item", "Details"),
                   col_widths=(5.0, 11.0))


def add_schedule_2(doc: Document) -> None:
    _page_break(doc)
    _h1(doc, "Schedule 2 — Technical and Organisational Security Measures")

    _p(doc,
       "This Schedule 2 is issued pursuant to Clause 5.3 and describes the "
       "minimum technical and organisational measures that the Processor "
       "shall implement and maintain. Measures shall be benchmarked to "
       "ISO/IEC 27001:2022 controls and to the DPDP Rules (as notified).",
       italic=True)

    measures = [
        ("Information Security Governance",
         "Documented Information Security Management System (ISMS); "
         "appointed information security officer; annual management review; "
         "ISO/IEC 27001 certification (current or planned within [X] months)."),
        ("Encryption in Transit",
         "TLS 1.2 or higher for all network communications; HSTS where "
         "applicable; mutual TLS for service-to-service communication where "
         "feasible."),
        ("Encryption at Rest",
         "AES-256 (or equivalent) for Personal Data at rest in databases, "
         "object storage and backups; key management via a managed KMS with "
         "rotation at least annually."),
        ("Access Control",
         "Role-Based Access Control (RBAC); least-privilege principle; "
         "Multi-Factor Authentication (MFA) for all administrative and "
         "remote access; quarterly access reviews; timely de-provisioning "
         "on role change or exit."),
        ("Logging and Monitoring",
         "Tamper-evident logs of access to Personal Data retained for not "
         "less than 180 days (or such longer period as CERT-In Directions "
         "2022 may require); 24x7 security monitoring with alerting; "
         "SIEM / equivalent."),
        ("Vulnerability and Patch Management",
         "Monthly vulnerability scanning; critical patches within 14 days; "
         "annual penetration testing by an independent assessor; "
         "remediation tracking."),
        ("Secure Software Development",
         "Secure SDLC with code review, static and dynamic application "
         "security testing, dependency scanning; environment segregation "
         "(dev/test/prod); no production data in non-production "
         "environments without masking."),
        ("Backup and Business Continuity",
         "Encrypted, access-controlled backups; documented Business "
         "Continuity Plan and Disaster Recovery Plan; defined RTO and "
         "RPO; annual DR test."),
        ("Incident Response",
         "Documented Personal Data Breach response plan; trained incident "
         "response team; runbooks; 24x7 contact; integration with the "
         "notification obligations in Clause 6; CERT-In reporting "
         "procedure."),
        ("Personnel Training and Screening",
         "Pre-engagement background verification (to the extent permitted "
         "by law); confidentiality undertakings; annual data-protection "
         "and security training; awareness programmes."),
        ("Physical Security",
         "Secure data-centre facilities with access controls, CCTV, "
         "visitor logs, environmental controls; secure disposal of media "
         "using NIST 800-88 or equivalent standards."),
        ("Vendor / Sub-Processor Risk Management",
         "Due diligence and risk-rating of Sub-Processors; back-to-back "
         "data-protection obligations; periodic review; right to audit."),
        ("Data Minimisation and Segregation",
         "Processing limited to Personal Data necessary for the specified "
         "purpose; logical segregation of Fiduciary's data from other "
         "customers' data."),
        ("Record-keeping",
         "Records of Processing Activities (ROPA), Data Protection Impact "
         "Assessment register (where applicable), breach register, consent "
         "log, sub-processor register."),
    ]
    _two_col_table(doc, measures, header=("Control Area", "Measures"),
                   col_widths=(5.0, 11.0))


def add_schedule_3(doc: Document) -> None:
    _page_break(doc)
    _h1(doc, "Schedule 3 — Approved Sub-Processors")

    _p(doc,
       "This Schedule 3 is issued pursuant to Clause 5.4 and sets out the "
       "Sub-Processors engaged by the Processor as at the Effective Date. "
       "The Processor shall update this Schedule in accordance with "
       "Clause 5.4 on addition or replacement of a Sub-Processor.",
       italic=True)

    headers = ("S. No.", "Sub-Processor (Legal Name)",
               "Processing Activity",
               "Location(s) of Processing",
               "Cross-Border?",
               "Security Certifications")

    table = doc.add_table(rows=6, cols=6)
    table.style = "Light Grid Accent 1"
    table.autofit = True

    for i, h in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], h, bold=True, size=10)
        _shade_cell(table.rows[0].cells[i])

    # Five fillable rows
    for r in range(1, 6):
        for c in range(6):
            _set_cell_text(table.rows[r].cells[c],
                           f"[{r}.{c+1}]" if c == 0 else "[…]", size=10)

    _p(doc, "")
    _p(doc,
       "Notes: \n"
       "1. \"Cross-Border?\" shall indicate whether Personal Data is "
       "transferred outside India, with destination country/territory.\n"
       "2. The Processor shall provide the Fiduciary with documentary "
       "evidence of Sub-Processor data-protection commitments on request.\n"
       "3. The Fiduciary may object to any new or replacement Sub-Processor "
       "on reasonable data-protection grounds within thirty (30) days of "
       "notice under Clause 5.4.",
       italic=True)


def add_schedule_4(doc: Document) -> None:
    _page_break(doc)
    _h1(doc, "Schedule 4 — Personal Data Breach Notification Template")

    _p(doc,
       "This Schedule 4 is issued pursuant to Clause 6 and sets out the "
       "information to be included in a Personal Data Breach notification "
       "from the Processor to the Fiduciary. The template shall be "
       "completed to the extent information is available at the time of "
       "initial notification; additional information shall be provided in "
       "phases without further undue delay.",
       italic=True)

    breach_rows = [
        ("Incident Reference Number", "[PDB-YYYY-NNNN]"),
        ("Date and Time of Notification", "[DD/MM/YYYY HH:MM IST]"),
        ("Date and Time Processor Became Aware", "[DD/MM/YYYY HH:MM IST]"),
        ("Date and Time of Occurrence (estimated)",
         "[DD/MM/YYYY HH:MM IST]"),
        ("Nature of the Personal Data Breach",
         "[Describe — unauthorised access / accidental disclosure / "
         "ransomware / system failure / loss of device / insider / other]"),
        ("Categories of Personal Data Affected",
         "[Identity / contact / financial / health / authentication / "
         "location / other]"),
        ("Approximate Number of Data Principals Affected", "[Number or range]"),
        ("Approximate Number of Personal Data Records Affected",
         "[Number or range]"),
        ("Categories of Data Principals Affected",
         "[End customers / employees / minors / other]"),
        ("Geographic Scope", "[India / specify foreign jurisdictions]"),
        ("Likely Consequences of the Personal Data Breach",
         "[Describe — identity theft, financial loss, reputational harm, "
         "discrimination, loss of service, other]"),
        ("Containment and Mitigation Measures Taken or Proposed",
         "[Describe technical, organisational, communication and legal "
         "measures]"),
        ("Root-Cause (interim hypothesis / confirmed)", "[Describe]"),
        ("Sub-Processor(s) Involved, if any", "[Specify]"),
        ("Regulatory Notifications Made / Pending",
         "[DPB — date and reference / CERT-In — date and reference / "
         "sectoral regulator — details / other]"),
        ("Law-Enforcement Notifications, if any", "[Specify]"),
        ("Processor Point of Contact",
         "Name: [NAME]\nDesignation: [DESIGNATION]\nEmail: [EMAIL]\n"
         "Phone: [PHONE]"),
        ("Follow-up Timeline",
         "Root-cause analysis and remediation plan to be provided within "
         "seventy-two (72) hours of initial notification."),
        ("Signed", "For the Processor: __________________________"),
    ]
    _two_col_table(doc, breach_rows, header=("Field", "Entry"),
                   col_widths=(6.0, 10.0))


# --------------------------------------------------------------------------- #
# Build
# --------------------------------------------------------------------------- #

def build() -> str:
    doc = Document()
    add_brand_header(doc)
    _set_page(doc)
    _set_base_style(doc)
    _build_footer(doc)

    add_title(doc)
    add_parties_and_recitals(doc)
    add_definitions(doc)
    add_scope(doc)
    add_processor_obligations(doc)
    add_breach(doc)
    add_audits(doc)
    add_cross_border(doc)
    add_rights(doc)
    add_sdf_and_sectoral(doc)
    add_termination(doc)
    add_liability(doc)
    add_misc(doc)
    add_execution(doc)

    add_schedule_1(doc)
    add_schedule_2(doc)
    add_schedule_3(doc)
    add_schedule_4(doc)

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_FILE)
    return OUTPUT_FILE


if __name__ == "__main__":
    path = build()
    size_kb = os.path.getsize(path) / 1024.0
    print(f"Generated: {path}")
    print(f"Size: {size_kb:.2f} KB")
