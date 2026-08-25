"""
build_03_ipr.py
================
Generates the combined Intellectual Property Rights Assignment Deed (Part A)
and AnantaSutra Intellectual Property Policy (Part B) as a single .docx file.

Target: c:/Users/duneja8515/Desktop/bhavya/prsnl/sutranet/legal-documents/03_IPR_Assignment_and_Policy.docx

Rendered with python-docx 1.2.0. Formatting:
  - A4 paper, 1-inch margins
  - Calibri 11pt, justified, 1.15 line-spacing
  - Part A title centered bold 16pt
  - Part B begins on a new page, centered bold 16pt
  - Numbered headings using Word Heading 1 / Heading 2 styles
  - Footer: "Confidential — AnantaSutra" + page number
  - Signature block as a two-column table
  - Schedule (Invention Disclosure Form) rendered as a fillable-looking table
"""

from __future__ import annotations

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.section import WD_SECTION, WD_ORIENTATION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
OUTPUT_PATH = (
    r"c:/Users/duneja8515/Desktop/bhavya/prsnl/sutranet/legal-documents/"
    r"03_IPR_Assignment_and_Policy.docx"
)

BASE_FONT = "Calibri"


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Brand header (AnantaSutra) — injected by rebrand
# ---------------------------------------------------------------------------
LOGO_PATH = r"c:/Users/duneja8515/Desktop/bhavya/prsnl/sutranet/public/Favicon/logo-nobg.png"


def add_brand_header(doc):
    """Insert centered logo + brand lines at top of page 1. Never crashes."""
    try:
        from docx.shared import Inches as _Inches, Pt as _Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH as _CENTERED
    except Exception:
        return
    p = doc.add_paragraph()
    p.alignment = _CENTERED.CENTER
    try:
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=_Inches(1.2))
    except Exception:
        pass
    p2 = doc.add_paragraph()
    p2.alignment = _CENTERED.CENTER
    r2 = p2.add_run("AnantaSutra — अनन्तसूत्र")
    r2.italic = True
    r2.font.size = _Pt(10)
    p3 = doc.add_paragraph()
    p3.alignment = _CENTERED.CENTER
    r3 = p3.add_run("https://anantasutra.com  •  contact@anantasutra.com")
    r3.italic = True
    r3.font.size = _Pt(9)

def set_cell_border(cell, **kwargs):
    """Set borders on a single table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            tag = OxmlElement(f"w:{edge}")
            for k, v in kwargs[edge].items():
                tag.set(qn(f"w:{k}"), v)
            existing = tcBorders.find(qn(f"w:{edge}"))
            if existing is not None:
                tcBorders.remove(existing)
            tcBorders.append(tag)


def add_page_number(paragraph):
    """Insert a PAGE field into the paragraph."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def style_run(run, *, size=11, bold=False, italic=False, font=BASE_FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)


def set_paragraph_format(paragraph, *, justify=True, space_after=6, space_before=0):
    pf = paragraph.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if justify:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


# ---------------------------------------------------------------------------
# Convenience writers
# ---------------------------------------------------------------------------
def add_body_paragraph(doc, text, *, italic=False, bold=False, justify=True, indent=None):
    p = doc.add_paragraph()
    set_paragraph_format(p, justify=justify)
    if indent is not None:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    style_run(run, size=11, bold=bold, italic=italic)
    return p


def add_mixed_paragraph(doc, segments, *, justify=True, indent=None):
    """segments is a list of (text, {bold:bool, italic:bool}) tuples."""
    p = doc.add_paragraph()
    set_paragraph_format(p, justify=justify)
    if indent is not None:
        p.paragraph_format.left_indent = Inches(indent)
    for text, opts in segments:
        run = p.add_run(text)
        style_run(
            run,
            size=11,
            bold=opts.get("bold", False),
            italic=opts.get("italic", False),
        )
    return p


def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    run = p.add_run(text)
    style_run(run, size=13, bold=True)
    run.font.color.rgb = RGBColor(0x1F, 0x2D, 0x5C)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.15
    run = p.add_run(text)
    style_run(run, size=12, bold=True)
    run.font.color.rgb = RGBColor(0x2E, 0x4A, 0x7B)
    return p


def add_bullets(doc, items, *, indent=0.3):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_format(p, justify=True)
        p.paragraph_format.left_indent = Inches(indent)
        run = p.add_run(item)
        style_run(run, size=11)


def add_numbered_sub(doc, items, *, indent=0.3):
    """Manually numbered sub-list a, b, c... as justified paragraphs."""
    letters = "abcdefghijklmnopqrstuvwxyz"
    for i, item in enumerate(items):
        p = doc.add_paragraph()
        set_paragraph_format(p, justify=True, space_after=4)
        p.paragraph_format.left_indent = Inches(indent)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        run = p.add_run(f"({letters[i]})  ")
        style_run(run, size=11, bold=True)
        run2 = p.add_run(item)
        style_run(run2, size=11)


def add_centered_title(doc, text, *, size=16, space_before=0, space_after=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.15
    run = p.add_run(text)
    style_run(run, size=size, bold=True)
    return p


def add_blank(doc, count=1):
    for _ in range(count):
        doc.add_paragraph()


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ---------------------------------------------------------------------------
# Document construction
# ---------------------------------------------------------------------------
def configure_document(doc: Document) -> None:
    """Set page size, margins, base style."""
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.orientation = WD_ORIENTATION.PORTRAIT

    normal = doc.styles["Normal"]
    normal.font.name = BASE_FONT
    normal.font.size = Pt(11)
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), BASE_FONT)
    rFonts.set(qn("w:hAnsi"), BASE_FONT)
    rFonts.set(qn("w:cs"), BASE_FONT)
    pf = normal.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)


def configure_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer_para.runs:
            run.text = ""
        run = footer_para.add_run("Confidential — AnantaSutra    |    Page ")
        style_run(run, size=9, italic=True)
        add_page_number(footer_para)
        tail = footer_para.add_run("")
        style_run(tail, size=9, italic=True)


# ---------------------------------------------------------------------------
# PART A — IP Assignment Deed
# ---------------------------------------------------------------------------
def build_part_a(doc: Document) -> None:
    add_centered_title(
        doc,
        "DEED OF ASSIGNMENT OF INTELLECTUAL PROPERTY RIGHTS",
        size=16,
        space_before=0,
        space_after=14,
    )

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(16)
    run = sub.add_run(
        "(Executed pursuant to the Indian Contract Act, 1872, the Copyright Act, 1957, "
        "the Patents Act, 1970, the Trade Marks Act, 1999, the Designs Act, 2000 and "
        "the Semiconductor Integrated Circuits Layout-Design Act, 2000)"
    )
    style_run(run, size=10, italic=True)

    # Deed reference line
    add_mixed_paragraph(
        doc,
        [
            ("This Deed of Assignment of Intellectual Property Rights (\"", {}),
            ("Deed", {"bold": True}),
            ("\") is made and executed at ", {}),
            ("New Delhi, Delhi, India", {"bold": True}),
            (" on this ", {}),
            ("[DAY]", {"bold": True}),
            (" day of ", {}),
            ("[MONTH]", {"bold": True}),
            (", ", {}),
            ("[YEAR]", {"bold": True}),
            (" (\"", {}),
            ("Effective Date", {"bold": True}),
            ("\")", {}),
            (".", {}),
        ],
    )

    # Parties
    add_heading1(doc, "BETWEEN")
    add_mixed_paragraph(
        doc,
        [
            ("(1)  ", {"bold": True}),
            ("[FULL LEGAL NAME OF ASSIGNOR]", {"bold": True}),
            (
                ", son/daughter/spouse of [PARENT/SPOUSE NAME], aged about "
                "[AGE] years, holding PAN [PAN NO.], residing at "
                "[COMPLETE RESIDENTIAL ADDRESS] (hereinafter referred to as the ",
                {},
            ),
            ("\"Assignor\"", {"bold": True}),
            (
                ", which expression shall, unless repugnant to the context or "
                "meaning thereof, include his/her/their heirs, legal representatives, "
                "executors, administrators and permitted assigns) of the ",
                {},
            ),
            ("ONE PART", {"bold": True}),
            (";", {}),
        ],
    )
    add_centered_title(doc, "AND", size=11, space_before=4, space_after=4)
    add_mixed_paragraph(
        doc,
        [
            ("(2)  ", {"bold": True}),
            ("ANANTASUTRA", {"bold": True}),
            (
                ", a business concern operated and represented by its Founder & "
                "CEO, Mr. Himanshu Mishra, carrying on business under the trade "
                "name \"AnantaSutra\", having its principal place of business at "
                "Delhi, India, with contact address contact@anantasutra.com "
                "(hereinafter referred to as the ",
                {},
            ),
            ("\"Assignee\"", {"bold": True}),
            (
                " or \"",
                {},
            ),
            ("AnantaSutra", {"bold": True}),
            (
                "\", which expression shall, unless repugnant to the context or "
                "meaning thereof, include its successors, assigns, and the "
                "person(s) for the time being in control of the business), acting "
                "through its Founder & CEO, Mr. Himanshu Mishra, of the ",
                {},
            ),
            ("OTHER PART", {"bold": True}),
            (".", {}),
        ],
    )

    add_body_paragraph(
        doc,
        "The Assignor and the Assignee are hereinafter referred to individually as a "
        "\"Party\" and collectively as the \"Parties\".",
    )

    # Recitals
    add_heading1(doc, "WHEREAS:")
    recitals = [
        "The Assignee is a private limited company engaged, inter alia, in the "
        "business of research, design, development, deployment, licensing and support "
        "of software, platforms, algorithms, networked solutions, hardware-adjacent "
        "components, data products and related technology services.",
        "The Assignor is engaged or proposed to be engaged by the Assignee in the "
        "capacity of [EMPLOYEE / CONSULTANT / INTERN / FOUNDER / DIRECTOR] pursuant "
        "to a [EMPLOYMENT AGREEMENT / CONSULTANCY AGREEMENT / LETTER OF ENGAGEMENT] "
        "dated [DATE] (the \"Engagement\"), and in the course of such Engagement the "
        "Assignor has created, is creating and/or may create intellectual property "
        "of value to the Assignee.",
        "The Parties acknowledge that while Section 17(c) of the Copyright Act, 1957 "
        "vests first ownership of copyright in works made in the course of "
        "employment under a contract of service in the employer, (i) such vesting "
        "is subject to any agreement to the contrary, (ii) does not operate in "
        "respect of consultants or independent contractors, and (iii) does not "
        "extend to rights under the Patents Act, 1970, the Trade Marks Act, 1999, "
        "the Designs Act, 2000 or the Semiconductor Integrated Circuits "
        "Layout-Design Act, 2000. The Parties therefore wish to record, confirm "
        "and effect by this Deed an express, present, irrevocable and unequivocal "
        "assignment of all such intellectual property rights in favour of the "
        "Assignee.",
        "The Parties intend that this Deed shall satisfy the requirements of "
        "Section 18 and Section 19 of the Copyright Act, 1957, Section 68 of the "
        "Patents Act, 1970, Sections 37 to 45 of the Trade Marks Act, 1999 and "
        "Section 30 of the Designs Act, 2000, and shall stand as a deed of "
        "assignment for all purposes under the Indian Stamp Act, 1899 and the "
        "Indian Stamp Act, 1899 (as applicable to the NCT of Delhi).",
        "The Parties have agreed to execute this Deed on the terms and conditions "
        "set out below.",
    ]
    for i, r in enumerate(recitals, 1):
        add_mixed_paragraph(
            doc,
            [(f"({chr(64+i)})  ", {"bold": True}), (r, {})],
        )

    add_heading1(doc, "NOW THIS DEED WITNESSETH AS FOLLOWS:")

    # 1. Definitions
    add_heading1(doc, "1.  DEFINITIONS AND INTERPRETATION")
    add_body_paragraph(
        doc,
        "1.1  In this Deed, unless the context otherwise requires, capitalised "
        "terms shall bear the meanings ascribed to them below:",
    )

    definitions = [
        ("\"Affiliate\"",
         "means, in relation to the Assignee, any entity that directly or "
         "indirectly controls, is controlled by or is under common control with "
         "the Assignee, where \"control\" means beneficial ownership of not less "
         "than fifty per cent (50%) of the voting equity of the entity concerned."),
        ("\"Background IP\"",
         "means any Intellectual Property Rights owned, licensed-in or otherwise "
         "controlled by the Assignor (a) prior to the commencement of the "
         "Engagement, or (b) developed entirely independently of the Engagement "
         "and without reference to or use of any Confidential Information, "
         "resources, systems, facilities or data of the Assignee, and which is "
         "specifically identified and disclosed by the Assignor in "
         "Schedule I hereto."),
        ("\"Confidential Information\"",
         "has the meaning assigned to it under the non-disclosure agreement "
         "executed between the Parties and, in addition, includes all non-public "
         "technical, commercial, financial, customer, employee, source-code, "
         "algorithmic, model, data-set and product information of the Assignee "
         "disclosed to or accessed by the Assignor in any form."),
        ("\"Foreground IP\"",
         "means all Intellectual Property Rights in any Work Product and in any "
         "Invention created, conceived, developed, reduced to practice, authored, "
         "fixed, recorded or otherwise brought into existence by the Assignor, "
         "alone or jointly with others, during, in the course of, or as a result "
         "of, the Engagement or through the use of the Assignee's time, "
         "resources, facilities, Confidential Information or Background IP."),
        ("\"Intellectual Property Rights\" or \"IP\"",
         "means all intellectual and industrial property rights of whatsoever "
         "nature, whether registered or unregistered and whether presently known "
         "to law or hereafter coming into existence, including without limitation "
         "(i) copyright and neighbouring rights, including rights in computer "
         "programmes, databases and compilations, (ii) patents, patent "
         "applications, petty patents, utility models and rights in inventions, "
         "(iii) trade marks, service marks, trade names, domain names, logos, "
         "get-up and goodwill, (iv) registered designs, design rights and rights "
         "in industrial designs, (v) semiconductor integrated circuit layout "
         "designs, (vi) rights in Confidential Information, know-how and trade "
         "secrets, (vii) rights in data and databases, (viii) moral rights, and "
         "(ix) all rights to apply for, prosecute, maintain, renew, extend, "
         "revive, enforce and recover damages, profits and other relief in "
         "respect of any of the foregoing, in each case throughout the world and "
         "for the entire term of such rights."),
        ("\"Invention\"",
         "means any idea, conception, discovery, invention, improvement, "
         "modification, method, process, algorithm, technique, formula, know-how, "
         "trade secret, design, work of authorship, mask work, software, source "
         "code, object code, architecture, user interface, data model, schema, "
         "database, documentation, specification, report or deliverable, whether "
         "or not patentable, copyrightable or otherwise protectable by law."),
        ("\"Moral Rights\"",
         "means the rights conferred on the author of a work under Section 57 of "
         "the Copyright Act, 1957 (including the right of paternity and the right "
         "of integrity) and any analogous rights conferred under the laws of any "
         "other jurisdiction."),
        ("\"Open-Source Software\"",
         "means any software that is made available under a licence approved by "
         "the Open Source Initiative or any similar licence, including but not "
         "limited to the MIT Licence, the BSD Licences, the Apache Licence 2.0, "
         "the GNU General Public Licence (all versions), the GNU Lesser General "
         "Public Licence, the Affero General Public Licence, the Mozilla Public "
         "Licence and the Eclipse Public Licence."),
        ("\"Prior IP\"",
         "means, collectively, the Background IP and any other Intellectual "
         "Property Rights existing prior to the Engagement and specifically "
         "disclosed in Schedule I hereto."),
        ("\"Work Product\"",
         "means all tangible and intangible materials, deliverables, outputs and "
         "artefacts embodying any Invention, including without limitation source "
         "code, object code, design documents, specifications, drawings, "
         "flowcharts, diagrams, notebooks, data-sets, trained machine-learning "
         "models, model weights, prompts, evaluation sets, reports, presentations, "
         "marketing materials and any derivative works of the foregoing."),
    ]
    for term, meaning in definitions:
        add_mixed_paragraph(
            doc,
            [(f"{term}  ", {"bold": True}), (meaning, {})],
        )

    add_body_paragraph(
        doc,
        "1.2  References to any statute, statutory provision, rule or regulation "
        "shall include any amendment, modification, re-enactment or replacement "
        "thereof in force from time to time. The headings used in this Deed are "
        "for convenience only and shall not affect its interpretation. Words "
        "importing the singular include the plural and vice versa; words importing "
        "any gender include all genders.",
    )

    # 2. Present-tense assignment
    add_heading1(doc, "2.  PRESENT AND IRREVOCABLE ASSIGNMENT OF FOREGROUND IP")

    add_heading2(doc, "2.1  Assignment")
    add_body_paragraph(
        doc,
        "In consideration of the Engagement, the remuneration, fees, salary, "
        "benefits, equity, stock options, training, access to Confidential "
        "Information and other valuable consideration provided and to be provided "
        "by the Assignee to the Assignor (the sufficiency and adequacy of which "
        "the Assignor hereby acknowledges), and further in consideration of the "
        "sum of Indian Rupees One (INR 1) paid by the Assignee to the Assignor "
        "(the receipt of which is hereby acknowledged), the Assignor HEREBY "
        "IRREVOCABLY, ABSOLUTELY AND UNCONDITIONALLY ASSIGNS, TRANSFERS, CONVEYS "
        "AND MAKES OVER unto the Assignee, with full title guarantee and free "
        "from all encumbrances, all right, title and interest (whether vested, "
        "contingent or future) in and to the Foreground IP, together with all "
        "Intellectual Property Rights subsisting or capable of subsisting therein.",
    )

    add_heading2(doc, "2.2  Scope of Assignment")
    add_body_paragraph(
        doc,
        "The assignment effected by Clause 2.1 shall be (a) in respect of all "
        "modes, media, forms and means of expression, exploitation and use, "
        "whether now known or hereafter devised; (b) worldwide in territorial "
        "extent; (c) for the entire term of each such Intellectual Property "
        "Right, including all extensions, renewals, revivals and reversions "
        "thereof; (d) perpetual and irrevocable; (e) royalty-free, no further "
        "consideration being payable by the Assignee beyond that recited herein; "
        "and (f) accompanied by the unrestricted right to grant sub-licences, "
        "further assignments, security interests and other dispositions in "
        "favour of any third party without the further consent of or reference "
        "to the Assignor.",
    )

    add_heading2(doc, "2.3  Compliance with Section 19, Copyright Act, 1957")
    add_body_paragraph(
        doc,
        "In satisfaction of the mandatory particulars prescribed by "
        "Section 19 of the Copyright Act, 1957, the Parties specifically record "
        "that:",
    )
    add_numbered_sub(
        doc,
        [
            "the \"work\" assigned is every work created by the Assignor in the "
            "course of or pursuant to the Engagement, including all literary "
            "works (including computer programmes and source code), artistic "
            "works, musical works, cinematograph films, sound recordings, "
            "dramatic works and compilations;",
            "the \"rights\" assigned are each and every economic right "
            "enumerated in Section 14 of the Copyright Act, 1957, and all "
            "equivalent rights in every other jurisdiction, without reservation;",
            "the \"duration\" of the assignment is the entire term of copyright "
            "subsisting in each work, together with all extensions, revivals "
            "and renewals thereof;",
            "the \"territorial extent\" of the assignment is worldwide; and",
            "the \"royalty\" payable is Indian Rupees One (INR 1) together with "
            "the Engagement consideration, which the Assignor confirms is fair, "
            "adequate and sufficient consideration for the rights assigned.",
        ],
    )

    add_heading2(doc, "2.4  Assignment of Future-Arising Foreground IP")
    add_body_paragraph(
        doc,
        "To the extent that any Foreground IP comes into existence after the "
        "Effective Date and cannot, as a matter of law, vest in the Assignee by "
        "virtue of the present-tense assignment in Clause 2.1 alone, the "
        "Assignor hereby assigns the same to the Assignee with effect from the "
        "moment of its creation, without the need for any further act, deed or "
        "instrument. The Assignor shall, upon written request by the Assignee, "
        "execute a confirmatory deed of assignment in respect of any such "
        "Foreground IP.",
    )

    add_heading2(doc, "2.5  Vesting at Source")
    add_body_paragraph(
        doc,
        "It is the express intention of the Parties that title to all Foreground "
        "IP shall vest in the Assignee at the point of creation and shall never "
        "vest (even momentarily) in the Assignor. The Assignor shall hold in "
        "trust for the Assignee any Foreground IP or rights therein that may, "
        "for any reason, accrue to the Assignor notwithstanding this Deed, and "
        "shall account to the Assignee for the same.",
    )

    # 3. Patents
    add_heading1(doc, "3.  PATENTS AND INVENTIONS")

    add_heading2(doc, "3.1  Assignment of Inventions")
    add_body_paragraph(
        doc,
        "Without prejudice to the generality of Clause 2, the Assignor hereby "
        "assigns to the Assignee all right, title and interest in and to every "
        "Invention forming part of the Foreground IP, including the exclusive "
        "right to apply for, prosecute, obtain, maintain, enforce and exploit "
        "patents, utility models and other inventor's rights in respect thereof "
        "in any and all jurisdictions worldwide, including before the Indian "
        "Patent Office, the United States Patent and Trademark Office (USPTO), "
        "the European Patent Office (EPO), the World Intellectual Property "
        "Organization (WIPO) under the Patent Cooperation Treaty, and any other "
        "national or regional intellectual-property office.",
    )

    add_heading2(doc, "3.2  Inventor Identification")
    add_body_paragraph(
        doc,
        "The Assignor acknowledges that Section 28 of the Patents Act, 1970 "
        "entitles the true and first inventor to be mentioned as the inventor "
        "in the patent and in the patent register. The Assignor agrees to be so "
        "identified as the inventor (jointly where applicable) but irrevocably "
        "confirms that the right to apply for and to be granted the patent and "
        "all economic rights flowing therefrom vest solely in the Assignee.",
    )

    add_heading2(doc, "3.3  Registration under Section 68, Patents Act, 1970")
    add_body_paragraph(
        doc,
        "The Assignor acknowledges that Section 68 of the Patents Act, 1970 "
        "requires an assignment of a patent or of a share in a patent to be in "
        "writing and, to be valid against third parties, to be registered with "
        "the Controller of Patents on Form 16 within a reasonable time and in "
        "any event within six (6) months of execution to avoid late-filing "
        "surcharges. The Assignor shall promptly, and without demanding any "
        "further consideration, execute and deliver all such forms, affidavits "
        "and documents as may be necessary to enable the Assignee to effect "
        "such registration.",
    )

    # 4. Trade Marks, Designs and SICLD
    add_heading1(doc, "4.  TRADE MARKS, DESIGNS AND SEMICONDUCTOR LAYOUT-DESIGNS")

    add_heading2(doc, "4.1  Trade Marks")
    add_body_paragraph(
        doc,
        "The Assignor hereby assigns to the Assignee, together with the "
        "goodwill of the business in which such marks are used, all rights in "
        "any trade marks, service marks, trade names, logos, get-up or domain "
        "names created, coined, designed, selected or first used by the "
        "Assignor in the course of the Engagement. The Assignor shall render "
        "all assistance required by the Assignee to secure registration of "
        "such assignments under Sections 37 to 45 of the Trade Marks Act, 1999, "
        "it being acknowledged that, under Section 45 of that Act, an assignee "
        "is not entitled to sue for infringement unless the assignment is "
        "entered on the register.",
    )

    add_heading2(doc, "4.2  Designs")
    add_body_paragraph(
        doc,
        "The Assignor hereby assigns to the Assignee all right, title and "
        "interest in any design (whether registrable or registered) within the "
        "meaning of the Designs Act, 2000, created in the course of the "
        "Engagement, and shall execute and deliver such documents as are "
        "required under Section 30 of the Designs Act, 2000 and the Designs "
        "Rules, 2001 to register the assignment with the Controller of Designs.",
    )

    add_heading2(doc, "4.3  Semiconductor Integrated Circuit Layout-Designs")
    add_body_paragraph(
        doc,
        "The Assignor hereby assigns to the Assignee all rights in any layout-"
        "design within the meaning of the Semiconductor Integrated Circuits "
        "Layout-Design Act, 2000 created by the Assignor in the course of the "
        "Engagement, together with the exclusive right to apply for, prosecute "
        "and maintain registration thereof under that Act.",
    )

    # 5. Moral Rights Waiver
    add_heading1(doc, "5.  MORAL RIGHTS")

    add_heading2(doc, "5.1  Waiver to the Maximum Extent Permitted")
    add_body_paragraph(
        doc,
        "The Assignor irrevocably and unconditionally waives, to the maximum "
        "extent permitted by applicable law, all Moral Rights (including the "
        "right of paternity and the right of integrity under Section 57 of the "
        "Copyright Act, 1957, and any equivalent rights in any other "
        "jurisdiction) in respect of every Foreground IP and Work Product, in "
        "favour of the Assignee and any person authorised by the Assignee.",
    )

    add_heading2(doc, "5.2  Undertaking Not to Assert")
    add_body_paragraph(
        doc,
        "The Assignor further undertakes that the Assignor shall not, directly "
        "or indirectly, institute, maintain or support any claim, action, "
        "proceeding or objection against the Assignee or any person claiming "
        "through or under the Assignee, asserting any Moral Rights, including "
        "any objection based on alleged distortion, mutilation, modification "
        "or other act prejudicial to honour or reputation. Where any such "
        "right is incapable of full waiver under applicable law, the Assignor "
        "consents to all acts or omissions of the Assignee and its licensees "
        "that would otherwise constitute an infringement of such right.",
    )

    add_heading2(doc, "5.3  Acknowledgement of Statutory Limits")
    add_body_paragraph(
        doc,
        "The Parties acknowledge that certain aspects of the Moral Rights "
        "(including protection against distortion, mutilation or modification "
        "prejudicial to the author's honour or reputation) may not be wholly "
        "waivable under Section 57 of the Copyright Act, 1957. This Clause 5 "
        "shall operate to the maximum extent legally permissible, and the "
        "unenforceability of any part of it shall not affect the remainder.",
    )

    # 6. Consultant carve-out
    add_heading1(doc, "6.  CONSULTANT AND NON-EMPLOYEE CARVE-OUT")
    add_body_paragraph(
        doc,
        "Where the Assignor is engaged by the Assignee otherwise than under a "
        "contract of service (including as an independent contractor, "
        "consultant, adviser, intern, founder, director or seconded personnel), "
        "the Parties acknowledge that Section 17(c) of the Copyright Act, 1957 "
        "does not operate to vest first ownership of copyright in the Assignee "
        "by default. In such cases, the Parties expressly rely on the "
        "assignment effected by Clause 2 of this Deed, and the Assignor "
        "irrevocably affirms that the present-tense assignment contained herein "
        "is intended to and does transfer all right, title and interest in the "
        "Foreground IP to the Assignee notwithstanding the absence of any "
        "contract of service.",
    )

    # 7. Prior IP Disclosure and Licence-Back
    add_heading1(doc, "7.  PRIOR IP DISCLOSURE AND LICENCE-BACK")

    add_heading2(doc, "7.1  Disclosure")
    add_body_paragraph(
        doc,
        "The Assignor represents and warrants that Schedule I contains a "
        "complete and accurate list of all Prior IP that the Assignor wishes "
        "to exclude from the scope of this Deed. Any Intellectual Property "
        "Right not listed in Schedule I shall be conclusively deemed not to "
        "constitute Prior IP and shall, if incorporated in or necessary for the "
        "exploitation of any Foreground IP, be assigned to the Assignee under "
        "Clause 2.",
    )

    add_heading2(doc, "7.2  Licence-Back of Prior IP")
    add_body_paragraph(
        doc,
        "To the extent that any Prior IP is incorporated in, combined with, or "
        "necessary for the lawful use, exploitation, modification or further "
        "development of any Foreground IP or any product, service or "
        "deliverable of the Assignee, the Assignor hereby grants the Assignee "
        "a perpetual, irrevocable, worldwide, royalty-free, fully paid-up, "
        "non-exclusive, transferable and sub-licensable licence to use, "
        "reproduce, modify, create derivative works of, distribute, perform, "
        "display, make, have made, sell, import and otherwise exploit such "
        "Prior IP for any purpose.",
    )

    add_heading2(doc, "7.3  No Incorporation Without Disclosure")
    add_body_paragraph(
        doc,
        "The Assignor shall not knowingly incorporate any Prior IP in any "
        "Foreground IP, Work Product or deliverable without first disclosing "
        "such incorporation in writing to the Assignee and obtaining the "
        "Assignee's written consent.",
    )

    # 8. Further Assurance
    add_heading1(doc, "8.  FURTHER ASSURANCE AND POWER OF ATTORNEY")

    add_heading2(doc, "8.1  Further Assurance")
    add_body_paragraph(
        doc,
        "The Assignor shall, promptly upon the written request of the Assignee "
        "and at the Assignee's reasonable cost, do and execute, or procure to "
        "be done and executed, all such further acts, deeds, confirmations, "
        "assignments, declarations, oaths, affidavits, inventor's declarations, "
        "forms (including, as applicable, Form 16 under the Patents Rules, 2003, "
        "Form TM-P under the Trade Marks Rules, 2017, Form 12 under the Designs "
        "Rules, 2001 and the form prescribed under Rule 19 of the Copyright "
        "Rules, 2013 for registration of assignment), applications, powers of "
        "attorney and other documents as may be reasonably required by the "
        "Assignee to perfect, confirm, register, protect, enforce or otherwise "
        "exploit the Assignee's rights in and to the Foreground IP, whether in "
        "India or in any other jurisdiction.",
    )

    add_heading2(doc, "8.2  Power of Attorney")
    add_body_paragraph(
        doc,
        "The Assignor hereby irrevocably appoints the Assignee (and any "
        "officer nominated by the Board of Directors of the Assignee from time "
        "to time) as the Assignor's true and lawful attorney, such appointment "
        "being coupled with an interest, to execute, sign, swear, file and do "
        "all acts, deeds, forms and documents necessary or desirable to give "
        "full effect to this Deed and to perfect, register, prosecute, defend "
        "and enforce the Assignee's rights in and to the Foreground IP, in "
        "each case where the Assignor is, after reasonable written notice, "
        "unable, unwilling or has failed to do so within fifteen (15) days of "
        "being requested by the Assignee.",
    )

    # 9. Warranties
    add_heading1(doc, "9.  REPRESENTATIONS AND WARRANTIES")
    add_body_paragraph(doc, "The Assignor represents and warrants to the Assignee that:")
    add_numbered_sub(
        doc,
        [
            "the Assignor has full legal capacity, power and authority to enter "
            "into this Deed and to effect the assignment contemplated herein;",
            "the Assignor is the sole and exclusive author and original creator "
            "of every Foreground IP, save where expressly identified as "
            "jointly created with other personnel of the Assignee;",
            "the Foreground IP is, and shall be, original to the Assignor and "
            "does not and shall not infringe, misappropriate or otherwise "
            "violate the Intellectual Property Rights, privacy rights, "
            "publicity rights or other legal rights of any third party;",
            "the Assignor has not granted, and shall not grant, any licence, "
            "assignment, security interest, lien, charge, encumbrance or other "
            "right in or to any Foreground IP to any person other than the "
            "Assignee;",
            "the Assignor has not used and shall not use any Open-Source "
            "Software in the development of any Foreground IP in a manner that "
            "(i) subjects any Foreground IP or any proprietary software of the "
            "Assignee to any copyleft obligation, (ii) requires the public "
            "distribution of source code, or (iii) restricts the commercial "
            "exploitation of any Foreground IP, without in each case first "
            "disclosing such use in writing to the Assignee and obtaining the "
            "Assignee's prior written approval, and in particular no component "
            "licensed under the GNU General Public Licence, the GNU Lesser "
            "General Public Licence or the GNU Affero General Public Licence "
            "has been introduced without such approval;",
            "no Foreground IP has been developed using funding, facilities or "
            "resources of any government, academic institution, former employer "
            "or third-party grant-making body (including DSIR, BIRAC or similar) "
            "in a manner that would impose any residual rights on such third "
            "party, save as disclosed in Schedule I;",
            "the Assignor is not subject to any agreement, court order, "
            "settlement or other obligation that conflicts with, restricts or "
            "would be breached by the Assignor's performance of this Deed; and",
            "all statements made and information furnished by the Assignor in "
            "this Deed, including Schedule I, are true, complete and not "
            "misleading in any material respect.",
        ],
    )

    # 10. Indemnity
    add_heading1(doc, "10.  INDEMNITY")
    add_body_paragraph(
        doc,
        "10.1  The Assignor shall indemnify, defend and hold harmless the "
        "Assignee, its Affiliates and their respective directors, officers, "
        "employees and agents from and against any and all losses, damages, "
        "liabilities, costs and expenses (including reasonable legal fees) "
        "suffered or incurred by reason of any breach by the Assignor of any "
        "representation, warranty, covenant or obligation contained in this "
        "Deed, including any third-party claim alleging that any Foreground IP "
        "infringes the Intellectual Property Rights of such third party.",
    )
    add_body_paragraph(
        doc,
        "10.2  The aggregate liability of the Assignor under Clause 10.1 "
        "(except in the case of fraud, wilful misconduct or breach of Moral "
        "Rights undertakings) shall not exceed an amount equal to "
        "[INDEMNITY CAP - e.g., twelve (12) months' total consideration paid "
        "by the Assignee to the Assignor under the Engagement]. For the "
        "avoidance of doubt, no indemnity cap shall apply to liability arising "
        "from the Assignor's fraud or wilful misrepresentation.",
    )

    # 11. Stamp Duty
    add_heading1(doc, "11.  STAMP DUTY AND REGISTRATION")

    add_heading2(doc, "11.1  Stamp Duty Acknowledgement")
    add_body_paragraph(
        doc,
        "The Parties acknowledge that this Deed is chargeable to stamp duty "
        "under the Indian Stamp Act, 1899 (as applicable to the NCT of Delhi) and the Indian Stamp Act, 1899 as "
        "applicable to the National Capital Territory of Delhi. The Assignee shall bear the "
        "cost of the stamp duty payable on this Deed and shall ensure that the "
        "Deed is duly stamped prior to, or contemporaneously with, execution, "
        "whether through e-stamp issued under the e-Stamp procedure "
        "(administered by the Stock Holding Corporation of India Limited) or "
        "through such other permissible mode. The Parties note that under "
        "Section 35 of the Indian Stamp Act, 1899, an instrument that is not "
        "duly stamped shall not be admitted in evidence until the deficient "
        "duty (together with penalty of up to ten (10) times the deficient "
        "duty) has been paid.",
    )

    add_heading2(doc, "11.2  Registration with IP Offices")
    add_body_paragraph(
        doc,
        "The Parties further acknowledge that, in addition to stamp duty:",
    )
    add_numbered_sub(
        doc,
        [
            "an assignment of copyright may be recorded with the Registrar of "
            "Copyrights in accordance with the Copyright Rules, 2013;",
            "an assignment of a patent shall be registered with the Controller "
            "of Patents under Section 68 of the Patents Act, 1970 and is not "
            "valid as against a person acquiring a later interest in the "
            "patent unless so registered;",
            "an assignment of a registered trade mark shall be recorded with "
            "the Registrar of Trade Marks under Section 45 of the Trade Marks "
            "Act, 1999, without which the assignee is not entitled to sue for "
            "infringement; and",
            "an assignment of a registered design shall be recorded with the "
            "Controller of Designs under Section 30 of the Designs Act, 2000.",
        ],
    )
    add_body_paragraph(
        doc,
        "11.3  The Assignee shall be responsible for effecting all such "
        "registrations at its own cost, and the Assignor shall render all "
        "reasonable assistance required for that purpose pursuant to Clause 8.",
    )

    # 12. Confidentiality (link)
    add_heading1(doc, "12.  CONFIDENTIALITY")
    add_body_paragraph(
        doc,
        "The Parties' respective rights and obligations with respect to "
        "Confidential Information shall be governed by the non-disclosure "
        "agreement executed between the Parties (\"NDA\") and, in addition, by "
        "such obligations as are set out in the Engagement documentation. "
        "Nothing in this Deed shall derogate from, or be construed to waive, "
        "such obligations, which survive termination of the Engagement and of "
        "this Deed in accordance with their terms.",
    )

    # 13. Term
    add_heading1(doc, "13.  TERM AND SURVIVAL")
    add_body_paragraph(
        doc,
        "This Deed takes effect on the Effective Date and shall be perpetual "
        "and irrevocable. Termination, expiry or rescission of the Engagement "
        "shall not affect the assignment effected by this Deed. All warranties, "
        "indemnities, waivers, licences and obligations hereunder shall survive "
        "any such termination, expiry or rescission.",
    )

    # 14. Governing Law and Dispute Resolution
    add_heading1(doc, "14.  GOVERNING LAW AND DISPUTE RESOLUTION")

    add_heading2(doc, "14.1  Governing Law")
    add_body_paragraph(
        doc,
        "This Deed shall be governed by, and construed in accordance with, the "
        "substantive laws of the Republic of India, without reference to its "
        "conflict-of-laws principles.",
    )

    add_heading2(doc, "14.2  Arbitration")
    add_body_paragraph(
        doc,
        "Any dispute, controversy or claim arising out of, in connection with, "
        "or in relation to this Deed, including any question regarding its "
        "existence, validity, interpretation, breach or termination, shall be "
        "referred to and finally resolved by arbitration under the Arbitration "
        "and Conciliation Act, 1996 (as amended). The arbitral tribunal shall "
        "consist of a sole arbitrator to be appointed by mutual agreement of "
        "the Parties, or, failing such agreement within thirty (30) days, in "
        "accordance with Section 11 of that Act. The seat and venue of "
        "arbitration shall be New Delhi, Delhi, India. The language of the "
        "arbitration shall be English. The arbitral award shall be final and "
        "binding on the Parties.",
    )

    add_heading2(doc, "14.3  Jurisdiction")
    add_body_paragraph(
        doc,
        "Subject to Clause 14.2, the courts at New Delhi, Delhi shall have "
        "exclusive jurisdiction in respect of any matter arising from or in "
        "connection with this Deed, including for the grant of any interim or "
        "conservatory relief under Section 9 of the Arbitration and "
        "Conciliation Act, 1996.",
    )

    # 15. Miscellaneous
    add_heading1(doc, "15.  MISCELLANEOUS")

    add_heading2(doc, "15.1  Entire Agreement")
    add_body_paragraph(
        doc,
        "This Deed, together with the Engagement documentation and the NDA, "
        "constitutes the entire agreement between the Parties with respect to "
        "its subject matter and supersedes all prior understandings, "
        "communications and agreements.",
    )

    add_heading2(doc, "15.2  Severability")
    add_body_paragraph(
        doc,
        "If any provision of this Deed is held to be invalid, illegal or "
        "unenforceable, the remaining provisions shall continue in full force "
        "and effect, and the Parties shall negotiate in good faith a valid, "
        "legal and enforceable substitute provision that most nearly reflects "
        "the original intent.",
    )

    add_heading2(doc, "15.3  No Waiver")
    add_body_paragraph(
        doc,
        "No failure or delay by either Party to exercise any right or remedy "
        "under this Deed shall operate as a waiver thereof. A waiver of any "
        "right or remedy shall be effective only if given in writing and "
        "shall not be deemed a waiver of any subsequent breach or default.",
    )

    add_heading2(doc, "15.4  Amendment")
    add_body_paragraph(
        doc,
        "No amendment to or variation of this Deed shall be valid unless made "
        "in writing and signed by or on behalf of each of the Parties.",
    )

    add_heading2(doc, "15.5  Counterparts and Electronic Execution")
    add_body_paragraph(
        doc,
        "This Deed may be executed in two or more counterparts, each of which "
        "shall be deemed an original and all of which together shall "
        "constitute one and the same instrument. Execution by Aadhaar-based "
        "e-signature or by a digital signature certificate issued under "
        "Section 3A of the Information Technology Act, 2000 shall be deemed "
        "valid execution, provided that stamp-duty requirements are separately "
        "complied with.",
    )

    # Execution block
    add_heading1(doc, "IN WITNESS WHEREOF")
    add_body_paragraph(
        doc,
        "the Parties have executed this Deed of Assignment of Intellectual "
        "Property Rights on the day, month and year first written above.",
    )

    add_blank(doc)

    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.autofit = True
    left, right = sig_table.rows[0].cells

    def populate_sig_cell(cell, title, name_label, lines):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(title)
        style_run(r, size=11, bold=True)
        cell.add_paragraph("")
        cell.add_paragraph("_________________________________")
        p2 = cell.add_paragraph()
        r2 = p2.add_run(name_label)
        style_run(r2, size=11, bold=True)
        for line in lines:
            q = cell.add_paragraph(line)
            set_paragraph_format(q, justify=False, space_after=2)
            for r_ in q.runs:
                style_run(r_, size=11)

    populate_sig_cell(
        left,
        "SIGNED AND DELIVERED AS A DEED BY THE ASSIGNOR",
        "Name: [FULL LEGAL NAME OF ASSIGNOR]",
        [
            "PAN: [PAN NUMBER]",
            "Address: [ADDRESS]",
            "Date: _______________________",
            "Place: New Delhi, Delhi",
        ],
    )
    populate_sig_cell(
        right,
        "SIGNED AND DELIVERED AS A DEED BY THE ASSIGNEE",
        "For and on behalf of ANANTASUTRA",
        [
            "Name: Mr. Himanshu Mishra",
            "Designation: Founder & CEO",
            "Email: contact@anantasutra.com",
            "Date: [DATE]",
            "Place: Delhi",
        ],
    )

    add_blank(doc)
    add_heading2(doc, "WITNESSES:")
    witness_table = doc.add_table(rows=1, cols=2)
    w1, w2 = witness_table.rows[0].cells

    def populate_witness_cell(cell, idx):
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(f"Witness {idx}")
        style_run(r, size=11, bold=True)
        cell.add_paragraph("_________________________________")
        for line in (
            "Name: _________________________________",
            "Address: _______________________________",
            "___________________________________________",
            "Signature: ______________________________",
            "Date: _______________________",
        ):
            q = cell.add_paragraph(line)
            set_paragraph_format(q, justify=False, space_after=2)
            for r_ in q.runs:
                style_run(r_, size=11)

    populate_witness_cell(w1, 1)
    populate_witness_cell(w2, 2)

    # Schedule I
    add_page_break(doc)
    add_centered_title(doc, "SCHEDULE I - PRIOR IP / BACKGROUND IP DISCLOSURE", size=13)
    add_body_paragraph(
        doc,
        "The Assignor hereby identifies the following items as Prior IP / "
        "Background IP to be excluded from the scope of the assignment under "
        "this Deed (attach additional sheets if required). If this Schedule "
        "is left blank or marked \"NIL\", the Assignor confirms that no Prior "
        "IP exists.",
    )
    prior_table = doc.add_table(rows=6, cols=5)
    prior_table.style = "Table Grid"
    prior_headers = [
        "S. No.",
        "Title / Description",
        "Type of IP (copyright / patent / TM / design / trade secret)",
        "Date of Creation",
        "Jurisdiction / Registration No. (if any)",
    ]
    for i, h in enumerate(prior_headers):
        cell = prior_table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        style_run(r, size=10, bold=True)
    for r_idx in range(1, 6):
        for c_idx in range(5):
            cell = prior_table.rows[r_idx].cells[c_idx]
            if c_idx == 0:
                cell.text = str(r_idx)
                for run in cell.paragraphs[0].runs:
                    style_run(run, size=10)
            else:
                cell.text = ""

    add_blank(doc)
    add_body_paragraph(
        doc,
        "Signed by the Assignor in confirmation of the contents of this "
        "Schedule:    ______________________________    Date: ________________",
    )


# ---------------------------------------------------------------------------
# PART B — IP Policy
# ---------------------------------------------------------------------------
def build_part_b(doc: Document) -> None:
    add_page_break(doc)
    add_centered_title(
        doc,
        "ANANTASUTRA INTELLECTUAL PROPERTY POLICY",
        size=16,
        space_before=0,
        space_after=10,
    )

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "Policy Owner: Mr. Himanshu Mishra, Founder & CEO, AnantaSutra\n"
        "Approved by the Founder & CEO on: [DATE OF APPROVAL]\n"
        "Version: 1.0    |    Effective Date: [EFFECTIVE DATE]    |    "
        "Next Review: [REVIEW DATE]"
    )
    style_run(run, size=10, italic=True)

    add_blank(doc)

    # 1. Purpose
    add_heading1(doc, "1.  PURPOSE AND SCOPE")

    add_heading2(doc, "1.1  Purpose")
    add_body_paragraph(
        doc,
        "This Intellectual Property Policy (\"Policy\") sets out the principles, "
        "standards and procedures governing the creation, ownership, "
        "protection, management, commercialisation and enforcement of "
        "intellectual property (\"IP\") at AnantaSutra (\"AnantaSutra\" "
        "or the \"Company\"). The Policy is intended to ensure that (a) "
        "AnantaSutra's chain of title to all IP used in its business is complete "
        "and defensible, (b) statutory obligations under applicable Indian and "
        "international IP laws are satisfied, and (c) the Company's products "
        "and services do not infringe the IP rights of any third party.",
    )

    add_heading2(doc, "1.2  Scope")
    add_body_paragraph(
        doc,
        "This Policy applies to all personnel of AnantaSutra, including (a) "
        "full-time and part-time employees, (b) fixed-term and probationary "
        "employees, (c) independent contractors, consultants and advisers, "
        "(d) interns, trainees and apprentices, (e) founders, promoters and "
        "directors, and (f) seconded, outsourced or agency personnel (each a "
        "\"Personnel\"). Every Personnel is required to comply with this Policy "
        "as a condition of engagement.",
    )

    add_heading2(doc, "1.3  Statutory Framework")
    add_body_paragraph(
        doc,
        "The Policy is designed to be consistent with the Copyright Act, 1957, "
        "the Patents Act, 1970, the Trade Marks Act, 1999, the Designs Act, "
        "2000, the Semiconductor Integrated Circuits Layout-Design Act, 2000, "
        "the Indian Contract Act, 1872, the Information Technology Act, 2000, "
        "the Indian Stamp Act, 1899, the Companies Act, 2013 and the Digital "
        "Personal Data Protection Act, 2023, and shall be interpreted "
        "accordingly.",
    )

    # 2. Ownership principle
    add_heading1(doc, "2.  OWNERSHIP PRINCIPLE")

    add_heading2(doc, "2.1  Default Ownership")
    add_body_paragraph(
        doc,
        "All Intellectual Property Rights in any Invention, Work Product or "
        "material created by any Personnel (a) in the course of, or as a "
        "result of, engagement with the Company, (b) using the Company's "
        "time, resources, facilities, Confidential Information or Background "
        "IP, or (c) otherwise within the scope of the business of the "
        "Company, shall be the sole and exclusive property of AnantaSutra.",
    )

    add_heading2(doc, "2.2  Mechanism of Vesting")
    add_body_paragraph(
        doc,
        "Vesting of IP in AnantaSutra shall be effected by a combination of (a) "
        "the operation of Section 17(c) of the Copyright Act, 1957 in the case "
        "of employees engaged under a contract of service, and (b) the "
        "executed Deed of Assignment of Intellectual Property Rights (the "
        "\"IP Assignment Deed\") referred to in Part A of this document, which "
        "every Personnel is required to execute. Where a Personnel is not an "
        "employee, the IP Assignment Deed is the primary and indispensable "
        "instrument of vesting.",
    )

    add_heading2(doc, "2.3  No Joint Ownership by Default")
    add_body_paragraph(
        doc,
        "AnantaSutra's default position is sole ownership. Joint ownership of IP "
        "with any third party shall not be created without the prior written "
        "approval of the IP Officer and the Board of Directors. Where joint "
        "ownership is unavoidable, the arrangement shall be documented so as "
        "to mitigate the statutory consequences under Section 50 of the "
        "Patents Act, 1970 and comparable provisions.",
    )

    # 3. Invention Disclosure
    add_heading1(doc, "3.  INVENTION DISCLOSURE PROCESS")

    add_heading2(doc, "3.1  Mandatory Disclosure")
    add_body_paragraph(
        doc,
        "Every Personnel shall promptly, and in any event within fifteen (15) "
        "working days of conception or recognition, disclose to the IP "
        "Officer any Invention, creative work, design or other item that may "
        "constitute IP, whether or not the Personnel considers it to be "
        "patentable, registrable or otherwise protectable. Disclosure shall "
        "be made using the Invention Disclosure Form set out in Schedule II.",
    )

    add_heading2(doc, "3.2  Review by IP Officer")
    add_body_paragraph(
        doc,
        "The IP Officer, in consultation with the Chief Technology Officer and "
        "external counsel where appropriate, shall review every disclosure and "
        "determine (a) whether the Invention is to be maintained as a trade "
        "secret, (b) whether to pursue patent, design or trade mark "
        "registration, (c) whether copyright registration is warranted under "
        "Section 45 of the Copyright Act, 1957, (d) whether any defensive "
        "publication is advisable, and (e) the appropriate inventor-recognition "
        "reward under Clause 3.4.",
    )

    add_heading2(doc, "3.3  No Effect on Ownership")
    add_body_paragraph(
        doc,
        "A Personnel's failure to disclose an Invention shall not affect "
        "AnantaSutra's title thereto, which vests automatically under Clause 2. "
        "Such failure may however constitute grounds for disciplinary action "
        "under Clause 12.",
    )

    add_heading2(doc, "3.4  Inventor Recognition")
    add_body_paragraph(
        doc,
        "AnantaSutra may, at its sole discretion, grant recognition awards to "
        "inventors, including (a) an award of [INR 25,000] on filing of a "
        "patent application naming the Personnel as inventor, and (b) an "
        "award of [INR 1,00,000] on grant of a patent. Such awards are "
        "ex gratia and do not constitute or imply any residual ownership "
        "interest in the IP.",
    )

    # 4. Prior IP declaration
    add_heading1(doc, "4.  PRIOR IP DECLARATION REQUIREMENT")
    add_body_paragraph(
        doc,
        "Every Personnel shall, at the time of joining or engagement, submit a "
        "written declaration listing all pre-existing inventions, works or "
        "other IP that the Personnel wishes to exclude from the scope of "
        "AnantaSutra's ownership. Such declaration shall be made in Schedule I to "
        "the IP Assignment Deed. Any IP not listed in that declaration shall "
        "be conclusively deemed to be AnantaSutra's property if subsequently "
        "incorporated into the Company's products, services, codebase or "
        "other materials.",
    )

    # 5. Open-Source Software
    add_heading1(doc, "5.  OPEN-SOURCE SOFTWARE USE POLICY")

    add_heading2(doc, "5.1  Classification of Licences")
    add_body_paragraph(
        doc,
        "Open-source software licences are classified for internal purposes "
        "into three categories:",
    )
    add_numbered_sub(
        doc,
        [
            "Approved (permissive): the MIT Licence, the BSD 2-Clause and "
            "3-Clause Licences, the Apache Licence 2.0, the ISC Licence, the "
            "Python Software Foundation Licence and similar permissive "
            "licences. These may be used without individual approval, subject "
            "to attribution and notice requirements;",
            "Flagged (weak copyleft): the Mozilla Public Licence 2.0, the GNU "
            "Lesser General Public Licence (v2.1 / v3), the Eclipse Public "
            "Licence and similar licences. These require written approval by "
            "the IP Officer and a segregation plan to ensure that the "
            "Company's proprietary code does not become subject to the "
            "copyleft obligation;",
            "Restricted (strong copyleft): the GNU General Public Licence "
            "(v2 / v3), the GNU Affero General Public Licence (v3) and "
            "similar licences. These shall not be introduced into any "
            "product or codebase without the prior written approval of both "
            "the IP Officer and the Chief Technology Officer, which approval "
            "shall be granted, if at all, only after a documented "
            "contamination-risk assessment. AGPL components, in particular, "
            "are presumptively incompatible with AnantaSutra's SaaS model.",
        ],
    )

    add_heading2(doc, "5.2  Software Bill of Materials")
    add_body_paragraph(
        doc,
        "The engineering team shall maintain, for every product and every "
        "release, a Software Bill of Materials (\"SBOM\") identifying each "
        "open-source component, its version and its licence. The SBOM shall "
        "be updated on each release and audited by the IP Officer at least "
        "quarterly.",
    )

    add_heading2(doc, "5.3  Contamination Checks")
    add_body_paragraph(
        doc,
        "Automated licence-scanning tools approved by the IP Officer (such as "
        "scanning tools integrated into the CI/CD pipeline) shall be run on "
        "every merge to the main branch. Any component triggering a flagged "
        "or restricted licence match shall block the build pending review.",
    )

    add_heading2(doc, "5.4  Contribution to Open-Source Projects")
    add_body_paragraph(
        doc,
        "No Personnel shall contribute any code, documentation or artefact to "
        "an external open-source project on behalf of AnantaSutra, or in a "
        "manner identifying AnantaSutra as the contributor, without the prior "
        "written approval of the IP Officer. Personal open-source activity "
        "not involving AnantaSutra's Confidential Information or Foreground IP "
        "is permitted subject to Clause 9.",
    )

    # 6. Third-Party IP
    add_heading1(doc, "6.  USE OF THIRD-PARTY INTELLECTUAL PROPERTY")

    add_heading2(doc, "6.1  Licensing Discipline")
    add_body_paragraph(
        doc,
        "All third-party software, libraries, fonts, stock images, music, "
        "videos, data sets and AI-model outputs used in AnantaSutra's products, "
        "services, marketing materials or internal systems shall be used "
        "only under a valid licence that covers the intended use. Pirated, "
        "cracked or unlicensed software is strictly prohibited.",
    )

    add_heading2(doc, "6.2  AI-Generated Content")
    add_body_paragraph(
        doc,
        "Where generative AI tools are used to produce code, text, images, "
        "designs or other content for inclusion in AnantaSutra's products or "
        "deliverables, Personnel shall (a) use only AI tools approved by the "
        "IP Officer, (b) not input any Confidential Information or personal "
        "data into public AI tools, (c) mark AI-assisted outputs as such in "
        "the repository, and (d) ensure, to a reasonable extent, that the "
        "output does not infringe any third-party IP. Personnel are "
        "reminded that the Indian Copyright Office's present position is "
        "that copyright subsists only in works of human authorship, and "
        "AI-only outputs may not qualify for registration.",
    )

    add_heading2(doc, "6.3  Third-Party Brand and Trade Marks")
    add_body_paragraph(
        doc,
        "No Personnel shall use any third-party trade mark, logo or brand "
        "element in AnantaSutra materials except (a) in factual references that "
        "comply with the doctrine of nominative fair use under Indian trade "
        "marks law, or (b) under a written licence. Customer logos shall be "
        "used in case studies and marketing only with the customer's written "
        "consent.",
    )

    # 7. Trade Secrets and Confidentiality
    add_heading1(doc, "7.  TRADE SECRETS AND CONFIDENTIALITY")
    add_body_paragraph(
        doc,
        "AnantaSutra relies on trade-secret protection for significant portions "
        "of its technology stack, including proprietary algorithms, model "
        "weights, prompts, heuristics, customer lists, pricing models and "
        "non-public roadmaps. Every Personnel shall comply with (a) the "
        "AnantaSutra Employee Confidentiality and Non-Disclosure Agreement, (b) "
        "any applicable Client Non-Disclosure Agreement, (c) Section 72 and "
        "Section 72A of the Information Technology Act, 2000, and (d) "
        "internal access-control, tagging and clean-desk protocols issued by "
        "the IP Officer. Trade-secret materials shall be stored only in "
        "systems designated for restricted access.",
    )

    # 8. Publications, Talks, Demos
    add_heading1(doc, "8.  PUBLICATIONS, CONFERENCE TALKS, DEMONSTRATIONS")

    add_heading2(doc, "8.1  Pre-Clearance")
    add_body_paragraph(
        doc,
        "No Personnel shall publish, present, demonstrate or otherwise "
        "disclose in a public forum any information that (a) describes or "
        "discloses Foreground IP, (b) references AnantaSutra's customers, "
        "internal data or non-public roadmap, or (c) identifies the Personnel "
        "as a AnantaSutra representative, without first obtaining written "
        "clearance from the IP Officer.",
    )

    add_heading2(doc, "8.2  Patent-Bar Considerations")
    add_body_paragraph(
        doc,
        "Personnel shall ensure, in particular, that any publication or "
        "public demonstration does not prejudice the novelty of any Invention "
        "in respect of which AnantaSutra may wish to file a patent application. "
        "Where patent filing is under consideration, the publication shall be "
        "deferred until the earliest applicable priority date is secured.",
    )

    # 9. IP Protection Strategy
    add_heading1(doc, "9.  IP PROTECTION STRATEGY")
    add_body_paragraph(
        doc,
        "AnantaSutra's portfolio strategy shall be reviewed by the IP Officer at "
        "least annually and shall consider the following instruments of "
        "protection, selected according to the nature of the underlying IP:",
    )
    add_numbered_sub(
        doc,
        [
            "Patents - for technical inventions exhibiting a technical effect "
            "or being embodied with hardware, having regard to the scope of "
            "Section 3(k) of the Patents Act, 1970 and the Guidelines for "
            "Examination of Computer Related Inventions issued by the Indian "
            "Patent Office;",
            "Trade marks and service marks - for brands, product names, "
            "logos and taglines, including defensive filings in key "
            "jurisdictions and classes;",
            "Copyrights - for source code, documentation, marketing "
            "materials and databases, with optional registration under "
            "Section 45 of the Copyright Act, 1957 where evidentiary value is "
            "material;",
            "Registered designs - for distinctive user-interface layouts "
            "and industrial designs under the Designs Act, 2000;",
            "Semiconductor layout-design registrations - where applicable "
            "under the Semiconductor Integrated Circuits Layout-Design Act, "
            "2000;",
            "Trade-secret protection - for algorithms, model weights, "
            "training data, prompts and other know-how where registration is "
            "inappropriate or would be counter-productive; and",
            "Defensive publications - to create prior art where protection "
            "is not pursued but freedom-to-operate requires public "
            "disclosure.",
        ],
    )

    # 10. IP Registers
    add_heading1(doc, "10.  IP REGISTERS MAINTAINED BY ANANTASUTRA")
    add_body_paragraph(
        doc,
        "The IP Officer shall maintain, in a secure repository, the following "
        "registers, each updated at least quarterly:",
    )
    add_numbered_sub(
        doc,
        [
            "Register of Background IP - Prior IP declared by each "
            "Personnel at the time of engagement;",
            "Register of Foreground IP - all Inventions and Work Products "
            "arising from the Company's operations, indexed by disclosure "
            "date, contributors and status (trade-secret / filed / granted / "
            "published);",
            "Register of Registered IP - patents, trade marks, designs, "
            "copyrights and layout-designs filed and/or granted, with "
            "renewal diary;",
            "Register of Executed IP Assignment Deeds - one entry per "
            "Personnel, with the date of execution, stamp-duty particulars "
            "and scanned copy;",
            "Register of Open-Source Components - derived from the SBOM;",
            "Register of Third-Party Licences - in-bound and out-bound.",
        ],
    )

    # 11. Exit Obligations
    add_heading1(doc, "11.  EXIT OBLIGATIONS")
    add_body_paragraph(
        doc,
        "Upon cessation of engagement (whether by resignation, termination, "
        "completion of the consultancy term, retirement or otherwise) every "
        "Personnel shall:",
    )
    add_numbered_sub(
        doc,
        [
            "return to AnantaSutra all property of the Company, including all "
            "devices, documents, source code, data, drawings, credentials, "
            "access cards, and Confidential Information in any form;",
            "permanently delete, to the extent practicable, all copies of "
            "AnantaSutra Confidential Information and Foreground IP held on "
            "personal devices or cloud accounts, and certify such deletion "
            "in writing;",
            "disclose in writing to the IP Officer all Inventions conceived "
            "or being developed as at the date of cessation, whether or not "
            "previously disclosed;",
            "attend an exit interview with the IP Officer, at which the "
            "foregoing shall be reviewed;",
            "execute any further confirmatory assignment, affidavit or form "
            "reasonably required by the Company to perfect its rights in any "
            "Foreground IP (Further Assurance); and",
            "respect, for the survival period set out in the Employee NDA, "
            "all obligations of confidentiality and non-solicitation, "
            "remembering that under Section 27 of the Indian Contract Act, "
            "1872 post-employment non-compete obligations are not "
            "enforceable in India and shall accordingly not be imposed by "
            "the Company.",
        ],
    )

    # 12. Violations
    add_heading1(doc, "12.  VIOLATIONS AND CONSEQUENCES")
    add_body_paragraph(
        doc,
        "Violation of this Policy may result in (a) disciplinary action up "
        "to and including termination of engagement, (b) recovery of costs "
        "and losses incurred by AnantaSutra (including costs of remediation, "
        "legal fees and third-party damages), (c) forfeiture of discretionary "
        "bonuses, incentives or unvested equity, and (d) civil and criminal "
        "proceedings under the Copyright Act, 1957 (including Section 63), "
        "the Patents Act, 1970, the Trade Marks Act, 1999, the Information "
        "Technology Act, 2000 and other applicable law. The Company may also "
        "report suspected IP theft to law-enforcement authorities.",
    )

    # 13. Review and Amendment
    add_heading1(doc, "13.  REVIEW AND AMENDMENT")
    add_body_paragraph(
        doc,
        "This Policy shall be reviewed by the IP Officer at least once every "
        "twelve (12) months, or sooner upon any material change in law or in "
        "the Company's business. Amendments shall be approved by the Board "
        "of Directors and circulated to all Personnel. The version "
        "in force from time to time, as published on the Company's internal "
        "repository, shall prevail.",
    )

    # 14. Interpretation
    add_heading1(doc, "14.  INTERPRETATION AND GOVERNING LAW")
    add_body_paragraph(
        doc,
        "Capitalised terms used but not defined in this Policy shall bear "
        "the meanings assigned to them in the IP Assignment Deed. This "
        "Policy shall be interpreted and applied in accordance with the laws "
        "of India, and any dispute arising in respect of its application "
        "shall, in the first instance, be referred to the IP Officer and the "
        "Board of Directors.",
    )

    # Schedule II - Invention Disclosure Form
    add_page_break(doc)
    add_centered_title(
        doc,
        "SCHEDULE II - INVENTION DISCLOSURE FORM",
        size=14,
        space_after=6,
    )
    add_body_paragraph(
        doc,
        "This Form is to be completed by every Personnel disclosing an "
        "Invention to the Founder & CEO of AnantaSutra, as "
        "required by Clause 3 of the AnantaSutra Intellectual Property Policy. "
        "Each Invention should be disclosed on a separate form. Continuation "
        "sheets may be attached.",
        italic=True,
    )

    add_blank(doc)

    # Disclosure identification
    disclosure_rows = [
        ("Disclosure Reference No.", "[To be assigned by IP Officer]"),
        ("Date of Disclosure", "_______________________________"),
        ("Title of the Invention", "_______________________________________________________________"),
        ("Primary Inventor - Full Name", "_______________________________________________________________"),
        ("Primary Inventor - Employee / Consultant ID", "_______________________________"),
        ("Primary Inventor - Email", "_______________________________"),
        ("Co-Inventors (if any) - Names and IDs", "_______________________________________________________________"),
        ("Department / Team", "_______________________________"),
        ("Project or Product Concerned", "_______________________________"),
        ("Date / Period of Conception", "_______________________________"),
        ("Date of First Written Record (Notebook / Repository)", "_______________________________"),
        ("Date of First Reduction to Practice", "_______________________________"),
    ]
    t = doc.add_table(rows=len(disclosure_rows), cols=2)
    t.style = "Table Grid"
    for i, (label, blank) in enumerate(disclosure_rows):
        left_cell = t.rows[i].cells[0]
        right_cell = t.rows[i].cells[1]
        left_cell.text = ""
        p = left_cell.paragraphs[0]
        r = p.add_run(label)
        style_run(r, size=10, bold=True)
        right_cell.text = ""
        p2 = right_cell.paragraphs[0]
        r2 = p2.add_run(blank)
        style_run(r2, size=10)

    add_blank(doc)
    add_heading2(doc, "PART A - TECHNICAL DESCRIPTION")

    tech_rows = [
        "1. Brief description of the Invention (2-3 sentences, for index purposes):",
        "2. Problem solved / technical or commercial need addressed:",
        "3. Existing solutions / known prior art of which the Inventor is aware:",
        "4. Novel or inventive features (identify what is different from the prior art):",
        "5. Technical effect or advantage (for evaluation against Section 3(k), Patents Act, 1970):",
        "6. Detailed description of the Invention (attach additional sheets):",
        "7. Drawings / diagrams / flowcharts / screenshots attached? (Y / N):",
        "8. Source-code repository path and commit hash (if applicable):",
    ]
    tech_table = doc.add_table(rows=len(tech_rows), cols=1)
    tech_table.style = "Table Grid"
    for i, prompt in enumerate(tech_rows):
        cell = tech_table.rows[i].cells[0]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(prompt)
        style_run(r, size=10, bold=True)
        for _ in range(3):
            blank_p = cell.add_paragraph("")
            set_paragraph_format(blank_p, justify=False, space_after=0)

    add_blank(doc)
    add_heading2(doc, "PART B - USE OF THIRD-PARTY AND OPEN-SOURCE MATERIAL")

    third_party_rows = [
        ("Has any third-party or open-source code, data, model or library been used in the Invention?", "Yes / No"),
        ("If Yes, list each component, its licence and version:", "___________________________________"),
        ("Has any Prior IP of the Inventor been incorporated? (If Yes, refer to Schedule I of the IP Assignment Deed.)", "Yes / No"),
        ("Has any Confidential Information of a customer been used?", "Yes / No"),
        ("Has the Invention been disclosed to any person outside AnantaSutra?", "Yes / No"),
        ("If Yes, to whom, when and under what confidentiality arrangements?", "___________________________________"),
        ("Has the Invention been published, presented or demonstrated publicly?", "Yes / No"),
        ("If Yes, provide particulars (venue, date, audience):", "___________________________________"),
    ]
    tp_table = doc.add_table(rows=len(third_party_rows), cols=2)
    tp_table.style = "Table Grid"
    for i, (label, blank) in enumerate(third_party_rows):
        left_cell = tp_table.rows[i].cells[0]
        right_cell = tp_table.rows[i].cells[1]
        left_cell.text = ""
        p = left_cell.paragraphs[0]
        r = p.add_run(label)
        style_run(r, size=10, bold=True)
        right_cell.text = ""
        p2 = right_cell.paragraphs[0]
        r2 = p2.add_run(blank)
        style_run(r2, size=10)

    add_blank(doc)
    add_heading2(doc, "PART C - INVENTOR'S DECLARATION")
    add_body_paragraph(
        doc,
        "I, the undersigned, hereby certify that the information set out in "
        "this Form is true, complete and not misleading to the best of my "
        "knowledge and belief. I confirm that the Invention was conceived "
        "and developed in the course of my engagement with AnantaSutra Private "
        "Limited or using AnantaSutra's resources, and I acknowledge that all "
        "right, title and interest in and to the Invention vest in AnantaSutra "
        "Private Limited pursuant to the Deed of Assignment of Intellectual "
        "Property Rights executed by me and the AnantaSutra Intellectual "
        "Property Policy.",
    )

    add_blank(doc)

    sig = doc.add_table(rows=1, cols=2)
    sig.autofit = True
    l, r = sig.rows[0].cells

    def fill_disclosure_sig(cell, who, lines):
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(who)
        style_run(run, size=11, bold=True)
        cell.add_paragraph("")
        cell.add_paragraph("_________________________________")
        for line in lines:
            q = cell.add_paragraph(line)
            set_paragraph_format(q, justify=False, space_after=2)
            for rr in q.runs:
                style_run(rr, size=11)

    fill_disclosure_sig(
        l,
        "Inventor",
        [
            "Name: _________________________________",
            "Employee / Consultant ID: _____________",
            "Date: _______________________",
            "Place: _______________________",
        ],
    )
    fill_disclosure_sig(
        r,
        "Acknowledged - IP Officer",
        [
            "Name: _________________________________",
            "Designation: __________________________",
            "Date of Receipt: _______________________",
            "Disclosure Reference No.: ______________",
        ],
    )

    add_blank(doc)
    add_heading2(doc, "FOR INTERNAL USE BY IP OFFICER")
    internal_rows = [
        ("Classification", "[ ] Trade secret  [ ] Patent candidate  [ ] Copyright  [ ] Design  [ ] Other: ______"),
        ("Freedom-to-Operate review required?", "Yes / No"),
        ("Open-source contamination review required?", "Yes / No"),
        ("Recommended action", "_______________________________________________________________"),
        ("Target filing jurisdictions (if any)", "_______________________________"),
        ("Decision of IP Officer / Board", "_______________________________________________________________"),
        ("Date of decision", "_______________________________"),
        ("Inventor-recognition award (if any)", "_______________________________"),
    ]
    int_table = doc.add_table(rows=len(internal_rows), cols=2)
    int_table.style = "Table Grid"
    for i, (label, blank) in enumerate(internal_rows):
        left_cell = int_table.rows[i].cells[0]
        right_cell = int_table.rows[i].cells[1]
        left_cell.text = ""
        p = left_cell.paragraphs[0]
        r_ = p.add_run(label)
        style_run(r_, size=10, bold=True)
        right_cell.text = ""
        p2 = right_cell.paragraphs[0]
        r2_ = p2.add_run(blank)
        style_run(r2_, size=10)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main() -> None:
    doc = Document()
    add_brand_header(doc)
    configure_document(doc)

    # Tune built-in Heading styles slightly so they render consistently.
    for name, size in (("Heading 1", 13), ("Heading 2", 12)):
        style = doc.styles[name]
        style.font.name = BASE_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x1F, 0x2D, 0x5C)

    build_part_a(doc)
    build_part_b(doc)

    # Footer must be configured after all sections are finalised.
    configure_footer(doc)

    out_dir = os.path.dirname(OUTPUT_PATH)
    os.makedirs(out_dir, exist_ok=True)
    doc.save(OUTPUT_PATH)

    size_bytes = os.path.getsize(OUTPUT_PATH)
    size_kb = size_bytes / 1024.0
    print(f"Written: {OUTPUT_PATH}")
    print(f"Size:    {size_bytes} bytes ({size_kb:.2f} KB)")


if __name__ == "__main__":
    main()
