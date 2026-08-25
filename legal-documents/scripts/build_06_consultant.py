"""
build_06_consultant.py

Generates: 06_Consultant_Agreement.docx
India-law-compliant Consultant / Independent Contractor Agreement for
AnantaSutra (New Delhi, Delhi).

Drafting anchors:
  - Indian Contract Act, 1872
  - Copyright Act, 1957 (ss.17, 18, 19, 19A, 57) -- present-tense assignment
  - Patents Act, 1970; Trade Marks Act, 1999; Designs Act, 2000
  - Information Technology Act, 2000
  - Digital Personal Data Protection Act, 2023
  - Income-tax Act, 1961 (s.194J); CGST Act, 2017
  - Specific Relief Act, 1963
  - Arbitration & Conciliation Act, 1996
  - Indian Stamp Act, 1899 (as applicable to NCT of Delhi)

Requires python-docx v1.2.0.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = (
    r"c:/Users/duneja8515/Desktop/bhavya/prsnl/sutranet/"
    r"legal-documents/06_Consultant_Agreement.docx"
)

BODY_FONT = "Calibri"
BODY_SIZE = 11

# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------


# ---------------------------------------------------------------------------
# Brand header (AnantaSutra) — injected by rebrand
# ---------------------------------------------------------------------------
LOGO_PATH = r"c:/Users/duneja8515/Desktop/bhavya/prsnl/sutranet/public/Favicon/logo-nobg.png"


def add_brand_header(doc):
    """Insert centered logo + brand lines at top of page 1. Never crashes."""
    try:
        from docx.shared import Inches as _Inches, Pt as _Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH as _CENTERED
    except Exception:
        return
    p = doc.add_paragraph()
    p.alignment = _CENTERED.CENTER
    try:
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=_Inches(1.2))
    except Exception:
        pass
    p2 = doc.add_paragraph()
    p2.alignment = _CENTERED.CENTER
    r2 = p2.add_run("AnantaSutra — अनन्तसूत्र")
    r2.italic = True
    r2.font.size = _Pt(10)
    p3 = doc.add_paragraph()
    p3.alignment = _CENTERED.CENTER
    r3 = p3.add_run("https://anantasutra.com  •  contact@anantasutra.com")
    r3.italic = True
    r3.font.size = _Pt(9)

def set_cell_border(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def style_run(run, *, bold=False, italic=False, size=BODY_SIZE, font=BODY_FONT,
              color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    # Ensure east-asian font also set to avoid Calibri fallback glitches
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)


def add_para(doc, text="", *, bold=False, italic=False, align=None,
             size=BODY_SIZE, space_after=6, space_before=0, left_indent=None,
             style=None):
    if style is not None:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    if align is not None:
        p.alignment = align
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if left_indent is not None:
        p.paragraph_format.left_indent = Inches(left_indent)
    if text:
        run = p.add_run(text)
        style_run(run, bold=bold, italic=italic, size=size)
    return p


def add_heading1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    style_run(run, bold=True, size=13)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    style_run(run, bold=True, size=12)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return p


def add_sub(doc, label, text):
    """Numbered sub-clause paragraph like '6.1 Label. Body text...'"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(f"{label} ")
    style_run(r1, bold=True)
    r2 = p.add_run(text)
    style_run(r2)
    return p


def add_bullet(doc, text, *, indent=0.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    style_run(run)
    return p


def add_lettered(doc, letter, text, *, indent=0.4):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(f"({letter}) ")
    style_run(r1, bold=True)
    r2 = p.add_run(text)
    style_run(r2)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break()
    run._r.append(OxmlElement("w:br"))
    run._r[-1].set(qn("w:type"), "page")


# ---------------------------------------------------------------------------
# Document configuration
# ---------------------------------------------------------------------------

def configure_document(doc):
    # A4, 1-inch margins
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base Normal style
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(BODY_SIZE)
    pf = style.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)

    # Footer: "Confidential -- AnantaSutra" + page number
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = fp.add_run("Confidential — AnantaSutra   |   Page ")
        style_run(r, size=9, italic=True)
        add_page_number_field(fp)


# ---------------------------------------------------------------------------
# Document content
# ---------------------------------------------------------------------------

def add_title_block(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("CONSULTANT / INDEPENDENT CONTRACTOR AGREEMENT")
    style_run(r, bold=True, size=16)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(12)
    r2 = p2.add_run("(Governed by the laws of the Republic of India)")
    style_run(r2, italic=True, size=11)


def add_parties(doc):
    add_heading1(doc, "1. PARTIES")
    add_para(
        doc,
        "This Consultant / Independent Contractor Agreement (the “Agreement”) "
        "is made and entered into at New Delhi, Delhi, on this "
        "[DAY] day of [MONTH], [YEAR] (the “Effective Date”), "
        "BY AND BETWEEN:",
    )
    add_para(
        doc,
        "(1) ANANTASUTRA, a business concern operated and represented by "
        "its Founder & CEO, Mr. Himanshu Mishra, carrying on business "
        "under the trade name “AnantaSutra”, having its principal place "
        "of business at Delhi, India, with contact address "
        "contact@anantasutra.com "
        "(hereinafter referred to as “AnantaSutra” or the “Company”, which "
        "expression shall, unless repugnant to the context or meaning "
        "thereof, be deemed to include its successors, assigns, and the "
        "person(s) for the time being in control of the business), acting "
        "through its Founder & CEO, Mr. Himanshu Mishra, of the ONE PART;",
    )
    add_para(doc, "AND")
    add_para(
        doc,
        "(2) [CONSULTANT NAME], "
        "[an individual Indian citizen aged [AGE] years, residing at "
        "[ADDRESS] / a sole proprietorship concern / a partnership firm "
        "registered under the Indian Partnership Act, 1932 / a limited "
        "liability partnership registered under the Limited Liability "
        "Partnership Act, 2008 / a company incorporated under the "
        "Companies Act, 2013], bearing Permanent Account Number (PAN) "
        "[CONSULTANT PAN], Goods and Services Tax Identification "
        "Number (GSTIN) [CONSULTANT GSTIN — if registered; else "
        "write “Not Registered”], having its principal place of "
        "business / residence at [CONSULTANT ADDRESS] (hereinafter "
        "referred to as the “Consultant”, which expression shall, "
        "unless repugnant to the context or meaning thereof, be deemed "
        "to include its / his / her legal heirs, successors-in-interest "
        "and permitted assigns), of the OTHER PART.",
    )
    add_para(
        doc,
        "AnantaSutra and the Consultant are hereinafter referred to individually "
        "as a “Party” and collectively as the “Parties”.",
    )


def add_recitals(doc):
    add_heading1(doc, "2. RECITALS")
    add_para(
        doc,
        "WHEREAS, the Company is engaged in the business of "
        "[DESCRIBE ANANTASUTRA BUSINESS — e.g., software products, "
        "technology consulting, digital platforms] and, in connection "
        "therewith, requires specialised professional services from "
        "time to time;",
    )
    add_para(
        doc,
        "WHEREAS, the Consultant represents that the Consultant possesses "
        "the requisite skill, qualifications, experience, infrastructure "
        "and capability to render professional consultancy services in the "
        "field of [CONSULTANT'S DOMAIN / AREA OF EXPERTISE] on a "
        "“contract for services” (principal-to-principal) basis;",
    )
    add_para(
        doc,
        "WHEREAS, the Company is desirous of engaging the Consultant as an "
        "independent contractor, and the Consultant is desirous of being "
        "so engaged, upon the terms and conditions set out in this Agreement;",
    )
    add_para(
        doc,
        "WHEREAS, the Parties expressly acknowledge and agree that this "
        "Agreement is executed on a principal-to-principal basis, and "
        "nothing herein shall be construed as creating a relationship of "
        "employer and employee, master and servant, principal and agent, "
        "joint venture, partnership, or association between the Parties.",
    )
    add_para(
        doc,
        "NOW, THEREFORE, in consideration of the mutual covenants, promises "
        "and undertakings contained herein, and for other good and valuable "
        "consideration, the receipt and sufficiency whereof are hereby "
        "acknowledged, the Parties agree as follows:",
    )


def add_definitions(doc):
    add_heading1(doc, "3. DEFINITIONS AND INTERPRETATION")
    add_sub(
        doc, "3.1",
        "In this Agreement, unless the context otherwise requires, the "
        "following capitalised terms shall have the meanings assigned to "
        "them below:",
    )
    defs = [
        ("“Affiliate”",
         "in relation to a Party, means any entity that directly or "
         "indirectly controls, is controlled by, or is under common "
         "control with such Party, where “control” means beneficial "
         "ownership of fifty per cent (50%) or more of the voting equity "
         "or the power to direct the management and policies of such entity."),
        ("“Agreement”",
         "means this Consultant / Independent Contractor Agreement together "
         "with all schedules, exhibits, annexures and statements of work "
         "attached hereto or executed pursuant hereto, as amended in "
         "writing from time to time."),
        ("“Applicable Laws”",
         "means all applicable statutes, enactments, ordinances, rules, "
         "regulations, notifications, guidelines, directions, orders or "
         "decrees of any Governmental Authority having jurisdiction over "
         "the Parties or the subject matter of this Agreement, including "
         "but not limited to the Indian Contract Act, 1872, the Copyright "
         "Act, 1957, the Patents Act, 1970, the Trade Marks Act, 1999, "
         "the Designs Act, 2000, the Information Technology Act, 2000, "
         "the Digital Personal Data Protection Act, 2023, the "
         "Income-tax Act, 1961, the Central Goods and Services Tax Act, "
         "2017 and State GST legislations, the Indian Stamp Act, 1899, "
         "the Arbitration and Conciliation Act, 1996, and the Specific "
         "Relief Act, 1963."),
        ("“Background IP”",
         "means Intellectual Property owned, controlled or licensed "
         "(with the right to sub-licence) by a Party prior to the "
         "Effective Date, or developed independently of this Agreement "
         "and outside the scope of the Services, as further identified "
         "for the Consultant in Exhibit B."),
        ("“Confidential Information”",
         "means all non-public information, in any form or medium, "
         "whether or not marked or designated as “confidential” or "
         "“proprietary”, disclosed by or on behalf of the Disclosing "
         "Party to the Receiving Party in connection with this Agreement, "
         "including but not limited to technical data, source code, object "
         "code, algorithms, architectures, designs, trade secrets, "
         "know-how, business plans, customer lists, pricing, financial "
         "information, employee and consultant information, Personal Data, "
         "and all derivatives, analyses, summaries and extracts thereof."),
        ("“Deliverables”",
         "means all work product, materials, reports, designs, software, "
         "source code, object code, documentation, data, drawings, "
         "specifications, plans, and other tangible or intangible outputs "
         "created, developed, produced, generated or delivered by the "
         "Consultant (alone or with others) in the course of or as a "
         "result of performing the Services, as further described in the "
         "applicable Statement of Work."),
        ("“DPDP Act”",
         "means the Digital Personal Data Protection Act, 2023, together "
         "with the rules, regulations and notifications made thereunder, "
         "as amended from time to time."),
        ("“Effective Date”",
         "means the date first written at the head of this Agreement."),
        ("“Fees”",
         "means the professional fees payable by the Company to the "
         "Consultant for the Services, as specified in the applicable "
         "Statement of Work or Exhibit A."),
        ("“Foreground IP”",
         "means all Intellectual Property created, conceived, developed, "
         "made, reduced to practice or first fixed in a tangible medium by "
         "the Consultant (alone or jointly with others, including employees, "
         "contractors or personnel of the Company), in the course of, "
         "arising out of, or in connection with the performance of the "
         "Services or the Deliverables, whether during or outside the "
         "Consultant's working hours and whether at the Consultant's or "
         "the Company's premises."),
        ("“Governmental Authority”",
         "means the Government of India, any State Government, any "
         "statutory, regulatory, judicial or quasi-judicial authority, "
         "tribunal, body, commission or agency having jurisdiction over "
         "the Parties or the subject matter of this Agreement."),
        ("“Intellectual Property” or “IP”",
         "means any and all intellectual and industrial property rights "
         "recognised under Applicable Laws, including but not limited to "
         "copyrights and related rights, patents and patent applications, "
         "trade marks and service marks (whether registered or "
         "unregistered), designs (whether registered or unregistered), "
         "trade secrets, know-how, database rights, topography rights, "
         "rights in domain names, moral rights, rights of publicity, "
         "and all applications, registrations, renewals and extensions "
         "thereof, anywhere in the world."),
        ("“Personal Data”",
         "has the meaning ascribed to it under Section 2(t) of the DPDP Act."),
        ("“Services”",
         "means the professional consultancy services to be provided by "
         "the Consultant to the Company as described in Clause 5 and the "
         "applicable Statement of Work set out in Exhibit A."),
        ("“Statement of Work” or “SOW”",
         "means a written document executed by both Parties, "
         "substantially in the form of Exhibit A, describing the specific "
         "scope, Deliverables, timelines, Fees, acceptance criteria and "
         "other particulars of a specific engagement under this Agreement."),
        ("“Term”",
         "has the meaning ascribed to it in Clause 6."),
        ("“Third Party”",
         "means any person other than the Parties and their respective "
         "Affiliates."),
    ]
    for term, meaning in defs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r1 = p.add_run(term + " ")
        style_run(r1, bold=True)
        r2 = p.add_run("means " + meaning if not meaning.startswith(("means ", "has ", "in "))
                       else meaning)
        style_run(r2)

    add_sub(
        doc, "3.2 Interpretation.",
        "In this Agreement, unless the context otherwise requires: "
        "(a) the singular includes the plural and vice versa; "
        "(b) references to clauses, schedules and exhibits are to clauses, "
        "schedules and exhibits of this Agreement; "
        "(c) headings are for convenience only and shall not affect "
        "interpretation; "
        "(d) references to statutes include any modifications or "
        "re-enactments thereof; "
        "(e) the words “including”, “includes” and “include” shall be "
        "deemed to be followed by the words “without limitation”; and "
        "(f) all references to time are to Indian Standard Time.",
    )


def add_engagement(doc):
    add_heading1(doc, "4. ENGAGEMENT AND INDEPENDENT-CONTRACTOR STATUS")
    add_sub(
        doc, "4.1 Engagement.",
        "The Company hereby engages the Consultant, and the Consultant "
        "hereby accepts such engagement, on a non-exclusive, "
        "principal-to-principal, contract-for-services basis, to render "
        "the Services upon the terms and conditions set out in this "
        "Agreement and the applicable Statement of Work.",
    )
    add_sub(
        doc, "4.2 Not an employee.",
        "The Parties expressly acknowledge and agree that the Consultant "
        "is, and shall at all times remain, an independent contractor. "
        "Nothing in this Agreement shall be construed as creating between "
        "the Parties, whether directly or indirectly, a relationship of "
        "employer and employee, master and servant, principal and agent, "
        "partnership, joint venture, association of persons, or any "
        "similar legal relationship. The Consultant is not, and shall "
        "not hold itself out as, an employee, officer, director, agent "
        "or representative of the Company.",
    )
    add_sub(
        doc, "4.3 No employee benefits.",
        "The Consultant shall not be entitled to, and shall not claim, "
        "any benefits available to the employees of the Company, "
        "including but not limited to: (a) provident fund contributions "
        "under the Employees' Provident Funds and Miscellaneous "
        "Provisions Act, 1952; (b) benefits under the Employees' State "
        "Insurance Act, 1948; (c) gratuity under the Payment of Gratuity "
        "Act, 1972; (d) bonus under the Payment of Bonus Act, 1965; "
        "(e) paid leave, sick leave, casual leave, earned leave or "
        "public holiday pay; (f) medical insurance, life insurance, or "
        "group health cover; (g) employee stock option plans (ESOPs) or "
        "restricted stock units, save and except under a separate, "
        "expressly written Consultant stock plan (if any); "
        "(h) maternity, paternity or parental benefits; or "
        "(i) any statutory notice pay, retrenchment compensation or "
        "severance.",
    )
    add_sub(
        doc, "4.4 No fixed hours; own tools and control.",
        "The Consultant shall determine the manner, method, means and "
        "sequence of performing the Services so as to achieve the "
        "Deliverables and meet the acceptance criteria set out in the "
        "Statement of Work. The Consultant shall not be required to "
        "observe fixed office hours, punch attendance, apply for leave, "
        "or adhere to any shift roster of the Company. The Consultant "
        "shall provide and use its own tools, equipment, hardware, "
        "software, licences and infrastructure to perform the Services, "
        "save for Company-provided resources (if any) expressly "
        "identified in the Statement of Work.",
    )
    add_sub(
        doc, "4.5 No authority to bind.",
        "The Consultant shall have no authority, express or implied, to "
        "bind, represent, act for, or incur any obligation or liability "
        "on behalf of the Company, and shall not hold itself out as "
        "having any such authority. The Consultant shall not sign any "
        "contract, issue any purchase order, accept any offer, or make "
        "any commitment on behalf of the Company.",
    )
    add_sub(
        doc, "4.6 Non-exclusive engagement.",
        "Subject to Clause 4.7 (Conflict), the Consultant may, during "
        "the Term, render services to other clients on its own account, "
        "provided that such engagements do not (i) interfere with the "
        "timely performance of the Services, or (ii) create a conflict "
        "of interest with the Company or its customers.",
    )
    add_sub(
        doc, "4.7 Conflict of interest.",
        "The Consultant represents that, as on the Effective Date, no "
        "conflict of interest exists between its engagements and the "
        "Services. The Consultant shall, within five (5) Working Days of "
        "becoming aware of any actual or potential conflict, notify the "
        "Company in writing, and the Company may, at its option, (a) "
        "waive the conflict in writing; (b) require the Consultant to "
        "implement mitigating measures; or (c) terminate this Agreement "
        "(or the affected Statement of Work) without liability save for "
        "Fees accrued for Services properly rendered up to the date of "
        "termination.",
    )
    add_sub(
        doc, "4.8 Tax, PF/ESI and statutory-classification indemnity.",
        "The Consultant acknowledges and agrees that the Consultant is "
        "solely responsible for the discharge of all direct taxes "
        "(including income-tax and advance tax), indirect taxes "
        "(including GST, subject to Clause 8), and all contributions, "
        "dues and statutory obligations payable in respect of the "
        "Consultant and the Consultant's own employees, contractors, "
        "agents or personnel (including, where applicable, provident "
        "fund, employees' state insurance, gratuity, bonus, professional "
        "tax, labour welfare fund and contract-labour compliances). The "
        "Consultant shall indemnify, defend and hold harmless the Company "
        "and its Affiliates, officers, directors and employees from and "
        "against any and all claims, demands, assessments, orders, "
        "penalties, interest, costs and expenses (including reasonable "
        "legal fees) arising out of or in connection with (i) any "
        "re-characterisation by any Governmental Authority of the "
        "Consultant or its personnel as employees of the Company, "
        "(ii) any claim by the Consultant's personnel for employment "
        "benefits from the Company, or (iii) any non-payment or "
        "short-payment of taxes or statutory dues for which the "
        "Consultant is responsible under this Clause 4.8.",
    )


def add_scope(doc):
    add_heading1(doc, "5. SCOPE OF SERVICES")
    add_sub(
        doc, "5.1 Services.",
        "The Consultant shall provide the Services described in the "
        "Statement(s) of Work agreed between the Parties from time to "
        "time in the form set out in Exhibit A. Each SOW, once signed by "
        "both Parties, shall be deemed incorporated into, and shall form "
        "an integral part of, this Agreement. In the event of any "
        "conflict between this Agreement and a SOW, this Agreement shall "
        "prevail, save and except in relation to any commercial terms "
        "(such as Fees, specific Deliverables and timelines) expressly "
        "set out in the SOW.",
    )
    add_sub(
        doc, "5.2 Standard of performance.",
        "The Consultant shall perform the Services: (a) with due skill, "
        "care and diligence, in accordance with professional standards "
        "prevailing in the Consultant's field of expertise; "
        "(b) in compliance with all Applicable Laws and the Company's "
        "reasonable written policies notified to the Consultant; "
        "(c) in a timely manner so as to meet the milestones and "
        "deadlines specified in the SOW; and (d) in a manner that does "
        "not infringe the Intellectual Property or other rights of any "
        "Third Party.",
    )
    add_sub(
        doc, "5.3 Key personnel.",
        "If the Consultant is not an individual, the Consultant shall "
        "identify key personnel in the SOW. The Consultant shall not "
        "substitute any key personnel without the prior written consent "
        "of the Company, such consent not to be unreasonably withheld. "
        "Proposed substitutes shall possess equivalent or superior "
        "qualifications and experience.",
    )
    add_sub(
        doc, "5.4 Acceptance.",
        "The Company shall, within the acceptance period specified in "
        "the SOW (or, if not specified, within fifteen (15) Working "
        "Days of delivery), either (i) accept the Deliverable in "
        "writing, or (ii) provide written details of any non-conformity "
        "with the acceptance criteria. The Consultant shall rectify "
        "notified non-conformities at no additional cost within a "
        "reasonable period. A Deliverable shall be deemed accepted if "
        "the Company fails to respond within the acceptance period, or "
        "puts the Deliverable into productive use.",
    )


def add_term_termination(doc):
    add_heading1(doc, "6. TERM AND TERMINATION")
    add_sub(
        doc, "6.1 Term.",
        "This Agreement shall commence on the Effective Date and, "
        "unless terminated earlier in accordance with this Clause 6, "
        "shall continue for an initial period of [TERM — e.g., twelve "
        "(12) months] (the “Initial Term”), and thereafter shall "
        "automatically renew for successive periods of [RENEWAL "
        "PERIOD] each, unless either Party gives the other not less "
        "than thirty (30) days' prior written notice of non-renewal "
        "(the Initial Term together with all renewals being the "
        "“Term”).",
    )
    add_sub(
        doc, "6.2 Termination for convenience.",
        "Either Party may terminate this Agreement, or any Statement of "
        "Work, for convenience by giving the other Party not less than "
        "thirty (30) days' prior written notice.",
    )
    add_sub(
        doc, "6.3 Termination for cause.",
        "Either Party may terminate this Agreement with immediate effect "
        "by written notice if the other Party: (a) commits a material "
        "breach of this Agreement and fails to cure such breach within "
        "fifteen (15) days of receipt of written notice specifying the "
        "breach (or, in the case of non-payment, within seven (7) days); "
        "(b) commits a breach that is not capable of cure (including any "
        "breach of Clause 9 (Confidentiality), Clause 10 (Intellectual "
        "Property) or Clause 11 (Data Protection)); (c) becomes insolvent, "
        "makes an assignment for the benefit of creditors, or becomes "
        "the subject of any winding-up, insolvency or bankruptcy "
        "proceedings under Applicable Laws (including the Insolvency and "
        "Bankruptcy Code, 2016); or (d) engages in fraud, wilful "
        "misconduct, or any act that brings the other Party into "
        "disrepute.",
    )
    add_sub(
        doc, "6.4 Consequences of termination.",
        "Upon termination or expiry of this Agreement for any reason: "
        "(a) the Consultant shall immediately cease the performance of "
        "the Services, save where the Company requests continued "
        "performance to effect an orderly transition; (b) the Company "
        "shall pay the Consultant for Services properly rendered and "
        "accepted (and, where applicable, for authorised reimbursable "
        "expenses incurred) up to the effective date of termination; "
        "(c) the Consultant shall deliver to the Company all Deliverables "
        "(whether complete or in-progress), source files, working "
        "papers, Company Confidential Information and Company property in "
        "the Consultant's possession or control; and (d) the Consultant "
        "shall, upon request and at no additional cost, provide up to "
        "five (5) Working Days of transition assistance, with any "
        "additional transition assistance to be provided at the rates "
        "agreed in the SOW.",
    )
    add_sub(
        doc, "6.5 Survival.",
        "Clauses 3 (Definitions), 4.8 (Tax / Classification Indemnity), "
        "6.4 (Consequences), 6.5 (Survival), 9 (Confidentiality), "
        "10 (Intellectual Property), 11 (Data Protection), 13 "
        "(Non-Solicitation), 14 (Representations and Warranties), "
        "15 (Indemnity), 16 (Limitation of Liability), 20 (Dispute "
        "Resolution), 21 (Governing Law), and all other provisions which "
        "by their nature are intended to survive, shall survive the "
        "termination or expiry of this Agreement.",
    )


def add_fees(doc):
    add_heading1(doc, "7. FEES AND PAYMENT")
    add_sub(
        doc, "7.1 Fees.",
        "In consideration of the Services rendered and the Deliverables "
        "provided (including the assignment of Intellectual Property "
        "under Clause 10), the Company shall pay the Consultant the "
        "professional Fees specified in the applicable Statement of "
        "Work. The Fees shall be (as specified in the SOW) "
        "[a fixed lump-sum per Deliverable / a monthly retainer of "
        "INR [AMOUNT] / an hourly / daily rate of INR [AMOUNT] per "
        "hour / day, subject to the cap set out in the SOW].",
    )
    add_sub(
        doc, "7.2 Invoicing.",
        "The Consultant shall raise a tax invoice on the Company in "
        "respect of Fees due, as per the cadence specified in the SOW "
        "(typically monthly in arrears or on completion of a "
        "Deliverable). Each invoice shall contain the Consultant's "
        "name, address, PAN, GSTIN (if registered), invoice number and "
        "date, HSN / SAC code, description of Services, period covered, "
        "Fees (exclusive of GST), GST component (if applicable), total "
        "amount payable, bank account details, and any such other "
        "particulars as may be required under Applicable Laws.",
    )
    add_sub(
        doc, "7.3 GST.",
        "The Fees are exclusive of Goods and Services Tax. If the "
        "Consultant is registered under the Central Goods and Services "
        "Tax Act, 2017 and applicable State GST legislations, the "
        "Consultant shall charge GST on each invoice at the applicable "
        "rate. The Consultant shall (a) ensure timely filing of GST "
        "returns and remittance of GST collected; (b) upload invoices "
        "such that they are reflected in the Company's GSTR-2A / 2B; "
        "and (c) be liable to the Company for any loss of input tax "
        "credit arising from the Consultant's default. If the Consultant "
        "is not required to be registered under GST legislation, the "
        "Consultant shall so state on every invoice, and any "
        "reverse-charge liability arising on the Company shall be "
        "discharged by the Company in accordance with law.",
    )
    add_sub(
        doc, "7.4 TDS under Section 194J.",
        "The Company shall deduct tax at source at the rate of ten per "
        "cent (10%) on professional fees payable under this Agreement "
        "pursuant to Section 194J of the Income-tax Act, 1961 (or at "
        "such other rate as may be applicable from time to time, "
        "including the rate of two per cent (2%) where the Services "
        "qualify as “fees for technical services” and are so taxed, or "
        "the rate under Section 194C for works contracts if applicable). "
        "The Company shall deposit the tax so deducted with the Central "
        "Government within the prescribed time, and issue Form 16A to "
        "the Consultant for each quarter. Where the Consultant is a "
        "non-resident, tax shall be deducted under Section 195 and the "
        "Consultant shall furnish the tax residency certificate (TRC) "
        "and Form 10F in advance.",
    )
    add_sub(
        doc, "7.5 Payment terms.",
        "Subject to Clause 5.4 (Acceptance) and deduction of TDS, the "
        "Company shall make payment of undisputed invoices within "
        "thirty (30) days of receipt of a valid invoice, by electronic "
        "transfer to the Consultant's bank account. Any disputed portion "
        "of an invoice shall be notified in writing within fifteen "
        "(15) days of receipt of the invoice; undisputed portions shall "
        "be paid in full on the due date. Delay in payment of undisputed "
        "amounts shall attract interest at the rate of one and a half "
        "per cent (1.5%) per month or part thereof.",
    )
    add_sub(
        doc, "7.6 Reimbursable expenses.",
        "Save as expressly set out in the SOW, the Fees are inclusive of "
        "all costs and expenses incurred by the Consultant in performing "
        "the Services. Any out-of-pocket expenses (such as pre-approved "
        "travel or third-party licences) shall be reimbursed at actuals "
        "only if (a) pre-approved in writing by the Company; and "
        "(b) supported by original receipts / tax invoices in the name of "
        "the Company. Domestic travel shall be in accordance with the "
        "Company's travel policy notified to the Consultant.",
    )
    add_sub(
        doc, "7.7 No other consideration.",
        "The Fees (together with any authorised reimbursable expenses) "
        "represent the full and final consideration for all Services "
        "rendered and for the assignment of all Foreground IP under "
        "Clause 10. The Consultant acknowledges that the royalty "
        "contemplated under Section 19 of the Copyright Act, 1957 is "
        "included in, and fully discharged by, the Fees.",
    )


def add_confidentiality(doc):
    add_heading1(doc, "8. CONFIDENTIALITY")
    add_sub(
        doc, "8.1 Obligation.",
        "The Consultant shall (a) hold all Confidential Information of "
        "the Company in strict confidence, with a standard of care no "
        "less than that which the Consultant applies to its own "
        "confidential information of like importance, and in no event "
        "less than reasonable care; (b) use Confidential Information "
        "solely for the purpose of performing the Services under this "
        "Agreement (the “Permitted Purpose”); (c) not disclose, "
        "transmit, make available, reverse-engineer, decompile, "
        "reproduce or otherwise use Confidential Information, in whole "
        "or in part, except as expressly permitted under this Agreement; "
        "and (d) promptly notify the Company of any actual or suspected "
        "unauthorised disclosure, loss or access of Confidential "
        "Information.",
    )
    add_sub(
        doc, "8.2 Permitted disclosures.",
        "The Consultant may disclose Confidential Information only to "
        "those of its personnel who need to know such information for "
        "the Permitted Purpose and who are bound by written "
        "confidentiality obligations no less stringent than those "
        "contained herein. The Consultant shall remain primarily liable "
        "for any breach of confidentiality by its personnel.",
    )
    add_sub(
        doc, "8.3 Exclusions.",
        "The obligations in this Clause 8 shall not apply to information "
        "which the Consultant can demonstrate by contemporaneous written "
        "evidence: (a) was already in its possession without a duty of "
        "confidentiality prior to disclosure; (b) is or becomes "
        "publicly available other than through a breach of this "
        "Agreement; (c) is lawfully received from a Third Party without "
        "restriction; or (d) is independently developed by the "
        "Consultant without use of or reference to any Confidential "
        "Information.",
    )
    add_sub(
        doc, "8.4 Compelled disclosure.",
        "If the Consultant is required by any Governmental Authority, "
        "court or by operation of Applicable Laws to disclose "
        "Confidential Information, the Consultant shall (a) give the "
        "Company prompt prior written notice (where lawfully "
        "permissible) to enable the Company to seek a protective "
        "order; (b) disclose only so much of the Confidential "
        "Information as is strictly required; and (c) cooperate "
        "(at the Company's expense) with the Company's efforts to "
        "preserve the confidentiality of such information.",
    )
    add_sub(
        doc, "8.5 Return / destruction.",
        "On the earlier of (a) expiry or termination of this Agreement "
        "or the relevant SOW, or (b) written request by the Company, "
        "the Consultant shall promptly, and in any event within seven "
        "(7) Working Days, return to the Company or (at the Company's "
        "option) securely destroy all Confidential Information in its "
        "possession or control, including all copies, notes and "
        "derivatives, and certify such destruction in writing, save "
        "that the Consultant may retain one (1) copy in its legal "
        "archives solely to the extent required by Applicable Laws, "
        "subject to continuing confidentiality obligations.",
    )
    add_sub(
        doc, "8.6 Survival.",
        "The obligations in this Clause 8 shall survive termination or "
        "expiry of this Agreement for a period of five (5) years, "
        "provided that in respect of Confidential Information that "
        "constitutes a trade secret under Applicable Laws, the "
        "obligations shall survive for so long as the information "
        "retains its trade-secret status.",
    )
    add_sub(
        doc, "8.7 Equitable remedies.",
        "The Parties acknowledge that a breach of this Clause 8 may "
        "cause irreparable harm for which monetary damages may be "
        "inadequate. Accordingly, the Company shall be entitled to seek "
        "interim, interlocutory and permanent injunctive relief and "
        "specific performance under the Specific Relief Act, 1963, in "
        "addition to all other remedies available at law or in equity.",
    )


def add_ip(doc):
    add_heading1(doc, "9. INTELLECTUAL PROPERTY")

    # Drafting callout — critical warning
    callout = doc.add_paragraph()
    callout.paragraph_format.space_before = Pt(4)
    callout.paragraph_format.space_after = Pt(10)
    callout.paragraph_format.left_indent = Inches(0.2)
    callout.paragraph_format.right_indent = Inches(0.2)
    callout.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = callout.add_run(
        "DRAFTING CALLOUT — CRITICAL CLAUSE: This Clause 9 is the single "
        "most important provision of this Agreement. Unlike employees, a "
        "consultant is engaged on a “contract for services” and, under "
        "Section 17(c) of the Copyright Act, 1957 (read with the proviso "
        "thereto), copyright in consultant-authored works DOES NOT "
        "automatically vest in the engaging entity. Without an express, "
        "present-tense assignment compliant with Sections 18 and 19 of "
        "the Copyright Act, 1957 (and mirror provisions of the Patents "
        "Act, 1970, the Designs Act, 2000 and the Trade Marks Act, 1999), "
        "AnantaSutra will not own the Deliverables it has paid for. The "
        "clause below is therefore drafted in the present tense (“hereby "
        "assigns”), covers all categories of IP, specifies all Section 19 "
        "parameters (work, rights, duration, territory and royalty), and "
        "includes a waiver of moral rights to the extent permissible in "
        "law, a power of attorney and a further-assurances covenant. DO "
        "NOT WEAKEN THIS CLAUSE."
    )
    style_run(r, bold=True, italic=True, size=10,
              color=RGBColor(0x8B, 0x00, 0x00))

    add_sub(
        doc, "9.1 Present-tense assignment of Foreground IP.",
        "The Consultant hereby irrevocably, absolutely and unconditionally "
        "ASSIGNS and transfers to the Company, with full title guarantee "
        "and free from all encumbrances, all right, title and interest "
        "(including all Intellectual Property rights, whether existing now "
        "or arising in the future) in and to all Deliverables and all "
        "Foreground IP, including but not limited to all works, source "
        "code, object code, documentation, algorithms, inventions, "
        "discoveries, improvements, designs, drawings, data, databases, "
        "reports, photographs, audio-visual works, trade marks, trade "
        "dress and other materials created, conceived, developed or "
        "first reduced to practice by the Consultant (alone or jointly) "
        "in the course of, arising out of, or in connection with the "
        "performance of the Services. This assignment operates as a "
        "present assignment of all such Foreground IP created after the "
        "Effective Date, effective automatically upon the coming into "
        "existence of the relevant work or right, without the need for "
        "any further act, deed or instrument.",
    )
    add_sub(
        doc, "9.2 Section 19 parameters (Copyright Act, 1957).",
        "In satisfaction of the requirements of Sections 18 and 19 of "
        "the Copyright Act, 1957, the Parties expressly record:",
    )
    add_lettered(
        doc, "a",
        "Identification of the work: all copyrightable works comprised "
        "in the Deliverables and Foreground IP, including all literary "
        "works (including computer programs), artistic works, dramatic "
        "works, musical works, sound recordings and cinematograph films "
        "created by the Consultant in the course of the Services.",
    )
    add_lettered(
        doc, "b",
        "Rights assigned: all economic rights recognised under the "
        "Copyright Act, 1957, including the rights under Section 14 of "
        "the Copyright Act, 1957, and all corresponding rights under "
        "the Patents Act, 1970, the Designs Act, 2000, the Trade Marks "
        "Act, 1999, the Semiconductor Integrated Circuits Layout-Design "
        "Act, 2000, the Protection of Plant Varieties and Farmers' "
        "Rights Act, 2001, and any corresponding or analogous laws "
        "anywhere in the world.",
    )
    add_lettered(
        doc, "c",
        "Duration: the entire term of protection afforded to the "
        "relevant Intellectual Property under Applicable Laws, "
        "including all renewals, revivals and extensions thereof.",
    )
    add_lettered(
        doc, "d",
        "Territory: worldwide, without limitation.",
    )
    add_lettered(
        doc, "e",
        "Royalty / consideration: the Fees payable under this Agreement "
        "are agreed by the Parties to be inclusive of, and to fully "
        "discharge, any royalty that would otherwise be payable in "
        "respect of this assignment. The Consultant hereby acknowledges "
        "receipt of consideration for the assignment and waives any "
        "claim to further royalty under Section 19(3) of the Copyright "
        "Act, 1957.",
    )
    add_lettered(
        doc, "f",
        "Reversion: the Consultant expressly waives any right of "
        "reversion under Section 19(4) and Section 19A of the Copyright "
        "Act, 1957, to the maximum extent permissible in law.",
    )
    add_sub(
        doc, "9.3 Waiver of moral rights.",
        "The Consultant hereby waives, to the maximum extent permissible "
        "under Section 57 of the Copyright Act, 1957 and any analogous "
        "legislation anywhere in the world, all moral rights, rights of "
        "paternity, rights of integrity, rights of attribution and "
        "rights against false attribution in respect of the Deliverables "
        "and Foreground IP, and undertakes not to assert any such rights "
        "against the Company, its Affiliates, licensees, successors or "
        "assigns. The Consultant acknowledges that the right against "
        "“distortion, mutilation, modification or other act in relation "
        "to the said work which would be prejudicial to [the author's] "
        "honour or reputation” may not be fully waivable under Indian "
        "law, and such carve-out is expressly recognised.",
    )
    add_sub(
        doc, "9.4 Power of attorney.",
        "The Consultant hereby irrevocably appoints the Company (and its "
        "duly authorised officers) as its attorney-in-fact, coupled with "
        "interest, with full power to execute, deliver, file and record "
        "any and all documents, applications, declarations, assignments "
        "and forms, and to do any and all acts and things, that the "
        "Company reasonably considers necessary or desirable to apply "
        "for, prosecute, obtain, maintain, defend, enforce, perfect or "
        "record the Company's ownership of the Foreground IP anywhere in "
        "the world (including filings before the Indian Patent Office, "
        "Copyright Office, Trade Marks Registry, Designs Office and any "
        "equivalent foreign authority). This power of attorney is granted "
        "for consideration, is irrevocable, and shall survive termination "
        "of this Agreement.",
    )
    add_sub(
        doc, "9.5 Further assurances.",
        "The Consultant shall, at the Company's reasonable written "
        "request and at the Company's cost, promptly execute such "
        "further documents, instruments of assignment, declarations, "
        "affidavits and powers, and do such further acts, as may be "
        "required to vest, perfect, record, protect or enforce the "
        "Company's ownership of the Foreground IP. This obligation "
        "shall survive termination of this Agreement.",
    )
    add_sub(
        doc, "9.6 Background IP.",
        "Each Party shall retain all right, title and interest in and "
        "to its respective Background IP. The Consultant's Background IP "
        "as on the Effective Date is set out in Exhibit B. The "
        "Consultant shall not incorporate any Background IP (of the "
        "Consultant or of any Third Party) into a Deliverable without "
        "prior written disclosure to, and consent of, the Company. To "
        "the extent any Consultant Background IP is incorporated into, "
        "or is reasonably necessary to use, a Deliverable, the "
        "Consultant hereby grants to the Company and its Affiliates a "
        "perpetual, irrevocable, worldwide, royalty-free, "
        "fully paid-up, non-exclusive licence (with the right to "
        "sub-licence through multiple tiers) to use, reproduce, modify, "
        "create derivative works of, distribute, sub-licence, and "
        "otherwise exploit such Background IP solely in connection with "
        "the Deliverable and the business of the Company.",
    )
    add_sub(
        doc, "9.7 Open-source and third-party components.",
        "The Consultant shall not incorporate any open-source software "
        "or Third-Party Intellectual Property into any Deliverable "
        "without prior written disclosure to, and consent of, the "
        "Company. The Consultant shall provide a complete bill of "
        "materials for each Deliverable, identifying every open-source "
        "or third-party component, its version and its licence. The "
        "Consultant shall not use any component licensed under the GNU "
        "Affero General Public License (AGPL), the Server Side Public "
        "License (SSPL), or any other licence with “copyleft” effect "
        "that could require the licensing or disclosure of the "
        "Company's proprietary source code, without prior written "
        "consent of the Company.",
    )
    add_sub(
        doc, "9.8 Warranty of originality and non-infringement.",
        "The Consultant represents and warrants that: (a) the "
        "Deliverables and Foreground IP shall be original works of the "
        "Consultant, save for portions expressly disclosed and permitted "
        "under Clauses 9.6 and 9.7; (b) the Consultant has full right, "
        "title and authority to make the assignments and grants under "
        "this Clause 9; (c) the Deliverables, and the Company's use "
        "thereof as contemplated by this Agreement, shall not infringe "
        "or misappropriate any Intellectual Property, privacy, publicity "
        "or other right of any Third Party; and (d) the Consultant has "
        "not assigned, licensed, encumbered or otherwise committed to "
        "any Third Party any right that would conflict with the rights "
        "granted to the Company under this Agreement.",
    )


def add_dpdp(doc):
    add_heading1(doc, "10. DATA PROTECTION")
    add_sub(
        doc, "10.1 Applicability.",
        "Where, in the course of providing the Services, the Consultant "
        "processes Personal Data on behalf of the Company, the Company "
        "shall be the Data Fiduciary and the Consultant shall be a "
        "Data Processor (as defined under the DPDP Act). The terms of "
        "this Clause 10 are supplemental to, and shall be read together "
        "with, any data processing agreement executed between the "
        "Parties pursuant to Section 8(5) of the DPDP Act.",
    )
    add_sub(
        doc, "10.2 Processor obligations.",
        "The Consultant shall: (a) process Personal Data only on the "
        "documented instructions of the Company and for the purposes "
        "specified in the SOW; (b) implement and maintain reasonable "
        "technical and organisational security safeguards (including "
        "encryption, access controls, logging and personnel training) "
        "commensurate with the nature, scope, and risk of processing; "
        "(c) ensure that personnel authorised to process Personal Data "
        "are bound by confidentiality obligations; (d) not engage any "
        "sub-processor without the prior written consent of the Company, "
        "and shall flow down all obligations of this Clause 10 to such "
        "sub-processor; (e) assist the Company with responses to "
        "requests from Data Principals and with compliance with the "
        "Company's obligations under the DPDP Act; (f) notify the "
        "Company of any personal data breach, without undue delay and in "
        "any event within twenty-four (24) hours of becoming aware, "
        "with sufficient detail to enable the Company to notify the "
        "Data Protection Board of India within the statutorily "
        "prescribed timeline; and (g) upon termination or on request, "
        "return or securely erase all Personal Data (including copies) "
        "in its possession.",
    )
    add_sub(
        doc, "10.3 IT Act compliance.",
        "The Consultant shall, in addition, comply with the Information "
        "Technology Act, 2000 and the rules made thereunder, including "
        "the Information Technology (Reasonable Security Practices and "
        "Procedures and Sensitive Personal Data or Information) Rules, "
        "2011, to the extent applicable.",
    )
    add_sub(
        doc, "10.4 Cross-border transfers.",
        "The Consultant shall not transfer Personal Data outside India "
        "except in compliance with the DPDP Act and only with the prior "
        "written consent of the Company.",
    )


def add_reps_warranties(doc):
    add_heading1(doc, "11. REPRESENTATIONS AND WARRANTIES")
    add_sub(
        doc, "11.1 Mutual representations.",
        "Each Party represents and warrants to the other that: (a) it "
        "has full power, capacity and authority to enter into and "
        "perform this Agreement; (b) the execution, delivery and "
        "performance of this Agreement have been duly authorised by all "
        "necessary corporate or personal action; (c) this Agreement "
        "constitutes its legal, valid and binding obligation, "
        "enforceable against it in accordance with its terms; and (d) "
        "the execution and performance of this Agreement does not and "
        "will not violate any Applicable Laws or any contractual or "
        "other obligation owed to any Third Party.",
    )
    add_sub(
        doc, "11.2 Consultant representations.",
        "The Consultant additionally represents and warrants that: "
        "(a) it possesses the requisite qualifications, skill, "
        "experience and resources to perform the Services to a "
        "professional standard; (b) it is not bound by any "
        "non-compete, confidentiality, or other restriction that would "
        "prevent or restrict performance of the Services; (c) it shall "
        "perform the Services in compliance with all Applicable Laws, "
        "including anti-bribery and anti-corruption laws (including the "
        "Prevention of Corruption Act, 1988), sanctions laws, anti-money-"
        "laundering laws, and applicable labour laws in respect of its "
        "own personnel; (d) it is not disqualified or debarred from "
        "rendering the Services; and (e) it shall maintain all licences, "
        "registrations and approvals required for rendering the "
        "Services.",
    )


def add_nonsolicit(doc):
    add_heading1(doc, "12. NON-SOLICITATION")
    add_sub(
        doc, "12.1 Non-solicitation of personnel.",
        "During the Term and for a period of twelve (12) months "
        "thereafter, the Consultant shall not, directly or indirectly, "
        "solicit for employment or engagement, or induce to leave the "
        "Company, any employee or consultant of the Company with whom "
        "the Consultant had material dealings during the Term; provided "
        "that a general advertisement not targeted at the Company's "
        "personnel, and the hiring of any person who responds thereto "
        "without any other targeted solicitation, shall not constitute "
        "a breach of this Clause 12.1.",
    )
    add_sub(
        doc, "12.2 Non-solicitation of customers.",
        "During the Term and for a period of twelve (12) months "
        "thereafter, the Consultant shall not, directly or indirectly, "
        "solicit any customer, prospect or business partner of the "
        "Company with whom the Consultant had dealings or learned "
        "Confidential Information, for the purpose of diverting "
        "business away from the Company.",
    )
    add_sub(
        doc, "12.3 Section 27 acknowledgment.",
        "The Parties confirm that the restrictions in this Clause 12 "
        "are reasonable, necessary to protect the legitimate business "
        "interests of the Company, and are not in unreasonable "
        "restraint of trade within the meaning of Section 27 of the "
        "Indian Contract Act, 1872. The Parties expressly record that "
        "no post-engagement non-compete obligation is imposed on the "
        "Consultant.",
    )


def add_indemnity(doc):
    add_heading1(doc, "13. INDEMNITY")
    add_sub(
        doc, "13.1 Consultant indemnity.",
        "The Consultant shall indemnify, defend and hold harmless the "
        "Company, its Affiliates and their respective directors, officers, "
        "employees, agents and successors (the “Company Indemnitees”) "
        "from and against any and all claims, demands, suits, proceedings, "
        "liabilities, losses, damages, penalties, fines, costs and "
        "expenses (including reasonable legal fees) (collectively, "
        "“Losses”) incurred by the Company Indemnitees arising out of "
        "or in connection with: (a) any breach by the Consultant of "
        "Clauses 8 (Confidentiality), 9 (Intellectual Property), 10 "
        "(Data Protection), 11 (Representations and Warranties) or 12 "
        "(Non-Solicitation); (b) any claim that the Deliverables or "
        "Foreground IP infringe, misappropriate or violate any "
        "Intellectual Property or other right of any Third Party; "
        "(c) any bodily injury (including death) to, or damage to "
        "property of, any person caused by the Consultant or its "
        "personnel; (d) any claim, assessment or proceeding arising "
        "under Clause 4.8 (tax / classification indemnity); and "
        "(e) any gross negligence, wilful misconduct or fraud of the "
        "Consultant or its personnel.",
    )
    add_sub(
        doc, "13.2 Indemnity procedure.",
        "The Company Indemnitees shall (a) promptly notify the "
        "Consultant in writing of any claim for which indemnity is "
        "sought; (b) permit the Consultant to assume the defence and "
        "settlement of the claim with counsel reasonably acceptable to "
        "the Company, provided that no settlement may be entered into "
        "without the Company's prior written consent if it requires any "
        "admission, payment or restriction by the Company; and (c) "
        "provide reasonable cooperation to the Consultant (at the "
        "Consultant's cost) in the defence of the claim. Failure to "
        "give prompt notice shall not relieve the Consultant of its "
        "obligations except to the extent of material prejudice.",
    )


def add_liability(doc):
    add_heading1(doc, "14. LIMITATION OF LIABILITY")
    add_sub(
        doc, "14.1 Exclusion of indirect losses.",
        "Save as provided in Clause 14.3, neither Party shall be liable "
        "to the other for any indirect, incidental, special, punitive, "
        "exemplary or consequential loss or damage, including loss of "
        "profits, loss of revenue, loss of goodwill, loss of business, "
        "or loss of anticipated savings, whether arising in contract, "
        "tort (including negligence), under statute, or otherwise, even "
        "if advised of the possibility thereof.",
    )
    add_sub(
        doc, "14.2 Cap.",
        "Save as provided in Clause 14.3, the total aggregate liability "
        "of either Party under or in connection with this Agreement, "
        "whether arising in contract, tort or otherwise, shall not "
        "exceed the total Fees paid and payable by the Company to the "
        "Consultant under this Agreement during the twelve (12) months "
        "immediately preceding the event giving rise to the claim.",
    )
    add_sub(
        doc, "14.3 Carve-outs.",
        "The exclusions and caps in Clauses 14.1 and 14.2 shall not "
        "apply to: (a) the Consultant's obligations under Clause 9 "
        "(Intellectual Property); (b) either Party's breach of Clause 8 "
        "(Confidentiality); (c) either Party's breach of Clause 10 "
        "(Data Protection); (d) the Consultant's indemnity obligations "
        "under Clauses 4.8 and 13; (e) any liability arising from "
        "fraud, gross negligence or wilful misconduct; and (f) the "
        "Company's obligation to pay undisputed Fees.",
    )


def add_insurance(doc):
    add_heading1(doc, "15. INSURANCE")
    add_sub(
        doc, "15.1 Insurance cover.",
        "Where the value of any Statement of Work exceeds INR 50 "
        "(fifty) lakh, or where the Services involve the processing of "
        "Personal Data or access to sensitive Company systems, the "
        "Consultant shall procure and maintain, at its own cost, "
        "throughout the Term and for a period of twelve (12) months "
        "thereafter: (a) professional indemnity insurance of not less "
        "than INR [AMOUNT — e.g., 1 crore] per occurrence; and "
        "(b) cyber-liability insurance of not less than INR [AMOUNT — "
        "e.g., 50 lakh] per occurrence. The Consultant shall, on "
        "request, furnish certificates of insurance to the Company.",
    )


def add_ic_confirmation(doc):
    add_heading1(doc, "16. INDEPENDENT-CONTRACTOR STATUS: ADDITIONAL CONFIRMATIONS")
    add_sub(
        doc, "16.1 Anti-misclassification.",
        "Without prejudice to Clause 4, the Consultant additionally "
        "acknowledges and confirms that: (a) the Consultant is engaged "
        "as an independent professional on a principal-to-principal "
        "basis; (b) the Consultant shall not be included in the "
        "Company's organisation chart or internal reporting lines, "
        "shall not be designated as a manager or officer of the Company, "
        "and shall not receive performance appraisals of the kind "
        "administered to the Company's employees; (c) the Consultant "
        "shall use its own email and communication tools, save where "
        "Company-provided access is strictly necessary for the Services "
        "and is clearly labelled as consultant / contractor access; "
        "(d) the Consultant may, subject to Clause 4.7 (Conflict), serve "
        "other clients; and (e) the Consultant bears its own business "
        "risk, supplies its own equipment, and operates its own business.",
    )
    add_sub(
        doc, "16.2 Annual confirmation.",
        "The Consultant shall, upon the Company's request (not more "
        "than once per calendar year), sign a written acknowledgment "
        "confirming the Consultant's continuing independent-contractor "
        "status and the absence of any employment relationship with the "
        "Company.",
    )


def add_tax_compliance(doc):
    add_heading1(doc, "17. TAX AND STATUTORY COMPLIANCE")
    add_sub(
        doc, "17.1 Consultant's direct taxes.",
        "The Consultant shall be solely responsible for the payment of "
        "all direct taxes (including income-tax and advance tax) on the "
        "Fees received under this Agreement, and shall indemnify the "
        "Company against any claims arising from the Consultant's "
        "non-payment or short-payment of such taxes.",
    )
    add_sub(
        doc, "17.2 PAN and TDS.",
        "The Consultant shall furnish its PAN at execution. The Company "
        "shall deduct TDS at the rate applicable under the Income-tax "
        "Act, 1961 (principally Section 194J) and shall issue Form 16A "
        "quarterly. Where the Consultant's aggregate receipts require "
        "audit under Section 44AB, the Consultant shall ensure timely "
        "audit and filing.",
    )
    add_sub(
        doc, "17.3 GST registration and invoicing.",
        "The Consultant shall obtain and maintain GST registration if "
        "required under the Central Goods and Services Tax Act, 2017 or "
        "applicable State GST legislation, and shall raise GST-compliant "
        "tax invoices in accordance with Clause 7.2. Any loss of input "
        "tax credit suffered by the Company on account of the "
        "Consultant's default (including delayed filing of GSTR-1) shall "
        "be recoverable from the Consultant.",
    )
    add_sub(
        doc, "17.4 Professional tax and other dues.",
        "The Consultant shall be responsible for payment of any "
        "professional tax, labour welfare fund contribution, or similar "
        "state-level levies applicable to the Consultant or the "
        "Consultant's personnel.",
    )
    add_sub(
        doc, "17.5 Contract Labour (R&A) Act.",
        "Where the Contract Labour (Regulation & Abolition) Act, 1970 "
        "is applicable to the engagement (by reason of the Consultant "
        "deploying twenty (20) or more workmen at the Company's "
        "premises, or such lower threshold as may be applicable under "
        "the relevant State rules), the Consultant shall obtain and "
        "maintain a valid labour licence, comply with all statutory "
        "obligations (including wage, PF, ESI, welfare and safety "
        "obligations in respect of its own workmen), and indemnify the "
        "Company as “principal employer” against any resulting claims.",
    )


def add_subcontracting(doc):
    add_heading1(doc, "18. SUBCONTRACTING")
    add_sub(
        doc, "18.1 No subcontracting without consent.",
        "The Consultant shall not subcontract, assign or delegate any "
        "of its obligations under this Agreement or any Statement of "
        "Work to any Third Party without the prior written consent of "
        "the Company.",
    )
    add_sub(
        doc, "18.2 Flow-down and primary liability.",
        "Where the Company consents to a subcontracting, the Consultant "
        "shall (a) ensure that the subcontractor is bound by written "
        "obligations of confidentiality, IP assignment, data protection, "
        "and compliance with Applicable Laws no less stringent than "
        "those contained in this Agreement; and (b) remain primarily "
        "liable to the Company for all acts, omissions, defaults and "
        "breaches of the subcontractor as if they were the acts, "
        "omissions, defaults and breaches of the Consultant.",
    )


def add_dispute_resolution(doc):
    add_heading1(doc, "19. DISPUTE RESOLUTION")
    add_sub(
        doc, "19.1 Amicable settlement.",
        "Any dispute, controversy or claim arising out of or in "
        "connection with this Agreement, including any question "
        "regarding its existence, validity, interpretation, breach or "
        "termination (a “Dispute”), shall first be attempted to be "
        "resolved amicably by good-faith discussion between duly "
        "authorised representatives of the Parties. If the Dispute is "
        "not resolved within thirty (30) days of written notice by one "
        "Party to the other, it shall be referred to arbitration under "
        "Clause 19.2.",
    )
    add_sub(
        doc, "19.2 Arbitration.",
        "Any Dispute not resolved under Clause 19.1 shall be finally "
        "settled by arbitration conducted in accordance with the "
        "provisions of the Arbitration and Conciliation Act, 1996, as "
        "amended from time to time. The arbitral tribunal shall consist "
        "of a sole arbitrator mutually appointed by the Parties, "
        "failing which the arbitrator shall be appointed in accordance "
        "with Section 11 of the said Act. The seat of arbitration shall "
        "be New Delhi, Delhi (which shall determine the curial "
        "law); the venue shall be New Delhi, Delhi; the language of "
        "arbitration shall be English; and the arbitral award shall be "
        "final and binding on the Parties.",
    )
    add_sub(
        doc, "19.3 Interim reliefs (Section 9 carve-out).",
        "Nothing in this Clause 19 shall preclude either Party from "
        "seeking interim, interlocutory or protective relief from the "
        "competent courts at New Delhi under Section 9 of the "
        "Arbitration and Conciliation Act, 1996, or equitable relief "
        "under the Specific Relief Act, 1963, in aid of arbitration or "
        "for the protection of Confidential Information or "
        "Intellectual Property.",
    )
    add_sub(
        doc, "19.4 Costs.",
        "The costs of arbitration shall be borne as directed by the "
        "arbitral tribunal. Each Party shall bear its own legal costs "
        "unless the tribunal orders otherwise.",
    )


def add_governing_law(doc):
    add_heading1(doc, "20. GOVERNING LAW AND JURISDICTION")
    add_sub(
        doc, "20.1 Governing law.",
        "This Agreement shall be governed by, and construed in "
        "accordance with, the laws of the Republic of India.",
    )
    add_sub(
        doc, "20.2 Exclusive jurisdiction.",
        "Subject to Clause 19 (Dispute Resolution), the courts at "
        "New Delhi, Delhi, shall have exclusive jurisdiction over "
        "any proceedings arising out of or in connection with this "
        "Agreement.",
    )


def add_force_majeure(doc):
    add_heading1(doc, "21. FORCE MAJEURE")
    add_sub(
        doc, "21.1 FM event.",
        "Neither Party shall be liable for any failure or delay in "
        "performing its obligations under this Agreement (other than "
        "payment obligations for Services already rendered and "
        "accepted) to the extent such failure or delay is caused by an "
        "event beyond its reasonable control, including acts of God, "
        "earthquakes, floods, fires, pandemics, epidemics, "
        "governmental lockdowns or curfews, wars, terrorism, riots, "
        "civil disturbances, regional internet or telecommunications "
        "outages, or binding orders of Governmental Authorities (a "
        "“Force Majeure Event”).",
    )
    add_sub(
        doc, "21.2 Mitigation and termination.",
        "The affected Party shall (a) promptly notify the other in "
        "writing of the Force Majeure Event and its expected duration; "
        "(b) use reasonable endeavours to mitigate the effects; and "
        "(c) resume performance as soon as reasonably practicable. If "
        "the Force Majeure Event continues for more than sixty (60) "
        "consecutive days, either Party may terminate the affected "
        "Statement of Work (and, where appropriate, this Agreement) "
        "upon written notice, without liability other than for Services "
        "properly rendered and accepted up to the date of termination.",
    )


def add_notices(doc):
    add_heading1(doc, "22. NOTICES")
    add_sub(
        doc, "22.1 Mode of service.",
        "Any notice or other communication required or permitted to be "
        "given under this Agreement shall be in writing and shall be "
        "delivered: (a) by hand; (b) by registered post / speed post "
        "with acknowledgment due; (c) by nationally recognised courier "
        "with proof of delivery; or (d) by email to the email "
        "address(es) specified below (with a copy by any of the methods "
        "in (a)–(c) for notices of breach, indemnity or termination).",
    )
    add_sub(
        doc, "22.2 Addresses.",
        "Notices to AnantaSutra shall be marked for the attention of "
        "Mr. Himanshu Mishra, Founder & CEO, at Delhi, India, by email to "
        "contact@anantasutra.com. Notices to the Consultant shall be sent "
        "to the address set out at the head of this Agreement, email "
        "[CONSULTANT EMAIL]. Either Party may change its notice "
        "address by written notice to the other.",
    )


def add_boilerplate(doc):
    add_heading1(doc, "23. MISCELLANEOUS")
    add_sub(
        doc, "23.1 Assignment.",
        "The Consultant shall not assign, transfer, charge or otherwise "
        "dispose of any of its rights or obligations under this "
        "Agreement without the prior written consent of the Company. "
        "The Company may assign or novate this Agreement to an Affiliate, "
        "or to a successor in a merger, amalgamation, demerger or sale "
        "of substantially all its assets, upon written notice to the "
        "Consultant.",
    )
    add_sub(
        doc, "23.2 Waiver.",
        "No failure or delay by a Party in exercising any right or "
        "remedy under this Agreement shall operate as a waiver thereof, "
        "nor shall any single or partial exercise preclude any further "
        "exercise thereof.",
    )
    add_sub(
        doc, "23.3 Severability.",
        "If any provision of this Agreement is held to be invalid, "
        "illegal or unenforceable by any competent authority, the "
        "remaining provisions shall continue in full force and effect, "
        "and the Parties shall negotiate in good faith a substitute "
        "provision that most nearly reflects the original intent.",
    )
    add_sub(
        doc, "23.4 Entire agreement.",
        "This Agreement, together with all Statements of Work and "
        "Exhibits hereto, constitutes the entire agreement between the "
        "Parties in relation to its subject matter and supersedes all "
        "prior discussions, negotiations, arrangements and "
        "understandings, whether written or oral. No Party has relied "
        "on any representation not expressly set out in this Agreement.",
    )
    add_sub(
        doc, "23.5 Amendment.",
        "No amendment or modification of this Agreement shall be valid "
        "or binding unless made in writing and signed by authorised "
        "representatives of both Parties.",
    )
    add_sub(
        doc, "23.6 Counterparts and electronic execution.",
        "This Agreement may be executed in counterparts, each of which "
        "when executed shall be an original and all of which together "
        "shall constitute one and the same instrument. The Parties "
        "agree that this Agreement may be executed by way of "
        "electronic signature (including Aadhaar e-Sign or digital "
        "signature certificate) under the Information Technology Act, "
        "2000, and that such electronic execution shall be valid and "
        "enforceable as between the Parties.",
    )
    add_sub(
        doc, "23.7 No partnership or agency.",
        "Nothing in this Agreement shall constitute a partnership, "
        "joint venture, or agency relationship between the Parties, and "
        "neither Party shall have authority to bind the other except as "
        "expressly provided herein.",
    )
    add_sub(
        doc, "23.8 Third-party rights.",
        "A person who is not a Party to this Agreement shall have no "
        "rights under it, save as expressly provided herein.",
    )
    add_sub(
        doc, "23.9 Stamp duty.",
        "The Parties acknowledge that appropriate stamp duty is "
        "chargeable on this Agreement under the Delhi Stamp Act, "
        "1957 (read with Article 5 of the Schedule thereto, being an "
        "“agreement not otherwise provided for”), and shall be borne by "
        "[the Company / equally by the Parties], and that the original "
        "of this Agreement shall be e-stamped through the Stock Holding "
        "Corporation of India Ltd. (SHCIL). The Parties acknowledge that "
        "an inadequately stamped document may be inadmissible in "
        "evidence under Section 35 of the Indian Stamp Act, 1899 until "
        "stamped together with applicable penalty.",
    )


def add_execution_block(doc):
    add_heading1(doc, "24. EXECUTION")
    add_para(
        doc,
        "IN WITNESS WHEREOF, the Parties have executed this Agreement on "
        "the date first above written, in the presence of the witnesses "
        "whose names appear below.",
    )

    # Two-column signature table
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    col_w = Inches(3.15)
    hdr = table.rows[0].cells
    hdr[0].width = col_w
    hdr[1].width = col_w

    sutranet_lines = [
        "For and on behalf of",
        "ANANTASUTRA",
        "",
        "",
        "Signature: ____________________________",
        "Name:   Mr. Himanshu Mishra",
        "Designation: Founder & CEO",
        "Email:  contact@anantasutra.com",
        "Date:   [DATE]",
        "Place:  Delhi",
    ]
    consultant_lines = [
        "For and on behalf of",
        "[CONSULTANT NAME]",
        "",
        "",
        "Signature: ____________________________",
        "Name:   [CONSULTANT / AUTHORISED SIGNATORY]",
        "Designation: [DESIGNATION, if firm/company]",
        "Date:   ____________________________",
        "Place:  ____________________________",
    ]

    for cell, lines in ((hdr[0], sutranet_lines), (hdr[1], consultant_lines)):
        set_cell_border(cell)
        cell.paragraphs[0].text = ""
        for i, line in enumerate(lines):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(line)
            bold = i < 2
            style_run(run, bold=bold)

    add_para(doc, "")
    add_para(doc, "WITNESSES:", bold=True, space_before=10)

    wit_table = doc.add_table(rows=1, cols=2)
    wit_table.autofit = False
    for cell in wit_table.rows[0].cells:
        cell.width = col_w
        set_cell_border(cell)

    wit_lines_1 = [
        "Witness 1",
        "",
        "Signature: ____________________________",
        "Name:  ____________________________",
        "Address: ____________________________",
        "PAN:   ____________________________",
    ]
    wit_lines_2 = [
        "Witness 2",
        "",
        "Signature: ____________________________",
        "Name:  ____________________________",
        "Address: ____________________________",
        "PAN:   ____________________________",
    ]
    for cell, lines in ((wit_table.rows[0].cells[0], wit_lines_1),
                        (wit_table.rows[0].cells[1], wit_lines_2)):
        cell.paragraphs[0].text = ""
        for i, line in enumerate(lines):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(line)
            bold = i == 0
            style_run(run, bold=bold)


# ---------------------------------------------------------------------------
# Exhibits
# ---------------------------------------------------------------------------

def add_exhibit_a(doc):
    add_page_break(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("EXHIBIT A — STATEMENT OF WORK (TEMPLATE)")
    style_run(r, bold=True, size=14)

    add_para(
        doc,
        "This Statement of Work (“SOW No. [•]”) is entered into pursuant "
        "to the Consultant / Independent Contractor Agreement dated "
        "[EFFECTIVE DATE] between AnantaSutra (the "
        "“Company”) and [CONSULTANT NAME] (the “Consultant”). Capitalised "
        "terms used but not defined herein shall have the meanings given "
        "to them in the Agreement. In the event of conflict between the "
        "Agreement and this SOW, the Agreement shall prevail, save in "
        "respect of commercial terms (Fees, Deliverables, timelines) set "
        "out below.",
    )

    rows = [
        ("1. SOW Number and Date",
         "SOW No. [•]; effective [DATE]."),
        ("2. Description of Services",
         "[Describe the Services in sufficient detail. E.g., development "
         "of a [module/feature], technical advisory on [area], "
         "preparation of [report/design].]"),
        ("3. Deliverables",
         "[List each Deliverable with a unique reference; include format, "
         "acceptance criteria, and whether partial deliveries are "
         "permitted.]"),
        ("4. Timeline / Milestones",
         "[Milestone 1 — [Date]; Milestone 2 — [Date]; Final delivery — "
         "[Date]. Indicate any contingent dates.]"),
        ("5. Fees",
         "[Fixed fee of INR [AMOUNT] payable as follows: [●]% on "
         "signature; [●]% on Milestone 1; [●]% on Milestone 2; balance on "
         "final acceptance. OR Hourly rate of INR [AMOUNT] per hour, "
         "estimated [HOURS] hours, capped at INR [AMOUNT]. Fees "
         "exclusive of GST.]"),
        ("6. Out-of-pocket expenses",
         "[Travel outside New Delhi, subject to pre-approval, at actuals; "
         "Third-party software licences reimbursed at cost plus zero "
         "margin; other expenses not reimbursable.]"),
        ("7. Acceptance criteria",
         "[Define objective, measurable acceptance criteria for each "
         "Deliverable; reference any test plan, specification or "
         "industry standard.]"),
        ("8. Key personnel",
         "[List key personnel with role and % of time committed.]"),
        ("9. Location of Services",
         "[Primarily remote / at Consultant's premises / with periodic "
         "visits to Company premises at [ADDRESS].]"),
        ("10. Background IP of Consultant used",
         "[List from Exhibit B or mark “None”.]"),
        ("11. Third-party / open-source components used",
         "[List with licence or mark “None”.]"),
        ("12. Access to Personal Data",
         "[Yes — scope and categories: [•]; sub-processors: [•] / No.]"),
        ("13. Insurance required",
         "[Professional indemnity: INR [AMOUNT]; Cyber: INR [AMOUNT] / "
         "Not required.]"),
        ("14. Term of this SOW",
         "[From [DATE] to [DATE] unless terminated earlier.]"),
        ("15. Special provisions",
         "[Any SOW-specific variations from the Agreement. If none, "
         "mark “None”.]"),
    ]
    tab = doc.add_table(rows=len(rows), cols=2)
    tab.autofit = False
    for i, (head, body) in enumerate(rows):
        for c in tab.rows[i].cells:
            set_cell_border(c)
        tab.rows[i].cells[0].width = Inches(2.0)
        tab.rows[i].cells[1].width = Inches(4.3)
        c0 = tab.rows[i].cells[0].paragraphs[0]
        c0.paragraph_format.space_after = Pt(2)
        r0 = c0.add_run(head)
        style_run(r0, bold=True)
        c1 = tab.rows[i].cells[1].paragraphs[0]
        c1.paragraph_format.space_after = Pt(2)
        r1 = c1.add_run(body)
        style_run(r1)

    add_para(doc, "", space_before=10)
    add_para(
        doc,
        "Agreed and accepted:",
        bold=True,
    )

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    for cell in table.rows[0].cells:
        cell.width = Inches(3.15)
        set_cell_border(cell)
    left_lines = [
        "For AnantaSutra",
        "",
        "Signature: ______________________",
        "Name:  Mr. Himanshu Mishra",
        "Designation: Founder & CEO",
        "Date:  [DATE]",
    ]
    right_lines = [
        "For [CONSULTANT NAME]",
        "",
        "Signature: ______________________",
        "Name:  [•]",
        "Designation: [•]",
        "Date:  [•]",
    ]
    for cell, lines in ((table.rows[0].cells[0], left_lines),
                        (table.rows[0].cells[1], right_lines)):
        cell.paragraphs[0].text = ""
        for i, line in enumerate(lines):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(line)
            style_run(run, bold=(i == 0))


def add_exhibit_b(doc):
    add_page_break(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("EXHIBIT B — PRIOR IP DISCLOSURE SCHEDULE")
    style_run(r, bold=True, size=14)

    add_para(
        doc,
        "The Consultant discloses the following items of Background IP / "
        "pre-existing Intellectual Property owned or controlled by the "
        "Consultant as on the Effective Date. The items listed below are "
        "NOT assigned to the Company under Clause 9 of the Agreement, but "
        "shall be subject to the licence-back in Clause 9.6 of the "
        "Agreement to the extent embedded in, or reasonably necessary to "
        "use, any Deliverable.",
    )
    add_para(
        doc,
        "If no Background IP is to be disclosed, the Consultant shall "
        "mark the table below “NONE” and initial the same. The absence "
        "of disclosure shall be deemed a warranty by the Consultant "
        "that no Background IP of the Consultant shall be incorporated "
        "into any Deliverable.",
        italic=True,
    )

    headers = ["S. No.", "Title / Description",
               "Type (Copyright / Patent / TM / Design / Know-how / Other)",
               "Date of creation / Registration no. (if any)",
               "Owner (Consultant / Third Party with consent)",
               "Notes / Intended use in Deliverables"]
    rows_n = 6  # template rows
    tab = doc.add_table(rows=rows_n + 1, cols=len(headers))
    tab.autofit = False
    # Headers
    for j, head in enumerate(headers):
        c = tab.rows[0].cells[j]
        set_cell_border(c)
        c.paragraphs[0].text = ""
        r = c.paragraphs[0].add_run(head)
        style_run(r, bold=True, size=10)
    for i in range(1, rows_n + 1):
        for j in range(len(headers)):
            c = tab.rows[i].cells[j]
            set_cell_border(c)
            c.paragraphs[0].text = ""
            run = c.paragraphs[0].add_run(" ")
            style_run(run, size=10)

    add_para(doc, "", space_before=10)
    add_para(
        doc,
        "Declaration: I / We hereby declare that the information set out "
        "above is true, complete and correct as on the Effective Date, "
        "and that the Consultant has full right and authority to grant "
        "the licence-back contemplated under Clause 9.6 of the Agreement "
        "in respect of the items so disclosed.",
    )

    add_para(doc, "", space_before=10)
    add_para(doc, "Signed for the Consultant:", bold=True)
    add_para(doc, "Signature: ______________________     Date: __________________")
    add_para(doc, "Name:   [CONSULTANT NAME / AUTHORISED SIGNATORY]")


# ---------------------------------------------------------------------------
# Build pipeline
# ---------------------------------------------------------------------------

def build():
    doc = Document()
    add_brand_header(doc)
    configure_document(doc)

    add_title_block(doc)
    add_parties(doc)
    add_recitals(doc)
    add_definitions(doc)
    add_engagement(doc)
    add_scope(doc)
    add_term_termination(doc)
    add_fees(doc)
    add_confidentiality(doc)
    add_ip(doc)
    add_dpdp(doc)
    add_reps_warranties(doc)
    add_nonsolicit(doc)
    add_indemnity(doc)
    add_liability(doc)
    add_insurance(doc)
    add_ic_confirmation(doc)
    add_tax_compliance(doc)
    add_subcontracting(doc)
    add_dispute_resolution(doc)
    add_governing_law(doc)
    add_force_majeure(doc)
    add_notices(doc)
    add_boilerplate(doc)
    add_execution_block(doc)
    add_exhibit_a(doc)
    add_exhibit_b(doc)

    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    size_kb = os.path.getsize(OUTPUT_PATH) / 1024
    print(f"Saved: {OUTPUT_PATH}")
    print(f"Size:  {size_kb:.2f} KB")
    return OUTPUT_PATH, size_kb


if __name__ == "__main__":
    build()
