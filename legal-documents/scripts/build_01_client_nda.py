"""
Build Mutual Non-Disclosure Agreement (Client <-> AnantaSutra)
India-law-compliant, research-backed, execution-ready.

Output: c:/Users/duneja8515/Desktop/bhavya/prsnl/sutranet/legal-documents/01_Client_Company_NDA.docx
"""

from __future__ import annotations

import os
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, Inches


OUTPUT_PATH = (
    r"c:/Users/duneja8515/Desktop/bhavya/prsnl/sutranet/legal-documents/"
    r"01_Client_Company_NDA.docx"
)

BODY_FONT = "Calibri"
BODY_SIZE = 11


# ---------------------------------------------------------------------------
# Low-level formatting helpers
# ---------------------------------------------------------------------------


# ---------------------------------------------------------------------------
# Brand header (AnantaSutra) — injected by rebrand
# ---------------------------------------------------------------------------
LOGO_PATH = r"c:/Users/duneja8515/Desktop/bhavya/prsnl/sutranet/public/Favicon/logo-nobg.png"


def add_brand_header(doc):
    """Insert centered logo + brand lines at top of page 1. Never crashes."""
    try:
        from docx.shared import Inches as _Inches, Pt as _Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH as _CENTERED
    except Exception:
        return
    p = doc.add_paragraph()
    p.alignment = _CENTERED.CENTER
    try:
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=_Inches(1.2))
    except Exception:
        pass
    p2 = doc.add_paragraph()
    p2.alignment = _CENTERED.CENTER
    r2 = p2.add_run("AnantaSutra — अनन्तसूत्र")
    r2.italic = True
    r2.font.size = _Pt(10)
    p3 = doc.add_paragraph()
    p3.alignment = _CENTERED.CENTER
    r3 = p3.add_run("https://anantasutra.com  •  contact@anantasutra.com")
    r3.italic = True
    r3.font.size = _Pt(9)

def set_a4_and_margins(document: Document) -> None:
    """Set all sections to A4 with 1-inch margins."""
    for section in document.sections:
        # A4 is 21.0 cm x 29.7 cm
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def configure_base_styles(document: Document) -> None:
    """Normal, Heading 1, Heading 2 styles tuned for legal body text."""
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    h1 = styles["Heading 1"]
    h1.font.name = BODY_FONT
    h1.font.size = Pt(13)
    h1.font.bold = True
    h1.font.color.rgb = None  # default black per template
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    h2.font.name = BODY_FONT
    h2.font.size = Pt(11)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after = Pt(3)
    h2.paragraph_format.keep_with_next = True


def add_page_number_field(paragraph) -> None:
    """Add a PAGE field to a paragraph."""
    run = paragraph.add_run()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_sep)
    run._r.append(fldChar_end)


def configure_footer(document: Document) -> None:
    """Two-line footer: confidentiality notice and page number."""
    for section in document.sections:
        footer = section.footer
        # First paragraph: confidentiality notice (left)
        p1 = footer.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p1.text = ""  # clear
        run = p1.add_run("Confidential \u2014 AnantaSutra")
        run.font.name = BODY_FONT
        run.font.size = Pt(9)
        run.italic = True

        # Second paragraph: page X of Y (centered)
        p2 = footer.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run("Page ")
        run2.font.name = BODY_FONT
        run2.font.size = Pt(9)
        add_page_number_field(p2)


def add_title(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(text)
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(16)


def add_centered_line(document: Document, text: str, bold: bool = False,
                      italic: bool = False, size: int = BODY_SIZE) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = BODY_FONT
    run.font.size = Pt(size)


def add_h1(document: Document, number: str, text: str) -> None:
    p = document.add_paragraph(style="Heading 1")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"{number}. {text.upper()}")
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(13)


def add_clause(document: Document, number: str, text: str,
               first_line_indent_cm: float = 0.0) -> None:
    """A numbered sub-clause in body style, e.g. '1.1 Text...'."""
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    if first_line_indent_cm:
        p.paragraph_format.first_line_indent = Cm(first_line_indent_cm)
    run_num = p.add_run(f"{number} ")
    run_num.bold = True
    run_num.font.name = BODY_FONT
    run_num.font.size = Pt(BODY_SIZE)
    run_body = p.add_run(text)
    run_body.font.name = BODY_FONT
    run_body.font.size = Pt(BODY_SIZE)


def add_para(document: Document, text: str, bold: bool = False,
             italic: bool = False, indent_cm: float = 0.0) -> None:
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    if indent_cm:
        p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = BODY_FONT
    run.font.size = Pt(BODY_SIZE)


def add_lettered(document: Document, letter: str, text: str) -> None:
    """Indented lettered sub-point like '(a) Text...'."""
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"({letter}) ")
    r1.bold = True
    r1.font.name = BODY_FONT
    r1.font.size = Pt(BODY_SIZE)
    r2 = p.add_run(text)
    r2.font.name = BODY_FONT
    r2.font.size = Pt(BODY_SIZE)


def add_page_break(document: Document) -> None:
    document.add_page_break()


# ---------------------------------------------------------------------------
# Document build
# ---------------------------------------------------------------------------

def build_document() -> Document:
    doc = Document()
    add_brand_header(doc)
    set_a4_and_margins(doc)
    configure_base_styles(doc)
    configure_footer(doc)

    # -----------------------------------------------------------------------
    # TITLE / HEADING BLOCK
    # -----------------------------------------------------------------------
    add_title(doc, "MUTUAL NON-DISCLOSURE AGREEMENT")
    add_centered_line(doc, "(Executed in New Delhi, Delhi, India)",
                      italic=True)
    add_centered_line(doc,
                      "This Agreement is executed on this [DAY] day of "
                      "[MONTH], [YEAR] (the \u201cEffective Date\u201d).")
    doc.add_paragraph()  # spacer

    # -----------------------------------------------------------------------
    # PARTIES
    # -----------------------------------------------------------------------
    add_h1(doc, "1", "Parties")
    add_clause(
        doc, "1.1",
        "ANANTASUTRA, a business concern operated and represented by its "
        "Founder & CEO, Mr. Himanshu Mishra, carrying on business under the "
        "trade name \u201cAnantaSutra\u201d, having its principal place of "
        "business at Delhi, India, with contact address "
        "contact@anantasutra.com "
        "(hereinafter referred to as \u201cAnantaSutra\u201d or \u201cthe "
        "Company\u201d, which expression shall, unless repugnant to the "
        "context or meaning thereof, be deemed to include its successors, "
        "assigns, and the person(s) for the time being in control of the "
        "business), acting through its Founder & CEO, Mr. Himanshu Mishra;"
    )
    add_para(doc, "AND", bold=True)
    add_clause(
        doc, "1.2",
        "[CLIENT LEGAL NAME], a [company incorporated under the "
        "Companies Act, 2013 / limited liability partnership / "
        "partnership firm / sole proprietorship / foreign entity "
        "incorporated under the laws of [JURISDICTION]], having its "
        "registered office / principal place of business at "
        "[CLIENT ADDRESS], bearing [CIN / LLPIN / registration number] "
        "[NUMBER] and Permanent Account Number [PAN] (hereinafter "
        "referred to as the \u201cClient\u201d, which expression shall, "
        "unless repugnant to the context or meaning thereof, be deemed "
        "to include its successors-in-interest and permitted assigns), "
        "acting through its duly authorised signatory [NAME], "
        "[DESIGNATION]."
    )
    add_para(
        doc,
        "AnantaSutra and the Client are hereinafter individually referred to "
        "as a \u201cParty\u201d and collectively as the \u201cParties\u201d."
    )

    # -----------------------------------------------------------------------
    # RECITALS
    # -----------------------------------------------------------------------
    add_h1(doc, "2", "Recitals")
    add_para(
        doc,
        "WHEREAS, AnantaSutra is engaged in the business of providing "
        "[software development / information technology services / "
        "design / consultancy / data engineering] services and owns "
        "certain proprietary methodologies, know-how, software, tools, "
        "technical documentation and commercial information relating "
        "thereto;"
    )
    add_para(
        doc,
        "AND WHEREAS, the Client is engaged in the business of "
        "[CLIENT BUSINESS DESCRIPTION] and possesses certain proprietary, "
        "commercial, financial, technical and customer information "
        "relating thereto;"
    )
    add_para(
        doc,
        "AND WHEREAS, the Parties wish to explore, negotiate and / or "
        "undertake a potential business relationship in relation to "
        "[DESCRIBE PURPOSE \u2014 e.g., evaluation of a proposed services "
        "engagement, proof-of-concept, request-for-proposal, or the "
        "performance of services under a definitive agreement] "
        "(the \u201cPurpose\u201d), and in furtherance thereof, each Party "
        "may be required to disclose to the other certain non-public "
        "information of a confidential or proprietary nature;"
    )
    add_para(
        doc,
        "AND WHEREAS, the Parties desire to record the terms upon which "
        "such Confidential Information (as defined below) shall be "
        "disclosed, received, used and protected."
    )
    add_para(
        doc,
        "NOW, THEREFORE, in consideration of the mutual covenants, "
        "promises and undertakings set forth herein and for other good "
        "and valuable consideration, the receipt and sufficiency of which "
        "are hereby acknowledged, the Parties hereby agree as follows:",
        bold=True
    )

    # -----------------------------------------------------------------------
    # DEFINITIONS
    # -----------------------------------------------------------------------
    add_h1(doc, "3", "Definitions and Interpretation")
    add_clause(
        doc, "3.1",
        "In this Agreement, unless the context otherwise requires, the "
        "following capitalised terms shall have the meanings ascribed to "
        "them below:"
    )
    add_clause(
        doc, "3.1.1",
        "\u201cAffiliate\u201d means, in relation to a Party, any entity "
        "that directly or indirectly controls, is controlled by, or is "
        "under common control with such Party, where \u201ccontrol\u201d "
        "means the beneficial ownership of not less than fifty per cent "
        "(50%) of the voting equity, or the power to direct the management "
        "and policies of the entity, whether by contract or otherwise."
    )
    add_clause(
        doc, "3.1.2",
        "\u201cApplicable Law\u201d means all applicable statutes, "
        "enactments, ordinances, rules, regulations, notifications, "
        "guidelines, directions, judgments, orders and policies of any "
        "central, state or local governmental, statutory, regulatory or "
        "judicial authority of India, having jurisdiction over the "
        "relevant subject matter, including without limitation the "
        "Indian Contract Act, 1872; the Specific Relief Act, 1963; the "
        "Information Technology Act, 2000; the Digital Personal Data "
        "Protection Act, 2023; the Copyright Act, 1957; the Arbitration "
        "and Conciliation Act, 1996; and the Indian Stamp Act, 1899 "
        "(read with the Indian Stamp Act, 1899 (as applicable to the NCT of Delhi))."
    )
    add_clause(
        doc, "3.1.3",
        "\u201cConfidential Information\u201d means all non-public "
        "technical, commercial, financial, strategic, customer, employee, "
        "source code, object code, algorithmic, architectural, product, "
        "pricing, process, research and development, trade secret, "
        "know-how and any other information \u2014 whether disclosed orally, "
        "visually, electronically, in writing or in any other form, and "
        "whether or not marked or designated as \u201cconfidential\u201d "
        "\u2014 that is disclosed by or on behalf of the Disclosing Party "
        "(including through its Representatives) to the Receiving Party "
        "(including its Representatives) before or after the Effective "
        "Date in connection with the Purpose, and that a reasonable "
        "person would consider to be confidential in the circumstances. "
        "Confidential Information shall include all copies, extracts, "
        "summaries, analyses, notes and other derivative works prepared "
        "by or for the Receiving Party that contain or are derived from "
        "such information. Oral or visual disclosures that are identified "
        "as confidential at the time of disclosure shall be treated as "
        "Confidential Information whether or not reduced to writing."
    )
    add_clause(
        doc, "3.1.4",
        "\u201cData Principal\u201d, \u201cData Fiduciary\u201d, \u201cData "
        "Processor\u201d and \u201cPersonal Data\u201d shall have the "
        "meanings respectively ascribed to them under section 2 of the "
        "Digital Personal Data Protection Act, 2023 (\u201cDPDP Act\u201d)."
    )
    add_clause(
        doc, "3.1.5",
        "\u201cDisclosing Party\u201d means the Party, or its Affiliates or "
        "Representatives, disclosing Confidential Information hereunder; "
        "and \u201cReceiving Party\u201d means the Party, or its Affiliates "
        "or Representatives, receiving such Confidential Information. Each "
        "Party may act in either capacity at different times during the "
        "Term."
    )
    add_clause(
        doc, "3.1.6",
        "\u201cEffective Date\u201d means the later of (a) the date of "
        "last signature of this Agreement by both Parties, and (b) the "
        "date on which applicable stamp duty has been paid in accordance "
        "with Clause 22 (Stamp Duty) hereof."
    )
    add_clause(
        doc, "3.1.7",
        "\u201cPurpose\u201d means the purpose described in the recitals "
        "and more particularly set out in Schedule 1 hereto, and shall "
        "not be construed to extend to any other purpose without the prior "
        "written consent of the Disclosing Party."
    )
    add_clause(
        doc, "3.1.8",
        "\u201cRepresentatives\u201d means, in relation to a Party, its "
        "directors, officers, employees, Affiliates, and its (and its "
        "Affiliates\u2019) professional advisers (including legal, "
        "financial, tax and accounting advisers), auditors and approved "
        "sub-contractors, in each case on a strict \u201cneed-to-know\u201d "
        "basis for the Purpose."
    )
    add_clause(
        doc, "3.1.9",
        "\u201cTerm\u201d has the meaning ascribed to it in Clause 8.1."
    )
    add_clause(
        doc, "3.1.10",
        "\u201cTrade Secret\u201d means any Confidential Information that "
        "derives independent economic value, actual or potential, from not "
        "being generally known to, and not being readily ascertainable "
        "through proper means by, another person who can obtain economic "
        "value from its disclosure or use, including without limitation "
        "source code, object code, cryptographic keys, proprietary "
        "algorithms, formulae, and customer lists of the Disclosing Party."
    )
    add_clause(
        doc, "3.2",
        "Interpretation. In this Agreement, unless the context otherwise "
        "requires: (a) headings are for convenience only and shall not "
        "affect interpretation; (b) the singular includes the plural and "
        "vice versa; (c) references to statutes are to such statutes as "
        "amended, modified or re-enacted from time to time; (d) the words "
        "\u201cincluding\u201d and \u201cinclude\u201d shall be deemed to "
        "be followed by the words \u201cwithout limitation\u201d; and (e) "
        "references to a \u201cperson\u201d include any individual, body "
        "corporate, partnership, limited liability partnership, trust, "
        "governmental authority or other legal entity."
    )

    # -----------------------------------------------------------------------
    # OBLIGATIONS OF RECEIVING PARTY
    # -----------------------------------------------------------------------
    add_h1(doc, "4", "Obligations of the Receiving Party")
    add_clause(
        doc, "4.1",
        "Standard of Care. The Receiving Party shall hold all Confidential "
        "Information of the Disclosing Party in strict confidence and shall "
        "exercise at least the same degree of care to prevent the "
        "unauthorised disclosure, use, copying, dissemination or "
        "publication of such Confidential Information as it exercises to "
        "protect its own confidential information of a similar nature, and "
        "in no event less than a reasonable degree of care."
    )
    add_clause(
        doc, "4.2",
        "Use Restriction. The Receiving Party shall use the Confidential "
        "Information solely and exclusively for the Purpose and shall not "
        "use, reproduce, exploit or commercialise, directly or indirectly, "
        "any Confidential Information for any other purpose, including, "
        "without limitation, for its own benefit or the benefit of any "
        "third party, or in any manner adverse to the Disclosing Party."
    )
    add_clause(
        doc, "4.3",
        "Need-to-Know Disclosure to Representatives. The Receiving Party "
        "may disclose Confidential Information only to those of its "
        "Representatives who (a) have a bona fide need to know such "
        "information for the Purpose; (b) have been informed of the "
        "confidential nature of such information; and (c) are bound by "
        "written obligations of confidentiality and non-use no less "
        "restrictive than those set forth herein (or, in the case of "
        "professional advisers, by equivalent professional duties of "
        "confidentiality)."
    )
    add_clause(
        doc, "4.4",
        "Responsibility for Representatives. The Receiving Party shall "
        "be responsible for any act or omission of its Representatives in "
        "breach of this Agreement as if such act or omission were that of "
        "the Receiving Party itself, and shall indemnify the Disclosing "
        "Party in respect thereof in accordance with Clause 12 "
        "(Indemnity)."
    )
    add_clause(
        doc, "4.5",
        "Security Measures. The Receiving Party shall implement and "
        "maintain commercially reasonable administrative, physical and "
        "technical safeguards (including, where appropriate, encryption, "
        "access controls, secure storage, audit logs and incident response "
        "procedures) designed to protect the confidentiality, integrity "
        "and availability of the Confidential Information, consistent with "
        "the reasonable security practices and procedures contemplated by "
        "section 43A of the Information Technology Act, 2000 and the "
        "rules framed thereunder."
    )
    add_clause(
        doc, "4.6",
        "No Reverse Engineering. Notwithstanding anything contained "
        "herein, the Receiving Party shall not, and shall procure that "
        "its Representatives shall not, decompile, disassemble, reverse "
        "engineer or otherwise attempt to derive the source code, "
        "underlying ideas, algorithms, file formats or non-public "
        "application programming interfaces of any software or technology "
        "comprised in the Confidential Information, except to the "
        "limited extent expressly permitted by Applicable Law."
    )
    add_clause(
        doc, "4.7",
        "Notice of Unauthorised Disclosure. The Receiving Party shall "
        "promptly, and in any event within five (5) Working Days of "
        "becoming aware thereof, notify the Disclosing Party in writing "
        "of any actual or suspected loss, unauthorised access, disclosure, "
        "use or misappropriation of any Confidential Information and shall "
        "cooperate with the Disclosing Party, at the Receiving Party\u2019s "
        "cost, in taking all reasonable steps to mitigate, contain and "
        "remedy such occurrence."
    )

    # -----------------------------------------------------------------------
    # PERMITTED DISCLOSURES / EXCLUSIONS
    # -----------------------------------------------------------------------
    add_h1(doc, "5", "Exclusions and Permitted Disclosures")
    add_clause(
        doc, "5.1",
        "Exclusions. The obligations of confidentiality and non-use set "
        "forth in Clause 4 shall not apply to any information that the "
        "Receiving Party can demonstrate, by reasonable contemporaneous "
        "written evidence:"
    )
    add_lettered(
        doc, "a",
        "was lawfully in its possession, free of any obligation of "
        "confidence, prior to the time of disclosure by the Disclosing "
        "Party;"
    )
    add_lettered(
        doc, "b",
        "is or becomes generally known to the public through no act or "
        "omission of the Receiving Party or its Representatives in breach "
        "of this Agreement;"
    )
    add_lettered(
        doc, "c",
        "was independently developed by the Receiving Party without "
        "reference to, use of or reliance upon the Confidential "
        "Information of the Disclosing Party;"
    )
    add_lettered(
        doc, "d",
        "was lawfully received by the Receiving Party from a third party "
        "having, to the Receiving Party\u2019s knowledge, the right to "
        "disclose the same and without any obligation of confidentiality; "
        "or"
    )
    add_lettered(
        doc, "e",
        "is expressly released from confidential treatment by the prior "
        "written consent of the Disclosing Party."
    )
    add_clause(
        doc, "5.2",
        "Compelled Disclosure. In the event that the Receiving Party or "
        "any of its Representatives is required by Applicable Law, by a "
        "court of competent jurisdiction, or by any regulatory, judicial "
        "or governmental authority to disclose any Confidential "
        "Information, the Receiving Party shall, to the extent legally "
        "permissible, (a) provide the Disclosing Party with prompt prior "
        "written notice of such requirement (and in any event not later "
        "than five (5) Working Days after becoming aware thereof or as "
        "soon as legally permissible, whichever is earlier); (b) cooperate, "
        "to the extent of commercially reasonable steps and at the "
        "Disclosing Party\u2019s cost, with any efforts by the Disclosing "
        "Party to obtain a protective order or other appropriate remedy; "
        "and (c) disclose only that portion of the Confidential "
        "Information which is legally required to be disclosed and shall "
        "use reasonable efforts to obtain assurances that confidential "
        "treatment will be accorded to such Confidential Information."
    )
    add_clause(
        doc, "5.3",
        "Whistleblowing Carve-out. Nothing in this Agreement shall "
        "restrict or prohibit either Party or its Representatives from "
        "making a good-faith report or disclosure of any suspected "
        "violation of Applicable Law to any appropriate governmental, "
        "regulatory or judicial authority, or from participating in any "
        "investigation or proceeding conducted by such authority, it being "
        "acknowledged that the lawful object of this Agreement does not, "
        "and shall not, extend to the suppression of information "
        "pertaining to unlawful acts."
    )

    # -----------------------------------------------------------------------
    # RETURN / DESTRUCTION
    # -----------------------------------------------------------------------
    add_h1(doc, "6", "Return or Destruction of Confidential Information")
    add_clause(
        doc, "6.1",
        "Upon (a) the earlier of expiry or termination of this Agreement; "
        "(b) completion or abandonment of the Purpose; or (c) written "
        "request by the Disclosing Party at any time, the Receiving Party "
        "shall, within thirty (30) days of such event or request, at the "
        "Disclosing Party\u2019s election, either return to the Disclosing "
        "Party or irrevocably destroy all Confidential Information in its "
        "possession or control (and that of its Representatives), "
        "including all copies, extracts, summaries and derivative works "
        "thereof, in whatever medium held."
    )
    add_clause(
        doc, "6.2",
        "The Receiving Party shall, within the said thirty (30)-day period, "
        "furnish to the Disclosing Party a written certificate signed by "
        "an authorised officer of the Receiving Party confirming full "
        "compliance with Clause 6.1."
    )
    add_clause(
        doc, "6.3",
        "Archival Retention. Notwithstanding Clause 6.1, the Receiving "
        "Party may retain (a) one (1) archival copy of the Confidential "
        "Information solely to the extent required for compliance with "
        "Applicable Law, bona fide internal record-keeping, or legal or "
        "regulatory audit requirements; and (b) Confidential Information "
        "contained in routine electronic back-ups that are not readily "
        "accessible to ordinary users, provided that any such retained "
        "Confidential Information shall remain subject to the obligations "
        "of confidentiality and non-use set out in this Agreement for so "
        "long as it is retained."
    )

    # -----------------------------------------------------------------------
    # TERM & SURVIVAL
    # -----------------------------------------------------------------------
    add_h1(doc, "7", "Term and Survival")
    add_clause(
        doc, "8.1" if False else "7.1",
        "Term. This Agreement shall commence on the Effective Date and "
        "shall, unless earlier terminated in accordance with Clause 7.2, "
        "continue in force for a period of two (2) years (the "
        "\u201cTerm\u201d). The Parties may extend the Term by mutual "
        "written agreement."
    )
    add_clause(
        doc, "7.2",
        "Termination. Either Party may terminate this Agreement (a) for "
        "convenience, by giving the other Party not less than thirty (30) "
        "days\u2019 prior written notice; or (b) forthwith, upon written "
        "notice, in the event of a material breach by the other Party "
        "which, if capable of being cured, remains uncured for fifteen "
        "(15) days after written notice thereof."
    )
    add_clause(
        doc, "7.3",
        "Survival of Confidentiality Obligations. Notwithstanding anything "
        "contained herein, the obligations of confidentiality and non-use "
        "set forth in this Agreement shall survive the expiry or "
        "termination of this Agreement for a period of three (3) years, "
        "in respect of general Confidential Information; and shall "
        "continue in perpetuity (or for the maximum period permitted "
        "under Applicable Law) in respect of Trade Secrets, source code, "
        "cryptographic keys and Personal Data, for so long as such "
        "information retains the character of a Trade Secret or remains "
        "subject to statutory confidentiality under Applicable Law."
    )
    add_clause(
        doc, "7.4",
        "Survival of Clauses. Clauses 4 (Obligations), 5 (Exclusions), "
        "6 (Return or Destruction), 7 (Term and Survival), 9 (No Licence), "
        "10 (Intellectual Property), 11 (Remedies and Injunctive Relief), "
        "12 (Indemnity), 13 (Limitation of Liability), 14 (Dispute "
        "Resolution), 15 (Governing Law and Jurisdiction), 16 (Notices), "
        "and 18 (Miscellaneous) shall survive the expiry or termination "
        "of this Agreement to the extent necessary to give effect to "
        "their terms."
    )

    # -----------------------------------------------------------------------
    # NO LICENCE / NO WARRANTY
    # -----------------------------------------------------------------------
    add_h1(doc, "8", "No Licence and No Warranty")
    add_clause(
        doc, "8.1",
        "No Licence. Nothing in this Agreement shall be construed as "
        "granting or conferring, expressly, impliedly, by estoppel or "
        "otherwise, any right, title, interest, licence or sub-licence, "
        "by either Party to the other, in or to any Confidential "
        "Information or any intellectual property, including without "
        "limitation any patent, copyright, trade mark, design, trade "
        "secret or know-how, save and except the limited right to use "
        "Confidential Information solely for the Purpose in accordance "
        "with the terms of this Agreement."
    )
    add_clause(
        doc, "8.2",
        "No Warranty. All Confidential Information is provided on an "
        "\u201cas is\u201d basis. Neither Party makes any representation "
        "or warranty, whether express or implied, as to the accuracy, "
        "completeness, fitness for any particular purpose, or "
        "non-infringement of any Confidential Information disclosed "
        "hereunder, and each Party expressly disclaims all such "
        "representations and warranties. Neither Party shall have any "
        "liability to the other resulting from the use of Confidential "
        "Information by the other Party."
    )
    add_clause(
        doc, "8.3",
        "No Obligation to Proceed. Neither this Agreement nor the "
        "disclosure or receipt of Confidential Information shall "
        "constitute or imply any promise, commitment or obligation on "
        "the part of either Party to enter into any further agreement, "
        "transaction or business relationship with the other."
    )

    # -----------------------------------------------------------------------
    # INTELLECTUAL PROPERTY
    # -----------------------------------------------------------------------
    add_h1(doc, "9", "Intellectual Property Preservation")
    add_clause(
        doc, "9.1",
        "All right, title and interest in and to the Confidential "
        "Information, and all intellectual property rights subsisting "
        "therein (including all rights under the Copyright Act, 1957; the "
        "Patents Act, 1970; the Trade Marks Act, 1999; the Designs Act, "
        "2000; and equivalent laws in other jurisdictions), shall remain "
        "the sole and exclusive property of the Disclosing Party."
    )
    add_clause(
        doc, "9.2",
        "No Transfer. Save as expressly set out herein, nothing in this "
        "Agreement shall operate as a transfer, assignment, licence or "
        "waiver of any intellectual property rights or any other "
        "proprietary rights of the Disclosing Party."
    )
    add_clause(
        doc, "9.3",
        "Feedback. Any suggestions, ideas, improvements or feedback "
        "provided by the Receiving Party to the Disclosing Party "
        "concerning the Disclosing Party\u2019s products, services or "
        "Confidential Information (\u201cFeedback\u201d) may be freely "
        "used by the Disclosing Party without attribution or compensation, "
        "provided that the Disclosing Party shall not identify the "
        "Receiving Party as the source of such Feedback without the "
        "Receiving Party\u2019s prior written consent."
    )

    # -----------------------------------------------------------------------
    # DATA PROTECTION
    # -----------------------------------------------------------------------
    add_h1(doc, "10", "Data Protection and DPDP Act 2023 Compliance")
    add_clause(
        doc, "10.1",
        "Where the Confidential Information disclosed under this Agreement "
        "includes or comprises Personal Data of any Data Principal, each "
        "Party shall, in respect of its role as a Data Fiduciary or Data "
        "Processor (as the case may be), comply with all applicable "
        "provisions of the Digital Personal Data Protection Act, 2023 and "
        "the rules, regulations, notifications and directions of the Data "
        "Protection Board of India issued thereunder, including the "
        "principles of lawful purpose, purpose limitation, data "
        "minimisation, accuracy, storage limitation, security and "
        "accountability."
    )
    add_clause(
        doc, "10.2",
        "Role Allocation. For the avoidance of doubt, the disclosure of "
        "Personal Data under this Agreement does not, by itself, "
        "constitute the Receiving Party a Data Processor for the "
        "Disclosing Party. Where the Parties intend to engage in "
        "processing of Personal Data on behalf of one another, they "
        "shall execute a separate data processing agreement setting out "
        "the matters required under section 8(5) of the DPDP Act, 2023."
    )
    add_clause(
        doc, "10.3",
        "Security Standard. Each Party shall implement appropriate "
        "technical and organisational measures to protect Personal Data "
        "against unauthorised or unlawful processing, accidental loss, "
        "destruction, damage, alteration or disclosure, including the "
        "\u201creasonable security practices and procedures\u201d "
        "contemplated by section 43A of the Information Technology Act, "
        "2000."
    )
    add_clause(
        doc, "10.4",
        "Breach Notification. In the event of any \u201cpersonal data "
        "breach\u201d (as understood under section 8(6) of the DPDP Act, "
        "2023) affecting Personal Data received from the other Party, "
        "the affected Party shall notify the other Party in writing "
        "without undue delay, and in any event within seventy-two (72) "
        "hours of becoming aware of such breach, and shall cooperate with "
        "the other Party in making any notifications required to be made "
        "to the Data Protection Board of India and to affected Data "
        "Principals."
    )
    add_clause(
        doc, "10.5",
        "Cross-Border Transfer. Neither Party shall transfer any Personal "
        "Data outside India except in accordance with section 16 of the "
        "DPDP Act, 2023 and any notifications issued by the Central "
        "Government thereunder."
    )

    # -----------------------------------------------------------------------
    # REMEDIES & INJUNCTIVE RELIEF
    # -----------------------------------------------------------------------
    add_h1(doc, "11", "Remedies and Injunctive Relief")
    add_clause(
        doc, "11.1",
        "Acknowledgement. Each Party acknowledges and agrees that the "
        "Confidential Information is of a special, unique, unusual and "
        "extraordinary character, the misuse or unauthorised disclosure "
        "of which would cause irreparable harm to the Disclosing Party "
        "for which monetary damages would not be an adequate remedy."
    )
    add_clause(
        doc, "11.2",
        "Specific Performance and Injunctive Relief. Accordingly, in "
        "addition to any other remedies available at law or in equity, "
        "the Disclosing Party shall be entitled, as a matter of right "
        "and without the need to prove actual damage, to seek specific "
        "performance under sections 10 and 14 of the Specific Relief Act, "
        "1963, and temporary, interlocutory, ad-interim, perpetual and "
        "mandatory injunctive relief under sections 36 to 42 of the "
        "Specific Relief Act, 1963 and Order XXXIX Rules 1 and 2 of the "
        "Code of Civil Procedure, 1908, to prevent or restrain any actual "
        "or threatened breach of this Agreement, without prejudice to "
        "its right to claim damages, accounts of profits, delivery-up "
        "and any other relief."
    )
    add_clause(
        doc, "11.3",
        "Criminal Liability. The Parties acknowledge that any wilful "
        "disclosure of Confidential Information in breach of this "
        "Agreement may attract criminal liability under section 72A of "
        "the Information Technology Act, 2000, and other Applicable Law, "
        "and nothing herein shall prejudice the right of the aggrieved "
        "Party to pursue such remedies as are available under Applicable "
        "Law."
    )
    add_clause(
        doc, "11.4",
        "Cumulative Remedies. All rights and remedies under this "
        "Agreement shall be cumulative and in addition to, and not in "
        "derogation or exclusion of, any rights and remedies available "
        "at law or in equity."
    )

    # -----------------------------------------------------------------------
    # INDEMNITY
    # -----------------------------------------------------------------------
    add_h1(doc, "12", "Indemnity")
    add_clause(
        doc, "12.1",
        "Mutual Indemnity. Each Party (the \u201cIndemnifying Party\u201d) "
        "shall indemnify, defend and hold harmless the other Party, its "
        "Affiliates, and its and their respective directors, officers, "
        "employees and agents (collectively, the \u201cIndemnified "
        "Parties\u201d), from and against any and all direct losses, "
        "damages, liabilities, costs and expenses (including reasonable "
        "attorneys\u2019 fees) actually suffered or incurred by the "
        "Indemnified Parties arising out of or in connection with any "
        "breach by the Indemnifying Party (or its Representatives) of "
        "Clauses 4, 5, 6 or 10 of this Agreement."
    )
    add_clause(
        doc, "12.2",
        "Cap on Indemnity. Notwithstanding anything contained herein, the "
        "aggregate liability of the Indemnifying Party under this "
        "Clause 12 shall not exceed Indian Rupees One Crore "
        "(INR 1,00,00,000) or one (1) times the aggregate value of the "
        "commercial engagement (if any) between the Parties to which the "
        "Purpose relates, whichever is lower; provided that this cap "
        "shall not apply to (a) liability arising from wilful misconduct "
        "or fraud; (b) breach of Clause 10 (Data Protection) resulting "
        "in statutory penalties under the DPDP Act, 2023; or (c) "
        "intentional misappropriation of Trade Secrets."
    )
    add_clause(
        doc, "12.3",
        "Indemnification Procedure. The Indemnified Parties shall (a) "
        "promptly notify the Indemnifying Party in writing of any claim "
        "in respect of which indemnity is sought; (b) provide the "
        "Indemnifying Party with reasonable cooperation in the defence "
        "of the claim (at the Indemnifying Party\u2019s cost); and (c) "
        "not make any admission or settlement of the claim without the "
        "Indemnifying Party\u2019s prior written consent (such consent "
        "not to be unreasonably withheld)."
    )

    # -----------------------------------------------------------------------
    # LIMITATION OF LIABILITY
    # -----------------------------------------------------------------------
    add_h1(doc, "13", "Limitation of Liability")
    add_clause(
        doc, "13.1",
        "Exclusion of Indirect Damages. Subject to Clause 13.2, in no "
        "event shall either Party be liable to the other Party, whether "
        "in contract, tort (including negligence), under statute or "
        "otherwise, for any indirect, incidental, consequential, special, "
        "punitive or exemplary damages, or for any loss of profits, loss "
        "of revenue, loss of business, loss of goodwill, loss of "
        "opportunity, or loss of data, arising out of or in connection "
        "with this Agreement, even if such Party has been advised of the "
        "possibility of such damages."
    )
    add_clause(
        doc, "13.2",
        "Carve-outs. The exclusions and limitations in Clause 13.1 shall "
        "not apply to: (a) a Party\u2019s indemnification obligations under "
        "Clause 12; (b) breach of the confidentiality and non-use "
        "obligations in Clause 4; (c) infringement or misappropriation "
        "of the other Party\u2019s intellectual property rights; (d) "
        "breach of Clause 10 (Data Protection); (e) fraud, gross "
        "negligence or wilful misconduct; or (f) any liability which "
        "cannot, as a matter of Applicable Law, be excluded or limited."
    )
    add_clause(
        doc, "13.3",
        "Liquidated Damages Acknowledgement. Any amount specified as "
        "liquidated damages in any schedule to this Agreement or in any "
        "related document represents a genuine and reasonable pre-estimate "
        "of the loss likely to be suffered by the Disclosing Party in the "
        "event of breach, taking into account the cost of investigation, "
        "the loss of competitive advantage, reputational harm and the "
        "difficulty of quantifying actual damages, and shall be recoverable "
        "in accordance with section 74 of the Indian Contract Act, 1872."
    )

    # -----------------------------------------------------------------------
    # DISPUTE RESOLUTION
    # -----------------------------------------------------------------------
    add_h1(doc, "14", "Dispute Resolution")
    add_clause(
        doc, "14.1",
        "Amicable Resolution. The Parties shall use their best endeavours "
        "to resolve amicably, through good-faith negotiations between "
        "designated senior representatives, any dispute, controversy or "
        "claim arising out of, relating to, or in connection with this "
        "Agreement, including any question regarding its existence, "
        "validity, interpretation, performance, breach or termination "
        "(a \u201cDispute\u201d), within a period of thirty (30) days from "
        "the date on which written notice of the Dispute is served by "
        "one Party on the other."
    )
    add_clause(
        doc, "14.2",
        "Arbitration. Any Dispute not resolved under Clause 14.1 shall "
        "be referred to and finally resolved by arbitration conducted "
        "under the Arbitration and Conciliation Act, 1996 (as amended "
        "from time to time)."
    )
    add_clause(
        doc, "14.3",
        "Number of Arbitrators. The arbitral tribunal shall consist of a "
        "sole arbitrator, to be appointed by mutual consent of the "
        "Parties; failing such consent within thirty (30) days of the "
        "request for arbitration, the sole arbitrator shall be appointed "
        "in accordance with section 11 of the Arbitration and "
        "Conciliation Act, 1996."
    )
    add_clause(
        doc, "14.4",
        "Seat and Venue. The seat of arbitration shall be New Delhi, "
        "Delhi, India, and the venue of the arbitration shall be "
        "New Delhi (or such other venue as the Parties may agree in "
        "writing). The courts at New Delhi shall have exclusive "
        "supervisory jurisdiction over the arbitration."
    )
    add_clause(
        doc, "14.5",
        "Language. The language of the arbitration shall be English. The "
        "award shall be reasoned, final and binding upon the Parties and "
        "may be enforced in any court of competent jurisdiction."
    )
    add_clause(
        doc, "14.6",
        "Governing Law of the Arbitration Agreement. The arbitration "
        "agreement contained in this Clause 14 shall be governed by the "
        "substantive laws of India."
    )
    add_clause(
        doc, "14.7",
        "Interim Relief \u2014 Section 9 Carve-out. Notwithstanding "
        "anything contained in this Clause 14, either Party shall be "
        "entitled, at any time (whether before, during or after the "
        "arbitration proceedings), to apply to a court of competent "
        "jurisdiction at New Delhi for interim measures of protection "
        "under section 9 of the Arbitration and Conciliation Act, 1996, "
        "including injunctions, preservation of assets and interim "
        "custody of Confidential Information, in respect of any actual "
        "or threatened breach of this Agreement."
    )
    add_clause(
        doc, "14.8",
        "Costs. The costs of the arbitration, including the fees of the "
        "arbitrator, shall be borne by the Parties as determined by the "
        "arbitrator in accordance with section 31A of the Arbitration and "
        "Conciliation Act, 1996."
    )
    add_clause(
        doc, "14.9",
        "Continuation of Performance. Pending resolution of any Dispute, "
        "the Parties shall continue to perform their respective "
        "obligations under this Agreement, save in respect of the subject "
        "matter of such Dispute."
    )

    # -----------------------------------------------------------------------
    # GOVERNING LAW & JURISDICTION
    # -----------------------------------------------------------------------
    add_h1(doc, "15", "Governing Law and Jurisdiction")
    add_clause(
        doc, "15.1",
        "This Agreement and any non-contractual obligations arising out "
        "of or in connection with it shall be governed by and construed "
        "in accordance with the substantive laws of the Republic of "
        "India, without regard to its conflict-of-laws principles."
    )
    add_clause(
        doc, "15.2",
        "Subject to Clause 14 (Dispute Resolution), the courts at "
        "New Delhi, Delhi, India shall have exclusive jurisdiction "
        "over any matter arising out of or in connection with this "
        "Agreement."
    )

    # -----------------------------------------------------------------------
    # NOTICES
    # -----------------------------------------------------------------------
    add_h1(doc, "16", "Notices")
    add_clause(
        doc, "16.1",
        "Form. All notices, requests, demands, claims and other "
        "communications under this Agreement shall be in writing in the "
        "English language and shall be sent to the addresses set out "
        "below (or to such other address as a Party may notify to the "
        "other Party from time to time in accordance with this Clause)."
    )
    add_clause(
        doc, "16.2",
        "Delivery. A notice shall be deemed to have been validly served: "
        "(a) if delivered by hand, on the date of delivery; (b) if sent "
        "by registered post or reputable courier, on the third (3rd) "
        "Working Day following the date of dispatch; or (c) if sent by "
        "electronic mail to the email address specified below, on the "
        "date of transmission, provided that no delivery-failure or "
        "bounce-back notification is received by the sender and a "
        "confirmatory copy is dispatched by any of the methods referred "
        "to in sub-clauses (a) or (b) within three (3) Working Days "
        "thereafter."
    )
    add_clause(
        doc, "16.3",
        "Addresses for service:"
    )
    add_para(doc, "If to AnantaSutra:", bold=True, indent_cm=1.0)
    add_para(
        doc,
        "AnantaSutra\n"
        "Attn: Mr. Himanshu Mishra, Founder & CEO\n"
        "Place: Delhi, India\n"
        "Email: contact@anantasutra.com\n"
        "Website: https://anantasutra.com",
        indent_cm=1.0
    )
    add_para(doc, "If to the Client:", bold=True, indent_cm=1.0)
    add_para(
        doc,
        "[CLIENT LEGAL NAME]\n"
        "Attention: [NAME], [DESIGNATION]\n"
        "Address: [CLIENT ADDRESS]\n"
        "Email: [EMAIL]\n"
        "Telephone: [TELEPHONE]",
        indent_cm=1.0
    )

    # -----------------------------------------------------------------------
    # FORCE MAJEURE
    # -----------------------------------------------------------------------
    add_h1(doc, "17", "Force Majeure")
    add_clause(
        doc, "17.1",
        "Neither Party shall be liable to the other for any failure or "
        "delay in the performance of its obligations under this Agreement "
        "(other than any obligation to make payment or to protect "
        "Confidential Information) to the extent that such failure or "
        "delay is caused by an event of Force Majeure, being any event "
        "beyond the reasonable control of the affected Party, including "
        "acts of God, earthquake, flood, fire, storm, epidemic, pandemic, "
        "governmental lockdown or quarantine, war, acts of terrorism, "
        "civil commotion, strikes, lockouts, failure of "
        "telecommunications or power infrastructure not attributable to "
        "the affected Party, or any change in Applicable Law that renders "
        "performance unlawful."
    )
    add_clause(
        doc, "17.2",
        "The affected Party shall (a) notify the other Party in writing "
        "without undue delay of the occurrence of the Force Majeure event "
        "and of its expected duration; and (b) use commercially "
        "reasonable efforts to mitigate the effects thereof and resume "
        "performance as soon as practicable. If a Force Majeure event "
        "continues for more than sixty (60) consecutive days, either "
        "Party may terminate this Agreement by written notice to the "
        "other, without liability (save in respect of accrued rights "
        "and obligations)."
    )

    # -----------------------------------------------------------------------
    # MISCELLANEOUS / BOILERPLATE
    # -----------------------------------------------------------------------
    add_h1(doc, "18", "Miscellaneous")
    add_clause(
        doc, "18.1",
        "Assignment. Neither Party shall assign, transfer, charge or "
        "otherwise deal with any of its rights or obligations under this "
        "Agreement without the prior written consent of the other Party, "
        "save that either Party may, upon prior written notice to the "
        "other, assign this Agreement to an Affiliate or to a successor "
        "in connection with a merger, amalgamation, corporate "
        "reorganisation or sale of substantially all of its assets, "
        "provided that the assignee agrees in writing to be bound by the "
        "terms of this Agreement."
    )
    add_clause(
        doc, "18.2",
        "Waiver. No failure or delay by either Party in exercising any "
        "right or remedy under this Agreement shall operate as a waiver "
        "thereof, nor shall any single or partial exercise preclude any "
        "other or further exercise thereof. Any waiver must be in writing "
        "and signed by the waiving Party to be effective."
    )
    add_clause(
        doc, "18.3",
        "Severability. If any provision of this Agreement is held by a "
        "court or arbitral tribunal of competent jurisdiction to be "
        "invalid, illegal or unenforceable, such provision shall be "
        "severed from this Agreement, and the remaining provisions shall "
        "continue in full force and effect. The Parties shall negotiate "
        "in good faith to replace the severed provision with a valid "
        "provision that most closely reflects the original commercial "
        "intent."
    )
    add_clause(
        doc, "18.4",
        "Entire Agreement. This Agreement constitutes the entire agreement "
        "and understanding between the Parties with respect to the "
        "subject matter hereof, and supersedes all prior and "
        "contemporaneous agreements, negotiations, representations, "
        "warranties and understandings, whether oral or written, in "
        "respect thereof. No variation of this Agreement shall be "
        "effective unless in writing and signed by duly authorised "
        "representatives of both Parties."
    )
    add_clause(
        doc, "18.5",
        "No Partnership or Agency. Nothing in this Agreement shall be "
        "deemed to constitute or create a partnership, joint venture, "
        "agency, employment or fiduciary relationship between the "
        "Parties, and neither Party shall have the authority to bind "
        "the other in any manner whatsoever."
    )
    add_clause(
        doc, "18.6",
        "No Publicity. Neither Party shall issue any press release, "
        "public announcement or marketing communication referring to the "
        "existence or subject matter of this Agreement, or use the name, "
        "logo or trade marks of the other Party, without the prior "
        "written consent of such other Party."
    )
    add_clause(
        doc, "18.7",
        "Non-Solicitation. During the Term and for a period of twelve "
        "(12) months thereafter, neither Party shall directly solicit for "
        "employment any employee of the other Party with whom it has had "
        "direct contact in connection with the Purpose, provided that "
        "nothing herein shall prohibit (a) general solicitations of "
        "employment not specifically directed at such employees; or (b) "
        "the hiring of any person who independently responds to such "
        "general solicitation or who approaches the soliciting Party on "
        "his or her own initiative."
    )
    add_clause(
        doc, "18.8",
        "Counterparts and Electronic Execution. This Agreement may be "
        "executed in any number of counterparts, each of which when "
        "executed shall constitute an original, and all of which together "
        "shall constitute one and the same instrument. The Parties agree "
        "that this Agreement may be validly executed by electronic "
        "signature, including by way of Aadhaar e-Sign or a digital "
        "signature certificate issued by a licensed Certifying Authority, "
        "and such electronic signature shall have the same legal force "
        "and effect as a physical signature, in accordance with sections "
        "3, 3A, 5 and 10A of the Information Technology Act, 2000. "
        "Electronic records of this Agreement shall be admissible in "
        "evidence in accordance with section 65B of the Indian Evidence "
        "Act, 1872 (and, upon its coming into force, section 63 of the "
        "Bharatiya Sakshya Adhiniyam, 2023)."
    )
    add_clause(
        doc, "18.9",
        "Third-Party Rights. No person other than a Party to this "
        "Agreement shall have any rights to enforce any term of this "
        "Agreement."
    )
    add_clause(
        doc, "18.10",
        "Further Assurances. Each Party shall, at its own cost, execute "
        "such further documents and take such further actions as may "
        "reasonably be required to give full effect to the provisions "
        "of this Agreement."
    )

    # -----------------------------------------------------------------------
    # STAMP DUTY ACKNOWLEDGMENT
    # -----------------------------------------------------------------------
    add_h1(doc, "19", "Stamp Duty")
    add_clause(
        doc, "19.1",
        "The Parties acknowledge that this Agreement is chargeable with "
        "stamp duty under Article 5 of the Schedule to the Delhi "
        "Stamp Act, 1957, as an agreement not otherwise provided for. "
        "The applicable stamp duty shall be paid by [ANANTASUTRA / THE "
        "CLIENT / THE PARTIES IN EQUAL SHARES] by way of e-stamp paper "
        "issued through the Stock Holding Corporation of India Limited "
        "(SHCIL) prior to the execution of this Agreement."
    )
    add_clause(
        doc, "19.2",
        "The Parties further acknowledge that, under section 35 of the "
        "Indian Stamp Act, 1899, an instrument not duly stamped shall not "
        "be admitted in evidence unless the stamp duty and any applicable "
        "penalty is paid, and that a Party seeking to rely on this "
        "Agreement shall bear the cost of any such deficiency to the "
        "extent attributable to its own default."
    )

    # -----------------------------------------------------------------------
    # EXECUTION BLOCK
    # -----------------------------------------------------------------------
    add_h1(doc, "20", "Execution")
    add_para(
        doc,
        "IN WITNESS WHEREOF, the Parties hereto have caused this Mutual "
        "Non-Disclosure Agreement to be executed by their respective duly "
        "authorised representatives as of the Effective Date first above "
        "written.",
        bold=True
    )
    doc.add_paragraph()  # spacer before signature block

    # --- Signature table: 2 columns AnantaSutra | Client ---
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    try:
        table.style = "Table Grid"
    except KeyError:
        pass

    hdr = table.rows[0].cells
    # AnantaSutra cell
    p = hdr[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("For and on behalf of\nANANTASUTRA")
    r.bold = True
    r.font.name = BODY_FONT
    r.font.size = Pt(BODY_SIZE)

    sutranet_lines = [
        "",
        "_________________________________",
        "Signature",
        "",
        "Name: Mr. Himanshu Mishra",
        "Designation: Founder & CEO",
        "Email: contact@anantasutra.com",
        "Date: [DATE]",
        "Place: Delhi",
        "",
        "Witness:",
        "Signature: _______________________",
        "Name: ___________________________",
        "Address: _________________________",
    ]
    for line in sutranet_lines:
        para = hdr[0].add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(line)
        run.font.name = BODY_FONT
        run.font.size = Pt(BODY_SIZE)

    # Client cell
    p = hdr[1].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("For and on behalf of\n[CLIENT LEGAL NAME]")
    r.bold = True
    r.font.name = BODY_FONT
    r.font.size = Pt(BODY_SIZE)

    client_lines = [
        "",
        "_________________________________",
        "Signature",
        "",
        "Name: [NAME]",
        "Designation: [DESIGNATION]",
        "Date: ______________________",
        "Place: _____________________",
        "",
        "(Common Seal, if adopted)",
        "",
        "Witness:",
        "Signature: _______________________",
        "Name: ___________________________",
        "Address: _________________________",
    ]
    for line in client_lines:
        para = hdr[1].add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(line)
        run.font.name = BODY_FONT
        run.font.size = Pt(BODY_SIZE)

    # -----------------------------------------------------------------------
    # SCHEDULE 1
    # -----------------------------------------------------------------------
    add_page_break(doc)
    add_centered_line(doc, "SCHEDULE 1", bold=True, size=13)
    add_centered_line(doc, "DESCRIPTION OF THE PURPOSE", bold=True, size=12)
    doc.add_paragraph()
    add_para(
        doc,
        "The \u201cPurpose\u201d for which Confidential Information shall "
        "be exchanged under this Agreement is:"
    )
    add_para(
        doc,
        "[DESCRIBE IN DETAIL THE SPECIFIC PURPOSE \u2014 e.g., evaluation "
        "of a proposed software development engagement; technical "
        "discovery and scoping workshops; proof-of-concept build; "
        "response to a request-for-proposal; negotiation of a Master "
        "Services Agreement and Statement of Work; due diligence in "
        "connection with a potential commercial transaction.]",
        indent_cm=0.5
    )
    add_para(
        doc,
        "Categories of Confidential Information reasonably expected to "
        "be exchanged include, without limitation: [e.g., technical "
        "architecture diagrams; source code extracts; proprietary "
        "algorithms; customer and employee lists; financial statements; "
        "pricing models; business plans; and Personal Data described in "
        "a separate data processing annexure, if any]."
    )

    return doc


def main() -> None:
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT_PATH)

    size_bytes = os.path.getsize(OUTPUT_PATH)
    size_kb = size_bytes / 1024
    print(f"Saved: {OUTPUT_PATH}")
    print(f"Size:  {size_bytes} bytes ({size_kb:.2f} KB)")


if __name__ == "__main__":
    main()
