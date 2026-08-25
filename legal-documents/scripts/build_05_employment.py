"""
Build the Employment Agreement DOCX for AnantaSutra

Generates: legal-documents/05_Employment_Agreement.docx

Complies with:
- Indian Contract Act 1872 (incl. s.27)
- Delhi Shops and Establishments Act 1954
- Payment of Wages Act 1936 / Code on Wages 2019
- Payment of Bonus Act 1965
- Payment of Gratuity Act 1972
- Employees' Provident Funds & Misc. Provisions Act 1952
- Employees' State Insurance Act 1948
- Maternity Benefit Act 1961 (as amended 2017)
- POSH Act 2013
- Industrial Employment (Standing Orders) Act 1946
- Four Labour Codes 2019-2020
- Copyright Act 1957 s.17(c) and Patents Act 1970
- IT Act 2000 and DPDP Act 2023
- Companies Act 2013 (s.62(1)(b) ESOP)
- Arbitration & Conciliation Act 1996
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUTPUT_PATH = r"c:/Users/duneja8515/Desktop/bhavya/prsnl/sutranet/legal-documents/05_Employment_Agreement.docx"

BASE_FONT = "Calibri"
BODY_SIZE = Pt(11)
HEADING1_SIZE = Pt(13)
HEADING2_SIZE = Pt(12)
TITLE_SIZE = Pt(16)


# -------------------------------------------------------------------
# Helpers
# -------------------------------------------------------------------


# ---------------------------------------------------------------------------
# Brand header (AnantaSutra) — injected by rebrand
# ---------------------------------------------------------------------------
LOGO_PATH = r"c:/Users/duneja8515/Desktop/bhavya/prsnl/sutranet/public/Favicon/logo-nobg.png"


def add_brand_header(doc):
    """Insert centered logo + brand lines at top of page 1. Never crashes."""
    try:
        from docx.shared import Inches as _Inches, Pt as _Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH as _CENTERED
    except Exception:
        return
    p = doc.add_paragraph()
    p.alignment = _CENTERED.CENTER
    try:
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=_Inches(1.2))
    except Exception:
        pass
    p2 = doc.add_paragraph()
    p2.alignment = _CENTERED.CENTER
    r2 = p2.add_run("AnantaSutra — अनन्तसूत्र")
    r2.italic = True
    r2.font.size = _Pt(10)
    p3 = doc.add_paragraph()
    p3.alignment = _CENTERED.CENTER
    r3 = p3.add_run("https://anantasutra.com  •  contact@anantasutra.com")
    r3.italic = True
    r3.font.size = _Pt(9)

def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'PAGE'
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_styles(doc):
    styles = doc.styles

    normal = styles['Normal']
    normal.font.name = BASE_FONT
    normal.font.size = BODY_SIZE
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    h1 = styles['Heading 1']
    h1.font.name = BASE_FONT
    h1.font.size = HEADING1_SIZE
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True

    h2 = styles['Heading 2']
    h2.font.name = BASE_FONT
    h2.font.size = HEADING2_SIZE
    h2.font.bold = True
    h2.font.italic = False
    h2.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.keep_with_next = True


def set_page_layout(doc):
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Confidential — AnantaSutra    |    Page ")
        run.font.name = BASE_FONT
        run.font.size = Pt(9)
        add_page_number_field(p)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = TITLE_SIZE
    run.font.name = BASE_FONT
    p.paragraph_format.space_after = Pt(12)


def add_h1(doc, text):
    p = doc.add_paragraph(text, style='Heading 1')
    return p


def add_h2(doc, text):
    p = doc.add_paragraph(text, style='Heading 2')
    return p


def add_para(doc, text, bold=False, italic=False, justify=True, indent_cm=None):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent_cm is not None:
        p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = BASE_FONT
    run.font.size = BODY_SIZE
    return p


def add_list_item(doc, text, level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.75)
    run = p.add_run(text)
    run.font.name = BASE_FONT
    run.font.size = BODY_SIZE
    return p


# -------------------------------------------------------------------
# Document construction
# -------------------------------------------------------------------

def build_document():
    doc = Document()
    add_brand_header(doc)
    configure_styles(doc)
    set_page_layout(doc)
    add_footer(doc)

    # Title
    add_title(doc, "EMPLOYMENT AGREEMENT")

    # ------------------------------------------------------------------
    # Parties and recitals
    # ------------------------------------------------------------------
    add_para(
        doc,
        "This Employment Agreement (\"Agreement\") is entered into on this [●] day of "
        "[●], [YEAR] (\"Effective Date\") at New Delhi, Delhi, India, by and between:",
    )

    add_para(
        doc,
        "ANANTASUTRA, a business concern operated and represented by its "
        "Founder & CEO, Mr. Himanshu Mishra, carrying on business under the trade "
        "name \"AnantaSutra\", having its principal place of business at Delhi, "
        "India, with contact address contact@anantasutra.com, acting through its "
        "Founder & CEO, Mr. Himanshu Mishra "
        "(hereinafter referred to as the "
        "\"Company\", \"Employer\" or \"AnantaSutra\", which expression shall, "
        "unless repugnant to the context or meaning thereof, be deemed to mean "
        "and include its successors, assigns, and the person(s) for the time "
        "being in control of the business) of the FIRST PART;",
    )

    add_para(doc, "AND", bold=True)

    add_para(
        doc,
        "[EMPLOYEE FULL NAME], son/daughter/spouse of [FATHER/GUARDIAN NAME], aged "
        "[AGE] years, Indian citizen holding PAN [PAN NO.] and Aadhaar reference "
        "[LAST FOUR DIGITS], residing at [EMPLOYEE RESIDENTIAL ADDRESS] (hereinafter "
        "referred to as the \"Employee\", which expression shall, unless repugnant to the "
        "context or meaning thereof, be deemed to mean and include his/her heirs, legal "
        "representatives, executors, administrators and permitted assigns) of the SECOND PART.",
    )

    add_para(
        doc,
        "The Company and the Employee are hereinafter individually referred to as a "
        "\"Party\" and collectively as the \"Parties\".",
    )

    # Recitals
    add_h1(doc, "RECITALS")
    add_para(
        doc,
        "A. The Company is engaged in the business of providing [DESCRIBE BUSINESS – e.g., "
        "information technology, software development, data analytics and allied "
        "services] (\"Business\").",
    )
    add_para(
        doc,
        "B. Pursuant to its letter of offer dated [OFFER LETTER DATE] (\"Offer Letter\"), "
        "the Company has offered employment to the Employee in the capacity of "
        "[DESIGNATION], and the Employee has accepted such offer on the terms and "
        "conditions set out in the Offer Letter and as more particularly described in this "
        "Agreement.",
    )
    add_para(
        doc,
        "C. The Parties now wish to record the terms and conditions governing the "
        "employment of the Employee with the Company in the manner set out herein.",
    )
    add_para(
        doc,
        "NOW THEREFORE, in consideration of the mutual covenants, promises and "
        "representations contained in this Agreement and for other good and valuable "
        "consideration, the receipt and sufficiency of which is hereby acknowledged, the "
        "Parties, intending to be legally bound, agree as follows:",
        bold=True,
    )

    # ------------------------------------------------------------------
    # 1. Definitions and interpretation
    # ------------------------------------------------------------------
    add_h1(doc, "1. DEFINITIONS AND INTERPRETATION")

    add_h2(doc, "1.1 Definitions")
    definitions = [
        ("\"Affiliate\"", "means, in relation to the Company, any entity that directly or "
         "indirectly controls, is controlled by, or is under common control with, the "
         "Company, where \"control\" means the beneficial ownership of not less than fifty "
         "per cent (50%) of the voting equity of such entity."),
        ("\"Applicable Law\"", "means all statutes, enactments, acts of legislature, laws, "
         "ordinances, rules, bye-laws, regulations, notifications, guidelines, directions, "
         "directives and orders of any governmental authority, statutory authority, court, "
         "or tribunal of competent jurisdiction in India, as may be in force from time to "
         "time."),
        ("\"CTC\"", "means the Cost to Company as set out in Annexure A."),
        ("\"Confidential Information\"", "shall have the meaning ascribed to it in the "
         "Employee Non-Disclosure Agreement executed or to be executed by the Employee "
         "contemporaneously with this Agreement (\"Employee NDA\")."),
        ("\"Intellectual Property\" or \"IP\"", "shall have the meaning ascribed to it in "
         "the Intellectual Property Assignment Deed executed or to be executed by the "
         "Employee contemporaneously with this Agreement (\"IP Deed\")."),
        ("\"Personal Data\"", "has the meaning given to it in section 2(t) of the Digital "
         "Personal Data Protection Act, 2023."),
        ("\"Policies\"", "means the policies, codes and standards adopted by the Company "
         "from time to time, including those referred to in Annexure C."),
        ("\"POSH Act\"", "means the Sexual Harassment of Women at Workplace (Prevention, "
         "Prohibition and Redressal) Act, 2013 and the rules made thereunder."),
        ("\"S&CE Act\"", "means the Delhi Shops and Commercial Establishments Act, "
         "1961 and the rules made thereunder."),
        ("\"Working Day\"", "means a day other than Saturday, Sunday or a public holiday "
         "declared by the Company or the Government of the NCT of Delhi."),
    ]
    for term, meaning in definitions:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.75)
        r1 = p.add_run(term + " ")
        r1.bold = True
        r1.font.name = BASE_FONT
        r1.font.size = BODY_SIZE
        r2 = p.add_run(meaning)
        r2.font.name = BASE_FONT
        r2.font.size = BODY_SIZE

    add_h2(doc, "1.2 Interpretation")
    add_para(
        doc,
        "(a) References to statutes or statutory provisions include references to such "
        "statutes or statutory provisions as amended, modified, re-enacted, consolidated "
        "or replaced from time to time (including, where applicable, the Code on Wages, "
        "2019; the Industrial Relations Code, 2020; the Code on Social Security, 2020; "
        "and the Occupational Safety, Health and Working Conditions Code, 2020).",
    )
    add_para(
        doc,
        "(b) Words importing the singular include the plural and vice versa; words "
        "importing one gender include all genders; and references to \"including\" shall be "
        "construed as \"including without limitation\".",
    )
    add_para(
        doc,
        "(c) The headings used in this Agreement are for convenience only and shall not "
        "affect the interpretation of any provision.",
    )

    # ------------------------------------------------------------------
    # 2. Appointment and designation
    # ------------------------------------------------------------------
    add_h1(doc, "2. APPOINTMENT AND DESIGNATION")
    add_para(
        doc,
        "2.1 The Company hereby appoints the Employee, and the Employee hereby accepts "
        "such appointment, in the position of [DESIGNATION] in the [DEPARTMENT/FUNCTION] of "
        "the Company, reporting to [REPORTING MANAGER NAME AND DESIGNATION], on and with "
        "effect from [DATE OF JOINING] (\"Date of Joining\").",
    )
    add_para(
        doc,
        "2.2 The Employee's duties and responsibilities are more particularly described in "
        "Annexure B (Job Description) and shall include such additional or ancillary "
        "duties as may be assigned to the Employee by the Company from time to time, "
        "consistent with the Employee's role, experience and seniority.",
    )
    add_para(
        doc,
        "2.3 The Company may, at its sole discretion and upon reasonable prior written "
        "notice, change the Employee's designation, reporting line, department, or scope "
        "of duties, provided that any such change shall not result in a material reduction "
        "of the Employee's CTC.",
    )

    # ------------------------------------------------------------------
    # 3. Commencement and probation
    # ------------------------------------------------------------------
    add_h1(doc, "3. COMMENCEMENT AND PROBATION")
    add_para(
        doc,
        "3.1 The Employee's employment under this Agreement shall commence on the Date of "
        "Joining and shall continue until terminated in accordance with Clause 13.",
    )
    add_para(
        doc,
        "3.2 The Employee shall be on probation for a period of six (6) months from the "
        "Date of Joining (\"Probation Period\"). The Company may, at its sole discretion, "
        "extend the Probation Period by a further period not exceeding three (3) months "
        "upon giving written notice to the Employee before expiry of the initial Probation "
        "Period.",
    )
    add_para(
        doc,
        "3.3 Unless the Employee is issued a written letter of confirmation by the "
        "Company, the Employee shall be deemed to continue on probation. Confirmation "
        "shall be subject to the Employee's satisfactory performance, attendance, "
        "conduct and successful completion of background verification under Clause 16.",
    )
    add_para(
        doc,
        "3.4 During the Probation Period, either Party may terminate this Agreement by "
        "giving the other Party thirty (30) days' prior written notice or salary in lieu "
        "thereof, without assigning any reason. Termination for misconduct during "
        "probation shall be governed by Clause 13.3.",
    )

    # ------------------------------------------------------------------
    # 4. Duties and responsibilities
    # ------------------------------------------------------------------
    add_h1(doc, "4. DUTIES AND RESPONSIBILITIES")
    add_para(
        doc,
        "4.1 The Employee shall devote the whole of his/her working time, attention, "
        "skills and abilities exclusively to the duties of his/her employment with the "
        "Company and shall faithfully and diligently perform such duties to the best of "
        "his/her ability.",
    )
    add_para(
        doc,
        "4.2 The Employee shall comply with all Applicable Law and with all lawful "
        "directions, Policies, codes and procedures of the Company in force from time to "
        "time, including those listed in Annexure C.",
    )
    add_para(
        doc,
        "4.3 Full-time and exclusive service. During the continuance of this Agreement, "
        "the Employee shall not, whether directly or indirectly, engage in or undertake any "
        "other employment, trade, profession, business, consultancy, or commercial activity "
        "(whether paid or unpaid) without the prior written consent of the Company. The "
        "Employee acknowledges that this restriction applies during the currency of this "
        "Agreement only and is a reasonable restraint permissible in law (being a restraint "
        "during employment) in accordance with the principles recognised under Indian law.",
    )
    add_para(
        doc,
        "4.4 Conflict of interest. The Employee shall promptly disclose in writing to the "
        "Company any actual or potential conflict of interest (including financial, "
        "familial or commercial interests in any supplier, vendor, customer, competitor, "
        "collaborator, or counterparty of the Company) that may affect, or be perceived to "
        "affect, the objective performance of the Employee's duties. The Employee shall "
        "abide by the Company's decision on the management or avoidance of such conflict.",
    )
    add_para(
        doc,
        "4.5 The Employee shall act with utmost good faith, honesty and integrity in all "
        "dealings on behalf of the Company and shall not do, or omit to do, any act that "
        "is likely to bring the Company or any of its Affiliates into disrepute.",
    )

    # ------------------------------------------------------------------
    # 5. Place of work
    # ------------------------------------------------------------------
    add_h1(doc, "5. PLACE OF WORK")
    add_para(
        doc,
        "5.1 The Employee's primary place of work shall be the Company's office situated "
        "at [WORK LOCATION ADDRESS], New Delhi, Delhi, or such other place in India as "
        "the Company may reasonably designate from time to time.",
    )
    add_para(
        doc,
        "5.2 The Employee may, at the discretion of the Company and in accordance with its "
        "remote / hybrid working Policy, be permitted or required to perform his/her "
        "duties from a remote location or on a hybrid basis. The Employee shall ensure a "
        "safe, ergonomic and secure remote working environment consistent with the "
        "Company's occupational safety and data security requirements.",
    )
    add_para(
        doc,
        "5.3 The Employee may be required, from time to time, to travel within India and "
        "abroad in connection with the Business, and the Company shall reimburse reasonable "
        "business-related travel expenses in accordance with the Company's travel and "
        "reimbursement Policy.",
    )
    add_para(
        doc,
        "5.4 The Company may require the Employee to be transferred or deputed to any "
        "office, branch, project, client site, subsidiary or Affiliate of the Company in "
        "India upon reasonable written notice. Such transfer shall not, of itself, "
        "constitute a material change to the terms of this Agreement.",
    )

    # ------------------------------------------------------------------
    # 6. Working hours and leave
    # ------------------------------------------------------------------
    add_h1(doc, "6. WORKING HOURS AND LEAVE")
    add_para(
        doc,
        "6.0 Statutory registration. The Company shall comply with the Delhi Shops "
        "and Establishments Act, 1954 (and, where applicable, the Delhi Shops and "
        "Commercial Establishments Act, 1961) and shall obtain registration "
        "thereunder within the statutory timeline from the date on which such "
        "registration becomes mandatory in respect of the Company. Until such time, "
        "the Company shall in good faith observe the substantive standards prescribed "
        "under the said legislation in respect of working hours, weekly off, leave "
        "and holidays.",
    )
    add_para(
        doc,
        "6.1 The Employee's normal hours of work shall be [HOURS – typically 9] hours per "
        "day (inclusive of meal and rest intervals) and shall not ordinarily exceed "
        "forty-eight (48) hours in any week, consistent with the S&CE Act. The Employee "
        "may, from time to time, be required to work beyond the normal hours to meet "
        "business exigencies, and the Employee agrees to do so as reasonably required.",
    )
    add_para(
        doc,
        "6.2 The weekly off shall be [WEEKLY OFF DAY, typically Sunday] and such other "
        "days as the Company may designate. The Employee shall be entitled to the public "
        "and festival holidays notified by the Company for New Delhi, Delhi, in "
        "accordance with the S&CE Act.",
    )
    add_para(
        doc,
        "6.3 The Employee shall be entitled to earned leave, casual leave, sick leave, "
        "maternity leave (under the Maternity Benefit Act, 1961, as amended), and other "
        "statutory and discretionary leaves in accordance with the Company's Leave Policy "
        "and the S&CE Act. The Company's Leave Policy shall not provide less than the "
        "statutory minima prescribed under Applicable Law.",
    )
    add_para(
        doc,
        "6.4 All leave shall be availed with prior written approval of the Employee's "
        "reporting manager save in the case of sudden illness or emergency, in which case "
        "the Employee shall notify the Company at the earliest reasonable opportunity.",
    )

    # ------------------------------------------------------------------
    # 7. Compensation and benefits
    # ------------------------------------------------------------------
    add_h1(doc, "7. COMPENSATION AND BENEFITS")

    add_h2(doc, "7.1 CTC and salary structure")
    add_para(
        doc,
        "In consideration of the services to be rendered by the Employee under this "
        "Agreement, the Company shall pay to the Employee a Cost to Company of "
        "INR [AMOUNT IN FIGURES] (Rupees [AMOUNT IN WORDS] only) per annum, structured "
        "into Basic, House Rent Allowance (HRA), Special Allowance, Leave Travel "
        "Allowance (LTA), statutory retirals and variable pay components, as more "
        "particularly described in Annexure A. Salary shall be paid monthly in arrears, "
        "on or before the [PAYROLL DATE] of each succeeding month, by direct credit to "
        "the Employee's bank account, in accordance with the Payment of Wages Act, 1936 "
        "(and, upon full notification, the Code on Wages, 2019).",
    )

    add_h2(doc, "7.2 Statutory contributions and retirals")
    add_para(
        doc,
        "(a) Provident Fund. EPF contributions under the Employees' Provident Funds "
        "and Miscellaneous Provisions Act, 1952 (and, upon full notification, the Code "
        "on Social Security, 2020) shall apply from the date on which the Company "
        "reaches the statutory threshold (currently 20 employees) or voluntarily "
        "registers, whichever is earlier. Upon such applicability, the Company shall "
        "deduct and contribute amounts towards the Employees' Provident Fund in "
        "accordance with the said Act and the schemes framed thereunder. The Company's "
        "approach to PF contribution (statutory cap or full basic, as the case may be) "
        "will be set out in Annexure A at that time.",
    )
    add_para(
        doc,
        "(b) Gratuity. The Employee shall be entitled to gratuity in accordance with "
        "the Payment of Gratuity Act, 1972, upon completion of five (5) years of "
        "continuous service and subject to the Act becoming applicable to the Company "
        "(typically on the Company employing ten (10) or more persons).",
    )
    add_para(
        doc,
        "(c) Employees' State Insurance. ESI contributions under the Employees' State "
        "Insurance Act, 1948 shall apply from the date on which the Company reaches "
        "the statutory threshold under that Act (currently ten (10) employees in "
        "notified areas) or voluntarily registers, whichever is earlier. Upon such "
        "applicability, contributions shall be deducted and remitted in accordance "
        "with the said Act.",
    )
    add_para(
        doc,
        "(d) Professional Tax. Professional Tax, if applicable under the laws of the "
        "State in which the Employee is based, shall be deducted at source. The "
        "Parties acknowledge that the National Capital Territory of Delhi does not "
        "currently levy professional tax on individuals.",
    )
    add_para(
        doc,
        "(e) Tax Deducted at Source. Tax shall be deducted at source from the Employee's "
        "salary and perquisites under Section 192 and other applicable provisions of the "
        "Income-tax Act, 1961. The Employee shall furnish such declarations, investment "
        "proofs and documents as may be required by the Company for this purpose.",
    )

    add_h2(doc, "7.3 Performance bonus and variable pay")
    add_para(
        doc,
        "The Employee shall be eligible for an annual performance-linked variable pay / "
        "bonus as set out in Annexure A, payable subject to (i) the Employee being in the "
        "Company's employment and not under notice of termination (by either Party) on the "
        "date of payment, (ii) the achievement of individual and Company performance "
        "targets, and (iii) the applicable bonus Policy. Statutory bonus, where applicable "
        "under the Payment of Bonus Act, 1965 (or the Code on Wages, 2019 once fully "
        "notified), shall be paid as required by law and shall be deemed subsumed in the "
        "variable pay to the extent permitted by law.",
    )

    add_h2(doc, "7.4 Reimbursements")
    add_para(
        doc,
        "The Employee shall be entitled to reimbursement of reasonable business expenses "
        "properly incurred in the course of employment, against original supporting "
        "documents and in accordance with the Company's reimbursement Policy.",
    )

    add_h2(doc, "7.5 Employee Stock Options")
    add_para(
        doc,
        "If the Employee is granted stock options, restricted stock units, or other "
        "equity-linked instruments, such grant shall be governed exclusively by the terms "
        "of the Company's employee stock option plan (\"ESOP Plan\") as adopted under "
        "Section 62(1)(b) of the Companies Act, 2013 read with the rules made thereunder, "
        "and, if and when the Company is listed, the Securities and Exchange Board of "
        "India (Share Based Employee Benefits and Sweat Equity) Regulations, 2021 (\"SBEB "
        "Regulations\"), together with the applicable grant letter. Nothing in this "
        "Agreement shall be construed as conferring any right or entitlement to equity in "
        "the Company otherwise than in accordance with the ESOP Plan and grant letter.",
    )

    add_h2(doc, "7.6 Review")
    add_para(
        doc,
        "The Employee's CTC may be reviewed by the Company once in each financial year. "
        "Any revision shall be at the sole discretion of the Company and shall be "
        "communicated to the Employee in writing.",
    )

    # ------------------------------------------------------------------
    # 8. Leave policy cross-reference
    # ------------------------------------------------------------------
    add_h1(doc, "8. LEAVE POLICY")
    add_para(
        doc,
        "Detailed entitlements for earned leave, casual leave, sick leave, maternity "
        "leave, paternity leave (if any), bereavement leave, and other categories of "
        "leave are set out in the Company's Leave Policy (forming part of the Policies "
        "listed in Annexure C). The Employee acknowledges having received or been given "
        "access to such Leave Policy. In the event of any conflict between the Leave "
        "Policy and Applicable Law, the more beneficial provision for the Employee shall "
        "prevail.",
    )

    # ------------------------------------------------------------------
    # 9. Confidentiality and intellectual property
    # ------------------------------------------------------------------
    add_h1(doc, "9. CONFIDENTIALITY AND INTELLECTUAL PROPERTY")
    add_para(
        doc,
        "9.1 Confidentiality. The Employee shall execute and comply with the Employee NDA "
        "contemporaneously with this Agreement. The obligations under the Employee NDA "
        "shall apply during the term of the Employee's employment and shall survive "
        "termination for the periods specified therein.",
    )
    add_para(
        doc,
        "9.2 Intellectual property. The Employee shall execute and comply with the IP "
        "Deed contemporaneously with this Agreement. Without prejudice to Section 17(c) of "
        "the Copyright Act, 1957 and the default employer-vesting rule thereunder, the "
        "Employee hereby irrevocably and unconditionally assigns to the Company all right, "
        "title and interest (including future rights, to the extent permissible under "
        "Section 18(1) of the Copyright Act, 1957) in and to all Intellectual Property "
        "created, conceived, developed or reduced to practice by the Employee, alone or "
        "jointly with others, in the course of employment or using the resources, "
        "facilities or Confidential Information of the Company. The Employee shall, at "
        "the Company's request and expense, execute such further documents and do such "
        "further acts as may be necessary to perfect the Company's title to such "
        "Intellectual Property, including filings under the Patents Act, 1970, the "
        "Trade Marks Act, 1999 and the Designs Act, 2000.",
    )
    add_para(
        doc,
        "9.3 Moral rights. To the extent permitted by Applicable Law, the Employee waives "
        "all moral rights in such Intellectual Property, save and except the right under "
        "Section 57 of the Copyright Act, 1957 to restrain or claim damages in respect of "
        "distortion, mutilation or other modification of the work which would be "
        "prejudicial to the Employee's honour or reputation.",
    )
    add_para(
        doc,
        "9.4 Prior IP. Any Intellectual Property owned or developed by the Employee prior "
        "to the Date of Joining that is, or may be, relevant to the Business shall be "
        "disclosed in Annexure D. Intellectual Property not so disclosed shall be "
        "conclusively presumed to have been created in the course of the Employee's "
        "employment with the Company.",
    )

    # ------------------------------------------------------------------
    # 10. Restrictive covenants
    # ------------------------------------------------------------------
    add_h1(doc, "10. RESTRICTIVE COVENANTS")

    add_h2(doc, "10.1 Non-solicitation of clients")
    add_para(
        doc,
        "For a period of twelve (12) months from the Effective Termination Date (defined "
        "in Clause 13.8), the Employee shall not, directly or indirectly, solicit, canvass "
        "or induce any client, customer or counterparty of the Company with whom the "
        "Employee had material dealings, or of whom the Employee had knowledge of "
        "Confidential Information, during the twelve (12) months preceding the Effective "
        "Termination Date, to cease doing business with the Company or to divert business "
        "away from the Company.",
    )

    add_h2(doc, "10.2 Non-solicitation of personnel")
    add_para(
        doc,
        "For a period of twelve (12) months from the Effective Termination Date, the "
        "Employee shall not, directly or indirectly, solicit, entice, induce or endeavour "
        "to entice away any employee, officer, consultant or contractor of the Company "
        "with whom the Employee had material professional interaction during the twelve "
        "(12) months preceding the Effective Termination Date, to leave the service of the "
        "Company or to accept employment or engagement elsewhere.",
    )

    add_h2(doc, "10.3 No post-termination non-compete")
    add_para(
        doc,
        "The Parties expressly acknowledge that Section 27 of the Indian Contract Act, "
        "1872 renders void any agreement by which any person is restrained from exercising "
        "a lawful profession, trade or business of any kind. Accordingly, this Agreement "
        "does not impose any post-termination non-competition restriction on the Employee. "
        "The obligation of full-time and exclusive service under Clause 4.3 applies only "
        "during the currency of the employment and shall cease upon the Effective "
        "Termination Date.",
    )

    add_h2(doc, "10.4 Reasonableness and severability")
    add_para(
        doc,
        "The Employee acknowledges that the covenants in Clauses 10.1 and 10.2 are "
        "reasonable in scope, duration and geography having regard to the nature of the "
        "Business, the confidential nature of the information entrusted to the Employee, "
        "and the market-standard practice in India for comparable roles. If any court or "
        "tribunal of competent jurisdiction finds any restriction to be unenforceable, "
        "such restriction shall be modified to the minimum extent necessary to render it "
        "enforceable, and the remaining provisions shall continue in full force.",
    )

    # ------------------------------------------------------------------
    # 11. Code of conduct, POSH, anti-bribery and conflict of interest
    # ------------------------------------------------------------------
    add_h1(doc, "11. CODE OF CONDUCT, POSH, ANTI-BRIBERY AND CONFLICT OF INTEREST")
    add_para(
        doc,
        "11.1 The Employee shall comply with the Company's Code of Conduct, Anti-Bribery "
        "and Anti-Corruption Policy, Whistleblower Policy, Insider Trading Policy (once "
        "applicable), and other Policies notified to the Employee from time to time.",
    )
    add_para(
        doc,
        "11.2 POSH compliance. The Employee shall comply with the Company's Policy on "
        "Prevention of Sexual Harassment at the Workplace framed under the POSH Act. "
        "The Company shall constitute an Internal Committee under the POSH Act, 2013 "
        "upon reaching the statutory threshold of ten (10) employees or earlier as "
        "deemed appropriate by the Company. Pending such constitution, complaints shall "
        "be addressed to the Founder & CEO, or escalated to the Local Committee "
        "constituted by the District Officer under Section 6 of the POSH Act, as "
        "applicable. The Employee acknowledges having received the Company's POSH Policy.",
    )
    add_para(
        doc,
        "11.3 Anti-bribery. The Employee shall not, directly or indirectly, offer, "
        "promise, give, solicit, or accept any bribe, kickback, illegal gratification or "
        "other undue advantage in connection with the Business, and shall comply with the "
        "Prevention of Corruption Act, 1988 and all other anti-bribery laws applicable to "
        "the Company and its Affiliates.",
    )
    add_para(
        doc,
        "11.4 The Employee shall comply with the Company's policies on non-discrimination "
        "and equal opportunity and shall treat all colleagues, clients and counterparties "
        "with dignity and respect, in line with the rights recognised under the "
        "Constitution of India, the Rights of Persons with Disabilities Act, 2016 and the "
        "Transgender Persons (Protection of Rights) Act, 2019.",
    )

    # ------------------------------------------------------------------
    # 12. Data protection and acceptable use
    # ------------------------------------------------------------------
    add_h1(doc, "12. DATA PROTECTION AND ACCEPTABLE USE")
    add_para(
        doc,
        "12.1 The Employee acknowledges receipt of the Company's Privacy Notice to "
        "employees issued under the Digital Personal Data Protection Act, 2023 (\"DPDP Act\") "
        "and consents, to the extent required under law, to the processing of his/her "
        "Personal Data by the Company for employment-related purposes, including payroll, "
        "statutory compliance, performance management, background verification, health "
        "and safety, workplace security and information-technology administration.",
    )
    add_para(
        doc,
        "12.2 The Employee shall handle Personal Data of other data principals (including "
        "colleagues, clients, vendors and end-users) strictly in accordance with the DPDP "
        "Act, the Information Technology Act, 2000 (including Sections 43, 43A, 65, 66 and "
        "72A thereof), the Company's DPDP and information-security Policies, and any "
        "instructions of the Company as data fiduciary or, as the case may be, data "
        "processor.",
    )
    add_para(
        doc,
        "12.3 The Employee shall use the Company's information-technology resources "
        "(including email, messaging platforms, devices, servers, cloud services and "
        "source-code repositories) solely for authorised business purposes, in accordance "
        "with the Company's IT / Acceptable Use Policy, and acknowledges that the Company "
        "may monitor the use of such resources to the extent permitted by Applicable Law "
        "and disclosed in the Privacy Notice.",
    )
    add_para(
        doc,
        "12.4 The Employee shall promptly report any actual or suspected information "
        "security incident or Personal Data breach to the Company so as to enable the "
        "Company to comply with its obligations, including any obligation to notify the "
        "Data Protection Board of India within the timelines prescribed under the DPDP "
        "Act.",
    )

    # ------------------------------------------------------------------
    # 13. Termination
    # ------------------------------------------------------------------
    add_h1(doc, "13. TERMINATION")

    add_h2(doc, "13.1 Termination by the Employee")
    add_para(
        doc,
        "After confirmation of employment, the Employee may terminate this Agreement by "
        "giving the Company [NOTICE PERIOD – typically sixty (60) or ninety (90)] days' "
        "prior written notice. The Company may, at its discretion, accept a shorter notice "
        "period on such terms as it deems fit.",
    )

    add_h2(doc, "13.2 Termination by the Company without cause")
    add_para(
        doc,
        "After confirmation of employment, the Company may terminate this Agreement, "
        "without assigning any reason, by giving the Employee [NOTICE PERIOD – typically "
        "sixty (60) or ninety (90)] days' prior written notice, or salary in lieu of such "
        "notice (calculated on the Employee's Basic and Dearness Allowance, if any) as "
        "set out in Annexure A.",
    )

    add_h2(doc, "13.3 Termination for cause")
    add_para(
        doc,
        "Notwithstanding anything to the contrary in this Agreement, the Company may "
        "terminate this Agreement immediately, without notice or payment in lieu of "
        "notice, upon the occurrence of any of the following events (each a \"Cause Event\"):",
    )
    cause_items = [
        "(a) any act of fraud, theft, dishonesty, embezzlement, or misappropriation by the Employee;",
        "(b) gross misconduct, gross negligence, wilful disobedience or insubordination;",
        "(c) material breach of this Agreement, the Employee NDA, the IP Deed, or any Policy of the Company, which breach, if capable of being remedied, is not remedied to the satisfaction of the Company within a reasonable period specified in a written notice;",
        "(d) conviction by a court of competent jurisdiction for any offence involving moral turpitude, fraud or dishonesty;",
        "(e) any finding by the Company's Internal Committee under the POSH Act or an equivalent body of sexual harassment, bullying, or other serious workplace misconduct;",
        "(f) any act causing material loss, damage or disrepute to the Company or any of its Affiliates;",
        "(g) the Employee being declared insolvent, or of unsound mind by a court of competent jurisdiction; or",
        "(h) any other ground recognised as a \"Cause Event\" under the Company's Standing Orders (once certified) or under Applicable Law.",
    ]
    for c in cause_items:
        add_list_item(doc, c, level=0)
    add_para(
        doc,
        "The Employee shall be given a reasonable opportunity to be heard before any "
        "termination for cause, in accordance with the principles of natural justice and "
        "Applicable Law.",
    )

    add_h2(doc, "13.4 Retirement")
    add_para(
        doc,
        "The Employee shall retire from service upon attaining the age of [RETIREMENT "
        "AGE – typically 58 or 60] years, or such other age as may be specified in the "
        "Company's retirement Policy from time to time.",
    )

    add_h2(doc, "13.5 Garden leave")
    add_para(
        doc,
        "The Company may, at its sole discretion, during all or any part of the notice "
        "period under Clause 13.1 or Clause 13.2, place the Employee on garden leave, "
        "during which the Employee shall continue to be an employee of the Company, "
        "receive his/her regular salary and benefits, and remain subject to the "
        "obligations under this Agreement, but may be required (i) not to attend the "
        "Company's premises, (ii) not to contact clients, employees or counterparties of "
        "the Company, and (iii) to return Company property. The Parties acknowledge that "
        "garden leave applies only during the notice period and not thereafter.",
    )

    add_h2(doc, "13.6 Payment on termination")
    add_para(
        doc,
        "Upon the Effective Termination Date, the Company shall pay to the Employee all "
        "unpaid salary up to such date, accrued and unavailed earned leave encashment, "
        "reimbursements for approved business expenses, statutory dues (including "
        "provident fund and, where applicable, gratuity), and any variable pay or bonus "
        "earned but unpaid as per the applicable Policy, less all sums due from the "
        "Employee to the Company. Such full and final settlement shall be effected within "
        "the period prescribed under the S&CE Act and the Code on Wages, 2019 (once fully "
        "notified), and in any event no later than the second ordinary wage period after "
        "the Effective Termination Date.",
    )

    add_h2(doc, "13.7 Effect of termination")
    add_para(
        doc,
        "Termination of this Agreement shall be without prejudice to any rights or "
        "obligations accrued prior to such termination. The provisions relating to "
        "confidentiality, intellectual property, non-solicitation, return of property, "
        "indemnity, dispute resolution, governing law and such other provisions which by "
        "their nature are intended to survive termination, shall so survive.",
    )

    add_h2(doc, "13.8 Effective Termination Date")
    add_para(
        doc,
        "The \"Effective Termination Date\" means the date on which the Employee ceases to "
        "be on the rolls of the Company, being (as applicable) the date of expiry of the "
        "notice period, the date of acceptance of resignation with immediate effect, the "
        "date specified in a notice of termination for cause, or the date of retirement.",
    )

    # ------------------------------------------------------------------
    # 14. Exit obligations
    # ------------------------------------------------------------------
    add_h1(doc, "14. EXIT OBLIGATIONS")
    add_para(
        doc,
        "14.1 Return of property. On or before the Effective Termination Date (or earlier "
        "if so required by the Company), the Employee shall return to the Company all "
        "property belonging to the Company or its Affiliates, including laptops, mobile "
        "devices, access cards, keys, documents, records, source code, client lists, "
        "marketing materials, correspondence, Confidential Information in any form, and "
        "all copies and extracts thereof, and shall securely delete (or enable the "
        "Company to securely delete) any Confidential Information or Personal Data held on "
        "any personal device or account of the Employee.",
    )
    add_para(
        doc,
        "14.2 Handover. The Employee shall provide a full and orderly handover of "
        "pending work, including documentation of ongoing projects, client relationships, "
        "source code repositories, access credentials, knowledge articles and other "
        "matters, to such person(s) as the Company may designate, and shall cooperate with "
        "any knowledge transfer required by the Company.",
    )
    add_para(
        doc,
        "14.3 Exit interview. The Employee shall attend an exit interview with the "
        "Company's human resources function, at which the Employee may raise any grievance "
        "or concern.",
    )
    add_para(
        doc,
        "14.4 Withdrawals and transfers. The Company shall, on receipt of requisite "
        "forms, effect the transfer or withdrawal of the Employee's provident fund "
        "account, disbursement of gratuity (if payable), issuance of Form 16 and "
        "experience / relieving letter, in accordance with Applicable Law and the "
        "Company's full-and-final settlement timelines.",
    )
    add_para(
        doc,
        "14.5 Post-termination cooperation. The Employee shall, at the Company's "
        "reasonable request and at the Company's expense, provide such assistance and "
        "testimony as may be required in connection with any investigation, arbitration, "
        "litigation or regulatory proceeding concerning matters within the Employee's "
        "knowledge during employment.",
    )

    # ------------------------------------------------------------------
    # 15. Background verification and representations
    # ------------------------------------------------------------------
    add_h1(doc, "15. BACKGROUND VERIFICATION AND REPRESENTATIONS")
    add_para(
        doc,
        "15.1 The Employee's continued employment is conditional upon the Company's "
        "satisfaction with the results of background verification checks (including "
        "education, prior employment, address, identity, criminal record and, where "
        "role-relevant, credit and reference checks), which the Company may carry out "
        "directly or through a reputed third-party service provider.",
    )
    add_para(
        doc,
        "15.2 The Employee hereby consents, for the purposes of the DPDP Act, to the "
        "processing of his/her Personal Data for such verification and to the disclosure "
        "thereof by relevant third parties (including prior employers, educational "
        "institutions, statutory authorities and credit bureaus) to the Company or its "
        "verification agency.",
    )
    add_para(
        doc,
        "15.3 The Employee represents and warrants that: (a) all information furnished by "
        "him/her to the Company (including on the resume and application form) is true, "
        "accurate, complete and not misleading; (b) he/she is legally entitled to work in "
        "India and is not subject to any restriction, injunction, restrictive covenant or "
        "subsisting obligation (including to any former employer) that would prevent or "
        "restrict the performance of his/her duties under this Agreement; (c) he/she has "
        "disclosed in Annexure D any pending or threatened claim, dispute or proceeding "
        "that may affect his/her employment; and (d) he/she has not been convicted of any "
        "offence involving moral turpitude.",
    )

    # ------------------------------------------------------------------
    # 16. Disciplinary procedure and grievance redressal
    # ------------------------------------------------------------------
    add_h1(doc, "16. DISCIPLINARY PROCEDURE AND GRIEVANCE REDRESSAL")
    add_para(
        doc,
        "16.1 Disciplinary procedure. Allegations of misconduct against the Employee "
        "shall be dealt with in accordance with the Company's disciplinary Policy and, "
        "where applicable, any certified Standing Orders under the Industrial Employment "
        "(Standing Orders) Act, 1946 or the Industrial Relations Code, 2020. The "
        "procedure shall include (i) issuance of a show-cause notice or charge-sheet "
        "setting out the allegations; (ii) an opportunity to submit a written reply; "
        "(iii) a domestic enquiry, where warranted, conducted by a person not directly "
        "involved in the matter, consistent with the principles of natural justice; and "
        "(iv) a written decision communicated to the Employee.",
    )
    add_para(
        doc,
        "16.2 Grievance redressal. The Employee may raise any grievance relating to "
        "his/her employment first with his/her reporting manager and, if not resolved, "
        "with the human resources function, in accordance with the Company's grievance "
        "redressal Policy. Nothing in this Clause shall affect the Employee's right to "
        "approach statutory forums under Applicable Law (including the POSH Internal "
        "Committee, labour authorities, or courts and tribunals of competent jurisdiction).",
    )
    add_para(
        doc,
        "16.3 POSH Internal Committee. Complaints of sexual harassment shall be dealt "
        "with by the Internal Committee once constituted under the POSH Act (on the "
        "Company reaching the statutory threshold of ten (10) employees or earlier), "
        "in accordance with that Act and the Company's POSH Policy. Until such "
        "Internal Committee is in place, complaints may be raised with the Founder & "
        "CEO or with the Local Committee constituted by the District Officer.",
    )

    # ------------------------------------------------------------------
    # 17. Indemnity and set-off
    # ------------------------------------------------------------------
    add_h1(doc, "17. INDEMNITY AND SET-OFF")
    add_para(
        doc,
        "17.1 The Employee shall indemnify and hold harmless the Company and its "
        "Affiliates from and against any direct loss, damage, cost, penalty or liability "
        "suffered by the Company as a direct result of (a) the Employee's gross negligence, "
        "fraud or wilful misconduct, (b) breach by the Employee of the Employee NDA, the "
        "IP Deed, or Clauses 9, 10, 11 or 12 of this Agreement, or (c) breach by the "
        "Employee of Applicable Law in the course of employment. The Employee's liability "
        "under this Clause shall not extend to consequential, indirect or punitive loss, "
        "save where mandated by Applicable Law.",
    )
    add_para(
        doc,
        "17.2 The Company shall be entitled to set off any sums due from the Employee to "
        "the Company (including any recovery under this Clause 17, notice-period short-"
        "fall, training-bond recovery or clawback) against any sums payable by the "
        "Company to the Employee, including in full and final settlement, to the extent "
        "permitted under the Payment of Wages Act, 1936 and the Code on Wages, 2019.",
    )

    # ------------------------------------------------------------------
    # 18. Dispute resolution
    # ------------------------------------------------------------------
    add_h1(doc, "18. DISPUTE RESOLUTION")
    add_para(
        doc,
        "18.1 Statutory forums. Any dispute arising out of or relating to the employment "
        "of the Employee that is within the exclusive jurisdiction of statutory forums "
        "under Indian labour or employment legislation (including the S&CE Act, the POSH "
        "Act, the Industrial Disputes Act, 1947 / Industrial Relations Code, 2020, the "
        "Employees' Provident Funds and Miscellaneous Provisions Act, 1952 / Code on "
        "Social Security, 2020, and the Payment of Gratuity Act, 1972) shall be pursued "
        "before such statutory forums, and nothing in this Agreement shall oust the "
        "jurisdiction of such forums.",
    )
    add_para(
        doc,
        "18.2 Contractual disputes. Any other dispute, difference or claim arising out "
        "of or in connection with this Agreement, including its existence, validity, "
        "interpretation, performance, breach or termination (a \"Dispute\"), shall first be "
        "sought to be resolved amicably between the Parties within thirty (30) days of "
        "written notice of the Dispute. Failing such resolution, the Dispute shall be "
        "referred to and finally resolved by arbitration under the Arbitration and "
        "Conciliation Act, 1996, by a sole arbitrator to be mutually appointed by the "
        "Parties (or, failing agreement, appointed in accordance with that Act). The seat "
        "and venue of arbitration shall be New Delhi, Delhi, India. The language of "
        "arbitration shall be English. The award shall be final and binding on the Parties.",
    )
    add_para(
        doc,
        "18.3 Interim relief. Nothing in this Clause shall prevent either Party from "
        "seeking urgent interim or injunctive relief from a court of competent "
        "jurisdiction at New Delhi, in aid of arbitration or to protect Confidential "
        "Information or Intellectual Property rights.",
    )

    # ------------------------------------------------------------------
    # 19. Governing law
    # ------------------------------------------------------------------
    add_h1(doc, "19. GOVERNING LAW AND JURISDICTION")
    add_para(
        doc,
        "This Agreement shall be governed by and construed in accordance with the laws of "
        "India. Subject to Clause 18, the courts at New Delhi, Delhi shall have "
        "exclusive jurisdiction over any matter arising out of or relating to this "
        "Agreement.",
    )

    # ------------------------------------------------------------------
    # 20. Notices
    # ------------------------------------------------------------------
    add_h1(doc, "20. NOTICES")
    add_para(
        doc,
        "20.1 Any notice, consent or other communication given under this Agreement shall "
        "be in writing and shall be delivered by hand, by registered post with "
        "acknowledgement due, by reputable courier, or by email to the following "
        "addresses (or such other address as a Party may notify in writing):",
    )
    add_para(
        doc,
        "(a) To the Company: [COMPANY NOTICE ADDRESS]; Attn.: [HR HEAD / COMPANY "
        "SECRETARY]; Email: [COMPANY EMAIL].",
    )
    add_para(
        doc,
        "(b) To the Employee: at the address and email set out in the preamble and "
        "records of the Company, as updated by the Employee from time to time.",
    )
    add_para(
        doc,
        "20.2 Notices shall be deemed received (i) upon delivery, if by hand or courier; "
        "(ii) seven (7) days after dispatch, if by registered post; and (iii) upon "
        "confirmed transmission, if by email on a Working Day before 5 p.m. IST, failing "
        "which on the next Working Day.",
    )

    # ------------------------------------------------------------------
    # 21. Miscellaneous
    # ------------------------------------------------------------------
    add_h1(doc, "21. MISCELLANEOUS")

    add_h2(doc, "21.1 Entire agreement")
    add_para(
        doc,
        "This Agreement, together with the Offer Letter, the Employee NDA, the IP Deed, "
        "the Annexures, the ESOP grant letter (if applicable), and the Policies referred "
        "to in Annexure C, constitutes the entire agreement between the Parties with "
        "respect to the Employee's employment and supersedes all prior discussions, "
        "understandings and agreements on the subject matter. In the event of any "
        "inconsistency between this Agreement and the Offer Letter, this Agreement shall "
        "prevail.",
    )

    add_h2(doc, "21.2 Amendment")
    add_para(
        doc,
        "No amendment to this Agreement shall be valid unless made in writing and signed "
        "by or on behalf of both Parties. The Company may, however, amend the Policies "
        "from time to time on reasonable prior notice to the Employee.",
    )

    add_h2(doc, "21.3 Severability")
    add_para(
        doc,
        "If any provision of this Agreement is held by any court or tribunal of competent "
        "jurisdiction to be invalid, illegal or unenforceable (whether in whole or in "
        "part), the validity, legality and enforceability of the remaining provisions "
        "shall not be affected, and such provision shall be modified or read down to the "
        "minimum extent necessary to render it valid, legal and enforceable, consistent "
        "with the Parties' original intent.",
    )

    add_h2(doc, "21.4 Waiver")
    add_para(
        doc,
        "No failure or delay by a Party to exercise any right under this Agreement shall "
        "operate as a waiver thereof; no waiver shall be effective unless made in writing "
        "and signed by the waiving Party.",
    )

    add_h2(doc, "21.5 Assignment")
    add_para(
        doc,
        "The Employee's obligations are personal and shall not be assigned. The Company "
        "may assign this Agreement to any Affiliate or successor in title without the "
        "Employee's consent, provided that such assignment does not adversely affect the "
        "Employee's material rights.",
    )

    add_h2(doc, "21.6 Counterparts and electronic execution")
    add_para(
        doc,
        "This Agreement may be executed in counterparts, each of which when executed and "
        "delivered shall be deemed an original, and all counterparts together shall "
        "constitute one and the same instrument. The Parties agree that this Agreement "
        "may be executed by electronic means (including Aadhaar e-Sign or digital "
        "signatures) in accordance with the Information Technology Act, 2000, and such "
        "execution shall be valid, binding and admissible in evidence in accordance with "
        "Section 65B of the Indian Evidence Act, 1872 (as re-enacted under Section 63 of "
        "the Bharatiya Sakshya Adhiniyam, 2023).",
    )

    add_h2(doc, "21.7 Stamp duty")
    add_para(
        doc,
        "This Agreement shall be stamped in accordance with Article 5 of the Schedule to "
        "the Indian Stamp Act, 1899 (as applicable to the NCT of Delhi) (agreement not otherwise provided for). The "
        "Employee acknowledges that the Company bears the cost of stamp duty and that the "
        "Agreement shall be executed on duly stamped paper or e-stamp generated through "
        "the Stock Holding Corporation of India Limited (SHCIL). Any deficiency in stamp "
        "duty shall be made good in accordance with the Indian Stamp Act, 1899 (as applicable to the NCT of Delhi).",
    )

    add_h2(doc, "21.8 Survival")
    add_para(
        doc,
        "Clauses 9 (Confidentiality and Intellectual Property), 10 (Restrictive "
        "Covenants), 12 (Data Protection and Acceptable Use), 14 (Exit Obligations), 17 "
        "(Indemnity and Set-off), 18 (Dispute Resolution), 19 (Governing Law) and 20 "
        "(Notices), together with any other provision which by its nature is intended to "
        "survive, shall survive the termination of this Agreement.",
    )

    # ------------------------------------------------------------------
    # In-witness-whereof
    # ------------------------------------------------------------------
    add_h1(doc, "IN WITNESS WHEREOF")
    add_para(
        doc,
        "The Parties have executed this Employment Agreement on the day, month and year "
        "first above written at New Delhi, Delhi, India.",
    )

    # Signature table
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.autofit = True
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    employee_cell = sig_table.rows[0].cells[0]
    company_cell = sig_table.rows[0].cells[1]

    def fill_sig_cell(cell, heading_lines, block_lines):
        cell.text = ""
        for i, line in enumerate(heading_lines):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(line)
            r.bold = True
            r.font.name = BASE_FONT
            r.font.size = BODY_SIZE
        cell.add_paragraph("")
        cell.add_paragraph("_______________________________")
        for line in block_lines:
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(line)
            r.font.name = BASE_FONT
            r.font.size = BODY_SIZE

    fill_sig_cell(
        employee_cell,
        ["FOR AND ON BEHALF OF THE EMPLOYEE"],
        [
            "Name: [EMPLOYEE FULL NAME]",
            "PAN: [PAN NO.]",
            "Date: ______________________",
            "Place: New Delhi, Delhi",
        ],
    )

    fill_sig_cell(
        company_cell,
        ["FOR AND ON BEHALF OF ANANTASUTRA"],
        [
            "Name: Mr. Himanshu Mishra",
            "Designation: Founder & CEO",
            "Email: contact@anantasutra.com",
            "Date: [DATE]",
            "Place: Delhi",
        ],
    )

    doc.add_paragraph("")
    add_para(doc, "WITNESSES:", bold=True)
    wit_table = doc.add_table(rows=1, cols=2)
    wit_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    w1 = wit_table.rows[0].cells[0]
    w2 = wit_table.rows[0].cells[1]

    def fill_witness(cell, n):
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(f"Witness {n}")
        r.bold = True
        r.font.name = BASE_FONT
        r.font.size = BODY_SIZE
        cell.add_paragraph("")
        cell.add_paragraph("_______________________________")
        for line in ["Name: ______________________",
                     "Address: ____________________",
                     "Signature: ___________________"]:
            p = cell.add_paragraph()
            r = p.add_run(line)
            r.font.name = BASE_FONT
            r.font.size = BODY_SIZE

    fill_witness(w1, 1)
    fill_witness(w2, 2)

    # ------------------------------------------------------------------
    # ANNEXURES
    # ------------------------------------------------------------------
    doc.add_page_break()
    add_h1(doc, "ANNEXURE A – COMPENSATION STRUCTURE")
    add_para(
        doc,
        "The Employee's Cost to Company (CTC) is structured as set out in the table "
        "below. All figures are in Indian Rupees (INR) and are indicative; actual "
        "deductions and net pay shall depend on Applicable Law, statutory deductions, "
        "and the Employee's declarations.",
    )

    comp_data = [
        ("Component", "Monthly (INR)", "Annual (INR)", "Notes"),
        ("Basic Salary", "[●]", "[●]", "40–50% of CTC"),
        ("House Rent Allowance (HRA)", "[●]", "[●]", "40% of Basic (non-metro) / 50% (metro)"),
        ("Special Allowance", "[●]", "[●]", "Balancing component"),
        ("Leave Travel Allowance (LTA)", "[●]", "[●]", "Paid annually, subject to claim"),
        ("Statutory Bonus / Ex-gratia", "[●]", "[●]", "Payment of Bonus Act, 1965 / Code on Wages"),
        ("Employer PF Contribution", "[●]", "[●]", "12% of Basic (subject to wage ceiling as per Policy)"),
        ("Gratuity (accrual)", "[●]", "[●]", "Payment of Gratuity Act, 1972"),
        ("Performance-linked Variable Pay", "—", "[●]", "Subject to Clause 7.3"),
        ("Gross Annual CTC", "—", "[●]", "Sum of the above"),
    ]

    comp_table = doc.add_table(rows=len(comp_data), cols=4)
    comp_table.style = 'Light Grid Accent 1'
    comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(comp_data):
        for j, value in enumerate(row_data):
            cell = comp_table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            run.font.name = BASE_FONT
            run.font.size = Pt(10)
            if i == 0:
                run.bold = True
                set_cell_shading(cell, "D9E2F3")

    add_para(doc, "")
    add_para(
        doc,
        "Statutory deductions (including employee PF contribution, professional tax, "
        "income-tax TDS under Section 192 of the Income-tax Act, 1961, and ESI where "
        "applicable) shall be deducted from the Employee's monthly salary at source. For "
        "the purposes of any pay-in-lieu-of-notice calculation under Clause 13.2, the "
        "relevant wage component shall be Basic Salary (and Dearness Allowance, if any).",
    )

    # Annexure B
    doc.add_page_break()
    add_h1(doc, "ANNEXURE B – JOB DESCRIPTION")
    add_para(doc, "Designation: [DESIGNATION]", bold=True)
    add_para(doc, "Department / Function: [DEPARTMENT]", bold=True)
    add_para(doc, "Reporting Manager: [REPORTING MANAGER]", bold=True)
    add_para(doc, "Grade / Level: [LEVEL]", bold=True)
    add_para(doc, "Location: [LOCATION], New Delhi, Delhi (with hybrid / remote working per Clause 5)", bold=True)
    add_para(doc, "")
    add_para(doc, "Key Responsibilities:", bold=True)
    for item in [
        "(a) [Describe primary responsibility 1];",
        "(b) [Describe primary responsibility 2];",
        "(c) [Describe primary responsibility 3];",
        "(d) participate in, and contribute to, project delivery, client communications, internal reviews and Company initiatives as reasonably required;",
        "(e) ensure compliance with all applicable Policies, Applicable Law, and contractual obligations owed by the Company to its clients and counterparties; and",
        "(f) any other duties reasonably assigned by the Reporting Manager consistent with the Employee's role and seniority.",
    ]:
        add_list_item(doc, item, level=0)

    add_para(doc, "")
    add_para(doc, "Key Performance Indicators (indicative):", bold=True)
    for item in [
        "(a) achievement of quarterly and annual objectives communicated in writing;",
        "(b) quality, timeliness and client-satisfaction metrics for deliverables;",
        "(c) adherence to Company Policies, Code of Conduct and ethical standards; and",
        "(d) contribution to team capability-building and knowledge management.",
    ]:
        add_list_item(doc, item, level=0)

    # Annexure C
    doc.add_page_break()
    add_h1(doc, "ANNEXURE C – LIST OF COMPANY POLICIES")
    add_para(
        doc,
        "The Employee agrees to comply with, and acknowledges having received or been "
        "given access to, the following Company Policies (as amended from time to time). "
        "Copies are maintained on the Company's human-resources portal and are available "
        "on request.",
    )
    policies = [
        "Code of Conduct and Business Ethics",
        "Policy on Prevention of Sexual Harassment at Workplace (POSH)",
        "Anti-Bribery and Anti-Corruption Policy",
        "Whistleblower / Protected Disclosure Policy",
        "Information Technology / Acceptable Use Policy",
        "Information Security Policy",
        "DPDP Policy and Employee Privacy Notice (under the Digital Personal Data Protection Act, 2023)",
        "Leave Policy (consistent with the Delhi Shops and Establishments Act, 1954 and the Maternity Benefit Act, 1961)",
        "Travel and Reimbursement Policy",
        "Remote / Hybrid Working Policy",
        "Performance Management Policy (including PIP procedure)",
        "Employee Grievance Redressal Policy",
        "Disciplinary Policy",
        "Diversity, Equity and Inclusion Policy (including compliance with the Rights of Persons with Disabilities Act, 2016 and the Transgender Persons (Protection of Rights) Act, 2019)",
        "Insider Trading Policy (once applicable under the SEBI framework)",
        "ESOP Plan and grant procedure (if applicable)",
    ]
    for i, pol in enumerate(policies, 1):
        add_list_item(doc, f"{i}. {pol}", level=0)

    # Annexure D
    doc.add_page_break()
    add_h1(doc, "ANNEXURE D – PRIOR IP AND PRIOR-EMPLOYER DISCLOSURE")
    add_para(
        doc,
        "The Employee hereby discloses the following for the purposes of Clause 9.4 and "
        "Clause 15.3 of the Agreement. Where no disclosure is made, the Employee confirms "
        "that there is no prior Intellectual Property, subsisting obligation or prior-"
        "employer restriction of the nature described.",
    )

    disclose_data = [
        ("S.No.", "Category", "Particulars"),
        ("1.", "Pre-existing Intellectual Property owned by the Employee relevant to the Business (title, date, rights holder, brief description)", "[●]"),
        ("2.", "Subsisting confidentiality / non-solicitation / other restrictive obligations owed to any prior employer or third party", "[●]"),
        ("3.", "Current directorships, partnerships, consultancies or other engagements (including moonlighting) to be disclosed under Clause 4.3 / 4.4", "[●]"),
        ("4.", "Pending or threatened civil, criminal, tax or regulatory proceedings, if any, against the Employee", "[●]"),
        ("5.", "Any other matter relevant to Clause 15.3", "[●]"),
    ]

    dtable = doc.add_table(rows=len(disclose_data), cols=3)
    dtable.style = 'Light Grid Accent 1'
    dtable.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(disclose_data):
        for j, val in enumerate(row):
            c = dtable.rows[i].cells[j]
            c.text = ""
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.name = BASE_FONT
            run.font.size = Pt(10)
            if i == 0:
                run.bold = True
                set_cell_shading(c, "D9E2F3")

    add_para(doc, "")
    add_para(
        doc,
        "Declaration by the Employee: I confirm that the information set out in this "
        "Annexure D is true, accurate and complete to the best of my knowledge and belief, "
        "and I understand that any omission or mis-statement may be treated as a material "
        "breach of this Agreement and/or a Cause Event under Clause 13.3.",
    )
    add_para(doc, "")
    add_para(doc, "Signature of Employee: ____________________________")
    add_para(doc, "Name: [EMPLOYEE FULL NAME]")
    add_para(doc, "Date: ______________________")
    add_para(doc, "Place: New Delhi, Delhi")

    # End marker
    add_para(doc, "")
    add_para(doc, "— End of Employment Agreement —", italic=True, justify=False)

    return doc


def main():
    out_dir = os.path.dirname(OUTPUT_PATH)
    os.makedirs(out_dir, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT_PATH)
    size_bytes = os.path.getsize(OUTPUT_PATH)
    size_kb = size_bytes / 1024
    print(f"File written: {OUTPUT_PATH}")
    print(f"Size: {size_kb:.2f} KB ({size_bytes} bytes)")


if __name__ == "__main__":
    main()
