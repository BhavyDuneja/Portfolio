"""
build_09_referral_partner_agreement.py

Generates a SHORT-FORM, PLAIN-LANGUAGE, PRO-ANANTASUTRA Referral Partner
(Profit-Share) Agreement for "collaborators" who bring clients to
AnantaSutra in exchange for an agreed share of profit.

Output: legal-documents/09_Referral_Partner_Agreement.docx

This is a reusable template: the collaborator's details and the agreed
profit-share percentage per client are left as fill-in fields / recorded in
the Referral Schedule (Annexure A).

Designed to protect AnantaSutra:
- Commission is a % of NET PROFIT (clearly defined), not of revenue.
- Paid only on amounts ACTUALLY RECEIVED from the client (no bad debt).
- No written % recorded => no commission (forces the negotiation in writing).
- Partner is an independent referrer with NO authority to bind AnantaSutra.
- AnantaSutra has sole discretion to accept/reject clients, set pricing/scope.
- Lead-registration mechanic + carve-outs prevent "who brought whom" disputes.
- Non-circumvention / non-interference with referred clients.
- Commission period capped (default 12 months) to avoid perpetual liability.

Indian-law anchors: Indian Contract Act 1872, Indian Partnership Act 1932
(no partnership created), Income-tax Act 1961 (TDS on commission),
Prevention of Corruption Act 1988 (no bribes/kickbacks), Arbitration &
Conciliation Act 1996, Specific Relief Act 1963.

Uses python-docx v1.2.0.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# -------------------------------------------------------------------
# Constants
# -------------------------------------------------------------------
OUTPUT_PATH = r"c:/Users/duneja8515/Desktop/bhavya/prsnl/sutranet/legal-documents/09_Referral_Partner_Agreement.docx"
LOGO_PATH = r"c:/Users/duneja8515/Desktop/bhavya/prsnl/sutranet/public/Favicon/logo-nobg.png"

BASE_FONT = "Calibri"
BODY_SIZE = Pt(11)
HEADING1_SIZE = Pt(13)
TITLE_SIZE = Pt(16)

# Company (AnantaSutra) — hardcoded
SP_DESCRIPTION = (
    "AnantaSutra, a business concern operated and represented by its "
    "Chief Executive Officer, Mr. Himanshu Mishra, having its principal "
    "place of business at Delhi, India"
)
SP_CONTACT = "contact@anantasutra.com"
SP_SIGNATORY_NAME = "Mr. Himanshu Mishra"
SP_SIGNATORY_DESIGNATION = "Chief Executive Officer (CEO)"

# Referral Partner — template placeholders
PARTNER_DESCRIPTION = (
    "[PARTNER FULL NAME], [son/daughter/proprietor of __________], an "
    "individual/entity having address at [PARTNER ADDRESS] and email "
    "[PARTNER EMAIL]"
)


# -------------------------------------------------------------------
# Helpers
# -------------------------------------------------------------------
def add_brand_header(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=Inches(1.2))
    except Exception:
        pass
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("AnantaSutra — अनन्तसूत्र")
    r2.italic = True
    r2.font.size = Pt(10)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("https://anantasutra.com  •  contact@anantasutra.com")
    r3.italic = True
    r3.font.size = Pt(9)


def _simple_field(paragraph, instr):
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = instr
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_end)


def configure_styles(doc):
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = BASE_FONT
    normal.font.size = BODY_SIZE
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    h1 = styles['Heading 1']
    h1.font.name = BASE_FONT
    h1.font.size = HEADING1_SIZE
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(4)
    h1.paragraph_format.keep_with_next = True


def set_page_layout(doc):
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Confidential — AnantaSutra    |    Page ")
        run.font.name = BASE_FONT
        run.font.size = Pt(9)
        _simple_field(p, "PAGE")
        mid = p.add_run(" of ")
        mid.font.name = BASE_FONT
        mid.font.size = Pt(9)
        _simple_field(p, "NUMPAGES")


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = TITLE_SIZE
    run.font.name = BASE_FONT
    p.paragraph_format.space_after = Pt(4)
    return p


def add_centered(doc, text, italic=False, size=None, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.name = BASE_FONT
    run.font.size = size or BODY_SIZE
    p.paragraph_format.space_after = Pt(6)
    return p


def add_h1(doc, text):
    return doc.add_paragraph(text, style='Heading 1')


def add_para(doc, text, bold=False, italic=False, justify=True):
    p = doc.add_paragraph()
    p.alignment = (WD_ALIGN_PARAGRAPH.JUSTIFY if justify
                   else WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = BASE_FONT
    run.font.size = BODY_SIZE
    return p


def add_sub(doc, lead, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.7)
    lr = p.add_run(lead + "  ")
    lr.bold = True
    lr.font.name = BASE_FONT
    lr.font.size = BODY_SIZE
    r = p.add_run(text)
    r.font.name = BASE_FONT
    r.font.size = BODY_SIZE
    return p


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "000000")
        tcBorders.append(border)
    tcPr.append(tcBorders)


# -------------------------------------------------------------------
# Document construction
# -------------------------------------------------------------------
def build_document():
    doc = Document()
    add_brand_header(doc)
    configure_styles(doc)
    set_page_layout(doc)
    add_footer(doc)

    # ---------------- Title ----------------
    add_title(doc, "REFERRAL PARTNER AGREEMENT")
    add_centered(
        doc,
        "(Client Introduction and Profit-Share Arrangement)",
        italic=True,
    )

    # ---------------- Date / execution ----------------
    add_para(
        doc,
        "This Referral Partner Agreement (the \"Agreement\") is made and "
        "executed at Delhi, India on this ______ day of ____________, 20____ "
        "(the \"Effective Date\"),",
    )

    # ---------------- Parties ----------------
    add_h1(doc, "BETWEEN")
    add_para(
        doc,
        SP_DESCRIPTION
        + " (hereinafter referred to as \"AnantaSutra\" or the \"Company\", "
        "which expression shall, unless repugnant to the context, include its "
        "successors and permitted assigns), of the ONE PART;",
    )
    add_h1(doc, "AND")
    add_para(
        doc,
        PARTNER_DESCRIPTION
        + " (hereinafter referred to as the \"Referral Partner\" or "
        "\"Partner\", which expression shall, unless repugnant to the "
        "context, include his/her/its heirs, successors, and permitted "
        "assigns), of the OTHER PART.",
    )
    add_para(
        doc,
        "AnantaSutra and the Referral Partner are each referred to as a "
        "\"Party\" and together as the \"Parties\".",
    )

    # ---------------- Background ----------------
    add_h1(doc, "BACKGROUND")
    add_para(
        doc,
        "A. AnantaSutra provides software development, website, technical "
        "support, and creative design services to its clients.",
    )
    add_para(
        doc,
        "B. The Referral Partner is willing to introduce potential clients "
        "to AnantaSutra, and AnantaSutra is willing to pay the Referral "
        "Partner an agreed share of the profit it actually earns from such "
        "introduced clients, on the terms set out below.",
    )
    add_para(
        doc,
        "NOW THEREFORE the Parties agree as follows:",
        bold=True,
    )

    # ---------------- 1. Definitions ----------------
    add_h1(doc, "1. Definitions")
    add_sub(
        doc,
        "1.1  Referral.",
        "An introduction by the Referral Partner of a potential client to "
        "AnantaSutra, made and registered in accordance with Clause 3.",
    )
    add_sub(
        doc,
        "1.2  Referred Client.",
        "A potential client introduced by the Referral Partner that "
        "AnantaSutra accepts in writing as a Qualifying Referral under Clause "
        "3, and that goes on to enter into a paid engagement with "
        "AnantaSutra.",
    )
    add_sub(
        doc,
        "1.3  Qualifying Referral.",
        "A Referral that meets all the conditions in Clause 3 and is not "
        "excluded under Clause 7.",
    )
    add_sub(
        doc,
        "1.4  Net Profit.",
        "In respect of a Referred Client, the amounts actually received and "
        "cleared by AnantaSutra from that client for the relevant work, LESS "
        "(a) all applicable taxes; (b) the direct cost of delivering that "
        "work (including the fees, salaries, or charges of the engineers, "
        "designers, and other personnel deployed on that client, third-party "
        "and asset costs, hosting/domain/licence costs, and payment-gateway "
        "or bank charges); and (c) any refund, credit, charge-back, or "
        "write-off relating to that client. AnantaSutra shall calculate Net "
        "Profit in good faith, and its calculation shall be final and binding "
        "save for manifest error.",
    )
    add_sub(
        doc,
        "1.5  Commission.",
        "The percentage share of Net Profit agreed for a particular Referred "
        "Client and recorded in the Referral Schedule (Annexure A) or in a "
        "separate writing signed by both Parties.",
    )
    add_sub(
        doc,
        "1.6  Commission Period.",
        "The period during which Commission is payable for a Referred Client, "
        "being twelve (12) months from the date AnantaSutra accepts the "
        "Referral, unless a different period is recorded for that client in "
        "the Referral Schedule.",
    )
    add_sub(
        doc,
        "1.7  Referral Schedule.",
        "Annexure A to this Agreement, in which each Referred Client, the "
        "agreed Commission percentage, and any client-specific terms are "
        "recorded.",
    )

    # ---------------- 2. Appointment and Scope ----------------
    add_h1(doc, "2. Appointment and Scope")
    add_sub(
        doc,
        "2.1",
        "AnantaSutra appoints the Referral Partner as a non-exclusive "
        "referrer to introduce potential clients to AnantaSutra. The "
        "Referral Partner accepts the appointment on the terms of this "
        "Agreement.",
    )
    add_sub(
        doc,
        "2.2",
        "The appointment is non-exclusive. AnantaSutra may appoint other "
        "referral partners, may obtain clients through its own or any other "
        "channel, and is under no obligation to use the Referral Partner for "
        "any client or business.",
    )
    add_sub(
        doc,
        "2.3",
        "The Referral Partner's role is limited to making introductions. The "
        "Referral Partner shall not negotiate, quote prices, make promises, "
        "sign documents, or otherwise commit or represent AnantaSutra, and "
        "has no authority to bind AnantaSutra in any way.",
    )

    # ---------------- 3. How to Refer (Lead Registration) ----------------
    add_h1(doc, "3. How to Refer a Client")
    add_sub(
        doc,
        "3.1",
        "To register a Referral, the Referral Partner shall send AnantaSutra "
        "a written notice (email is sufficient) giving the prospective "
        "client's name, contact details, and the nature of the opportunity, "
        "before any contact between AnantaSutra and that client.",
    )
    add_sub(
        doc,
        "3.2",
        "Within seven (7) days, AnantaSutra shall confirm in writing whether "
        "it accepts the Referral as a Qualifying Referral. AnantaSutra may "
        "decline any Referral at its sole discretion, including where the "
        "client is already known to AnantaSutra, is already in its pipeline, "
        "or has already been introduced by another person.",
    )
    add_sub(
        doc,
        "3.3",
        "Only a Referral expressly accepted in writing by AnantaSutra "
        "qualifies for Commission. AnantaSutra's records shall be conclusive "
        "as to whether, when, and by whom a client was first introduced.",
    )

    # ---------------- 4. AnantaSutra's Discretion ----------------
    add_h1(doc, "4. AnantaSutra's Discretion")
    add_para(
        doc,
        "AnantaSutra has sole discretion over whether to pursue or accept any "
        "Referred Client, the scope, pricing, and terms of any engagement, "
        "and whether and for how long to continue that engagement. "
        "AnantaSutra is not obliged to enter into, continue, or renew any "
        "engagement, and the Referral Partner shall have no claim if it does "
        "not.",
    )

    # ---------------- 5. Commission (Profit Share) ----------------
    add_h1(doc, "5. Commission (Profit Share)")
    add_sub(
        doc,
        "5.1",
        "For each Referred Client, AnantaSutra shall pay the Referral Partner "
        "the agreed Commission percentage of the Net Profit actually earned "
        "from that client during the Commission Period.",
    )
    add_sub(
        doc,
        "5.2  Percentage agreed per client.",
        "The Commission percentage is negotiated and agreed separately for "
        "each Referred Client and recorded in the Referral Schedule (or a "
        "separate signed writing). If no percentage is recorded in writing "
        "before AnantaSutra begins work for that client, no Commission is "
        "payable for that client.",
    )
    add_sub(
        doc,
        "5.3  Profit, not revenue.",
        "Commission is calculated on Net Profit as defined in Clause 1.4, "
        "and only on amounts actually received and cleared from the client. "
        "No Commission is payable on amounts invoiced but not received, on "
        "refunded or disputed amounts, or where a client engagement runs at "
        "no profit or a loss.",
    )

    # ---------------- 6. Payment of Commission ----------------
    add_h1(doc, "6. Payment of Commission")
    add_sub(
        doc,
        "6.1",
        "AnantaSutra shall calculate and pay Commission monthly in arrears, "
        "within fifteen (15) days after the end of each calendar month, in "
        "respect of cleared amounts received from Referred Clients during "
        "that month, together with a simple statement showing the basis of "
        "calculation.",
    )
    add_sub(
        doc,
        "6.2  Taxes.",
        "Commission is inclusive of all taxes for which the Referral Partner "
        "is liable. AnantaSutra may deduct tax at source (TDS) as required "
        "under the Income-tax Act, 1961 (including Section 194H), and the "
        "Referral Partner shall provide his/her/its PAN and any details "
        "AnantaSutra reasonably requires.",
    )
    add_sub(
        doc,
        "6.3",
        "The Referral Partner shall bear his/her/its own costs of generating "
        "referrals. AnantaSutra shall not reimburse any expense unless it has "
        "approved that expense in writing in advance.",
    )

    # ---------------- 7. When No Commission is Payable ----------------
    add_h1(doc, "7. When No Commission is Payable")
    add_para(
        doc,
        "No Commission is payable where: (a) the Referral was not registered "
        "and accepted under Clause 3; (b) the client was already known to or "
        "in discussion with AnantaSutra, or was first introduced by another "
        "person, as shown by AnantaSutra's records; (c) no Commission "
        "percentage was agreed in writing before work began (Clause 5.2); "
        "(d) the client does not pay, or amounts are refunded, charged back, "
        "disputed, or written off; (e) the relevant work falls outside the "
        "Commission Period; or (f) the Referral Partner is in breach of this "
        "Agreement.",
    )

    # ---------------- 8. Independent Contractor ----------------
    add_h1(doc, "8. Independent Contractor; No Partnership")
    add_para(
        doc,
        "The Referral Partner acts as an independent contractor. Nothing in "
        "this Agreement creates any employment, agency, partnership, or "
        "joint-venture relationship between the Parties, and no partnership "
        "arises under the Indian Partnership Act, 1932. The Referral Partner "
        "is not entitled to any salary, employee benefit, or fixed or minimum "
        "payment, and is responsible for his/her/its own taxes and "
        "compliances.",
    )

    # ---------------- 9. Non-Circumvention and Non-Interference ----------------
    add_h1(doc, "9. Non-Circumvention and Non-Interference")
    add_sub(
        doc,
        "9.1",
        "The Referral Partner shall not, directly or indirectly, during the "
        "Term and for twelve (12) months thereafter, (a) provide or arrange "
        "for any third party to provide to a Referred Client the same or "
        "similar services as AnantaSutra; (b) solicit, divert, or interfere "
        "with AnantaSutra's relationship with any Referred Client; or (c) "
        "contact a Referred Client about the engagement except as AnantaSutra "
        "authorises.",
    )
    add_sub(
        doc,
        "9.2",
        "AnantaSutra shall not, during the Commission Period, deliberately "
        "structure or route a Referred Client's work mainly to avoid paying "
        "Commission properly due under this Agreement.",
    )

    # ---------------- 10. Confidentiality ----------------
    add_h1(doc, "10. Confidentiality")
    add_para(
        doc,
        "Each Party shall keep confidential all non-public information of the "
        "other Party, and all information about Referred Clients, pricing, and "
        "terms, and shall use it only for the purposes of this Agreement. "
        "This obligation continues for three (3) years after termination. The "
        "Referral Partner shall not use AnantaSutra's name, logo, or "
        "materials, or make any representation about AnantaSutra, without "
        "AnantaSutra's prior written consent.",
    )

    # ---------------- 11. Partner's Conduct and Compliance ----------------
    add_h1(doc, "11. Partner's Conduct and Compliance")
    add_para(
        doc,
        "The Referral Partner shall: (a) act honestly and only make accurate "
        "statements about AnantaSutra, using only materials approved by "
        "AnantaSutra; (b) not give or accept any bribe, kickback, or unlawful "
        "inducement, in compliance with the Prevention of Corruption Act, "
        "1988; (c) not engage in any misleading, unlawful, or spam marketing; "
        "and (d) comply with all applicable laws in carrying out referral "
        "activities. Breach of this Clause is a material breach entitling "
        "AnantaSutra to terminate immediately and to withhold Commission.",
    )

    # ---------------- 12. Term and Termination ----------------
    add_h1(doc, "12. Term and Termination")
    add_sub(
        doc,
        "12.1",
        "This Agreement begins on the Effective Date and continues until "
        "terminated. Either Party may terminate for convenience by giving "
        "fifteen (15) days' prior written notice. AnantaSutra may terminate "
        "immediately if the Referral Partner breaches this Agreement.",
    )
    add_sub(
        doc,
        "12.2  Commission after termination.",
        "If this Agreement is terminated other than for the Referral "
        "Partner's breach, the Referral Partner shall continue to receive "
        "Commission for already-accepted Referred Clients for the remainder "
        "of their Commission Periods, subject to the terms of this Agreement. "
        "If this Agreement is terminated for the Referral Partner's breach, "
        "all entitlement to Commission ceases on termination.",
    )

    # ---------------- 13. Limitation of Liability ----------------
    add_h1(doc, "13. Limitation of Liability")
    add_para(
        doc,
        "To the maximum extent permitted by law, AnantaSutra shall not be "
        "liable to the Referral Partner for any indirect or consequential "
        "loss, or for any loss of profit or expected commission beyond "
        "Commission actually due and payable under this Agreement. "
        "AnantaSutra's total liability under this Agreement shall not exceed "
        "the total Commission properly payable to the Referral Partner in the "
        "three (3) months preceding the event giving rise to the claim.",
    )

    # ---------------- 14. Governing Law and Disputes ----------------
    add_h1(doc, "14. Governing Law and Dispute Resolution")
    add_para(
        doc,
        "This Agreement is governed by the laws of India. The Parties shall "
        "first try to resolve any dispute amicably; failing resolution within "
        "fifteen (15) days, the dispute shall be referred to and finally "
        "resolved by arbitration by a sole arbitrator mutually appointed by "
        "the Parties under the Arbitration and Conciliation Act, 1996. The "
        "seat and venue shall be Delhi and the language English. Subject to "
        "arbitration, the courts at Delhi shall have exclusive jurisdiction. "
        "Either Party may seek interim relief from a court under Section 9 of "
        "that Act and under the Specific Relief Act, 1963.",
    )

    # ---------------- 15. Notices ----------------
    add_h1(doc, "15. Notices")
    add_para(
        doc,
        "Notices shall be in writing and sent by email or recognised courier "
        "to the other Party's contact details set out in this Agreement or "
        "notified in writing. Notices to AnantaSutra shall be marked to Mr. "
        "Himanshu Mishra, CEO, at " + SP_CONTACT + ".",
    )

    # ---------------- 16. General ----------------
    add_h1(doc, "16. General")
    add_sub(
        doc,
        "16.1  Entire agreement.",
        "This Agreement (with its Annexure) is the entire agreement between "
        "the Parties on its subject matter and supersedes all prior "
        "discussions.",
    )
    add_sub(
        doc,
        "16.2  Amendment.",
        "No change is valid unless in writing and signed by both Parties.",
    )
    add_sub(
        doc,
        "16.3  Assignment.",
        "The Referral Partner shall not assign or transfer this Agreement or "
        "any right to Commission without AnantaSutra's prior written consent. "
        "AnantaSutra may assign to an affiliate or successor.",
    )
    add_sub(
        doc,
        "16.4  Severability and waiver.",
        "If any provision is held invalid, the rest continues in effect. A "
        "delay or failure to enforce a right is not a waiver of it.",
    )
    add_sub(
        doc,
        "16.5  Electronic execution.",
        "This Agreement may be signed in counterparts and executed "
        "electronically (including by scanned or digital signatures), which "
        "shall be valid and binding under the Information Technology Act, "
        "2000.",
    )

    # ---------------- Signature block ----------------
    doc.add_paragraph("")
    add_para(
        doc,
        "IN WITNESS WHEREOF, the Parties have signed this Agreement on the "
        "date first written above.",
        bold=True,
    )
    doc.add_paragraph("")

    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.autofit = True
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    company_cell = sig_table.rows[0].cells[0]
    partner_cell = sig_table.rows[0].cells[1]

    def fill_sig_cell(cell, heading_lines, block_lines):
        cell.text = ""
        for i, line in enumerate(heading_lines):
            p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(line)
            r.bold = True
            r.font.name = BASE_FONT
            r.font.size = BODY_SIZE
        cell.add_paragraph("")
        cell.add_paragraph("_______________________________")
        for line in block_lines:
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(line)
            r.font.name = BASE_FONT
            r.font.size = BODY_SIZE

    fill_sig_cell(
        company_cell,
        ["FOR AND ON BEHALF OF ANANTASUTRA"],
        [
            "Name: " + SP_SIGNATORY_NAME,
            "Designation: " + SP_SIGNATORY_DESIGNATION,
            "Email: " + SP_CONTACT,
            "Place: Delhi, India",
            "Date: ______________________",
        ],
    )
    fill_sig_cell(
        partner_cell,
        ["THE REFERRAL PARTNER"],
        [
            "Name: ______________________",
            "PAN: ______________________",
            "Signature: ______________________",
            "Place: ______________________",
            "Date: ______________________",
        ],
    )

    # =================================================================
    # ANNEXURE A — Referral Schedule
    # =================================================================
    doc.add_page_break()
    add_title(doc, "ANNEXURE A")
    add_centered(doc, "REFERRAL SCHEDULE", bold=True, size=Pt(13))
    add_para(
        doc,
        "Each Referred Client accepted by AnantaSutra is recorded below with "
        "the Commission percentage agreed for that client. A client not "
        "recorded here (or in a separate writing signed by both Parties) does "
        "not qualify for Commission.",
    )

    sched = doc.add_table(rows=1, cols=5)
    sched.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = (
        "Referred Client",
        "Date Accepted",
        "Commission % (of Net Profit)",
        "Commission Period",
        "Notes",
    )
    for idx, text in enumerate(headers):
        cell = sched.rows[0].cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.name = BASE_FONT
        r.font.size = Pt(10)
        set_cell_border(cell)
    # Example row + blank rows to fill
    rows_data = [
        ("[e.g. ABC Pvt Ltd]", "[DD/MM/YYYY]", "[e.g. 10%]", "[12 months / other]", "[any special terms]"),
    ]
    for _ in range(6):
        rows_data.append(("", "", "", "", ""))
    for row_data in rows_data:
        row = sched.add_row()
        for idx, text in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(text)
            r.italic = bool(text)
            r.font.name = BASE_FONT
            r.font.size = Pt(10)
            set_cell_border(cell)

    doc.add_paragraph("")
    add_para(
        doc,
        "Each entry in this Schedule shall be initialled or confirmed in "
        "writing (email is sufficient) by both Parties before AnantaSutra "
        "begins work for that client.",
        italic=True,
    )

    doc.add_paragraph("")
    # Mini illustration of how Commission is calculated
    add_h1(doc, "Illustration (example only)")
    add_para(
        doc,
        "If AnantaSutra receives Rs. 1,00,000 from a Referred Client for a "
        "project, and the direct cost of delivering it (personnel + assets + "
        "taxes + charges) is Rs. 70,000, the Net Profit is Rs. 30,000. At an "
        "agreed Commission of 10%, the Referral Partner receives Rs. 3,000 "
        "for that project. No Commission is payable on the cost portion or on "
        "any amount the client does not actually pay.",
        italic=True,
    )

    # Save
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    out = build_document()
    size_kb = os.path.getsize(out) / 1024.0
    print("REFERRAL PARTNER AGREEMENT COMPLETE")
    print("Output: " + out)
    print("Size: %.1f KB" % size_kb)
