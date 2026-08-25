"""
build_10_invoice.py

Generates a reusable, fillable INVOICE template for AnantaSutra.

Output: legal-documents/10_Invoice_Template.docx

Context / design choices:
  * AnantaSutra is NOT a registered company and is NOT registered for GST as
    at the date of issue. Accordingly this document is styled as a
    "Bill of Supply" (a plain invoice), NOT a "Tax Invoice", and no GST is
    charged. There is no GSTIN / CIN field.
  * The business is small and does not cross TDS-relevant thresholds on its
    side; the template simply notes that any TDS the CLIENT is required by law
    to deduct is the client's responsibility, with a certificate to be
    furnished. No tax is added or deducted on the AnantaSutra side.
  * The template is meant to be filled in per engagement: invoice number,
    dates, bill-to party, line items, amounts, and payment details all use
    blank placeholders.

Uses python-docx v1.2.0. Matches the house style of the other legal-documents
build scripts (Calibri, centered logo header, page numbers in footer).
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# -------------------------------------------------------------------
# Constants
# -------------------------------------------------------------------
OUTPUT_PATH = r"c:/Users/duneja8515/Desktop/bhavya/prsnl/sutranet/legal-documents/10_Invoice_Template.docx"
LOGO_PATH = r"c:/Users/duneja8515/Desktop/bhavya/prsnl/sutranet/public/Favicon/logo-nobg.png"

BASE_FONT = "Calibri"
BODY_SIZE = Pt(11)
HEADING1_SIZE = Pt(13)
TITLE_SIZE = Pt(20)

# Service Provider (AnantaSutra) — hardcoded, mirrors the other documents
SP_BRAND = "AnantaSutra"
SP_OWNER = "Bhavya Duneja"
SP_DESIGNATION = "Co-Founder & Owner"
SP_PLACE = "Delhi, India"
SP_CONTACT = "contact@anantasutra.com"
SP_WEBSITE = "https://anantasutra.com"

LIGHT_FILL = "F2F2F2"     # light grey for header / total rows
DARK_TEXT = RGBColor(0x00, 0x00, 0x00)


# -------------------------------------------------------------------
# Low-level helpers (shared house style)
# -------------------------------------------------------------------
def _simple_field(paragraph, instr):
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = instr
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_end)


def _shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)


def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "BFBFBF")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _fixed_layout(table):
    """Force a fixed table layout so explicit column widths are honoured by Word."""
    tblPr = table._tbl.tblPr
    # remove any existing layout element first
    for existing in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(existing)
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)


def _set_col_widths(table, widths):
    """widths: list of Cm() values, one per column.

    Printable width on A4 with the margins set in set_page_layout is ~16.4 cm,
    so the sum of widths must stay at or below that or the table overflows the
    page (and the last column gets clipped).
    """
    table.autofit = False
    table.allow_autofit = False
    _fixed_layout(table)
    for row in table.rows:
        for idx, w in enumerate(widths):
            row.cells[idx].width = w


def _cell_text(cell, text, bold=False, italic=False, size=None,
               align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = BASE_FONT
    r.font.size = size or BODY_SIZE
    if color is not None:
        r.font.color.rgb = color
    return p


def configure_styles(doc):
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = BASE_FONT
    normal.font.size = BODY_SIZE
    pf = normal.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT

    h1 = styles['Heading 1']
    h1.font.name = BASE_FONT
    h1.font.size = HEADING1_SIZE
    h1.font.bold = True
    h1.font.color.rgb = DARK_TEXT
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(4)
    h1.paragraph_format.keep_with_next = True


def set_page_layout(doc):
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(
            SP_BRAND + "  •  " + SP_CONTACT + "  •  " + SP_WEBSITE + "    |    Page ")
        run.font.name = BASE_FONT
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        _simple_field(p, "PAGE")
        mid = p.add_run(" of ")
        mid.font.name = BASE_FONT
        mid.font.size = Pt(9)
        mid.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        _simple_field(p, "NUMPAGES")


def add_h1(doc, text):
    return doc.add_paragraph(text, style='Heading 1')


def add_para(doc, text, bold=False, italic=False, size=None,
             align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = BASE_FONT
    r.font.size = size or BODY_SIZE
    return p


# -------------------------------------------------------------------
# Section builders
# -------------------------------------------------------------------
def add_masthead(doc):
    """Two-column masthead: logo + brand on the left, big INVOICE on the right."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    _set_col_widths(table, [Cm(10.0), Cm(6.2)])
    left, right = table.rows[0].cells

    # Left: logo + brand block
    left.text = ""
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    try:
        run = p.add_run()
        run.add_picture(LOGO_PATH, width=Inches(1.15))
    except Exception:
        pass

    p2 = left.add_paragraph()
    r2 = p2.add_run("AnantaSutra — अनन्तसूत्र")
    r2.bold = True
    r2.font.name = BASE_FONT
    r2.font.size = Pt(13)

    p3 = left.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)
    r3 = p3.add_run(
        "Proprietor: " + SP_OWNER + " (" + SP_DESIGNATION + ")\n"
        + SP_PLACE + "\n"
        + SP_CONTACT + "  •  " + SP_WEBSITE)
    r3.font.name = BASE_FONT
    r3.font.size = Pt(9.5)
    r3.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    # Right: INVOICE wordmark + sub-label
    right.text = ""
    pr = right.paragraphs[0]
    pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = pr.add_run("INVOICE")
    rr.bold = True
    rr.font.name = BASE_FONT
    rr.font.size = TITLE_SIZE

    pr2 = right.add_paragraph()
    pr2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr2 = pr2.add_run("Bill of Supply")
    rr2.italic = True
    rr2.font.name = BASE_FONT
    rr2.font.size = Pt(10.5)
    rr2.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    pr3 = right.add_paragraph()
    pr3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr3 = pr3.add_run("(Not a Tax Invoice — see note below)")
    rr3.italic = True
    rr3.font.name = BASE_FONT
    rr3.font.size = Pt(8.5)
    rr3.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # thin rule below masthead
    doc.add_paragraph("")


def add_meta_and_billto(doc):
    """Two side-by-side blocks: Bill To (left) and Invoice details (right)."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    _set_col_widths(table, [Cm(8.1), Cm(8.1)])
    left, right = table.rows[0].cells
    for c in (left, right):
        set_cell_border(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # --- Bill To ---
    left.text = ""
    _cell_text(left, "BILL TO", bold=True, size=Pt(10),
               color=RGBColor(0x59, 0x59, 0x59))
    for label in [
        "Name / Business: ______________________________",
        "Contact person: ________________________________",
        "Address: _______________________________________",
        "____________________________________________",
        "Email / Phone: _________________________________",
        "Client GSTIN (if any): __________________________",
    ]:
        _cell_text(left, label, size=Pt(10.5))

    # --- Invoice details ---
    right.text = ""
    _cell_text(right, "INVOICE DETAILS", bold=True, size=Pt(10),
               color=RGBColor(0x59, 0x59, 0x59))
    for label in [
        "Invoice No.: ANS/____/20__-__ /____",
        "Invoice Date: ____ / ____ / 20____",
        "Due Date: ____ / ____ / 20____",
        "Place of Supply: Delhi, India",
        "Reference / PO: _________________________",
        "Billing Period: _________________________",
    ]:
        _cell_text(right, label, size=Pt(10.5))

    doc.add_paragraph("")


def add_line_items(doc, blank_rows=6):
    """Main line-items table with header, blank rows, and totals."""
    cols = ("#", "Description of Goods / Services", "Qty", "Rate (Rs.)", "Amount (Rs.)")
    table = doc.add_table(rows=1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_col_widths(
        table,
        [Cm(0.8), Cm(7.9), Cm(2.0), Cm(2.6), Cm(2.9)])

    # Header row
    hdr = table.rows[0].cells
    aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
              WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT,
              WD_ALIGN_PARAGRAPH.RIGHT]
    for idx, text in enumerate(cols):
        _cell_text(hdr[idx], text, bold=True, align=aligns[idx], size=Pt(10.5))
        _shade_cell(hdr[idx], LIGHT_FILL)
        set_cell_border(hdr[idx])

    # Blank line-item rows
    for n in range(1, blank_rows + 1):
        row = table.add_row().cells
        _cell_text(row[0], str(n), align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(row[1], "")
        _cell_text(row[2], "", align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(row[3], "", align=WD_ALIGN_PARAGRAPH.RIGHT)
        _cell_text(row[4], "", align=WD_ALIGN_PARAGRAPH.RIGHT)
        for c in row:
            set_cell_border(c)

    # Totals rows: label spans first 4 columns, value in last column
    def total_row(label, value="", bold=False, shade=False):
        row = table.add_row().cells
        merged = row[0].merge(row[1]).merge(row[2]).merge(row[3])
        _cell_text(merged, label, bold=bold, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _cell_text(row[4], value, bold=bold, align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_border(merged)
        set_cell_border(row[4])
        if shade:
            _shade_cell(merged, LIGHT_FILL)
            _shade_cell(row[4], LIGHT_FILL)

    total_row("Subtotal", "")
    total_row("Discount (if any)", "( )")
    total_row("GST (not registered)", "Nil")
    total_row("TOTAL AMOUNT PAYABLE (Rs.)", "", bold=True, shade=True)

    # Amount in words
    doc.add_paragraph("")
    add_para(doc,
             "Amount in words: Rupees ______________________________________ only.",
             italic=True, size=Pt(10.5))


def add_payment_and_notes(doc):
    # Payment details block
    add_h1(doc, "Payment Details")
    pay_table = doc.add_table(rows=1, cols=2)
    pay_table.autofit = False
    _set_col_widths(pay_table, [Cm(8.1), Cm(8.1)])
    left, right = pay_table.rows[0].cells
    for c in (left, right):
        set_cell_border(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    left.text = ""
    _cell_text(left, "Bank Transfer (NEFT / IMPS / RTGS)", bold=True, size=Pt(10.5))
    for label in [
        "Account Name: ______________________________",
        "Account No.: _________________________________",
        "Bank / Branch: ______________________________",
        "IFSC: ________________________________________",
    ]:
        _cell_text(left, label, size=Pt(10.5))

    right.text = ""
    _cell_text(right, "UPI", bold=True, size=Pt(10.5))
    for label in [
        "UPI ID: ______________________________________",
        "UPI Name: ___________________________________",
        "",
        "Kindly quote the Invoice No. in the payment reference.",
    ]:
        _cell_text(right, label, size=Pt(10.5))

    doc.add_paragraph("")

    # Notes / terms
    add_h1(doc, "Notes & Terms")
    notes = [
        "AnantaSutra is a small, unregistered business concern (sole "
        "proprietorship of " + SP_OWNER + ") and is not registered under the "
        "Companies Act or for Goods and Services Tax (GST). Accordingly, this "
        "document is a Bill of Supply / plain invoice, not a Tax Invoice, and "
        "NO GST is charged on the amounts above. There is no GSTIN or CIN.",
        "No tax is added or deducted by AnantaSutra. If the client is "
        "required by law to deduct tax at source (TDS) on this payment, the "
        "client shall deduct it as per applicable rates and furnish the TDS "
        "certificate within the statutory time; the net amount is otherwise "
        "payable in full without any other set-off or deduction.",
        "Payment is due by the due date shown above, by bank transfer or UPI "
        "to the details above. Please quote the Invoice No. in the payment "
        "reference.",
        "Any third-party or out-of-pocket costs (e.g. paid stock images, "
        "fonts, plugins, hosting, domain, or advertising spend) are billed at "
        "actuals and are shown as separate line items where applicable.",
        "All amounts are in Indian Rupees (INR / Rs.). This invoice is issued "
        "at " + SP_PLACE + ". Any dispute regarding this invoice must be "
        "raised in writing within seven (7) days of the invoice date.",
    ]
    for note in notes:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.7)
        for r0 in p.runs:
            r0.text = ""
        r = p.add_run(note)
        r.font.name = BASE_FONT
        r.font.size = Pt(10.5)

    # Signature / issued-by
    doc.add_paragraph("")
    add_para(doc, "For AnantaSutra", bold=True)
    doc.add_paragraph("")
    add_para(doc, "_______________________________")
    add_para(doc, SP_OWNER, size=Pt(10.5))
    add_para(doc, SP_DESIGNATION + ", AnantaSutra", size=Pt(10.5))
    add_para(doc, SP_CONTACT, size=Pt(10.5))
    add_para(doc, "Date: ______________________", size=Pt(10.5))

    doc.add_paragraph("")
    add_para(doc,
             "This is a computer-generated invoice and is valid without a "
             "physical signature.",
             italic=True, size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Thank you for your business.",
             italic=True, size=Pt(10),
             align=WD_ALIGN_PARAGRAPH.CENTER)


# -------------------------------------------------------------------
# Document construction
# -------------------------------------------------------------------
def build_document():
    doc = Document()
    configure_styles(doc)
    set_page_layout(doc)
    add_footer(doc)

    add_masthead(doc)
    add_meta_and_billto(doc)
    add_line_items(doc, blank_rows=6)
    add_payment_and_notes(doc)

    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    out = build_document()
    size_kb = os.path.getsize(out) / 1024.0
    print("ANANTASUTRA INVOICE TEMPLATE COMPLETE")
    print("Output: " + out)
    print("Size: %.1f KB" % size_kb)
